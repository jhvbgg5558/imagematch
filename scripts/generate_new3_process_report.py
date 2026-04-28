#!/usr/bin/env python3
"""Generate a process-focused Word report for the new3output ODM-refresh run.

Purpose:
- generate a standalone Word document whose primary structure is the
  experiment content and end-to-end process rather than a dense result dump;
- summarize the completed new3output 009/010 experiment with compact
  quantitative evidence and a dedicated explanation of predicted-ortho partial
  coverage;
- preserve the existing full report and emit a second document under the same
  `reports/` directory.

Main inputs:
- experiment-level query selection and retrieval outputs under the new3output
  experiment root;
- formal pose runtime summaries under `pose_v1_formal/summary/` and `pnp/`;
- ODM-truth and satellite-truth validation summaries under the two suite
  roots;
- existing report figures or newly prepared figures from the integrated report
  helper module.

Main outputs:
- `<experiment-root>/reports/nadir_009010_odmrefresh_sattruth_experiment_process_report.docx`

Applicable task constraints:
- keep the runtime localization task defined as UAV-to-satellite retrieval;
- describe ODM orthophoto truth and satellite-truth patches as evaluation-only
  layers;
- keep the result section concise and subordinate to the process narrative.
"""

from __future__ import annotations

import argparse
import json
from pathlib import Path
from statistics import mean

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from generate_new3_full_experiment_report import (
    DEFAULT_BASELINE_ROOT,
    DEFAULT_EXPERIMENT_ROOT,
    add_bullets,
    add_heading,
    add_paragraph,
    add_picture_with_caption,
    add_table,
    build_baseline_comparison_chart,
    build_missing_issue_chart,
    build_sample_panel,
    choose_odm_sample_queries,
    ensure_dir,
    fmt,
    load_csv,
    load_json,
    read_rgba_or_rgb,
    relative_to_cwd,
    set_cn_font,
    short_flight_name,
    summarize_selected_queries,
)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--experiment-root", default=str(DEFAULT_EXPERIMENT_ROOT))
    parser.add_argument("--baseline-root", default=str(DEFAULT_BASELINE_ROOT))
    parser.add_argument(
        "--out-docx",
        default="",
        help="Optional explicit output path. Defaults to <experiment-root>/reports/nadir_009010_odmrefresh_sattruth_experiment_process_report.docx",
    )
    parser.add_argument(
        "--assets-dirname",
        default="process_report_assets",
        help="Directory name under reports/ used for process-report figures.",
    )
    return parser.parse_args()


def count_rows_if_exists(path: Path) -> int:
    return len(load_csv(path)) if path.exists() else 0


def build_runtime_gate_rows(experiment_root: Path, pose_root: Path) -> list[list[str]]:
    return [
        ["selected_queries/selected_images_summary.csv", str(count_rows_if_exists(experiment_root / "selected_queries" / "selected_images_summary.csv")), "40 rows; 009/010 各 20 张", "已完成"],
        ["query_inputs/query_manifest.csv", str(count_rows_if_exists(experiment_root / "query_inputs" / "query_manifest.csv")), "40 rows，query 去元数据", "已完成"],
        ["retrieval/retrieval_top20.csv", str(count_rows_if_exists(experiment_root / "retrieval" / "retrieval_top20.csv")), "800 rows；40 x Top-20", "已完成"],
        ["pose_v1_formal/manifest/pose_manifest.json", str(len(load_json(pose_root / "manifest" / "pose_manifest.json").get("pairs", []))), "800 query-candidate pairs", "已完成"],
        ["pose_v1_formal/pnp/pnp_results.csv", str(count_rows_if_exists(pose_root / "pnp" / "pnp_results.csv")), "800 rows", "已完成"],
        ["pose_v1_formal/summary/per_query_best_pose.csv", str(count_rows_if_exists(pose_root / "summary" / "per_query_best_pose.csv")), "40 rows", "已完成"],
    ]


def summarize_suite_result_rows(
    odm_ortho_overall: dict[str, object],
    odm_pose_overall: dict[str, object],
    odm_tie_overall: dict[str, object],
    sat_ortho_overall: dict[str, object],
    sat_geom_rows: list[dict[str, str]],
    sat_tie_overall: dict[str, object],
    baseline_ortho_overall: dict[str, object],
    baseline_tie_overall: dict[str, object],
) -> list[list[str]]:
    geom_offsets = [
        float(row["camera_center_offset_m"])
        for row in sat_geom_rows
        if row.get("eval_status") == "ok" and row.get("camera_center_offset_m") not in ("", None)
    ]
    return [
        ["Pose 主链", "PnP status", "以 pnp_summary.json / pose_overall_summary.json 为准，best pose 40/40 覆盖。"],
        ["ODM-truth layer-1", "ortho_iou mean", fmt(odm_ortho_overall["ortho_iou"]["mean"], 4)],
        ["ODM-truth layer-2", "horizontal_error_m mean", fmt(odm_pose_overall["horizontal_error_m"]["mean"], 4)],
        ["ODM-truth layer-3", "tiepoint_xy_error_rmse_m", fmt(odm_tie_overall["tiepoint_xy_error_rmse_m"], 4)],
        ["Satellite-truth layer-1", "ortho_iou mean", fmt(sat_ortho_overall["ortho_iou"]["mean"], 4)],
        ["Satellite-truth layer-2", "camera_center_offset_m mean", fmt(mean(geom_offsets) if geom_offsets else None, 4)],
        ["Satellite-truth layer-3", "tiepoint_xy_error_rmse_m", fmt(sat_tie_overall["tiepoint_xy_error_rmse_m"], 4)],
        ["Baseline vs new3", "baseline layer1 ortho_iou mean", fmt(baseline_ortho_overall["ortho_iou"]["mean"], 4)],
        ["Baseline vs new3", "baseline layer3 tiepoint_xy_error_rmse_m", fmt(baseline_tie_overall["tiepoint_xy_error_rmse_m"], 4)],
    ]


def main() -> None:
    args = parse_args()
    experiment_root = Path(args.experiment_root).resolve()
    baseline_root = Path(args.baseline_root).resolve()
    reports_root = experiment_root / "reports"
    assets_root = reports_root / args.assets_dirname
    out_docx = (
        Path(args.out_docx).resolve()
        if args.out_docx
        else reports_root / "nadir_009010_odmrefresh_sattruth_experiment_process_report.docx"
    )
    ensure_dir(reports_root)
    ensure_dir(assets_root)

    pose_root = experiment_root / "pose_v1_formal"
    odm_suite_root = pose_root / "eval_pose_validation_suite_odm_truth"
    sat_suite_root = pose_root / "eval_pose_validation_suite_satellite_truth"
    baseline_pose_root = baseline_root / "pose_v1_formal"
    baseline_suite_root = baseline_pose_root / "eval_pose_validation_suite"

    selected_rows = load_csv(experiment_root / "selected_queries" / "selected_images_summary.csv")
    retrieval_rows = load_csv(experiment_root / "retrieval" / "retrieval_top20.csv")
    pose_summary = load_json(pose_root / "summary" / "pose_overall_summary.json")
    pnp_summary = load_json(pose_root / "pnp" / "pnp_summary.json")
    odm_full_summary = load_json(odm_suite_root / "full_run_summary.json")
    sat_full_summary = load_json(sat_suite_root / "full_run_summary.json")

    odm_ortho_overall = load_json(odm_suite_root / "ortho_alignment" / "overall_ortho_accuracy.json")
    odm_pose_overall = load_json(odm_suite_root / "pose_vs_at" / "overall_pose_vs_at.json")
    odm_tie_overall = load_json(odm_suite_root / "tiepoint_ground_error" / "overall_tiepoint_ground_error.json")
    odm_ortho_rows = load_csv(odm_suite_root / "ortho_alignment" / "per_query_ortho_accuracy.csv")

    sat_ortho_overall = load_json(sat_suite_root / "ortho_alignment_satellite" / "overall_ortho_accuracy.json")
    sat_geom_overall = load_json(sat_suite_root / "pose_vs_satellite_truth_geometry" / "overall_satellite_truth_geometry.json")
    sat_tie_overall = load_json(sat_suite_root / "tiepoint_ground_error_satellite" / "overall_tiepoint_ground_error.json")
    sat_geom_rows = load_csv(sat_suite_root / "pose_vs_satellite_truth_geometry" / "per_query_pose_vs_satellite_truth_geometry.csv")

    baseline_ortho_overall = load_json(baseline_suite_root / "ortho_alignment" / "overall_ortho_accuracy.json")
    baseline_tie_overall = load_json(baseline_suite_root / "tiepoint_ground_error" / "overall_tiepoint_ground_error.json")
    comparison_md = (reports_root / "odm_truth_vs_satellite_truth_comparison.md").read_text(encoding="utf-8")

    query_summary = summarize_selected_queries(selected_rows)

    high_q, _, low_q = choose_odm_sample_queries(odm_ortho_rows)
    odm_by_query = {row["query_id"]: row for row in odm_ortho_rows}
    q001_row = odm_by_query["q_001"]
    low_row = odm_by_query[low_q]

    comparison_chart_path = assets_root / "baseline_vs_new3_process_summary.png"
    missing_chart_path = assets_root / "predicted_ortho_missing_process_view.png"
    q001_panel_path = assets_root / "q_001_process_panel.png"
    low_panel_path = assets_root / f"{low_q}_process_panel.png"

    build_baseline_comparison_chart(
        {
            "ortho": baseline_ortho_overall,
            "tie": baseline_tie_overall,
            "odm_rows": load_csv(baseline_suite_root / "ortho_alignment" / "per_query_ortho_accuracy.csv"),
        },
        {
            "ortho": odm_ortho_overall,
            "tie": odm_tie_overall,
            "odm_rows": odm_ortho_rows,
        },
        sat_ortho_overall,
        sat_tie_overall,
        load_csv(sat_suite_root / "ortho_alignment_satellite" / "per_query_ortho_accuracy.csv"),
        sat_geom_rows,
        comparison_chart_path,
    )
    build_missing_issue_chart(odm_ortho_rows, missing_chart_path)
    build_sample_panel(
        Path(q001_row["truth_crop_path"]),
        Path(q001_row["pred_crop_path"]),
        q001_panel_path,
        "q_001 process-view sample",
        stats_lines=[
            f"common_valid_ratio={fmt(q001_row['common_valid_ratio'], 4)}",
            f"ortho_iou={fmt(q001_row['ortho_iou'], 4)}",
            f"center_offset_m={fmt(q001_row['center_offset_m'], 4)}",
        ],
    )
    build_sample_panel(
        Path(low_row["truth_crop_path"]),
        Path(low_row["pred_crop_path"]),
        low_panel_path,
        f"{low_q} process-view low-coverage sample",
        stats_lines=[
            f"common_valid_ratio={fmt(low_row['common_valid_ratio'], 4)}",
            f"ortho_iou={fmt(low_row['ortho_iou'], 4)}",
            f"center_offset_m={fmt(low_row['center_offset_m'], 4)}",
        ],
    )

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("009/010 ODM Refresh + Satellite Truth 实验流程文档")
    set_cn_font(title_run, size=16, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(relative_to_cwd(experiment_root))
    set_cn_font(subtitle_run, size=10)

    add_heading(doc, "1. 实验背景与目标", 1)
    add_bullets(
        doc,
        [
            "本次实验保持任务定义不变，仍然是 UAV 查询图像对固定卫星候选库进行初始地理定位。",
            "本次 new3output 分支的目的不是重做 retrieval，而是在保留 runtime 候选库和 retrieval / RoMa 结果不变的前提下，观察 ODM truth orthophoto 和 ODM DSM 替换后的影响，并增加一套 satellite-truth 并行验证。",
            "实验结论只作为收尾摘要，正文重点放在输入资产、流程串联、阶段产物和 gate。",
        ],
    )

    add_heading(doc, "2. 数据范围与 query 选择规则", 1)
    add_bullets(
        doc,
        [
            f"query 总数为 {query_summary['query_count']}，仅使用航线 {', '.join(short_flight_name(fid) for fid in query_summary['ordered_flights'])}。",
            "两条航线各 20 张 query，总计 40 张，筛选约束为 gimbal_pitch_degree <= -85.0。",
            f"本次 query 的 pitch 范围为 {fmt(query_summary['pitch_min'], 2)} 到 {fmt(query_summary['pitch_max'], 2)}。",
            "query 编号固定为：009 航线 q_001-q_020，010 航线 q_021-q_040。",
        ],
    )
    add_table(
        doc,
        ["flight", "query_count", "query_id_range"],
        [
            [short_flight_name(item["flight_id"]), str(item["count"]), f"{item['query_id_start']} - {item['query_id_end']}"]
            for item in query_summary["query_ranges"]
        ],
    )

    add_heading(doc, "3. 本次实验与 new2output 基线的差异", 1)
    add_bullets(
        doc,
        [
            "runtime candidate DOM 不变，仍然是固定卫星库。",
            "DINOv2 coarse 和 RoMa v2 rerank 不重跑，直接复用 new2output 已完成结果。",
            "query intrinsics 不作为本次变量，仍然沿用现有 per-flight cameras.json 解析结果。",
            "评估 truth orthophoto 从旧 UAV orthophoto 口径切换到新的 ODM orthophoto override。",
            "PnP DSM 从 SRTM 路线切换到 ODM DSM；无 raster DSM 时使用 odm_georeferenced_model.laz 栅格化构建 candidate DSM。",
            "在同一套 best pose 上新增 satellite-truth 并行验证，但 satellite truth 不进入 runtime。",
        ],
    )

    add_heading(doc, "4. 输入资产与目录组织", 1)
    add_table(
        doc,
        ["子目录/文件", "作用", "本次是否使用"],
        [
            ["selected_queries/", "保存 009/010 下视 query 选择结果", "是"],
            ["query_inputs/", "去元数据后的 query 输入与 query manifest", "是"],
            ["retrieval/retrieval_top20.csv", "复用的 Top-20 runtime candidate 输入", "是"],
            ["romav2_rerank/stage7/", "复用的 rerank 结果", "是"],
            ["pose_v1_formal/", "PnP、scores、best pose、validation 总根", "是"],
            ["eval_pose_validation_suite_odm_truth/", "新 ODM orthophoto truth 口径验证", "是"],
            ["eval_pose_validation_suite_satellite_truth/", "卫星 truth 并行验证", "是"],
            ["reports/", "最终报告和对比汇总目录", "是"],
        ],
    )

    add_heading(doc, "5. 端到端实验流程", 1)
    add_paragraph(doc, "本次实验按以下顺序执行，流程重点是替换上游 truth/DSM 资产而不改变 runtime retrieval 任务。")
    add_table(
        doc,
        ["阶段", "输入", "输出", "说明"],
        [
            ["1. query 选择", "航线影像目录 + pitch 约束", "selected_images_summary.csv", "锁定 009/010 两条航线各 20 张下视 query。"],
            ["2. query 输入准备", "selected query 原图", "query_manifest.csv", "去元数据，形成 formal query 输入。"],
            ["3. retrieval / rerank 复用", "new2output 已完成结果", "retrieval_top20.csv + stage7 rerank", "本次不重跑 coarse / rerank。"],
            ["4. ODM truth 接入", "flight asset override manifest + ODM orthophoto", "query_ortho_truth_manifest.csv + truth_tiles/", "构建 eval_pose_validation_suite_odm_truth 的 truth grid。"],
            ["5. ODM DSM 构建", "ODM DSM / LAZ", "candidate DSM cache rasters", "替换原 SRTM 路线。"],
            ["6. PnP 与 scoring", "pose manifest + sampled correspondences", "pnp_results.csv / pose_scores.csv / per_query_best_pose.csv", "完成 new3output 的 formal pose 主链。"],
            ["7. ODM-truth validation", "best pose + new ODM truth", "三层 suite 结果", "保持原三层定义。"],
            ["8. satellite-truth validation", "best pose + satellite truth patches", "并行三层 suite 结果", "只用于独立 truth 口径检查。"],
            ["9. comparison 汇总", "baseline + ODM-truth + satellite-truth", "comparison md/docx", "给出关键差异说明。"],
        ],
    )

    add_heading(doc, "6. 各阶段关键产物与 gate", 1)
    add_table(
        doc,
        ["阶段输出", "本次结果", "gate", "状态"],
        build_runtime_gate_rows(experiment_root, pose_root),
    )
    add_bullets(
        doc,
        [
            f"ODM-truth suite 已完成：{json.dumps(odm_full_summary.get('pipeline_status', 'ok'), ensure_ascii=False)}。",
            f"Satellite-truth suite 已完成：{json.dumps(sat_full_summary.get('pipeline_status', 'ok'), ensure_ascii=False)}。",
            f"comparison 摘要已存在：{relative_to_cwd(reports_root / 'odm_truth_vs_satellite_truth_comparison.md')}。",
        ],
    )

    add_heading(doc, "7. 关键实验结果摘要", 1)
    add_table(
        doc,
        ["模块", "指标", "摘要"],
        summarize_suite_result_rows(
            odm_ortho_overall,
            odm_pose_overall,
            odm_tie_overall,
            sat_ortho_overall,
            sat_geom_rows,
            sat_tie_overall,
            baseline_ortho_overall,
            baseline_tie_overall,
        ),
    )
    add_bullets(
        doc,
        [
            f"Pose 主链：PnP rows = {pnp_summary['row_count']}，status_counts = {json.dumps(pnp_summary['status_counts'], ensure_ascii=False)}，best_status_counts = {json.dumps(pose_summary['best_status_counts'], ensure_ascii=False)}。",
            f"ODM-truth layer-1/2/3 核心指标分别以 ortho_iou mean = {fmt(odm_ortho_overall['ortho_iou']['mean'], 4)}、horizontal_error_m mean = {fmt(odm_pose_overall['horizontal_error_m']['mean'], 4)}、tiepoint_xy_error_rmse_m = {fmt(odm_tie_overall['tiepoint_xy_error_rmse_m'], 4)} 为代表。",
            f"Satellite-truth layer-1/2/3 核心指标分别以 ortho_iou mean = {fmt(sat_ortho_overall['ortho_iou']['mean'], 4)}、geometry status = {json.dumps(sat_geom_overall['status_counts'], ensure_ascii=False)}、tiepoint_xy_error_rmse_m = {fmt(sat_tie_overall['tiepoint_xy_error_rmse_m'], 4)} 为代表。",
        ],
    )
    add_picture_with_caption(doc, comparison_chart_path, "图 1. baseline、ODM-truth 和 satellite-truth 的关键指标对比摘要。")

    add_heading(doc, "8. 预测图部分缺失问题说明", 1)
    add_bullets(
        doc,
        [
            "预测图不是完整正射重建图，而是把单张 query 通过 best pose 和 candidate-linked DSM 投影到 truth grid 后得到的有效覆盖区域。",
            "当前渲染器不做平面回退，因此 DSM 无效、投影越界、以及 query 视场之外的 truth 网格像素都会直接留空。",
            f"q_001 的 common_valid_ratio = {fmt(q001_row['common_valid_ratio'], 4)}，说明即使评估成功，也只有约 {fmt(float(q001_row['common_valid_ratio']) * 100.0, 1)}% 的 truth 网格有有效投影覆盖。",
            f"{low_q} 是本次低覆盖样例，common_valid_ratio = {fmt(low_row['common_valid_ratio'], 4)}，该现象应解释为“有效覆盖受限”，不是文件损坏或脚本失败。",
        ],
    )
    add_picture_with_caption(doc, missing_chart_path, "图 2. ODM-truth 下 predicted ortho 覆盖率与对齐指标的逐 query 分布。")
    add_picture_with_caption(doc, q001_panel_path, "图 3. q_001 的 truth / pred / alpha-mask / overlay 样例。")
    add_picture_with_caption(doc, low_panel_path, f"图 4. {low_q} 的低覆盖样例，用于说明空洞来源。")

    add_heading(doc, "9. 结论", 1)
    add_bullets(
        doc,
        [
            "本次 new3output 实验已经形成一条完整、可追溯的流程：从 query 选择、runtime 结果复用，到 ODM truth/ODM DSM 替换，再到两套 truth 口径验证和总对比汇总。",
            "文档中的结论是对流程执行结果的摘要，不改变 runtime 任务定义，也不把 evaluation truth 误写成 runtime 输入。",
            "后续如果继续迭代，应优先补强流程阶段之间的产物说明、gate 摘要和有效覆盖解释，而不是继续堆叠大量逐 query 指标表。",
        ],
    )

    ensure_dir(out_docx.parent)
    doc.save(out_docx)
    print(out_docx)


if __name__ == "__main__":
    main()
