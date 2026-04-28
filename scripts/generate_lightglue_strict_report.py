#!/usr/bin/env python3
"""Generate Word and Markdown report for LightGlue strict rerank results."""

from __future__ import annotations

import argparse
import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--baseline-result-dir", required=True)
    parser.add_argument("--lightglue-result-dir", required=True)
    parser.add_argument("--out-docx", required=True)
    parser.add_argument("--out-md", required=True)
    return parser.parse_args()


def load_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", newline="", encoding="utf-8-sig") as f:
        return list(csv.DictReader(f))


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def set_cn_font(run, size: int = 11, bold: bool = False) -> None:
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(size)
    run.bold = bold


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 10, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    set_cn_font(r, size=size, bold=bold)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.style = "Heading 1" if level == 1 else "Heading 2"
    r = p.add_run(text)
    set_cn_font(r, size=14 if level == 1 else 12, bold=True)


def add_paragraph(doc: Document, text: str, size: int = 11, center: bool = False) -> None:
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn_font(r, size=size)


def add_bullets(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(line)
        set_cn_font(r, size=11)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn_font(r, size=10)


def add_picture(doc: Document, path: Path, width_inch: float) -> None:
    doc.add_picture(str(path), width=Inches(width_inch))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def short_flight_name(flight_id: str) -> str:
    parts = flight_id.split("_")
    return parts[2] if len(parts) >= 3 else flight_id


def build_metric_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(["指标", "含义"]):
        set_cell_text(table.cell(0, i), h, bold=True)
        shade_cell(table.cell(0, i), "D9EAF7")
    rows = [
        ("Strict Recall@1", "正式主指标，衡量首位候选是否直接命中 strict truth。"),
        ("Strict Recall@5", "正式主指标，衡量前 5 名候选覆盖能力。"),
        ("Strict Recall@10", "正式主指标，衡量区域级初步定位能力。"),
        ("Strict Recall@20", "辅助诊断指标，用于观察 Top-20 coarse candidate 上限，不作为主结论。"),
        ("Strict MRR", "strict truth 排名倒数均值，衡量前排排序质量。"),
        ("Top-1 error mean (m)", "首位候选中心与 query 位置间的平均距离。"),
    ]
    for name, desc in rows:
        row = table.add_row().cells
        set_cell_text(row[0], name)
        row[1].text = ""
        p = row[1].paragraphs[0]
        r = p.add_run(desc)
        set_cn_font(r, size=10)


def build_overall_table(doc: Document, overall: dict) -> None:
    table = doc.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["方法", "R@1", "R@5", "R@10", "R@20", "MRR", "Top-1误差均值(m)"]
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, bold=True)
        shade_cell(table.cell(0, i), "D9EAF7")
    rows = [
        ("DINOv2 Baseline", overall["baseline_strict_recall@1"], overall["baseline_strict_recall@5"], overall["baseline_strict_recall@10"], "-", overall["baseline_strict_mrr"], overall["baseline_top1_error_m_mean"]),
        ("Coarse Top20", "-", "-", "-", overall["coarse_strict_recall@20"], "-", "-"),
        ("SuperPoint + LightGlue", overall["lightglue_strict_recall@1"], overall["lightglue_strict_recall@5"], overall["lightglue_strict_recall@10"], overall["lightglue_strict_recall@20"], overall["lightglue_strict_mrr"], overall["lightglue_top1_error_m_mean"]),
    ]
    for vals in rows:
        row = table.add_row().cells
        for i, val in enumerate(vals):
            if isinstance(val, float):
                set_cell_text(row[i], f"{val:.3f}")
            else:
                set_cell_text(row[i], str(val))


def build_per_flight_table(doc: Document, rows: list[dict[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["航线", "Baseline R@10", "LightGlue R@10", "Delta", "LightGlue R@20", "Baseline MRR", "LightGlue MRR"]
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, bold=True)
        shade_cell(table.cell(0, i), "D9EAF7")
    for item in rows:
        row = table.add_row().cells
        baseline_r10 = float(item["baseline_strict_recall@10"])
        lightglue_r10 = float(item["lightglue_strict_recall@10"])
        set_cell_text(row[0], item["flight_tag"])
        set_cell_text(row[1], f"{baseline_r10:.3f}")
        set_cell_text(row[2], f"{lightglue_r10:.3f}")
        set_cell_text(row[3], f"{lightglue_r10 - baseline_r10:+.3f}")
        set_cell_text(row[4], f"{float(item['lightglue_strict_recall@20']):.3f}")
        set_cell_text(row[5], f"{float(item['baseline_strict_mrr']):.3f}")
        set_cell_text(row[6], f"{float(item['lightglue_strict_mrr']):.3f}")


def choose_cases(comp_rows: list[dict[str, str]]) -> list[dict[str, str]]:
    promote_cases = [r for r in comp_rows if int(r["promoted_11_20_to_top10"]) == 1]
    promote_cases.sort(key=lambda r: r["query_id"])

    improve_cases = []
    for row in comp_rows:
        b = row["baseline_first_strict_truth_rank"]
        l = row["lightglue_first_strict_truth_rank"]
        if b != "" and l != "" and int(l) < int(b):
            improve_cases.append(row)
    improve_cases.sort(key=lambda r: int(r["baseline_first_strict_truth_rank"]) - int(r["lightglue_first_strict_truth_rank"]), reverse=True)

    fail_cases = [r for r in comp_rows if row_to_int(r["coarse_strict_hit@20"]) == 1 and row_to_int(r["lightglue_strict_hit@10"]) == 0]
    fail_cases.sort(key=lambda r: r["query_id"])

    selected = []
    if promote_cases:
        promote_cases[0]["title"] = "从 11..20 拉入 Top-10 的改进样例"
        selected.append(promote_cases[0])
    if improve_cases:
        candidate = next((r for r in improve_cases if r["query_id"] not in {x["query_id"] for x in selected}), improve_cases[0])
        candidate["title"] = "前排排名进一步提升样例"
        selected.append(candidate)
    if fail_cases:
        candidate = next((r for r in fail_cases if r["query_id"] not in {x["query_id"] for x in selected}), fail_cases[0])
        candidate["title"] = "Top-20 内仍未成功推入 Top-10 的样例"
        selected.append(candidate)
    return selected


def row_to_int(v: str) -> int:
    return int(v) if v != "" else 0


def add_case_section(doc: Document, case_idx: int, case: dict[str, str], baseline_dir: Path, lightglue_dir: Path) -> None:
    qid = case["query_id"]
    flight_id = case["flight_id"]
    add_heading(doc, f"6.{case_idx} {case['title']}：{short_flight_name(flight_id)} / {qid}", 2)
    add_paragraph(
        doc,
        f"Baseline 首个 strict truth 排名：{case['baseline_first_strict_truth_rank'] or '未命中'}；"
        f"Coarse Top20 首个 strict truth 排名：{case['coarse_first_strict_truth_rank'] or '未命中'}；"
        f"LightGlue 首个 strict truth 排名：{case['lightglue_first_strict_truth_rank'] or '未命中'}。",
    )
    add_picture(doc, baseline_dir / "figures" / flight_id / f"{qid}_top10.png", 5.8)
    add_caption(doc, f"图 {case_idx * 3 - 2}  Baseline 结果：{short_flight_name(flight_id)} / {qid}")
    add_picture(doc, lightglue_dir / "figures" / flight_id / f"{qid}_top10.png", 5.8)
    add_caption(doc, f"图 {case_idx * 3 - 1}  LightGlue 重排结果：{short_flight_name(flight_id)} / {qid}")
    add_picture(doc, lightglue_dir / "figures" / "_compare" / "baseline_vs_lightglue_compare.png", 5.8)
    add_caption(doc, f"图 {case_idx * 3}  Overall 对照图")


def main() -> None:
    args = parse_args()
    baseline_dir = Path(args.baseline_result_dir)
    lightglue_dir = Path(args.lightglue_result_dir)
    fig_dir = lightglue_dir / "figures"
    out_docx = Path(args.out_docx)
    out_md = Path(args.out_md)

    overall = load_json(lightglue_dir / "overall_summary.json")
    per_flight_rows = load_csv(lightglue_dir / "per_flight_comparison.csv")
    comp_rows = load_csv(lightglue_dir / "per_query_comparison.csv")
    cases = choose_cases(comp_rows)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("SuperPoint + LightGlue 在 Strict Truth 口径下的重排结果说明")
    set_cn_font(r, size=16, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("面向无人机单张图像初步地理定位的正式技术报告")
    set_cn_font(r, size=11)

    add_heading(doc, "1. 任务定义与实验设置", 1)
    add_paragraph(
        doc,
        "本组实验用于回答：在 `DINOv2 + FAISS` 已形成 strict truth 正式基线之后，若把粗检索候选窗口扩展到 `Top-20`，"
        "再在其上使用 `SuperPoint + LightGlue` 做局部匹配与融合重排，是否能够进一步提升正式主指标。"
    )
    add_bullets(
        doc,
        [
            "基线方法：DINOv2 pooler 全局特征 + FAISS 粗检索。",
            "重排方法：SuperPoint 局部特征 + LightGlue 匹配 + RANSAC 几何一致性 + 融合分数排序。",
            "执行窗口：Top-20 coarse candidates。",
            "正式真值口径：strict truth。",
        ],
    )

    add_heading(doc, "2. 指标定义", 1)
    build_metric_table(doc)

    add_heading(doc, "3. 方法说明", 1)
    add_paragraph(
        doc,
        "本轮并不改变 DINOv2 粗检索本身，只在其 `Top-20` 候选集合上增加局部几何验证与融合排序。"
        "更准确地说，这是一轮“扩窗后的学习型局部重排”实验。"
        "`Recall@20` 在本文中只用于描述 coarse candidate 上限，不替代正式主指标 `Recall@1/5/10`。"
    )

    add_heading(doc, "4. 定量结果", 1)
    add_paragraph(doc, "表 1 给出 strict baseline、coarse Top-20 上限和 LightGlue 重排后的 overall 对照结果。")
    build_overall_table(doc, overall)
    add_caption(doc, "表 1  Baseline 与 LightGlue 的 overall strict 指标对照")
    add_paragraph(doc, "表 2 给出四条航线上的 `Recall@10`、`Recall@20` 和 `MRR` 对照结果。")
    build_per_flight_table(doc, per_flight_rows)
    add_caption(doc, "表 2  四条航线上的 strict 指标对照")

    add_heading(doc, "5. 汇总图解读", 1)
    add_paragraph(
        doc,
        "图 1 展示 LightGlue 单轮 overall 指标，图 2 展示分航线 Recall，图 3 展示 `Top-20` 扩窗上限，图 4 直接对照 baseline 与 LightGlue 的 overall 差异。"
    )
    figs = [
        (fig_dir / "_aggregate" / "overall_metrics_bar.png", "图 1  LightGlue 的 overall strict 指标"),
        (fig_dir / "_aggregate" / "multi_flight_recall.png", "图 2  LightGlue 的分航线 strict Recall"),
        (fig_dir / "_aggregate" / "top20_upper_bound.png", "图 3  Top-20 扩窗上限与最终 Top-10 结果"),
        (fig_dir / "_compare" / "baseline_vs_lightglue_compare.png", "图 4  Baseline vs LightGlue overall 对照"),
    ]
    for path, caption in figs:
        add_picture(doc, path, 6.4)
        add_caption(doc, caption)
    add_paragraph(
        doc,
        "结果表明，LightGlue 重排确实把 `Strict Recall@10` 从 `0.425` 提升到了 `0.475`，说明 `Top-20` 扩窗后的局部重排能够把一部分原本排在 `11..20` 的真值拉回前 `Top-10`。"
        "但与此同时，`Strict Recall@1` 从 `0.175` 降到 `0.150`，`Strict MRR` 也略有下降，说明前排排序质量并未同步改善。"
        "因此，更准确的结论表述是：LightGlue 提升了前 `Top-10` 候选覆盖能力，但尚未稳定提升首位候选质量。"
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "6. 代表性案例", 1)
    for idx, case in enumerate(cases, start=1):
        add_case_section(doc, idx, case, baseline_dir, lightglue_dir)

    add_heading(doc, "7. 结论", 1)
    add_bullets(
        doc,
        [
            "在 strict truth 口径下，把候选窗口从 Top-10 扩展到 Top-20 是必要的，因为 coarse retrieval 中确有一部分真值排在 11..20。",
            "SuperPoint + LightGlue 融合重排把 Strict Recall@10 从 0.425 提升到 0.475，说明它对前十名真值召回有帮助。",
            "但该方法没有明显改善 Recall@1 与 MRR，说明当前局部匹配与融合策略尚未稳定提升前排排序质量。",
            "因此，更准确的结论是：LightGlue 目前更适合作为“扩大前十名覆盖能力”的重排方法，而不是已经成熟的最优候选提纯方法。",
        ],
    )

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))

    md_lines = [
        "# SuperPoint + LightGlue 在 Strict Truth 口径下的重排结果说明",
        "",
        "## 1. 任务定义与实验设置",
        "",
        "本组实验用于回答：在 `DINOv2 + FAISS` 已形成 strict truth 正式基线之后，若把粗检索候选窗口扩展到 `Top-20`，再在其上使用 `SuperPoint + LightGlue` 做局部匹配与融合重排，是否能够进一步提升正式主指标。",
        "",
        "## 2. 指标定义",
        "",
        "- 正式主指标：`Strict Recall@1/5/10`、`Strict MRR`",
        "- 辅助上限指标：`Strict Recall@20`",
        "",
        "## 3. 方法说明",
        "",
        "本轮不改变 DINOv2 粗检索本身，只在其 `Top-20` 候选集合上增加局部几何验证与融合排序。",
        "",
        "## 4. 定量结果",
        "",
        f"- Baseline：`R@1={overall['baseline_strict_recall@1']:.3f}`，`R@5={overall['baseline_strict_recall@5']:.3f}`，`R@10={overall['baseline_strict_recall@10']:.3f}`，`MRR={overall['baseline_strict_mrr']:.3f}`",
        f"- Coarse Top20 上限：`R@20={overall['coarse_strict_recall@20']:.3f}`",
        f"- LightGlue：`R@1={overall['lightglue_strict_recall@1']:.3f}`，`R@5={overall['lightglue_strict_recall@5']:.3f}`，`R@10={overall['lightglue_strict_recall@10']:.3f}`，`R@20={overall['lightglue_strict_recall@20']:.3f}`，`MRR={overall['lightglue_strict_mrr']:.3f}`",
        "",
        "## 5. 汇总图解读",
        "",
        "LightGlue 重排确实把 `Strict Recall@10` 从 `0.425` 提升到了 `0.475`，说明 `Top-20` 扩窗后的局部重排能够把一部分原本排在 `11..20` 的真值拉回前 `Top-10`。另一方面，`Strict Recall@1` 与 `Strict MRR` 没有同步提升，说明当前重排策略更像是在改善前十名覆盖，而不是稳定提升首位候选质量。",
        "",
        "## 6. 代表性案例",
        "",
    ]
    for case in cases:
        md_lines.extend(
            [
                f"- `{case['query_id']}` / `{short_flight_name(case['flight_id'])}`：{case['title']}",
                f"  Baseline rank=`{case['baseline_first_strict_truth_rank'] or 'miss'}`，Coarse Top20 rank=`{case['coarse_first_strict_truth_rank'] or 'miss'}`，LightGlue rank=`{case['lightglue_first_strict_truth_rank'] or 'miss'}`",
            ]
        )
    md_lines.extend(
        [
            "",
            "## 7. 结论",
            "",
            "- `Top-20` 扩窗是必要的，因为 coarse retrieval 中有一部分真值稳定落在 `11..20`。",
            "- `SuperPoint + LightGlue` 能提升 `Top-10` 真值召回，但还不能稳定提升 `Top-1` 候选质量。",
            "- 更准确的结论是：它当前更适合作为“扩大前十名覆盖能力”的重排方法，而不是已经成熟的首位候选提纯方法。",
        ]
    )
    out_md.write_text("\n".join(md_lines), encoding="utf-8")
    print(out_docx)
    print(out_md)


if __name__ == "__main__":
    main()
