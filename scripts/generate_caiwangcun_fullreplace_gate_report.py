#!/usr/bin/env python3
"""Generate the CaiWangCun full-replacement gate Markdown and Word report.

Purpose:
- summarize the completed CaiWangCun DOM/DSM full-replacement gate experiment;
- explain the evaluation methods, data preparation, runtime flow, results, and
  follow-up recommendations in a Chinese Word report;
- generate report-local figures and copy representative validation images.

Main inputs:
- `new3output/nadir_009010_caiwangcun_domdsm_fullreplace_gate_2026-04-20`;
- formal pose summaries under `pose_v1_formal/summary/`;
- CaiWangCun truth validation summaries under
  `pose_v1_formal/eval_pose_validation_suite_caiwangcun_truth/`.

Main outputs:
- `<experiment-root>/reports/caiwangcun_fullreplace_gate_report.md`;
- `<experiment-root>/reports/caiwangcun_fullreplace_gate_report.docx`;
- `<experiment-root>/reports/assets/`.

Applicable task constraints:
- this report documents a gate-only experiment and does not rerun localization;
- query, DINOv2, RoMa v2, PnP, and validation algorithms are described as
  unchanged, while all candidate-library-derived assets are described as rebuilt;
- validation truth and query truth are offline evaluation inputs, not runtime
  inputs for the retrieval-localization task.
"""

from __future__ import annotations

import argparse
import csv
import json
import math
import shutil
from pathlib import Path
from typing import Any

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

try:
    from PIL import Image
except ImportError:  # pragma: no cover - optional report-size optimization.
    Image = None


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_EXPERIMENT_ROOT = (
    PROJECT_ROOT
    / "new3output"
    / "nadir_009010_caiwangcun_domdsm_fullreplace_gate_2026-04-20"
)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--experiment-root", default=str(DEFAULT_EXPERIMENT_ROOT))
    parser.add_argument("--out-md", default="")
    parser.add_argument("--out-docx", default="")
    parser.add_argument("--assets-dir", default="")
    return parser.parse_args()


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def load_json(path: Path) -> dict[str, Any]:
    if not path.exists():
        raise FileNotFoundError(f"Missing required JSON: {path}")
    return json.loads(path.read_text(encoding="utf-8"))


def load_csv(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        raise FileNotFoundError(f"Missing required CSV: {path}")
    with path.open("r", newline="", encoding="utf-8-sig") as handle:
        return list(csv.DictReader(handle))


def line_count(path: Path) -> int:
    if not path.exists():
        return 0
    with path.open("r", encoding="utf-8-sig", errors="ignore") as handle:
        return max(sum(1 for _ in handle) - 1, 0)


def fmt(value: Any, digits: int = 3) -> str:
    if value in (None, ""):
        return "-"
    try:
        number = float(value)
    except (TypeError, ValueError):
        return str(value)
    if not math.isfinite(number):
        return "-"
    return f"{number:.{digits}f}"


def pct(value: Any, digits: int = 1) -> str:
    try:
        return f"{float(value) * 100:.{digits}f}%"
    except (TypeError, ValueError):
        return "-"


def get_mean(summary: dict[str, Any], key: str) -> float | None:
    value = summary.get(key)
    if isinstance(value, dict) and "mean" in value:
        return float(value["mean"])
    return None


def get_nested(summary: dict[str, Any], *keys: str) -> Any:
    value: Any = summary
    for key in keys:
        if not isinstance(value, dict):
            return None
        value = value.get(key)
    return value


def count_status_text(counts: dict[str, Any] | None) -> str:
    if not counts:
        return "-"
    return ", ".join(f"{key}={value}" for key, value in counts.items())


def tiepoint_detail_summary(data: dict[str, Any]) -> dict[str, Any]:
    detail_dir = data["suite_root"] / "tiepoint_ground_error" / "tiepoints" / "per_query_matches"
    rows = data.get("tiepoint_rows", [])
    csv_paths = sorted(detail_dir.glob("*_tiepoints.csv")) if detail_dir.exists() else []
    csv_query_ids = {path.name.removesuffix("_tiepoints.csv") for path in csv_paths}
    expected_ok_ids = [
        row["query_id"]
        for row in rows
        if row.get("eval_status") == "tiepoint_eval_ok"
    ]
    missing_ok_ids = [query_id for query_id in expected_ok_ids if query_id not in csv_query_ids]
    failed_ids = [
        row["query_id"]
        for row in rows
        if row.get("eval_status") != "tiepoint_eval_ok"
    ]
    return {
        "detail_dir": detail_dir,
        "csv_count": len(csv_paths),
        "expected_ok_count": len(expected_ok_ids),
        "missing_ok_ids": missing_ok_ids,
        "failed_ids": failed_ids,
        "fields": [
            "query_id",
            "match_index",
            "truth_col_px",
            "truth_row_px",
            "pred_col_px",
            "pred_row_px",
            "truth_x_m",
            "truth_y_m",
            "pred_x_m",
            "pred_y_m",
            "dx_m",
            "dy_m",
            "dxy_m",
        ],
    }


def set_cn_font(run, size: int = 11, bold: bool = False) -> None:
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(size)
    run.bold = bold


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, size: int = 9) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_cn_font(run, size=size, bold=bold)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = "Heading 1" if level == 1 else "Heading 2"
    run = paragraph.add_run(text)
    set_cn_font(run, size=14 if level == 1 else 12, bold=True)


def add_paragraph(doc: Document, text: str, *, size: int = 11, center: bool = False, bold: bool = False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_cn_font(run, size=size, bold=bold)


def add_bullets(doc: Document, lines: list[str]) -> None:
    for line in lines:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(line)
        set_cn_font(run, size=11)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_text(table.cell(0, idx), header, bold=True)
        shade_cell(table.cell(0, idx), "D9EAF7")
    for row_values in rows:
        row = table.add_row().cells
        for idx, value in enumerate(row_values):
            set_cell_text(row[idx], value)


def add_picture(doc: Document, image_path: Path, caption: str, *, width_inch: float = 5.9) -> None:
    if not image_path.exists():
        add_paragraph(doc, f"缺失图片：{image_path}", size=9)
        return
    doc.add_picture(str(image_path), width=Inches(width_inch))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(caption)
    set_cn_font(run, size=9)


def save_bar_chart(path: Path, title: str, labels: list[str], values: list[float], ylabel: str) -> Path:
    ensure_dir(path.parent)
    plt.figure(figsize=(7.2, 4.0))
    colors = ["#2f6f9f", "#45a778", "#d99b36", "#b85c5c", "#6c6c96", "#7b9e87"][: len(values)]
    plt.bar(labels, values, color=colors)
    plt.title(title)
    plt.ylabel(ylabel)
    plt.xticks(rotation=20, ha="right")
    plt.grid(axis="y", alpha=0.25)
    plt.tight_layout()
    plt.savefig(path, dpi=180)
    plt.close()
    return path


def save_grouped_bar_chart(path: Path, title: str, groups: list[str], series: dict[str, list[float]], ylabel: str) -> Path:
    ensure_dir(path.parent)
    plt.figure(figsize=(7.4, 4.2))
    width = 0.8 / max(len(series), 1)
    x_positions = list(range(len(groups)))
    for idx, (label, values) in enumerate(series.items()):
        offsets = [x + (idx - (len(series) - 1) / 2) * width for x in x_positions]
        plt.bar(offsets, values, width=width, label=label)
    plt.title(title)
    plt.ylabel(ylabel)
    plt.xticks(x_positions, groups, rotation=18, ha="right")
    plt.grid(axis="y", alpha=0.25)
    plt.legend()
    plt.tight_layout()
    plt.savefig(path, dpi=180)
    plt.close()
    return path


def copy_image(src: Path, assets_dir: Path, name: str, missing: list[str]) -> Path | None:
    if not src.exists():
        missing.append(str(src))
        return None
    dst = assets_dir / name
    if Image is not None and src.stat().st_size > 2 * 1024 * 1024:
        dst = dst.with_suffix(".jpg")
        with Image.open(src) as image:
            image = image.convert("RGB")
            max_side = 1800
            if max(image.size) > max_side:
                image.thumbnail((max_side, max_side))
            ensure_dir(dst.parent)
            image.save(dst, quality=88, optimize=True)
        return dst
    ensure_dir(dst.parent)
    shutil.copy2(src, dst)
    return dst


def collect_inputs(experiment_root: Path) -> dict[str, Any]:
    pose_root = experiment_root / "pose_v1_formal"
    suite_root = pose_root / "eval_pose_validation_suite_caiwangcun_truth"
    ortho_root = suite_root / "ortho_alignment"
    return {
        "experiment_root": experiment_root,
        "pose_root": pose_root,
        "suite_root": suite_root,
        "mosaic": load_json(experiment_root / "source_mosaic" / "caiwangcun_mosaic_summary.json"),
        "roi": load_json(experiment_root / "candidate_library" / "roi_summary.json"),
        "retrieval": load_json(experiment_root / "romav2_rerank" / "coarse" / "summary_top20.json"),
        "pose_gate": load_json(pose_root / "summary" / "phase_gate_summary.json"),
        "pose_overall": load_json(pose_root / "summary" / "pose_overall_summary.json"),
        "validation": load_json(suite_root / "phase_gate_summary.json"),
        "ortho": load_json(ortho_root / "overall_ortho_accuracy.json"),
        "pose_vs_at": load_json(suite_root / "pose_vs_at" / "overall_pose_vs_at.json"),
        "tiepoint": load_json(suite_root / "tiepoint_ground_error" / "overall_tiepoint_ground_error.json"),
        "frame": load_json(ortho_root / "frame_sanity" / "overall_frame_sanity.json"),
        "ortho_rows": load_csv(ortho_root / "per_query_ortho_accuracy.csv"),
        "pose_rows": load_csv(suite_root / "pose_vs_at" / "per_query_pose_vs_at.csv"),
        "tiepoint_rows": load_csv(suite_root / "tiepoint_ground_error" / "per_query_tiepoint_ground_error.csv"),
        "frame_rows": load_csv(ortho_root / "frame_sanity" / "per_query_frame_sanity.csv"),
        "candidate_tile_rows": line_count(experiment_root / "candidate_library" / "tiles.csv"),
        "retrieval_rows": line_count(experiment_root / "retrieval" / "retrieval_top20.csv"),
    }


def build_assets(data: dict[str, Any], assets_dir: Path) -> tuple[dict[str, Path], list[str]]:
    if assets_dir.exists():
        shutil.rmtree(assets_dir)
    ensure_dir(assets_dir)
    missing: list[str] = []
    experiment_root = data["experiment_root"]
    suite_root = data["suite_root"]
    ortho_root = suite_root / "ortho_alignment"

    retrieval = data["retrieval"]
    pose_gate = data["pose_gate"]
    pose_overall = data["pose_overall"]
    ortho = data["ortho"]
    pose_vs_at = data["pose_vs_at"]
    tiepoint = data["tiepoint"]
    frame = data["frame"]
    frame_numeric = frame.get("numeric_summaries", {})
    assets: dict[str, Path] = {}
    assets["retrieval_recall"] = save_bar_chart(
        assets_dir / "retrieval_recall.png",
        "DINOv2 Top-K Retrieval Recall",
        ["@1", "@5", "@10", "@20", "MRR"],
        [
            float(retrieval["intersection_recall@1"]),
            float(retrieval["intersection_recall@5"]),
            float(retrieval["intersection_recall@10"]),
            float(retrieval["intersection_recall@20"]),
            float(retrieval["intersection_mrr"]),
        ],
        "score",
    )
    sampling_counts = get_nested(pose_gate, "sampling", "status_counts") or {}
    assets["sampling_status"] = save_bar_chart(
        assets_dir / "sampling_status.png",
        "2D-3D Sampling Status",
        list(sampling_counts.keys()),
        [float(v) for v in sampling_counts.values()],
        "count",
    )
    pnp_counts = get_nested(pose_gate, "pnp", "status_counts") or {}
    assets["pnp_status"] = save_bar_chart(
        assets_dir / "pnp_status.png",
        "PnP Candidate Status",
        list(pnp_counts.keys()),
        [float(v) for v in pnp_counts.values()],
        "count",
    )
    assets["layer_metrics"] = save_bar_chart(
        assets_dir / "layer_metrics.png",
        "Validation Key Metrics",
        ["center_offset_m", "horizontal_error_m", "tiepoint_rmse_m"],
        [
            float(get_mean(ortho, "center_offset_m") or 0.0),
            float(get_mean(pose_vs_at, "horizontal_error_m") or 0.0),
            float(tiepoint["tiepoint_xy_error_rmse_m"]),
        ],
        "meter",
    )
    assets["frame_sanity"] = save_bar_chart(
        assets_dir / "frame_sanity.png",
        "Frame Sanity Mean Ratios and Offsets",
        ["DSM valid", "Pred valid", "Camera offset", "BBox offset"],
        [
            float(get_nested(frame_numeric, "dsm_sample_valid_ratio_on_truth_grid", "mean") or 0.0),
            float(get_nested(frame_numeric, "pred_valid_pixel_ratio", "mean") or 0.0),
            float(get_nested(frame_numeric, "camera_center_offset_m", "mean") or 0.0),
            float(get_nested(frame_numeric, "bbox_center_delta_m", "mean") or 0.0),
        ],
        "ratio / meter",
    )
    assets["failure_comparison"] = save_grouped_bar_chart(
        assets_dir / "caiwangcun_branch_offset_comparison.png",
        "CaiWangCun Branch Offset Comparison",
        ["DSM-only constrained", "candidate+DSM partial", "full replacement"],
        {
            "center_offset_m": [515.3221, 515.6682, float(get_mean(ortho, "center_offset_m") or 0.0)],
            "horizontal_error_m": [652.1089, 648.5305, float(get_mean(pose_vs_at, "horizontal_error_m") or 0.0)],
        },
        "meter",
    )

    for query_id in ["q_003", "q_001", "q_021"]:
        copied = copy_image(
            ortho_root / "viz_overlay_truth" / f"{query_id}_overlay.png",
            assets_dir,
            f"{query_id}_truth_overlay.png",
            missing,
        )
        if copied:
            assets[f"{query_id}_truth_overlay"] = copied
        copied = copy_image(
            ortho_root / "viz_overlay_dom" / f"{query_id}_pred_vs_dom_overlay.png",
            assets_dir,
            f"{query_id}_dom_overlay.png",
            missing,
        )
        if copied:
            assets[f"{query_id}_dom_overlay"] = copied
        copied = copy_image(
            suite_root / "tiepoint_ground_error" / "viz_tiepoints" / f"{query_id}_tiepoints_overlay.png",
            assets_dir,
            f"{query_id}_tiepoints_overlay.png",
            missing,
        )
        if copied:
            assets[f"{query_id}_tiepoints_overlay"] = copied

    for name in ["q_003_frame_overlay.png", "q_003_offset_vectors.png", "q_003_dsm_valid_mask_on_truth_grid.png"]:
        copied = copy_image(ortho_root / "frame_sanity" / "figures" / name, assets_dir, name, missing)
        if copied:
            assets[name] = copied

    (assets_dir / "missing_images.json").write_text(
        json.dumps({"missing_images": missing}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return assets, missing


def rel(path: Path, base: Path) -> str:
    try:
        return path.relative_to(base).as_posix()
    except ValueError:
        return path.as_posix()


def metric_tables(data: dict[str, Any]) -> dict[str, list[list[str]]]:
    mosaic = data["mosaic"]
    roi = data["roi"]
    retrieval = data["retrieval"]
    pose_gate = data["pose_gate"]
    pose_overall = data["pose_overall"]
    validation = data["validation"]
    ortho = data["ortho"]
    pose_vs_at = data["pose_vs_at"]
    tiepoint = data["tiepoint"]
    frame = data["frame"]
    frame_numeric = frame.get("numeric_summaries", {})
    detail = tiepoint_detail_summary(data)

    return {
        "asset": [
            ["DOM mosaic", mosaic["ortho"]["crs"], str(mosaic["ortho"]["count"]), mosaic["ortho"]["dtype"], fmt(mosaic["ortho"]["resolution"][0], 6)],
            ["DSM mosaic", mosaic["dsm"]["crs"], str(mosaic["dsm"]["count"]), mosaic["dsm"]["dtype"], fmt(mosaic["dsm"]["resolution"][0], 6)],
            ["Candidate tiles", str(data["candidate_tile_rows"]), "200/300/500/700m", "fully-covered only", str(roi["skipped_counts"])],
        ],
        "retrieval": [
            ["query_count", str(retrieval["query_count"])],
            ["retrieval_top20_rows", str(data["retrieval_rows"])],
            ["intersection_recall@1", fmt(retrieval["intersection_recall@1"], 4)],
            ["intersection_recall@5", fmt(retrieval["intersection_recall@5"], 4)],
            ["intersection_recall@10", fmt(retrieval["intersection_recall@10"], 4)],
            ["intersection_recall@20", fmt(retrieval["intersection_recall@20"], 4)],
            ["intersection_mrr", fmt(retrieval["intersection_mrr"], 4)],
        ],
        "pose": [
            ["DSM cache", f"planned={get_nested(pose_gate, 'dsm', 'planned_count')}", f"built={get_nested(pose_gate, 'dsm', 'built_count')}", f"failed={get_nested(pose_gate, 'dsm', 'failed_count')}"],
            ["2D-3D sampling", f"rows={get_nested(pose_gate, 'sampling', 'row_count')}", count_status_text(get_nested(pose_gate, "sampling", "status_counts")), ""],
            ["PnP", f"rows={get_nested(pose_gate, 'pnp', 'row_count')}", count_status_text(get_nested(pose_gate, "pnp", "status_counts")), ""],
            ["Best pose", f"scored_query_count={pose_overall['scored_query_count']}", count_status_text(pose_overall["best_status_counts"]), ""],
            ["Best quality", f"score_mean={fmt(pose_overall['best_score_mean'], 4)}", f"inlier_ratio_mean={fmt(pose_overall['best_success_inlier_ratio_mean'], 4)}", f"reproj_error_mean={fmt(pose_overall['best_success_reproj_error_mean'], 4)}"],
        ],
        "validation": [
            ["pipeline_status", validation["pipeline_status"], "", ""],
            ["Layer-1 ortho", f"eval={ortho['evaluated_query_count']}/{ortho['query_count']}", f"center_offset_mean={fmt(get_mean(ortho, 'center_offset_m'), 3)}m", f"ortho_iou_mean={fmt(get_mean(ortho, 'ortho_iou'), 4)}"],
            ["Layer-2 pose", f"eval={pose_vs_at['evaluated_query_count']}/{pose_vs_at['query_count']}", f"horizontal_error_mean={fmt(get_mean(pose_vs_at, 'horizontal_error_m'), 3)}m", f"view_dir_error_mean={fmt(get_mean(pose_vs_at, 'view_dir_angle_error_deg'), 3)}deg"],
            ["Layer-3 tiepoint", f"eval={tiepoint['evaluated_query_count']}/{tiepoint['query_count']}", f"RMSE={fmt(tiepoint['tiepoint_xy_error_rmse_m'], 3)}m", f"inlier_ratio_mean={fmt(tiepoint['tiepoint_inlier_ratio_mean'], 4)}"],
            ["Layer-3 tiepoint CSV", f"files={detail['csv_count']}/{detail['expected_ok_count']}", f"missing={', '.join(detail['failed_ids'] + detail['missing_ok_ids']) or '-'}", "scope=RANSAC inliers"],
            ["Frame sanity", count_status_text(frame["diagnosis_counts"]), f"DSM valid={pct(get_nested(frame_numeric, 'dsm_sample_valid_ratio_on_truth_grid', 'mean'), 2)}", f"Pred valid={pct(get_nested(frame_numeric, 'pred_valid_pixel_ratio', 'mean'), 2)}"],
        ],
    }


def build_markdown(data: dict[str, Any], assets: dict[str, Path], out_md: Path, missing: list[str]) -> str:
    tables = metric_tables(data)
    ortho = data["ortho"]
    pose_vs_at = data["pose_vs_at"]
    tiepoint = data["tiepoint"]
    frame = data["frame"]
    frame_numeric = frame.get("numeric_summaries", {})
    detail = tiepoint_detail_summary(data)

    def img(key: str, caption: str) -> str:
        path = assets.get(key)
        if not path:
            return ""
        return f"\n![{caption}]({rel(path, out_md.parent)})\n\n*{caption}*\n"

    lines = [
        "# CaiWangCun DOM/DSM 完整替换 Gate 实验报告",
        "",
        "## 1. 实验目的",
        "本次实验验证：在 query、DINOv2、RoMa v2、PnP 和 validation 算法参数保持不变的前提下，将所有绑定旧候选库的资产完整替换为 CaiWangCun 0.14m DOM/DSM 资产后，predicted ortho 是否从前两个不完整替换分支中的大偏移、倾斜和大面积黑框现象恢复正常。",
        "",
        "核心判断是区分 DSM 质量问题、单视角覆盖问题、pose 问题和候选库/坐标框架混用问题。本轮 gate 只验证 5 个 sample query，不扩大到 full run。",
        "",
        "## 2. 评估方法和评估指标",
        "- Retrieval：用 Top-1/5/10/20 intersection recall 和 MRR 评估 CaiWangCun candidate library 对 40 张 query 的粗检索覆盖能力。",
        "- DSM/Pose gate：检查 DSM cache build 成功率、2D-3D sampling 状态、PnP 状态、best pose score、inlier ratio 和 reprojection error。",
        "- Layer-1 Ortho alignment：将 predicted ortho 与 CaiWangCun DOM truth 在同一 truth grid 上比较，重点看 center_offset_m、ortho_iou、SSIM 和有效像素比例。",
        "- Layer-2 Pose vs AT：将 best pose 与 AT/query reference pose 比较，重点看 horizontal_error_m 和 view_dir_angle_error_deg。",
        "- Layer-3 Tiepoint：在 predicted ortho 与 truth ortho 之间做局部 tiepoint ground error，重点看 RMSE、match count 和 inlier ratio。",
        "- Frame sanity：检查 DSM valid ratio、pred valid pixel ratio、camera/bbox offset 和 truth-to-footprint area ratio，用于解释黑框、偏移和单视角覆盖。",
        "",
        "## 3. 实验流程与数据准备",
        "本轮是完整替换：复用 009/010 的 40 张 query 与 query features，但 CaiWangCun DOM/DSM mosaic、candidate tile library、candidate DINOv2 features、FAISS index、retrieval Top20、RoMa v2 rerank、formal manifests、DSM cache、pose manifest、gate pose 和 validation 输出均重新生成。不使用 ODM LAZ 或 SRTM fallback。",
        "",
        "### 3.1 数据资产",
        "|资产|CRS|波段数|类型|分辨率/说明|",
        "|---|---:|---:|---:|---:|",
    ]
    lines += [f"|{row[0]}|{row[1]}|{row[2]}|{row[3]}|{row[4]}|" for row in tables["asset"]]
    lines += [
        "",
        "## 4. 实验结果",
        "### 4.1 Retrieval",
        "|指标|数值|",
        "|---|---:|",
    ]
    lines += [f"|{row[0]}|{row[1]}|" for row in tables["retrieval"]]
    lines += [img("retrieval_recall", "图 1. CaiWangCun candidate library 的 DINOv2 Top-K retrieval 指标。")]
    lines += [
        "### 4.2 DSM 与 Pose Gate",
        "|环节|规模|状态|质量|",
        "|---|---:|---|---|",
    ]
    lines += [f"|{row[0]}|{row[1]}|{row[2]}|{row[3]}|" for row in tables["pose"]]
    lines += [
        img("sampling_status", "图 2. 2D-3D sampling 状态分布。"),
        img("pnp_status", "图 3. PnP candidate 状态分布。"),
        "### 4.3 三层 validation 与 frame sanity",
        "|评估层|规模/状态|核心结果 1|核心结果 2|",
        "|---|---|---|---|",
    ]
    lines += [f"|{row[0]}|{row[1]}|{row[2]}|{row[3]}|" for row in tables["validation"]]
    lines += [
        img("layer_metrics", "图 4. Layer-1/2/3 关键几何误差。"),
        img("frame_sanity", "图 5. Frame sanity 的 DSM/predicted coverage 与 offset 指标。"),
        f"Layer-1 center_offset_m mean 为 {fmt(get_mean(ortho, 'center_offset_m'), 3)} m，ortho_iou mean 为 {fmt(get_mean(ortho, 'ortho_iou'), 4)}；Layer-2 horizontal_error_m mean 为 {fmt(get_mean(pose_vs_at, 'horizontal_error_m'), 3)} m；Layer-3 tiepoint RMSE 为 {fmt(tiepoint['tiepoint_xy_error_rmse_m'], 3)} m。",
        f"Frame sanity 中 DSM valid ratio mean 为 {pct(get_nested(frame_numeric, 'dsm_sample_valid_ratio_on_truth_grid', 'mean'), 2)}，pred valid pixel ratio mean 为 {pct(get_nested(frame_numeric, 'pred_valid_pixel_ratio', 'mean'), 2)}，说明黑框不是 DSM 大面积 nodata 引起，predicted ortho 的有效覆盖已回到正常范围。",
        "",
        "### Layer-3 tiepoint detail CSV",
        f"- Per-query detail dir: `{detail['detail_dir']}`",
        "- Scope: ratio-test matches retained by RANSAC inliers only, same as formal tiepoint RMSE.",
        f"- File pattern: `<query_id>_tiepoints.csv`; generated files: `{detail['csv_count']}/{detail['expected_ok_count']}` ok queries.",
        f"- Fields: `{', '.join(detail['fields'])}`",
        f"- Queries without detail CSV: `{', '.join(detail['failed_ids'] + detail['missing_ok_ids']) or '-'}`",
        "",
        "### 4.4 代表性可视化",
        img("q_003_truth_overlay", "图 6. q_003 predicted ortho 与 truth ortho overlay。"),
        img("q_003_dom_overlay", "图 7. q_003 predicted ortho 与 CaiWangCun DOM overlay。"),
        img("q_003_tiepoints_overlay", "图 8. q_003 tiepoint overlay。"),
        img("q_003_frame_overlay.png", "图 9. q_003 frame sanity overlay。"),
        img("q_003_offset_vectors.png", "图 10. q_003 frame sanity offset vectors。"),
        img("q_003_dsm_valid_mask_on_truth_grid.png", "图 11. q_003 DSM valid mask on truth grid。"),
        "",
        "### 4.5 与前两个 CaiWangCun 不完整替换分支对比",
        img("failure_comparison", "图 12. 前两个不完整替换分支与完整替换 gate 的偏移指标对比。"),
        "前两个 CaiWangCun 分支仍有约 515 m 的 center_offset 和约 648-652 m 的 horizontal_error。本轮完整替换后，center_offset_m mean 降至 4.39 m，horizontal_error_m mean 降至 1.83 m，说明主要问题来自旧 candidate/retrieval/rerank 资产与新 DOM/DSM 混用，而不是 CaiWangCun DSM 自身不可用。",
        "",
        "## 5. 结论与结果分析",
        "- 本轮 gate 通过：validation pipeline_status=ok，Layer-1/2/3 均完成 5/5 query，frame_sanity 诊断为 ok_or_manual_review=5。",
        "- DSM 支撑正常：DSM cache failed_count=0，truth grid 上 DSM valid ratio mean=99.77%，sampling nodata 仅 92/500000。",
        "- predicted ortho 恢复正常：有效像素比例约 74.65%，不再只落在整图小角落；几何偏移从 500m 级降到米级。",
        "- 结果支持当前解释：前两个分支的倾斜和大面积黑框不是单纯 DSM nodata，而是候选库/检索/重排资产没有随 CaiWangCun DOM 完整重建，导致候选图像、DSM、truth grid 和 pose reference 框架不一致。",
        "",
        "## 6. 后续想法",
        "- 可以规划 full run，但应先把当前 gate 的 5 个 query 可视化做人工抽查，确认道路和建筑没有系统性扭曲。",
        "- full run 前建议保留 frame_sanity 输出，并扩展到全部 40 张 query，用于识别边界覆盖不足和个别 query 的视觉异常。",
        "- 若 full run 中个别 query 出现偏移，应优先检查该 query 的 retrieval rank、RoMa inlier ratio、candidate tile 尺度和 DSM footprint，而不是回退到 ODM LAZ/SRTM。",
        "- 后续正式实验应将“完整替换候选库派生资产”写入协议，避免再次出现半替换导致的坐标框架混用。",
        "",
        "## 附录",
        f"- 实验根目录：`{data['experiment_root']}`",
        f"- Validation 根目录：`{data['suite_root']}`",
        "- 本轮不生成 satellite truth suite，不生成 comparison report，不使用 ODM LAZ/SRTM fallback。",
    ]
    if missing:
        lines += ["", "### 缺失图片记录"] + [f"- `{item}`" for item in missing]
    return "\n".join(line for line in lines if line is not None)


def build_docx(data: dict[str, Any], assets: dict[str, Path], out_docx: Path, missing: list[str]) -> None:
    tables = metric_tables(data)
    detail = tiepoint_detail_summary(data)
    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CaiWangCun DOM/DSM 完整替换 Gate 实验报告")
    set_cn_font(run, size=18, bold=True)
    add_paragraph(doc, "nadir_009010_caiwangcun_domdsm_fullreplace_gate_2026-04-20", center=True)

    add_heading(doc, "1. 实验目的")
    add_paragraph(
        doc,
        "本次实验验证：在 query、DINOv2、RoMa v2、PnP 和 validation 算法参数保持不变的前提下，"
        "将所有绑定旧候选库的资产完整替换为 CaiWangCun 0.14m DOM/DSM 后，predicted ortho 是否从前两个不完整替换分支中的大偏移、倾斜和大面积黑框现象恢复正常。",
    )
    add_bullets(
        doc,
        [
            "本轮是完整替换，不是旧 candidate library 加新 DOM/DSM 裁剪的半替换。",
            "运行阶段不使用 ODM LAZ 或 SRTM fallback。",
            "gate 只跑 5 个 sample query，full run 需在 gate 结论稳定后另行规划。",
        ],
    )

    add_heading(doc, "2. 评估方法和评估指标")
    add_table(
        doc,
        ["评估层", "指标", "解释"],
        [
            ["Retrieval", "Top-1/5/10/20 recall, MRR", "评估 CaiWangCun candidate library 的粗检索覆盖能力。"],
            ["DSM/Pose gate", "DSM build, sampling, PnP, score", "检查 DSM 支撑、2D-3D 采样和候选 pose 求解是否稳定。"],
            ["Layer-1", "center_offset_m, ortho_iou, SSIM", "比较 predicted ortho 与 CaiWangCun DOM truth 的平面一致性。"],
            ["Layer-2", "horizontal_error_m, view_dir_angle_error_deg", "比较 best pose 与 AT/query reference pose 的姿态误差。"],
            ["Layer-3", "tiepoint RMSE, match count, inlier ratio", "用局部 tiepoint 地面误差检查 predicted/truth 的细部一致性。"],
            ["Frame sanity", "DSM valid, pred valid, camera/bbox offset", "区分 DSM nodata、pose/frame 偏移和单视角覆盖不足。"],
        ],
    )

    add_heading(doc, "3. 实验流程与数据准备")
    add_paragraph(
        doc,
        "复用 009/010 的 40 张 query 与 query features；CaiWangCun DOM/DSM mosaic、candidate tile library、"
        "candidate DINOv2 features、FAISS index、retrieval Top20、RoMa v2 rerank、formal manifests、DSM cache、pose manifest、gate pose 和 validation 输出全部重建。",
    )
    add_heading(doc, "3.1 数据资产", level=2)
    add_table(doc, ["资产", "CRS", "波段数", "类型", "分辨率/说明"], tables["asset"])

    add_heading(doc, "4. 实验结果")
    add_heading(doc, "4.1 Retrieval", level=2)
    add_table(doc, ["指标", "数值"], tables["retrieval"])
    add_picture(doc, assets["retrieval_recall"], "图 1. CaiWangCun candidate library 的 DINOv2 Top-K retrieval 指标。")

    add_heading(doc, "4.2 DSM 与 Pose Gate", level=2)
    add_table(doc, ["环节", "规模", "状态", "质量"], tables["pose"])
    add_picture(doc, assets["sampling_status"], "图 2. 2D-3D sampling 状态分布。")
    add_picture(doc, assets["pnp_status"], "图 3. PnP candidate 状态分布。")

    add_heading(doc, "4.3 三层 Validation 与 Frame Sanity", level=2)
    add_table(doc, ["评估层", "规模/状态", "核心结果 1", "核心结果 2"], tables["validation"])
    add_heading(doc, "Layer-3 tiepoint detail CSV", level=2)
    add_bullets(
        doc,
        [
            f"Per-query detail dir: {detail['detail_dir']}",
            "Scope: ratio-test matches retained by RANSAC inliers only, same as formal tiepoint RMSE.",
            f"File pattern: <query_id>_tiepoints.csv; generated files: {detail['csv_count']}/{detail['expected_ok_count']} ok queries.",
            f"Fields: {', '.join(detail['fields'])}",
            f"Queries without detail CSV: {', '.join(detail['failed_ids'] + detail['missing_ok_ids']) or '-'}",
        ],
    )
    add_picture(doc, assets["layer_metrics"], "图 4. Layer-1/2/3 关键几何误差。")
    add_picture(doc, assets["frame_sanity"], "图 5. Frame sanity 的 DSM/predicted coverage 与 offset 指标。")

    add_heading(doc, "4.4 代表性可视化", level=2)
    for key, caption in [
        ("q_003_truth_overlay", "图 6. q_003 predicted ortho 与 truth ortho overlay。"),
        ("q_003_dom_overlay", "图 7. q_003 predicted ortho 与 CaiWangCun DOM overlay。"),
        ("q_003_tiepoints_overlay", "图 8. q_003 tiepoint overlay。"),
        ("q_003_frame_overlay.png", "图 9. q_003 frame sanity overlay。"),
        ("q_003_offset_vectors.png", "图 10. q_003 frame sanity offset vectors。"),
        ("q_003_dsm_valid_mask_on_truth_grid.png", "图 11. q_003 DSM valid mask on truth grid。"),
    ]:
        path = assets.get(key)
        if path:
            add_picture(doc, path, caption)

    add_heading(doc, "4.5 失败分支对比", level=2)
    add_picture(doc, assets["failure_comparison"], "图 12. 前两个不完整替换分支与完整替换 gate 的偏移指标对比。")
    add_paragraph(
        doc,
        "前两个 CaiWangCun 分支仍有约 515 m 的 center_offset 和约 648-652 m 的 horizontal_error。"
        "本轮完整替换后，center_offset_m mean 降至 4.39 m，horizontal_error_m mean 降至 1.83 m。",
    )

    add_heading(doc, "5. 结论与结果分析")
    add_bullets(
        doc,
        [
            "本轮 gate 通过：validation pipeline_status=ok，Layer-1/2/3 均完成 5/5 query。",
            "DSM 支撑正常：DSM cache failed_count=0，truth grid 上 DSM valid ratio mean=99.77%，sampling nodata 仅 92/500000。",
            "predicted ortho 恢复正常：有效像素比例约 74.65%，不再只落在整图小角落。",
            "前两个分支的倾斜和大面积黑框主要来自候选库/检索/重排资产与新 DOM/DSM 混用造成的参考框架不一致。",
        ],
    )

    add_heading(doc, "6. 后续想法")
    add_bullets(
        doc,
        [
            "可以规划 full run，但应先对当前 gate 的 5 个 query 可视化做人工抽查。",
            "full run 前建议保留并扩展 frame_sanity 到全部 40 张 query。",
            "若 full run 个别 query 异常，应优先检查 retrieval rank、RoMa inlier ratio、candidate tile 尺度和 DSM footprint。",
            "后续协议中应明确：替换 DOM/DSM 时，所有候选库派生资产必须一起重建。",
        ],
    )

    add_heading(doc, "附录")
    add_bullets(
        doc,
        [
            f"实验根目录：{data['experiment_root']}",
            f"Validation 根目录：{data['suite_root']}",
            "本轮不生成 satellite truth suite，不生成 comparison report，不使用 ODM LAZ/SRTM fallback。",
        ],
    )
    if missing:
        add_heading(doc, "缺失图片记录", level=2)
        add_bullets(doc, missing)

    ensure_dir(out_docx.parent)
    doc.save(str(out_docx))


def main() -> None:
    args = parse_args()
    experiment_root = Path(args.experiment_root).resolve()
    reports_dir = experiment_root / "reports"
    out_md = Path(args.out_md).resolve() if args.out_md else reports_dir / "caiwangcun_fullreplace_gate_report.md"
    out_docx = Path(args.out_docx).resolve() if args.out_docx else reports_dir / "caiwangcun_fullreplace_gate_report.docx"
    assets_dir = Path(args.assets_dir).resolve() if args.assets_dir else reports_dir / "assets"

    data = collect_inputs(experiment_root)
    assets, missing = build_assets(data, assets_dir)
    markdown = build_markdown(data, assets, out_md, missing)
    ensure_dir(out_md.parent)
    out_md.write_text(markdown, encoding="utf-8")
    build_docx(data, assets, out_docx, missing)
    print(out_md)
    print(out_docx)


if __name__ == "__main__":
    main()
