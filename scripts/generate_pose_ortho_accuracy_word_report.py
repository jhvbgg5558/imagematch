#!/usr/bin/env python3
"""Generate a Word report for the formal pose orthophoto-truth experiment.

Purpose:
- write a formal experiment note for the current orthophoto-truth validation
  run under `pose_v1_formal/eval_ortho_truth`;
- place metric definitions before experiment purpose, assets, workflow,
  outputs, and result interpretation;
- generate a standalone `.docx` report under the experiment `reports/`
  directory.

Main inputs:
- `per_query_ortho_accuracy.csv`
- `overall_ortho_accuracy.json`
- `per_flight_ortho_accuracy.csv`
- `failure_buckets.csv`
- representative overlay PNGs from `viz_overlay_truth/`

Main outputs:
- `eval_ortho_truth/reports/pose_ortho_accuracy_report.docx`

Applicable task constraints:
- this report explains offline validation only;
- DOM overlays are diagnostic-only and must not be framed as the main
  accuracy evidence.
"""

from __future__ import annotations

import argparse
import csv
import json
import math
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--eval-root", required=True)
    parser.add_argument("--out-docx", default=None)
    return parser.parse_args()


def load_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", newline="", encoding="utf-8-sig") as handle:
        return list(csv.DictReader(handle))


def load_json(path: Path) -> dict[str, object]:
    return json.loads(path.read_text(encoding="utf-8"))


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def set_cn_font(run, size: int = 11, bold: bool = False) -> None:
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(size)
    run.bold = bold


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    size: int = 10,
    align=WD_ALIGN_PARAGRAPH.CENTER,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    run = paragraph.add_run(text)
    set_cn_font(run, size=size, bold=bold)


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = "Heading 1" if level == 1 else "Heading 2"
    run = paragraph.add_run(text)
    set_cn_font(run, size=14 if level == 1 else 12, bold=True)


def add_paragraph(doc: Document, text: str, *, size: int = 11, center: bool = False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_cn_font(run, size=size)


def add_bullets(doc: Document, lines: list[str]) -> None:
    for line in lines:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(line)
        set_cn_font(run, size=11)


def add_picture_with_caption(doc: Document, image_path: Path, caption: str, *, width_inch: float = 5.8) -> None:
    if not image_path.exists():
        return
    doc.add_picture(str(image_path), width=Inches(width_inch))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(caption)
    set_cn_font(run, size=10)


def fmt_float(value: float | None, digits: int = 3) -> str:
    if value is None or not math.isfinite(value):
        return "-"
    return f"{value:.{digits}f}"


def short_flight_name(flight_id: str) -> str:
    parts = flight_id.split("_")
    return parts[2] if len(parts) >= 3 else flight_id


def pick_cases(ok_rows: list[dict[str, str]]) -> dict[str, dict[str, str]]:
    selected: dict[str, dict[str, str]] = {}
    if not ok_rows:
        return selected
    selected["best_phase_corr"] = min(ok_rows, key=lambda row: float(row["phase_corr_error_m"]))
    selected["best_iou"] = max(ok_rows, key=lambda row: float(row["ortho_iou"]))
    selected["best_ssim"] = max(ok_rows, key=lambda row: float(row["ssim"]))
    selected["worst_phase_corr"] = max(ok_rows, key=lambda row: float(row["phase_corr_error_m"]))
    return selected


def build_metric_glossary_table(doc: Document) -> None:
    rows = [
        ("phase_corr_shift_x_m / phase_corr_shift_y_m", "预测正射相对真值正射的整体平移偏差分量，单位米。", "绝对值越小越好；接近 0 说明整体平移偏差更小。"),
        ("phase_corr_error_m", "整体平移误差主指标，由 X/Y 平移偏差合成。", "越小越好；当前最适合直接判断平面套合精度。"),
        ("center_offset_m", "预测正射有效区域中心与真值正射有效区域中心的偏移。", "越小越好；但它受有效覆盖形状影响较大，不能单独代表精度。"),
        ("ortho_iou", "预测正射有效区域与真值正射有效区域的重叠程度。", "越大越好；高值说明几何覆盖更一致。"),
        ("ortho_overlap_ratio", "预测正射覆盖到真值有效区域的比例。", "越大越好；低值常表示只在局部区域对齐。"),
        ("ncc", "灰度归一化互相关。", "越大越好；越接近 1 表示亮度纹理变化更一致。"),
        ("ssim", "结构相似性指标。", "越大越好；更关注结构和纹理是否一致。"),
        ("common_valid_ratio", "两张图可共同比较的有效区域比例。", "越大越好；太小会削弱局部高相似度的说服力。"),
        ("best_inlier_count", "best pose 对应 PnP 解算的内点数。", "通常越大越好；表示支持该 pose 的几何约束更多。"),
        ("best_inlier_ratio", "PnP 内点比例。", "通常越大越好；表示几何一致性更强。"),
        ("best_reproj_error", "PnP 最佳候选的重投影误差。", "越小越好；较大值往往意味着 pose 几何质量较差。"),
        ("best_score", "formal pose 主链的最终综合分数。", "通常越大越好；但它不是最终真值精度。"),
        ("eval_status", "该 query 的正射真值评估状态。", "`ok` 表示成功；其他值表示失败类型。"),
        ("eval_status_detail", "失败原因或补充说明。", "用于定位流程失败位置，不表示精度高低。"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["指标", "含义", "值大/值小代表什么"]
    for idx, header in enumerate(headers):
        set_cell_text(table.cell(0, idx), header, bold=True)
        shade_cell(table.cell(0, idx), "D9EAF7")
    for metric, meaning, interpretation in rows:
        row = table.add_row().cells
        set_cell_text(row[0], metric, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(row[1], meaning, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(row[2], interpretation, align=WD_ALIGN_PARAGRAPH.LEFT)


def build_overall_table(doc: Document, overall: dict[str, object]) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["指标", "均值", "中位数", "P90"]
    for idx, header in enumerate(headers):
        set_cell_text(table.cell(0, idx), header, bold=True)
        shade_cell(table.cell(0, idx), "D9EAF7")
    metric_names = ["phase_corr_error_m", "center_offset_m", "ortho_iou", "ssim"]
    for metric_name in metric_names:
        payload = overall[metric_name]
        row = table.add_row().cells
        set_cell_text(row[0], metric_name)
        set_cell_text(row[1], fmt_float(payload["mean"]))
        set_cell_text(row[2], fmt_float(payload["median"]))
        set_cell_text(row[3], fmt_float(payload["p90"]))


def build_per_flight_table(doc: Document, rows: list[dict[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Flight", "Query数", "phase_corr_error_m均值", "ortho_iou均值", "ssim均值"]
    for idx, header in enumerate(headers):
        set_cell_text(table.cell(0, idx), header, bold=True)
        shade_cell(table.cell(0, idx), "D9EAF7")
    for item in rows:
        row = table.add_row().cells
        set_cell_text(row[0], short_flight_name(item["flight_id"]))
        set_cell_text(row[1], item["query_count"])
        set_cell_text(row[2], fmt_float(float(item["phase_corr_error_m_mean"])))
        set_cell_text(row[3], fmt_float(float(item["ortho_iou_mean"])))
        set_cell_text(row[4], fmt_float(float(item["ssim_mean"])))


def build_case_table(doc: Document, title: str, row_data: dict[str, str]) -> None:
    add_heading(doc, title, 2)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_cell_text(table.cell(0, 0), "字段", bold=True)
    set_cell_text(table.cell(0, 1), "值", bold=True)
    shade_cell(table.cell(0, 0), "D9EAF7")
    shade_cell(table.cell(0, 1), "D9EAF7")
    fields = [
        ("query_id", row_data["query_id"]),
        ("flight_id", row_data["flight_id"]),
        ("best_candidate_id", row_data["best_candidate_id"]),
        ("phase_corr_error_m", row_data["phase_corr_error_m"]),
        ("center_offset_m", row_data["center_offset_m"]),
        ("ortho_iou", row_data["ortho_iou"]),
        ("ssim", row_data["ssim"]),
        ("common_valid_ratio", row_data["common_valid_ratio"]),
        ("best_inlier_count", row_data["best_inlier_count"]),
        ("best_inlier_ratio", row_data["best_inlier_ratio"]),
        ("best_reproj_error", row_data["best_reproj_error"]),
        ("best_score", row_data["best_score"]),
    ]
    for key, value in fields:
        row = table.add_row().cells
        set_cell_text(row[0], key)
        set_cell_text(row[1], str(value))


def build_flow_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["阶段", "输入", "处理", "输出"]
    for idx, header in enumerate(headers):
        set_cell_text(table.cell(0, idx), header, bold=True)
        shade_cell(table.cell(0, idx), "D9EAF7")
    rows = [
        (
            "1. Query输入",
            "单张 UAV query 图像",
            "读取 formal 任务锁定的 query 图像与内参",
            "可供检索和几何计算的 query 输入",
        ),
        (
            "2. 候选检索",
            "query 图像 + satellite library",
            "从卫星正射候选库中检索 Top-20 candidate DOM tiles",
            "每个 query 对应 20 个候选区域",
        ),
        (
            "3. 图像匹配",
            "query 图像 + 单个 candidate DOM",
            "使用 RoMa v2 提取 query 与 DOM 之间的 2D-2D 同名点",
            "query 像素点和 DOM 像素点配对",
        ),
        (
            "4. 2D-3D 构建",
            "DOM 同名点 + candidate DSM",
            "先把 DOM 像素点转成地面 XY，再从 DSM 采样 Z，得到地面 3D 点",
            "query 2D 点与地面 3D 点的对应关系",
        ),
        (
            "5. 位姿解算",
            "2D-3D 点对 + query 内参",
            "通过 PnP 解算相机旋转、平移和相机中心位置",
            "每个 query-candidate 的 pose 结果",
        ),
        (
            "6. Best pose 选择",
            "20 个 candidate pose",
            "根据 score / inlier / reprojection error 选择最终 best pose",
            "每个 query 的唯一最终位姿",
        ),
        (
            "7. 预测正射生成",
            "原始 query 图像 + best pose + best candidate DSM",
            "将原始 UAV 图像按位姿和 DSM 投影回地面",
            "预测正射 pred tile",
        ),
        (
            "8. 真值正射准备",
            "query_id + flight_id + ODM orthophoto",
            "从对应 flight 的无人机真值正射中裁出 query 对应区域",
            "真值正射 truth tile",
        ),
        (
            "9. 同网格评估",
            "pred tile + truth tile",
            "保证两张图同 CRS、同分辨率、同 transform、同 extent 后计算指标",
            "per-query 指标和 overall summary",
        ),
        (
            "10. 可视化诊断",
            "pred tile + truth tile + DOM",
            "输出 truth overlay、outline、DOM overlay 等图像",
            "用于人工检查和失败归因的可视化结果",
        ),
    ]
    for phase, inputs, process, outputs in rows:
        row = table.add_row().cells
        set_cell_text(row[0], phase, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(row[1], inputs, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(row[2], process, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(row[3], outputs, align=WD_ALIGN_PARAGRAPH.LEFT)


def build_text_flowchart(doc: Document) -> None:
    lines = [
        "原始 UAV query 图像",
        "↓",
        "Top-20 候选检索",
        "↓",
        "RoMa v2 匹配 query 与 candidate DOM",
        "↓",
        "DOM 同名点转地面 XY，DSM 采样得到 Z",
        "↓",
        "构建 2D-3D 对应关系",
        "↓",
        "PnP 解算每个 candidate 的 pose",
        "↓",
        "选择每个 query 的 best pose",
        "↓",
        "使用 best pose 把原始 UAV 图像投到地面，生成预测正射",
        "↓",
        "从对应 flight 的 ODM orthophoto 裁出真值正射",
        "↓",
        "在统一地理网格上比较预测正射与真值正射",
        "↓",
        "输出指标、失败清单和 overlay 图",
    ]
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for idx, line in enumerate(lines):
        run = paragraph.add_run(line)
        set_cn_font(run, size=11, bold=(idx % 2 == 0))
        if idx != len(lines) - 1:
            run.add_break()


def main() -> None:
    args = parse_args()
    eval_root = Path(args.eval_root)
    out_docx = Path(args.out_docx) if args.out_docx else eval_root / "reports" / "pose_ortho_accuracy_report.docx"

    per_query_rows = load_csv(eval_root / "per_query_ortho_accuracy.csv")
    overall = load_json(eval_root / "overall_ortho_accuracy.json")
    per_flight_rows = load_csv(eval_root / "per_flight_ortho_accuracy.csv")
    failure_rows = load_csv(eval_root / "failure_buckets.csv")
    viz_truth_root = eval_root / "viz_overlay_truth"

    ok_rows = [row for row in per_query_rows if row["eval_status"] == "ok"]
    cases = pick_cases(ok_rows)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Pose v1 正射真值验证实验说明")
    set_cn_font(title_run, size=16, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("基于 per_query_ortho_accuracy.csv 的正式实验说明与结果解读")
    set_cn_font(subtitle_run, size=11)

    add_heading(doc, "1. 评价指标介绍", 1)
    add_paragraph(
        doc,
        "本实验的正式主结论基于“预测正射 vs UAV 真值正射”的定量套合结果，而不是基于运行时参与定位的卫星 DOM"
        "（Digital Orthophoto Map，数字正射影像图）。其中 phase_corr_error_m 是当前最稳定的平面套合主指标，"
        "ortho_iou 和 SSIM（Structural Similarity，结构相似性）用于补充几何重叠与纹理一致性解释。",
    )
    build_metric_glossary_table(doc)

    add_heading(doc, "2. 实验目的", 1)
    add_bullets(
        doc,
        [
            "验证 formal Pose v1 估计出的相机位姿，是否能够支持把原始 UAV 图像正确正射到地面坐标系。",
            "验证“RoMa v2（一个稠密图像匹配模型）同名点 + DOM/DSM 物方点 + PnP（Perspective-n-Point，相机位姿解算）”得到的 pose，是否能在独立的 UAV 正射真值上表现出稳定套合精度。",
            "将“位姿是否解得出来”升级为“解出来的位姿能否把 UAV 图像正确放回地面”的验证口径。",
        ],
    )

    add_heading(doc, "3. 实验内容", 1)
    add_bullets(
        doc,
        [
            "输入 query 为 formal 任务锁定的 40 张 UAV query 图像。",
            "运行时 pose 结果来自 new2output/pose_v1_formal 下的 best pose 输出，其中 best pose 表示每个 query 在 20 个候选中最终选中的最佳位姿结果。",
            "真值正射来自原始 UAV flight 目录下现有的 ODM orthophoto，其中 ODM orthophoto 指 OpenDroneMap 工程生成的无人机正射成果，而不是由卫星 DOM 反推。",
            "最终在统一地理网格上比较预测正射与真值正射，并给出每 query 的定量指标与可视化结果。",
        ],
    )

    add_heading(doc, "4. 详细实验流程图", 1)
    add_paragraph(
        doc,
        "下面先给出适合直接写入 Word 的文字流程图。它对应的是本次实验从 query 输入到最终精度评估的完整主链。",
    )
    build_text_flowchart(doc)
    add_paragraph(
        doc,
        "如果把整条链拆成“输入什么、做了什么、输出什么”，则可以写成下面这张详细流程表。",
    )
    build_flow_table(doc)

    add_heading(doc, "5. 主要资产与路径", 1)
    add_bullets(
        doc,
        [
            "formal 位姿结果：new2output/pose_v1_formal/summary/per_query_best_pose.csv",
            "formal PnP 结果：new2output/pose_v1_formal/pnp/pnp_results.csv",
            "formal manifest：new2output/pose_v1_formal/manifest/pose_manifest.json，其中 manifest 是统一的输入清单与参数索引文件。",
            "query 输入：new1output/query_reselect_2026-03-26_v2/query_inputs/",
            "query truth 索引：new1output/query_reselect_2026-03-26_v2/query_truth/queries_truth_seed.csv",
            "UAV 真值正射根：D:/数据/武汉影像/无人机0.1m/<flight_id>/odm_orthophoto/odm_orthophoto.tif",
        ],
    )

    add_heading(doc, "6. 详细步骤解释", 1)
    add_bullets(
        doc,
        [
            "步骤 1：先输入单张 UAV query 图像。这里的 query 不是正射图，也没有地理坐标，因此不能直接拿来做地理套合。",
            "步骤 2：先做候选检索，从卫星正射库中为每个 query 找到 Top-20 个最可能对应的候选区域。这样做的目的，是把问题从全区域定位缩小到少量候选里精定位。",
            "步骤 3：对 query 和每个 candidate DOM 做 RoMa v2 匹配，得到 query 像素点与 DOM 像素点之间的 2D-2D 同名点。",
            "步骤 4：把 DOM 上的同名点转成真实地面点。做法是：先利用 DOM 的地理参考把 DOM 像素转成地面 XY，再利用该 candidate 对应的 DSM 采样高程 Z，最终得到地面 3D 点。",
            "步骤 5：用 query 图像上的 2D 点、地面 3D 点和 query 内参，调用 PnP 解算相机位姿，得到每个 candidate 的旋转、平移和相机中心位置。",
            "步骤 6：从 20 个 candidate 的 pose 结果里选出每个 query 的 best pose。这个 best pose 是后续生成预测正射时唯一使用的正式位姿。",
            "步骤 7：使用 best pose、原始 query 图像和 best candidate 对应的 DSM，把原始 UAV 图像投影回地面，生成预测正射。这里生成的是 pred tile。",
            "步骤 8：再从同一 query 所属 flight 的 ODM orthophoto 中裁出真值正射。这里裁出来的是 truth tile，它和 pred tile 必须对应同一个 query。",
            "步骤 9：让预测正射和真值正射共网格，即同 CRS、同分辨率、同 transform、同 extent，然后再计算 phase_corr_error_m、ortho_iou、ssim 等正式指标。",
            "步骤 10：最后输出数值表、失败清单和 overlay 可视化图，用来完成 formal Pose v1 的独立精度验证。",
        ],
    )

    add_heading(doc, "7. 输出内容", 1)
    add_bullets(
        doc,
        [
            "per_query_ortho_accuracy.csv：每个 query 一行，是主分析表。",
            "overall_ortho_accuracy.json：全量汇总统计。",
            "per_flight_ortho_accuracy.csv：按 flight 分组统计。",
            "failure_buckets.csv：失败样本和失败原因。",
            "viz_overlay_truth/：预测正射与真值正射叠加图。",
            "viz_overlay_dom/：预测正射与卫星 DOM 的辅助诊断图。",
        ],
    )

    add_heading(doc, "8. Full-40 结果概览", 1)
    add_paragraph(
        doc,
        f"当前 full-40 共评估 {overall['query_count']} 个 query，其中成功完成正射真值对比的 query 数为 {overall['evaluated_query_count']}，"
        f"失败 {overall['query_count'] - overall['evaluated_query_count']} 个。"
        f"状态分布为 {', '.join(f'{k}={v}' for k, v in overall['eval_status_counts'].items())}。",
    )
    add_paragraph(
        doc,
        "从当前结果看，phase_corr_error_m（通过 phase correlation，频域相位相关方法估计的整体平移误差）"
        "比 center_offset_m 更稳定，适合作为主平面套合指标；"
        "center_offset_m 主要作为有效覆盖形状和 footprint 支持情况的辅助解释项。",
    )
    build_overall_table(doc, overall)

    add_heading(doc, "9. 按 Flight 的稳定性", 1)
    add_paragraph(doc, "下表用于判断不同 flight 是否存在系统性难度差异。")
    build_per_flight_table(doc, per_flight_rows)

    add_heading(doc, "10. 代表性样本", 1)
    case_titles = {
        "best_phase_corr": "10.1 平移误差最小样本",
        "best_iou": "10.2 几何重叠最好样本",
        "best_ssim": "10.3 纹理一致性最好样本",
        "worst_phase_corr": "10.4 平移误差最大样本",
    }
    for key, heading_text in case_titles.items():
        row_data = cases.get(key)
        if row_data is None:
            continue
        build_case_table(doc, heading_text, row_data)
        query_id = row_data["query_id"]
        add_picture_with_caption(doc, viz_truth_root / f"{query_id}_overlay.png", f"{query_id} 预测正射与真值正射叠加图")
        add_picture_with_caption(doc, viz_truth_root / f"{query_id}_outline.png", f"{query_id} 预测正射与真值正射边界图")

    add_heading(doc, "11. 失败样本说明", 1)
    if failure_rows and failure_rows[0].get("query_id"):
        for failure in failure_rows:
            add_bullets(
                doc,
                [
                    f"{failure['query_id']} / {failure['best_candidate_id']} / {failure['failure_bucket']}",
                    f"原因：{failure['detail']}",
                ],
            )
    else:
        add_paragraph(doc, "当前没有失败样本。")

    add_heading(doc, "12. 当前结论", 1)
    add_bullets(
        doc,
        [
            "formal Pose v1 的“位姿解算 -> 正射纠正 -> UAV 真值正射验证”链路已经闭环。",
            "full-40 中已有 39 个 query 成功完成正射真值评估，说明这条链路具备较好的可运行性。",
            "当前 phase_corr_error_m 已经可以作为主平面精度指标使用。",
            "DOM overlay 只适合作为辅助诊断，不应替代 UAV 真值正射作为主证据。",
        ],
    )

    add_heading(doc, "13. 后续建议", 1)
    add_bullets(
        doc,
        [
            "先单独分析 q_022 的 dsm_intersection_failed 原因。",
            "再分析 best_score / best_inlier_count / best_reproj_error 与 phase_corr_error_m 的相关性。",
            "最终确定正式结论中应保留哪些指标作为主指标，哪些仅作为辅助解释。",
        ],
    )

    ensure_dir(out_docx.parent)
    doc.save(str(out_docx))
    print(out_docx)


if __name__ == "__main__":
    main()
