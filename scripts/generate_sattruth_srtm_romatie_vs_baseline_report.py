#!/usr/bin/env python3
"""Generate baseline-vs-satellite-truth-SRTM-RoMa tiepoint comparison reports."""

from __future__ import annotations

import argparse
import csv
import json
import statistics
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_BASELINE_ROOT = PROJECT_ROOT / "new2output" / "nadir_009010_dinov2_romav2_pose_2026-04-10"
DEFAULT_EXPERIMENT_ROOT = (
    PROJECT_ROOT
    / "new3output"
    / "nadir_009010_dinov2_romav2_pose_sattruth_srtm_romatie_2026-04-16"
)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--baseline-root", default=str(DEFAULT_BASELINE_ROOT))
    parser.add_argument("--experiment-root", default=str(DEFAULT_EXPERIMENT_ROOT))
    parser.add_argument("--out-md", default="")
    parser.add_argument("--out-docx", default="")
    return parser.parse_args()


def load_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", newline="", encoding="utf-8-sig") as handle:
        return list(csv.DictReader(handle))


def load_json(path: Path) -> dict[str, object]:
    return json.loads(path.read_text(encoding="utf-8"))


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def fmt(value: object, digits: int = 4) -> str:
    if value in ("", None):
        return "-"
    try:
        return f"{float(value):.{digits}f}"
    except (TypeError, ValueError):
        return str(value)


def set_cn_font(run, size: int = 11, bold: bool = False) -> None:
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(size)
    run.bold = bold


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_paragraph(doc: Document, text: str, *, bold: bool = False, size: int = 11) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_cn_font(run, size=size, bold=bold)


def add_heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_cn_font(run, size=14, bold=True)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        cell.text = ""
        run = cell.paragraphs[0].add_run(header)
        set_cn_font(run, size=10, bold=True)
        shade_cell(cell, "D9EAF7")
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].text = ""
            run = cells[idx].paragraphs[0].add_run(value)
            set_cn_font(run, size=10)


def summarize_pose_root(pose_root: Path) -> dict[str, object]:
    pnp_rows = load_csv(pose_root / "pnp" / "pnp_results.csv")
    best_rows = load_csv(pose_root / "summary" / "per_query_best_pose.csv")
    return {
        "pnp_row_count": len(pnp_rows),
        "best_row_count": len(best_rows),
        "pnp_status_counts": dict(Counter(row.get("status", "") for row in pnp_rows)),
        "best_status_counts": dict(Counter(row.get("best_status", "") for row in best_rows)),
    }


def mean_value(summary: dict[str, object], key: str) -> object:
    payload = summary.get(key, {})
    if isinstance(payload, dict):
        return payload.get("mean")
    return None


def median_of_counts(rows: list[dict[str, str]], key: str) -> float | None:
    vals = [float(row[key]) for row in rows if row.get(key) not in {"", None}]
    if not vals:
        return None
    return float(statistics.median(vals))


def build_low_match_rows(
    baseline_rows: list[dict[str, str]],
    current_rows: list[dict[str, str]],
) -> list[list[str]]:
    baseline_ok = [row for row in baseline_rows if row.get("eval_status") == "tiepoint_eval_ok" and row.get("tiepoint_match_count")]
    current_by_query = {row["query_id"]: row for row in current_rows}
    if not baseline_ok:
        return []
    threshold = statistics.quantiles([float(row["tiepoint_match_count"]) for row in baseline_ok], n=4)[0]
    low_rows = [row for row in baseline_ok if float(row["tiepoint_match_count"]) <= threshold]
    out: list[list[str]] = []
    for row in sorted(low_rows, key=lambda item: float(item["tiepoint_match_count"]))[:8]:
        current = current_by_query.get(row["query_id"], {})
        delta = "-"
        if current.get("tiepoint_match_count"):
            delta = fmt(float(current["tiepoint_match_count"]) - float(row["tiepoint_match_count"]), 0)
        out.append(
            [
                row["query_id"],
                fmt(row.get("tiepoint_match_count"), 0),
                fmt(current.get("tiepoint_match_count"), 0),
                delta,
                fmt(row.get("tiepoint_xy_error_rmse_m")),
                fmt(current.get("tiepoint_xy_error_rmse_m")),
            ]
        )
    return out


def main() -> None:
    args = parse_args()
    baseline_root = Path(args.baseline_root)
    experiment_root = Path(args.experiment_root)
    reports_root = experiment_root / "reports"
    ensure_dir(reports_root)
    out_md = Path(args.out_md) if args.out_md else reports_root / "sattruth_srtm_romatie_vs_baseline.md"
    out_docx = Path(args.out_docx) if args.out_docx else reports_root / "sattruth_srtm_romatie_vs_baseline.docx"

    baseline_pose_root = baseline_root / "pose_v1_formal"
    current_pose_root = experiment_root / "pose_v1_formal"
    baseline_suite_root = baseline_pose_root / "eval_pose_validation_suite"
    current_suite_root = current_pose_root / "eval_pose_validation_suite_sattruth_srtm"

    baseline_pose = summarize_pose_root(baseline_pose_root)
    current_pose = summarize_pose_root(current_pose_root)
    baseline_ortho = load_json(baseline_suite_root / "ortho_alignment" / "overall_ortho_accuracy.json")
    baseline_pose_vs_at = load_json(baseline_suite_root / "pose_vs_at" / "overall_pose_vs_at.json")
    baseline_tie = load_json(baseline_suite_root / "tiepoint_ground_error" / "overall_tiepoint_ground_error.json")
    current_ortho = load_json(current_suite_root / "ortho_alignment" / "overall_ortho_accuracy.json")
    current_pose_vs_at = load_json(current_suite_root / "pose_vs_at" / "overall_pose_vs_at.json")
    current_tie = load_json(current_suite_root / "tiepoint_ground_error" / "overall_tiepoint_ground_error.json")
    baseline_tie_rows = load_csv(baseline_suite_root / "tiepoint_ground_error" / "per_query_tiepoint_ground_error.csv")
    current_tie_rows = load_csv(current_suite_root / "tiepoint_ground_error" / "per_query_tiepoint_ground_error.csv")

    metric_rows = [
        ["baseline", "layer1 phase_corr_error_m mean", fmt(mean_value(baseline_ortho, "phase_corr_error_m"))],
        ["baseline", "layer1 ortho_iou mean", fmt(mean_value(baseline_ortho, "ortho_iou"))],
        ["baseline", "layer1 ssim mean", fmt(mean_value(baseline_ortho, "ssim"))],
        ["baseline", "layer2 horizontal_error_m mean", fmt(mean_value(baseline_pose_vs_at, "horizontal_error_m"))],
        ["baseline", "layer2 view_dir_angle_error_deg mean", fmt(mean_value(baseline_pose_vs_at, "view_dir_angle_error_deg"))],
        ["baseline", "layer3 tiepoint_match_count median", fmt(median_of_counts(baseline_tie_rows, "tiepoint_match_count"))],
        ["baseline", "layer3 tiepoint_inlier_count median", fmt(median_of_counts(baseline_tie_rows, "tiepoint_inlier_count"))],
        ["baseline", "layer3 tiepoint_xy_error_rmse_m", fmt(baseline_tie.get("tiepoint_xy_error_rmse_m"))],
        ["baseline", "layer3 tiepoint_xy_error_p90_m", fmt(baseline_tie.get("tiepoint_xy_error_p90_m"))],
        ["sattruth_srtm_romatie", "layer1 phase_corr_error_m mean", fmt(mean_value(current_ortho, "phase_corr_error_m"))],
        ["sattruth_srtm_romatie", "layer1 ortho_iou mean", fmt(mean_value(current_ortho, "ortho_iou"))],
        ["sattruth_srtm_romatie", "layer1 ssim mean", fmt(mean_value(current_ortho, "ssim"))],
        ["sattruth_srtm_romatie", "layer2 horizontal_error_m mean", fmt(mean_value(current_pose_vs_at, "horizontal_error_m"))],
        ["sattruth_srtm_romatie", "layer2 view_dir_angle_error_deg mean", fmt(mean_value(current_pose_vs_at, "view_dir_angle_error_deg"))],
        ["sattruth_srtm_romatie", "layer3 tiepoint_match_count median", fmt(median_of_counts(current_tie_rows, "tiepoint_match_count"))],
        ["sattruth_srtm_romatie", "layer3 tiepoint_inlier_count median", fmt(median_of_counts(current_tie_rows, "tiepoint_inlier_count"))],
        ["sattruth_srtm_romatie", "layer3 tiepoint_xy_error_rmse_m", fmt(current_tie.get("tiepoint_xy_error_rmse_m"))],
        ["sattruth_srtm_romatie", "layer3 tiepoint_xy_error_p90_m", fmt(current_tie.get("tiepoint_xy_error_p90_m"))],
    ]
    low_match_rows = build_low_match_rows(baseline_tie_rows, current_tie_rows)

    md_lines = [
        "# Satellite Truth + SRTM + RoMa Tiepoints vs Baseline",
        "",
        "## Scope",
        "",
        f"- Baseline root: `{baseline_root}`",
        f"- Current experiment root: `{experiment_root}`",
        "- Runtime retrieval / rerank / satellite candidate DOM library remained fixed.",
        "- Main changes were: satellite truth patches replaced UAV truth orthophotos, DSM returned to SRTM, and layer-3 matching switched to RoMa v2.",
        "",
        "## Pose Runtime",
        "",
        f"- Baseline PnP rows: `{baseline_pose['pnp_row_count']}`, status counts: `{baseline_pose['pnp_status_counts']}`",
        f"- Current PnP rows: `{current_pose['pnp_row_count']}`, status counts: `{current_pose['pnp_status_counts']}`",
        f"- Baseline best-pose rows: `{baseline_pose['best_row_count']}`, status counts: `{baseline_pose['best_status_counts']}`",
        f"- Current best-pose rows: `{current_pose['best_row_count']}`, status counts: `{current_pose['best_status_counts']}`",
        "",
        "## Validation Summary",
        "",
        "| Suite | Metric | Value |",
        "| --- | --- | --- |",
    ]
    for row in metric_rows:
        md_lines.append(f"| {row[0]} | {row[1]} | {row[2]} |")
    md_lines.extend(
        [
            "",
            "## Low-Match Query Review",
            "",
            "| Query | Baseline Matches | Current Matches | Delta | Baseline RMSE | Current RMSE |",
            "| --- | --- | --- | --- | --- | --- |",
        ]
    )
    for row in low_match_rows:
        md_lines.append(f"| {' | '.join(row)} |")
    md_lines.extend(
        [
            "",
            "## Interpretation",
            "",
            "- This route keeps the original runtime localization task unchanged and changes only the offline validation truth source plus the layer-3 matcher.",
            "- Layer-2 remains `pose_vs_at`, so changes in layer-2 should be interpreted as pose consistency changes rather than a change in evaluation semantics.",
            "- Low-match review is defined from the baseline layer-3 lower quartile of `tiepoint_match_count` and compares the same query IDs under the new route.",
            "",
        ]
    )
    out_md.write_text("\n".join(md_lines) + "\n", encoding="utf-8")

    doc = Document()
    add_heading(doc, "Satellite Truth + SRTM + RoMa Tiepoints vs Baseline")
    add_paragraph(doc, f"Baseline root: {baseline_root}")
    add_paragraph(doc, f"Current experiment root: {experiment_root}")
    add_paragraph(doc, "Runtime retrieval / rerank / candidate DOM remained fixed. Only the validation truth source, DSM source, and layer-3 matcher changed.")
    add_heading(doc, "Pose Runtime")
    add_table(
        doc,
        ["Run", "PnP rows", "PnP status counts", "Best rows", "Best status counts"],
        [
            ["baseline", str(baseline_pose["pnp_row_count"]), json.dumps(baseline_pose["pnp_status_counts"], ensure_ascii=False), str(baseline_pose["best_row_count"]), json.dumps(baseline_pose["best_status_counts"], ensure_ascii=False)],
            ["current", str(current_pose["pnp_row_count"]), json.dumps(current_pose["pnp_status_counts"], ensure_ascii=False), str(current_pose["best_row_count"]), json.dumps(current_pose["best_status_counts"], ensure_ascii=False)],
        ],
    )
    add_heading(doc, "Validation Summary")
    add_table(doc, ["Suite", "Metric", "Value"], metric_rows)
    add_heading(doc, "Low-Match Query Review")
    add_table(doc, ["Query", "Baseline Matches", "Current Matches", "Delta", "Baseline RMSE", "Current RMSE"], low_match_rows or [["-", "-", "-", "-", "-", "-"]])
    ensure_dir(out_docx.parent)
    doc.save(out_docx)
    print(out_md)
    print(out_docx)


if __name__ == "__main__":
    main()
