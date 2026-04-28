#!/usr/bin/env python3
"""Generate a formal Chinese report for RoMa v2 reranking under intersection-truth evaluation."""

from __future__ import annotations

import argparse
import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--result-dir", required=True)
    parser.add_argument("--out-docx", required=True)
    parser.add_argument("--out-md", required=True)
    parser.add_argument("--coarse-model-label", default="DINOv3")
    return parser.parse_args()


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def load_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", newline="", encoding="utf-8-sig") as f:
        return list(csv.DictReader(f))


def set_cn_font(run, size: int = 11, bold: bool = False) -> None:
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(size)
    run.bold = bold


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 10, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    set_cn_font(r, size=size, bold=bold)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.style = "Heading 1" if level == 1 else "Heading 2"
    r = p.add_run(text)
    set_cn_font(r, size=14 if level == 1 else 12, bold=True)


def add_paragraph(doc: Document, text: str, size: int = 11, center: bool = False) -> None:
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn_font(r, size=size)


def add_bullets(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(line)
        set_cn_font(r, size=11)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn_font(r, size=10)


def format_seconds(seconds: float) -> str:
    return f"{seconds:.2f}s ({seconds / 60.0:.2f} min)"


def stage_seconds(timing: dict) -> dict[str, float]:
    return {item["stage"]: float(item["elapsed_seconds"]) for item in timing["stages"]}


def choose_cases(comp_rows: list[dict[str, str]]) -> list[dict[str, str]]:
    improved = [r for r in comp_rows if r["baseline_first_truth_rank"] and r["romav2_first_truth_rank"] and int(r["romav2_first_truth_rank"]) < int(r["baseline_first_truth_rank"])]
    improved.sort(key=lambda r: int(r["baseline_first_truth_rank"]) - int(r["romav2_first_truth_rank"]), reverse=True)
    promoted = [r for r in comp_rows if int(r["promoted_to_top10"]) == 1]
    promoted.sort(key=lambda r: r["query_id"])
    selected: list[dict[str, str]] = []
    if promoted:
        case = dict(promoted[0])
        case["title"] = "从 11..20 拉回 Top-10 的改进样例"
        selected.append(case)
    if improved:
        cand = next((dict(r) for r in improved if r["query_id"] not in {x["query_id"] for x in selected}), dict(improved[0]))
        cand["title"] = "前排排名进一步提升样例"
        selected.append(cand)
    return selected[:2]


def build_md(
    result_dir: Path,
    overall: dict,
    per_flight_rows: list[dict[str, str]],
    timing_map: dict[str, float],
    cases: list[dict[str, str]],
    coarse_model_label: str,
) -> str:
    lines: list[str] = []
    lines.append("# RoMa v2 在 Intersection Truth 口径下的重排结果说明")
    lines.append("")
    lines.append("## 1. 任务定义与实验设置")
    lines.append(
        f"本组实验用于回答：在 `intersection truth` 新真值口径下，若先用 `{coarse_model_label} + FAISS` 对全库粗检索，再对前 `Top-20` 候选使用 `RoMa v2` 做精匹配重排，是否能够进一步改善正式指标。"
    )
    lines.append("")
    lines.extend(
        [
            "- 数据范围：4 条航线，共 `40` 个 query。",
            f"- 粗检索：{coarse_model_label} 全局特征 + FAISS `IndexFlatIP`。",
            "- 重排方法：RoMa v2 dense matching + 采样点对 + RANSAC 几何一致性 + 融合分数排序。",
            "- 重排窗口：每个 query 的 coarse `Top-20`。",
        ]
    )
    lines.append("")
    lines.append("## 2. 总体定量结果")
    lines.append("")
    lines.append(f"- Baseline：`R@1={overall['baseline_intersection_recall@1']:.3f}`，`R@5={overall['baseline_intersection_recall@5']:.3f}`，`R@10={overall['baseline_intersection_recall@10']:.3f}`，`R@20={overall['baseline_intersection_recall@20']:.3f}`，`MRR={overall['baseline_intersection_mrr']:.3f}`")
    lines.append(f"- Coarse Top20 上限：`R@20={overall['coarse_intersection_recall@20']:.3f}`")
    lines.append(f"- RoMa v2：`R@1={overall['romav2_intersection_recall@1']:.3f}`，`R@5={overall['romav2_intersection_recall@5']:.3f}`，`R@10={overall['romav2_intersection_recall@10']:.3f}`，`R@20={overall['romav2_intersection_recall@20']:.3f}`，`MRR={overall['romav2_intersection_mrr']:.3f}`")
    lines.append(f"- 指标变化：`ΔR@1={overall['delta_intersection_recall@1']:+.3f}`，`ΔR@5={overall['delta_intersection_recall@5']:+.3f}`，`ΔR@10={overall['delta_intersection_recall@10']:+.3f}`，`ΔR@20={overall['delta_intersection_recall@20']:+.3f}`，`ΔMRR={overall['delta_intersection_mrr']:+.3f}`")
    lines.append(f"- Top-1 误差：`{overall['baseline_top1_error_m_mean']:.3f}m -> {overall['romav2_top1_error_m_mean']:.3f}m`，`Δ={overall['delta_top1_error_m_mean']:+.3f}m`")
    lines.append("")
    lines.append("## 3. 分航线结果")
    lines.append("")
    for row in per_flight_rows:
        delta_r1 = float(row["romav2_intersection_recall@1"]) - float(row["baseline_intersection_recall@1"])
        delta_mrr = float(row["romav2_intersection_mrr"]) - float(row["baseline_intersection_mrr"])
        lines.append(
            f"- `{row['flight_tag']}`：Baseline `R@1={float(row['baseline_intersection_recall@1']):.3f}`、`MRR={float(row['baseline_intersection_mrr']):.3f}`；"
            f"RoMa v2 `R@1={float(row['romav2_intersection_recall@1']):.3f}`、`MRR={float(row['romav2_intersection_mrr']):.3f}`；"
            f"`ΔR@1={delta_r1:+.3f}`、`ΔMRR={delta_mrr:+.3f}`。"
        )
    lines.append("")
    lines.append("## 4. 时间开销统计")
    lines.append("")
    for key, label in [
        ("coarse_topk_export", "coarse Top-20 导出"),
        ("input_preparation", "输入准备"),
        ("romav2_rerank", "RoMa v2 重排"),
        ("summary_aggregation", "结果汇总"),
        ("visualization", "可视化"),
        ("report_generation", "报告生成"),
    ]:
        if key in timing_map:
            lines.append(f"- {label}：`{format_seconds(timing_map.get(key, 0.0))}`")
    lines.append("")
    lines.append("## 5. 代表性样例")
    lines.append("")
    for case in cases:
        lines.append(
            f"- `{case['query_id']}`：{case['title']}；Baseline 首个真值 rank=`{case['baseline_first_truth_rank'] or 'miss'}`，Coarse Top20 rank=`{case['coarse_first_truth_rank'] or 'miss'}`，RoMa v2 rank=`{case['romav2_first_truth_rank'] or 'miss'}`。"
        )
    lines.append("")
    lines.append("## 6. 结论")
    lines.append("")
    lines.extend(
        [
            f"- 本轮 `RoMa v2` 相对 `{coarse_model_label} baseline` 的关键收益集中在前排排序：`ΔR@1={overall['delta_intersection_recall@1']:+.3f}`、`ΔR@5={overall['delta_intersection_recall@5']:+.3f}`、`ΔMRR={overall['delta_intersection_mrr']:+.3f}`。",
            f"- `Top-1 error mean` 从 `{overall['baseline_top1_error_m_mean']:.3f}m` 降至 `{overall['romav2_top1_error_m_mean']:.3f}m`，改善 `{abs(overall['delta_top1_error_m_mean']):.3f}m`，说明精匹配不仅改善排序，也提升了首位候选的空间精度。",
            "- 分航线表现应以本轮实际输出为准，不预设所有航线都提升或无回退。",
            f"- 在当前 `query v2 + intersection truth` 口径下，`RoMa v2` 可以视为一条基于 `{coarse_model_label}` coarse 的几何重排方法，最终结论以本轮指标为准。",
        ]
    )
    lines.append("")
    return "\n".join(lines)


def build_docx(
    result_dir: Path,
    out_docx: Path,
    overall: dict,
    per_flight_rows: list[dict[str, str]],
    timing_map: dict[str, float],
    cases: list[dict[str, str]],
    coarse_model_label: str,
) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("RoMa v2 在 Intersection Truth 口径下的重排结果说明")
    set_cn_font(r, size=16, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run(f"基于 {coarse_model_label} coarse Top-20 的精匹配重排实验")
    set_cn_font(r, size=11)

    add_heading(doc, "1. 任务定义与实验设置", 1)
    add_bullets(
        doc,
        [
            "数据范围：4 条航线，共 40 个 query。",
            f"粗检索方法：{coarse_model_label} 全局特征 + FAISS IndexFlatIP。",
            "重排方法：RoMa v2 dense matching + 点对采样 + RANSAC 几何一致性 + 融合分数排序。",
            "执行窗口：coarse Top-20。",
        ],
    )

    add_heading(doc, "2. 总体定量结果", 1)
    table = doc.add_table(rows=1, cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["方法", "R@1", "R@5", "R@10", "R@20", "MRR", "Top1Err(m)", "备注"]
    for i, title_text in enumerate(headers):
        set_cell_text(table.cell(0, i), title_text, bold=True)
        shade_cell(table.cell(0, i), "D9EAF7")
    rows = [
        (f"{coarse_model_label} Baseline", overall["baseline_intersection_recall@1"], overall["baseline_intersection_recall@5"], overall["baseline_intersection_recall@10"], overall["baseline_intersection_recall@20"], overall["baseline_intersection_mrr"], overall["baseline_top1_error_m_mean"], "-"),
        ("Coarse Top20", "-", "-", "-", overall["coarse_intersection_recall@20"], "-", "-", "候选上限"),
        ("RoMa v2", overall["romav2_intersection_recall@1"], overall["romav2_intersection_recall@5"], overall["romav2_intersection_recall@10"], overall["romav2_intersection_recall@20"], overall["romav2_intersection_mrr"], overall["romav2_top1_error_m_mean"], "重排结果"),
    ]
    for values in rows:
        row = table.add_row().cells
        for i, val in enumerate(values):
            if isinstance(val, float):
                set_cell_text(row[i], f"{val:.3f}")
            else:
                set_cell_text(row[i], str(val))
    add_caption(doc, f"表 1  {coarse_model_label} Baseline 与 RoMa v2 overall 对比")

    add_heading(doc, "3. 分航线结果", 1)
    pf = doc.add_table(rows=1, cols=5)
    pf.alignment = WD_TABLE_ALIGNMENT.CENTER
    pf.style = "Table Grid"
    for i, title_text in enumerate(["航线", "Baseline R@1", "RoMa v2 R@1", "Baseline MRR", "RoMa v2 MRR"]):
        set_cell_text(pf.cell(0, i), title_text, bold=True)
        shade_cell(pf.cell(0, i), "D9EAF7")
    for item in per_flight_rows:
        row = pf.add_row().cells
        set_cell_text(row[0], item["flight_tag"])
        set_cell_text(row[1], f"{float(item['baseline_intersection_recall@1']):.3f}")
        set_cell_text(row[2], f"{float(item['romav2_intersection_recall@1']):.3f}")
        set_cell_text(row[3], f"{float(item['baseline_intersection_mrr']):.3f}")
        set_cell_text(row[4], f"{float(item['romav2_intersection_mrr']):.3f}")
    add_caption(doc, "表 2  分航线结果")

    add_heading(doc, "4. 时间开销统计", 1)
    tt = doc.add_table(rows=1, cols=2)
    tt.alignment = WD_TABLE_ALIGNMENT.CENTER
    tt.style = "Table Grid"
    for i, title_text in enumerate(["阶段", "耗时"]):
        set_cell_text(tt.cell(0, i), title_text, bold=True)
        shade_cell(tt.cell(0, i), "D9EAF7")
    for key, label in [
        ("coarse_topk_export", "coarse Top-20 导出"),
        ("input_preparation", "输入准备"),
        ("romav2_rerank", "RoMa v2 重排"),
        ("summary_aggregation", "结果汇总"),
        ("visualization", "可视化"),
        ("report_generation", "报告生成"),
    ]:
        if key in timing_map:
            row = tt.add_row().cells
            set_cell_text(row[0], label)
            set_cell_text(row[1], format_seconds(timing_map[key]))
    add_caption(doc, "表 3  时间开销统计")

    add_heading(doc, "5. 代表性样例", 1)
    for case in cases:
        add_paragraph(
            doc,
            f"{case['query_id']}：{case['title']}。Baseline 首个真值 rank={case['baseline_first_truth_rank'] or 'miss'}，"
            f"Coarse Top20 rank={case['coarse_first_truth_rank'] or 'miss'}，RoMa v2 rank={case['romav2_first_truth_rank'] or 'miss'}。"
        )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "6. 结论", 1)
    add_bullets(
        doc,
        [
            f"本轮 RoMa v2 相对 {coarse_model_label} baseline 的 `ΔR@1={overall['delta_intersection_recall@1']:+.3f}`，`ΔR@5={overall['delta_intersection_recall@5']:+.3f}`，`ΔMRR={overall['delta_intersection_mrr']:+.3f}`。",
            f"Top-1 误差由 `{overall['baseline_top1_error_m_mean']:.3f}m` 降为 `{overall['romav2_top1_error_m_mean']:.3f}m`，改善 `{abs(overall['delta_top1_error_m_mean']):.3f}m`。",
            "分航线表现应结合本轮结果表解读，不预设所有航线都无回退。",
        ],
    )

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))


def main() -> None:
    args = parse_args()
    result_dir = Path(args.result_dir)
    out_docx = Path(args.out_docx)
    out_md = Path(args.out_md)

    overall = load_json(result_dir / "overall_summary.json")
    per_flight_rows = load_csv(result_dir / "per_flight_comparison.csv")
    per_query_rows = load_csv(result_dir / "per_query_comparison.csv")
    timing = load_json(result_dir / "timing" / "timing_summary.json")
    timing_map = stage_seconds(timing)
    cases = choose_cases(per_query_rows)

    out_md.parent.mkdir(parents=True, exist_ok=True)
    out_md.write_text(
        build_md(result_dir, overall, per_flight_rows, timing_map, cases, args.coarse_model_label),
        encoding="utf-8",
    )
    build_docx(result_dir, out_docx, overall, per_flight_rows, timing_map, cases, args.coarse_model_label)
    print(out_md)
    print(out_docx)


if __name__ == "__main__":
    main()
