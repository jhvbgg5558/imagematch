#!/usr/bin/env python3
"""Generate a Word analysis report for LightGlue rerank failure modes."""

from __future__ import annotations

import argparse
import csv
import json
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--baseline-result-dir", required=True)
    parser.add_argument("--lightglue-result-dir", required=True)
    parser.add_argument("--out-docx", required=True)
    return parser.parse_args()


def load_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", newline="", encoding="utf-8-sig") as f:
        return list(csv.DictReader(f))


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def set_cn_font(run, size: int = 11, bold: bool = False) -> None:
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(size)
    run.bold = bold


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 10, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    set_cn_font(r, size=size, bold=bold)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.style = "Heading 1" if level == 1 else "Heading 2"
    r = p.add_run(text)
    set_cn_font(r, size=14 if level == 1 else 12, bold=True)


def add_paragraph(doc: Document, text: str, size: int = 11, center: bool = False) -> None:
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn_font(r, size=size)


def add_bullets(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(line)
        set_cn_font(r, size=11)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn_font(r, size=10)


def add_picture(doc: Document, path: Path, width_inch: float) -> None:
    doc.add_picture(str(path), width=Inches(width_inch))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def short_flight_name(flight_id: str) -> str:
    parts = flight_id.split("_")
    return parts[2] if len(parts) >= 3 else flight_id


def rank_to_int(v: str) -> int | None:
    return None if v == "" else int(v)


def load_reranked_by_query(lightglue_dir: Path) -> dict[str, list[dict[str, str]]]:
    out: dict[str, list[dict[str, str]]] = {}
    for csv_path in sorted((lightglue_dir / "stage7").glob("*/reranked_top20.csv")):
        for row in load_csv(csv_path):
            out.setdefault(row["query_id"], []).append(row)
    for rows in out.values():
        rows.sort(key=lambda x: int(x["rank"]))
    return out


def load_baseline_by_query(baseline_dir: Path) -> dict[str, list[dict[str, str]]]:
    out: dict[str, list[dict[str, str]]] = {}
    for row in load_csv(baseline_dir / "retrieval" / "retrieval_top10.csv"):
        out.setdefault(row["query_id"], []).append(row)
    for rows in out.values():
        rows.sort(key=lambda x: int(x["rank"]))
    return out


def compute_transition_stats(comp_rows: list[dict[str, str]]) -> dict[str, int]:
    stats = Counter()
    for row in comp_rows:
        b = rank_to_int(row["baseline_first_strict_truth_rank"])
        l = rank_to_int(row["lightglue_first_strict_truth_rank"])
        if b is None and l is None:
            stats["miss_to_miss"] += 1
        elif b is None and l is not None:
            stats["miss_to_hit"] += 1
            if l <= 10:
                stats["miss_to_top10"] += 1
            else:
                stats["miss_to_11_20"] += 1
        elif b is not None and l is None:
            stats["hit_to_miss"] += 1
        elif l < b:
            stats["improve"] += 1
        elif l > b:
            stats["worse"] += 1
        else:
            stats["same"] += 1
    stats["baseline_top1"] = sum(1 for r in comp_rows if r["baseline_first_strict_truth_rank"] == "1")
    stats["lightglue_top1"] = sum(1 for r in comp_rows if r["lightglue_first_strict_truth_rank"] == "1")
    stats["baseline_top10"] = sum(int(r["baseline_strict_hit@10"]) for r in comp_rows)
    stats["lightglue_top10"] = sum(int(r["lightglue_strict_hit@10"]) for r in comp_rows)
    stats["promoted_11_20_to_top10"] = sum(int(r["promoted_11_20_to_top10"]) for r in comp_rows)
    stats["coarse_hit20"] = sum(int(r["coarse_strict_hit@20"]) for r in comp_rows)
    stats["coarse_miss20"] = len(comp_rows) - stats["coarse_hit20"]
    return dict(stats)


def compute_top1_inlier_stats(reranked: dict[str, list[dict[str, str]]]) -> dict[str, object]:
    dist = Counter()
    hit_dist = Counter()
    for rows in reranked.values():
        top1 = rows[0]
        key = top1["inlier_count"]
        dist[key] += 1
        if top1["is_strict_truth_hit"] == "1":
            hit_dist[key] += 1
    return {"dist": dict(dist), "hit_dist": dict(hit_dist)}


def find_row(rows: list[dict[str, str]], query_id: str, tile_id: str) -> dict[str, str] | None:
    for row in rows:
        if row.get("query_id") == query_id and row.get("candidate_tile_id") == tile_id:
            return row
    return None


def choose_cases(comp_rows: list[dict[str, str]]) -> dict[str, dict[str, str] | None]:
    degraded_top1 = next(
        (
            r
            for r in comp_rows
            if r["baseline_first_strict_truth_rank"] == "1" and r["lightglue_first_strict_truth_rank"] != "1"
        ),
        None,
    )
    promoted = [r for r in comp_rows if int(r["promoted_11_20_to_top10"]) == 1]
    promoted.sort(key=lambda r: r["query_id"])
    unresolved = [
        r
        for r in comp_rows
        if int(r["coarse_strict_hit@20"]) == 1 and int(r["lightglue_strict_hit@10"]) == 0
    ]
    unresolved.sort(key=lambda r: (rank_to_int(r["coarse_first_strict_truth_rank"]) or 99, r["query_id"]))
    return {
        "degraded_top1": degraded_top1,
        "promoted_a": promoted[0] if len(promoted) >= 1 else None,
        "promoted_b": promoted[1] if len(promoted) >= 2 else None,
        "unresolved": unresolved[0] if unresolved else None,
    }


def build_overall_table(doc: Document, overall: dict) -> None:
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["方法", "R@1", "R@5", "R@10", "MRR", "Top-1误差均值(m)"]
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, bold=True)
        shade_cell(table.cell(0, i), "D9EAF7")
    rows = [
        ("Baseline", overall["baseline_strict_recall@1"], overall["baseline_strict_recall@5"], overall["baseline_strict_recall@10"], overall["baseline_strict_mrr"], overall["baseline_top1_error_m_mean"]),
        ("LightGlue", overall["lightglue_strict_recall@1"], overall["lightglue_strict_recall@5"], overall["lightglue_strict_recall@10"], overall["lightglue_strict_mrr"], overall["lightglue_top1_error_m_mean"]),
        ("Delta", overall["delta_strict_recall@1"], overall["delta_strict_recall@5"], overall["delta_strict_recall@10"], overall["delta_strict_mrr"], overall["lightglue_top1_error_m_mean"] - overall["baseline_top1_error_m_mean"]),
    ]
    for vals in rows:
        row = table.add_row().cells
        set_cell_text(row[0], str(vals[0]))
        for i, val in enumerate(vals[1:], start=1):
            if vals[0] == "Delta":
                set_cell_text(row[i], f"{float(val):+.3f}")
            else:
                set_cell_text(row[i], f"{float(val):.3f}")


def build_transition_table(doc: Document, stats: dict[str, int]) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["现象", "数量", "解释"]
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, bold=True)
        shade_cell(table.cell(0, i), "D9EAF7")
    rows = [
        ("Baseline Top-1 命中", stats["baseline_top1"], "说明原始粗检索已经有 7 个 query 首位正确。"),
        ("LightGlue Top-1 命中", stats["lightglue_top1"], "LightGlue 后首位正确样本减少到 6 个。"),
        ("Baseline Top-10 命中", stats["baseline_top10"], "Baseline 在 Top-10 内命中 strict truth 的 query 数。"),
        ("LightGlue Top-10 命中", stats["lightglue_top10"], "LightGlue 在 Top-10 内多救回了 2 个 query。"),
        ("11..20 拉回 Top-10", stats["promoted_11_20_to_top10"], "真正把真值从 11..20 推回 Top-10 的 query 数。"),
        ("Top-20 外无解", stats["coarse_miss20"], "这些 query 真值不在 coarse Top-20 中，LightGlue 无法挽救。"),
        ("miss -> 11..20", stats["miss_to_11_20"], "LightGlue 只把真值推进到 Top-20，但未进入正式 Top-10。"),
        ("排名变差", stats["worse"], "LightGlue 把已有命中的真值往后压的 query 数。"),
    ]
    for name, count, desc in rows:
        row = table.add_row().cells
        set_cell_text(row[0], name)
        set_cell_text(row[1], str(count))
        set_cell_text(row[2], desc, align=WD_ALIGN_PARAGRAPH.LEFT)


def build_inlier_table(doc: Document, inlier_stats: dict[str, object]) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Top-1 inlier_count", "样本数", "其中 strict truth Top-1 数"]
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, bold=True)
        shade_cell(table.cell(0, i), "D9EAF7")
    dist = inlier_stats["dist"]
    hit_dist = inlier_stats["hit_dist"]
    for key in sorted(dist, key=lambda x: int(x)):
        row = table.add_row().cells
        set_cell_text(row[0], str(key))
        set_cell_text(row[1], str(dist[key]))
        set_cell_text(row[2], str(hit_dist.get(key, 0)))


def explain_case(doc: Document, title: str, case: dict[str, str], baseline_dir: Path, lightglue_dir: Path, baseline_rows: dict[str, list[dict[str, str]]], reranked_rows: dict[str, list[dict[str, str]]]) -> None:
    qid = case["query_id"]
    flight_id = case["flight_id"]
    b_rank = case["baseline_first_strict_truth_rank"] or "未命中"
    c_rank = case["coarse_first_strict_truth_rank"] or "未命中"
    l_rank = case["lightglue_first_strict_truth_rank"] or "未命中"
    add_heading(doc, title, 2)
    add_paragraph(
        doc,
        f"{qid} / 航线 {short_flight_name(flight_id)}：Baseline 首个 strict truth 排名为 {b_rank}，"
        f"Coarse Top20 中的首个 strict truth 排名为 {c_rank}，LightGlue 后为 {l_rank}。",
    )
    if title.startswith("4.1"):
        truth_tile = baseline_rows[qid][0]["candidate_tile_id"]
        truth_row = find_row(reranked_rows[qid], qid, truth_tile)
        top1_row = reranked_rows[qid][0]
        if truth_row is not None:
            add_bullets(
                doc,
                [
                    f"Baseline 的 rank1 本身就是 strict truth，但 LightGlue 后掉到 rank {truth_row['rank']}。",
                    f"该真值 tile 的 global score={float(truth_row['global_score']):.3f}，fused score={float(truth_row['fused_score']):.3f}，inlier_count={truth_row['inlier_count']}。",
                    f"新的 rank1 候选不是 strict truth，但它拿到了 inlier_count={top1_row['inlier_count']} 和更高的 fused score={float(top1_row['fused_score']):.3f}。",
                    "这说明局部几何分支没有稳定保住原本已经正确的首位候选，反而让邻近假阳性候选获得了排序优势。",
                ],
            )
    elif title.startswith("4.2"):
        add_bullets(
            doc,
            [
                f"该样本在 baseline Top-10 内完全未命中，但 coarse Top20 已经把真值放到 rank {c_rank}。",
                f"LightGlue 后首个 strict truth 进入 rank {l_rank}，说明局部重排确实能把一部分 11..20 的真值拉回正式统计区间。",
                "这类样本解释了为什么 Recall@10 会提升，但它们数量很少，所以无法显著拉高整体主指标。",
            ],
        )
    else:
        add_bullets(
            doc,
            [
                f"该样本在 coarse Top20 内存在 strict truth（rank {c_rank}），但 LightGlue 后仍未进入 Top-10。",
                "说明当前局部匹配只能把真值从完全 miss 推进到 Top-20 或维持原状，但还不足以稳定压过前排的高相似度假阳性候选。",
                "这类样本解释了为什么 Recall@20 已有空间，但 Recall@1/5/10 提升依然有限。",
            ],
        )
    add_picture(doc, baseline_dir / "figures" / flight_id / f"{qid}_top10.png", 5.8)
    add_caption(doc, f"{qid} 的 Baseline Top-10 联系图")
    add_picture(doc, lightglue_dir / "figures" / flight_id / f"{qid}_top10.png", 5.8)
    add_caption(doc, f"{qid} 的 LightGlue Top-10 联系图")


def main() -> None:
    args = parse_args()
    baseline_dir = Path(args.baseline_result_dir)
    lightglue_dir = Path(args.lightglue_result_dir)
    out_docx = Path(args.out_docx)

    overall = load_json(lightglue_dir / "overall_summary.json")
    comp_rows = load_csv(lightglue_dir / "per_query_comparison.csv")
    transition_stats = compute_transition_stats(comp_rows)
    reranked_rows = load_reranked_by_query(lightglue_dir)
    baseline_rows = load_baseline_by_query(baseline_dir)
    inlier_stats = compute_top1_inlier_stats(reranked_rows)
    cases = choose_cases(comp_rows)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("LightGlue 重排为何 Recall@1 下降且整体提升有限：结果原因分析")
    set_cn_font(r, size=16, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("基于 strict truth 口径的逐 query 对照分析")
    set_cn_font(r, size=11)

    add_heading(doc, "1. 现象概览", 1)
    add_paragraph(
        doc,
        "本报告不重复描述实验设置，而是直接回答两个问题："
        "为什么 SuperPoint + LightGlue 重排后 `Recall@1` 不增反降，以及为什么其余主指标提升也不明显。"
    )
    build_overall_table(doc, overall)
    add_caption(doc, "表 1  Baseline 与 LightGlue 的主指标对照")
    add_picture(doc, lightglue_dir / "figures" / "_compare" / "baseline_vs_lightglue_compare.png", 6.3)
    add_caption(doc, "图 1  Baseline 与 LightGlue 的 overall 对照")
    add_bullets(
        doc,
        [
            f"Baseline Top-1 命中数为 {transition_stats['baseline_top1']}，LightGlue 后变为 {transition_stats['lightglue_top1']}，因此 `Recall@1` 从 {overall['baseline_strict_recall@1']:.3f} 降到 {overall['lightglue_strict_recall@1']:.3f}。",
            f"Baseline Top-10 命中数为 {transition_stats['baseline_top10']}，LightGlue 后变为 {transition_stats['lightglue_top10']}，所以 `Recall@10` 只净增 2 个 query。",
            f"Coarse Top20 的上限仅为 {overall['coarse_strict_recall@20']:.3f}，意味着仍有 {transition_stats['coarse_miss20']} 个 query 真值根本不在重排窗口内。",
        ],
    )

    add_heading(doc, "2. 逐 query 变化统计", 1)
    build_transition_table(doc, transition_stats)
    add_caption(doc, "表 2  LightGlue 重排后的 query 变化类型统计")
    add_paragraph(
        doc,
        "从统计上看，这轮收益主要来自“把 11..20 的真值拉回 Top-10”，但这种样本只有 2 个。"
        "与此同时，有 1 个原本已经 Top-1 命中的样本被压到后面，另外还有 3 个已命中样本排名变差，"
        "这就足以把 `Recall@1` 和 `MRR` 的收益抵消掉。",
    )

    add_heading(doc, "3. 为什么 Recall@1 会下降", 1)
    add_paragraph(
        doc,
        "这轮 `Recall@1` 下降并不是因为 LightGlue 全面失败，而是因为它没有稳定地保住已经正确的第一名。"
        "在 40 个 query 中，Baseline 原本有 7 个样本首位就是真值；LightGlue 后只剩 6 个。"
    )
    build_inlier_table(doc, inlier_stats)
    add_caption(doc, "表 3  LightGlue 最终 rank1 候选的 inlier_count 分布")
    add_bullets(
        doc,
        [
            "40 个 query 的最终 rank1 候选中，有 33 个只有 4 个 inlier，只有 7 个是 5 个 inlier，说明局部几何信号整体分辨率很低。",
            "所有成功的 strict-truth Top-1 都落在 4-inlier 桶里，而 5-inlier 并没有带来更可靠的首位正确率，说明当前 inlier 数本身并不足以稳定区分真值与假阳性。",
            "当真值 tile 没有拿到几何加分，或者拿到的加分与邻近假阳性几乎一致时，融合排序就可能把原本正确的第一名往后压。",
        ],
    )

    if cases["degraded_top1"] is not None:
        explain_case(doc, "4.1 Top-1 被压坏的代表样例", cases["degraded_top1"], baseline_dir, lightglue_dir, baseline_rows, reranked_rows)

    add_heading(doc, "5. 为什么其他指标提升也不明显", 1)
    add_bullets(
        doc,
        [
            f"首先，coarse Top20 的上限只有 {overall['coarse_strict_recall@20']:.3f}，还有 {transition_stats['coarse_miss20']} 个 query 完全无解，LightGlue 不可能在这些样本上带来任何收益。",
            f"其次，真正把真值从 11..20 拉回 Top-10 的 query 只有 {transition_stats['promoted_11_20_to_top10']} 个，因此 `Recall@10` 只能小幅提升。",
            f"再次，有 {transition_stats['miss_to_11_20']} 个 query 虽然在 LightGlue 后第一次进入了 Top-20，但仍停留在 11..20，不能进入正式 `Recall@10` 统计。",
            "最后，局部匹配分支更像是在改善“候选覆盖范围”，而不是稳定提纯第一名，因此 `Recall@1` 和 `MRR` 的改善会比 `Recall@10` 更难出现。",
        ],
    )
    add_picture(doc, lightglue_dir / "figures" / "_aggregate" / "top20_upper_bound.png", 6.2)
    add_caption(doc, "图 2  Coarse Top20 上限与 LightGlue 最终 Top-10 的差距")

    if cases["promoted_a"] is not None:
        explain_case(doc, "6.1 从 11..20 拉回 Top-10 的样例 A", cases["promoted_a"], baseline_dir, lightglue_dir, baseline_rows, reranked_rows)
    if cases["promoted_b"] is not None:
        explain_case(doc, "6.2 从 11..20 拉回 Top-10 的样例 B", cases["promoted_b"], baseline_dir, lightglue_dir, baseline_rows, reranked_rows)
    if cases["unresolved"] is not None:
        explain_case(doc, "6.3 Top-20 内仍救不回来的样例", cases["unresolved"], baseline_dir, lightglue_dir, baseline_rows, reranked_rows)

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "7. 结论", 1)
    add_bullets(
        doc,
        [
            "LightGlue 当前确实有用，但它的主要作用是把一小部分 11..20 的真值拉回 Top-10，而不是稳定提升首位候选质量。",
            "Recall@1 下降的直接原因是：少数原本已经正确的 Top-1 被局部重排压到了后面，典型样例是 q_006。",
            "整体提升有限的根本原因是：Top-20 可挽救空间本来就有限，而且局部几何分支的判别力不够强，许多候选只呈现同质化的 4-inlier 信号。",
            "如果后续继续优化，应优先处理两类问题：一是保护原本已经正确的 Top-1，二是提高局部几何分数对真值与邻近假阳性的区分度。",
        ],
    )

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))
    print(out_docx)


if __name__ == "__main__":
    main()
