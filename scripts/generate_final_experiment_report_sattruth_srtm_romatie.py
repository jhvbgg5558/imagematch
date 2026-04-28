#!/usr/bin/env python3
"""Generate the final Markdown + Word report for the satellite-truth SRTM route.

Purpose:
- write a standalone final experiment report for the completed
  `new3output/nadir_009010_dinov2_romav2_pose_sattruth_srtm_romatie_2026-04-16`
  full run;
- summarize the current runtime pose results, the three-layer validation suite,
  and the baseline comparison in one Chinese academic-style document;
- generate report-specific figures and sample-case panels under the experiment
  `reports/` directory without changing the runtime task definition.

Main inputs:
- current experiment outputs under `pose_v1_formal/summary/`,
  `pose_v1_formal/pnp/`, and
  `pose_v1_formal/eval_pose_validation_suite_sattruth_srtm/`;
- baseline comparison inputs under
  `new2output/nadir_009010_dinov2_romav2_pose_2026-04-10/pose_v1_formal/`;
- query selection manifests under `selected_queries/` and `query_inputs/`.

Main outputs:
- `<experiment-root>/reports/final_experiment_report_sattruth_srtm_romatie.md`;
- `<experiment-root>/reports/final_experiment_report_sattruth_srtm_romatie.docx`;
- `<experiment-root>/reports/final_experiment_report_assets/`.

Applicable task constraints:
- runtime query remains a single arbitrary UAV image without geographic
  metadata; the report must not describe validation truth as a runtime input;
- runtime retrieval/rerank/candidate library remain fixed to the baseline
  satellite-library pipeline;
- the report must keep layer-1 / layer-2 / layer-3 meanings aligned with the
  current formal suite: satellite truth alignment, pose-vs-AT, and RoMa v2
  tie-point ground error.
"""

from __future__ import annotations

import argparse
import csv
import json
import math
import shutil
import statistics
from collections import Counter
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from generate_new3_full_experiment_report import (
    mask_to_rgb,
    normalize_rgb,
    overlay_truth_and_pred,
    read_rgba_or_rgb,
)


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_EXPERIMENT_ROOT = (
    PROJECT_ROOT
    / "new3output"
    / "nadir_009010_dinov2_romav2_pose_sattruth_srtm_romatie_2026-04-16"
)
DEFAULT_BASELINE_ROOT = PROJECT_ROOT / "new2output" / "nadir_009010_dinov2_romav2_pose_2026-04-10"
DEFAULT_SUITE_ROOT = DEFAULT_EXPERIMENT_ROOT / "pose_v1_formal" / "eval_pose_validation_suite_sattruth_srtm"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--experiment-root", default=str(DEFAULT_EXPERIMENT_ROOT))
    parser.add_argument("--baseline-root", default=str(DEFAULT_BASELINE_ROOT))
    parser.add_argument("--suite-root", default=str(DEFAULT_SUITE_ROOT))
    parser.add_argument("--out-md", default="")
    parser.add_argument("--out-docx", default="")
    parser.add_argument("--assets-dir", default="")
    return parser.parse_args()


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def load_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", newline="", encoding="utf-8-sig") as handle:
        return list(csv.DictReader(handle))


def load_json(path: Path) -> dict[str, object]:
    return json.loads(path.read_text(encoding="utf-8"))


def write_text(path: Path, text: str) -> None:
    ensure_dir(path.parent)
    path.write_text(text, encoding="utf-8")


def fmt(value: object, digits: int = 4) -> str:
    if value in ("", None):
        return "-"
    try:
        numeric = float(value)
    except (TypeError, ValueError):
        return str(value)
    if not math.isfinite(numeric):
        return "-"
    return f"{numeric:.{digits}f}"


def parse_float(value: str | None) -> float | None:
    if value in ("", None):
        return None
    try:
        return float(value)
    except ValueError:
        return None


def mean_metric(summary: dict[str, object], key: str) -> float:
    payload = summary.get(key, {})
    if isinstance(payload, dict) and "mean" in payload:
        return float(payload.get("mean", float("nan")))
    for value in summary.values():
        if isinstance(value, dict) and key in value:
            nested = value.get(key)
            if isinstance(nested, dict) and "mean" in nested:
                return float(nested.get("mean", float("nan")))
    return float("nan")


def short_flight_name(flight_id: str) -> str:
    parts = flight_id.split("_")
    return parts[2] if len(parts) >= 3 else flight_id


def relative_to(root: Path, target: Path) -> str:
    try:
        return str(target.relative_to(root)).replace("\\", "/")
    except ValueError:
        return str(target).replace("\\", "/")


def relative_to_cwd(path: Path) -> str:
    try:
        return str(path.relative_to(Path.cwd()))
    except ValueError:
        return str(path)


def set_cn_font(run, size: int = 11, bold: bool = False) -> None:
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(size)
    run.bold = bold


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = "Heading 1" if level == 1 else "Heading 2"
    run = paragraph.add_run(text)
    set_cn_font(run, size=14 if level == 1 else 12, bold=True)


def add_paragraph(
    doc: Document,
    text: str,
    *,
    size: int = 11,
    center: bool = False,
    bold: bool = False,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_cn_font(run, size=size, bold=bold)


def add_bullets(doc: Document, lines: list[str]) -> None:
    for line in lines:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(line)
        set_cn_font(run, size=11)


def set_cell_text(cell, text: str, *, bold: bool = False, size: int = 10) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_cn_font(run, size=size, bold=bold)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], *, header_fill: str = "D9EAF7") -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_text(table.cell(0, idx), header, bold=True)
        shade_cell(table.cell(0, idx), header_fill)
    for values in rows:
        row = table.add_row().cells
        for idx, value in enumerate(values):
            set_cell_text(row[idx], value)


def add_picture_with_caption(doc: Document, image_path: Path, caption: str, *, width_inch: float = 6.2) -> None:
    if not image_path.exists():
        add_paragraph(doc, f"图像缺失：{image_path}", size=10)
        return
    doc.add_picture(str(image_path), width=Inches(width_inch))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(caption)
    set_cn_font(run, size=10)


def require_file(path: Path) -> Path:
    if not path.exists():
        raise FileNotFoundError(f"Missing required input: {path}")
    return path


def to_local_path(value: str | Path) -> Path:
    text = str(value)
    posix_path = Path(text)
    if text.startswith("/mnt/") or posix_path.exists():
        return posix_path
    if text.startswith("/mnt/") and len(text) > 6:
        drive = text[5].upper()
        remainder = text[6:].replace("/", "\\")
        return Path(f"{drive}:{remainder}")
    return Path(text)


def save_rgb_png(array: np.ndarray, path: Path) -> None:
    ensure_dir(path.parent)
    plt.imsave(path, np.clip(array, 0.0, 1.0))


def median_from_rows(rows: list[dict[str, str]], key: str) -> float:
    values = [parse_float(row.get(key)) for row in rows]
    clean = [value for value in values if value is not None and math.isfinite(value)]
    return float(statistics.median(clean)) if clean else float("nan")


def counter_to_text(counter: dict[str, object] | Counter) -> str:
    payload = dict(counter)
    ordered = sorted(payload.items(), key=lambda item: item[0])
    return ", ".join(f"{key}={value}" for key, value in ordered)


def tie_status_counts(payload: dict[str, object]) -> dict[str, object]:
    return dict(payload.get("status_counts") or payload.get("eval_status_counts") or {})


def load_required_inputs(
    experiment_root: Path,
    baseline_root: Path,
    suite_root: Path,
) -> dict[str, object]:
    current_pose_summary_path = require_file(experiment_root / "pose_v1_formal" / "summary" / "pose_overall_summary.json")
    current_suite_summary_path = require_file(suite_root / "full_run_summary.json")
    current_tie_overall_path = require_file(
        suite_root / "tiepoint_ground_error" / "overall_tiepoint_ground_error.json"
    )
    current_tie_csv_path = require_file(
        suite_root / "tiepoint_ground_error" / "per_query_tiepoint_ground_error.csv"
    )
    current_pose_csv_path = require_file(suite_root / "pose_vs_at" / "per_query_pose_vs_at.csv")
    current_query_manifest_path = require_file(experiment_root / "query_inputs" / "query_manifest.csv")
    current_selected_summary_path = require_file(experiment_root / "selected_queries" / "selected_images_summary.csv")
    comparison_md_path = require_file(experiment_root / "reports" / "sattruth_srtm_romatie_vs_baseline.md")

    baseline_pose_summary_path = require_file(baseline_root / "pose_v1_formal" / "summary" / "pose_overall_summary.json")
    baseline_suite_summary_path = require_file(baseline_root / "pose_v1_formal" / "eval_pose_validation_suite" / "full_run_summary.json")
    baseline_tie_overall_path = require_file(
        baseline_root / "pose_v1_formal" / "eval_pose_validation_suite" / "tiepoint_ground_error" / "overall_tiepoint_ground_error.json"
    )
    baseline_tie_csv_path = require_file(
        baseline_root / "pose_v1_formal" / "eval_pose_validation_suite" / "tiepoint_ground_error" / "per_query_tiepoint_ground_error.csv"
    )
    baseline_pose_csv_path = require_file(
        baseline_root / "pose_v1_formal" / "eval_pose_validation_suite" / "pose_vs_at" / "per_query_pose_vs_at.csv"
    )

    suite_summary = load_json(current_suite_summary_path)
    if suite_summary.get("phase") != "full":
        raise RuntimeError(f"Expected full suite summary, got phase={suite_summary.get('phase')!r}")

    return {
        "experiment_root": experiment_root,
        "baseline_root": baseline_root,
        "suite_root": suite_root,
        "current_pose_summary": load_json(current_pose_summary_path),
        "current_suite_summary": suite_summary,
        "current_tie_overall": load_json(current_tie_overall_path),
        "current_tie_rows": load_csv(current_tie_csv_path),
        "current_pose_rows": load_csv(current_pose_csv_path),
        "current_query_manifest": load_csv(current_query_manifest_path),
        "current_selected_summary": load_csv(current_selected_summary_path),
        "comparison_md": comparison_md_path.read_text(encoding="utf-8"),
        "baseline_pose_summary": load_json(baseline_pose_summary_path),
        "baseline_suite_summary": load_json(baseline_suite_summary_path),
        "baseline_tie_overall": load_json(baseline_tie_overall_path),
        "baseline_tie_rows": load_csv(baseline_tie_csv_path),
        "baseline_pose_rows": load_csv(baseline_pose_csv_path),
    }


def summarize_queries(
    selected_rows: list[dict[str, str]],
    query_rows: list[dict[str, str]],
) -> dict[str, object]:
    flight_counts = Counter(row["flight_id"] for row in query_rows)
    gimbal_values = [parse_float(row.get("gimbal_pitch_degree")) for row in selected_rows]
    abs_altitudes = [parse_float(row.get("absolute_altitude")) for row in selected_rows]
    rel_altitudes = [parse_float(row.get("relative_altitude")) for row in selected_rows]
    image_names = [row.get("image_name", "") for row in query_rows]
    return {
        "query_count": len(query_rows),
        "flight_counts": dict(sorted(flight_counts.items())),
        "gimbal_pitch_min": min(value for value in gimbal_values if value is not None),
        "gimbal_pitch_max": max(value for value in gimbal_values if value is not None),
        "absolute_altitude_mean": statistics.mean(value for value in abs_altitudes if value is not None),
        "relative_altitude_mean": statistics.mean(value for value in rel_altitudes if value is not None),
        "metadata_removed_count": sum(int(row.get("has_metadata_removed", "0")) for row in query_rows),
        "first_image_name": image_names[0] if image_names else "-",
        "last_image_name": image_names[-1] if image_names else "-",
    }


def index_rows(rows: list[dict[str, str]], key: str = "query_id") -> dict[str, dict[str, str]]:
    return {row[key]: row for row in rows if row.get(key)}


def build_comparison_rows(data: dict[str, object]) -> list[dict[str, object]]:
    baseline_tie_rows = data["baseline_tie_rows"]
    current_tie_rows = data["current_tie_rows"]
    baseline_pose_rows = data["baseline_pose_rows"]
    current_pose_rows = data["current_pose_rows"]
    baseline_tie_map = index_rows(baseline_tie_rows)
    current_tie_map = index_rows(current_tie_rows)
    baseline_pose_map = index_rows(baseline_pose_rows)
    current_pose_map = index_rows(current_pose_rows)
    rows: list[dict[str, object]] = []
    for query_id in sorted(current_tie_map):
        tie_current = current_tie_map[query_id]
        tie_baseline = baseline_tie_map.get(query_id, {})
        pose_current = current_pose_map.get(query_id, {})
        pose_baseline = baseline_pose_map.get(query_id, {})
        current_match = parse_float(tie_current.get("tiepoint_match_count")) or 0.0
        baseline_match = parse_float(tie_baseline.get("tiepoint_match_count")) or 0.0
        rows.append(
            {
                "query_id": query_id,
                "flight_id": tie_current.get("flight_id", ""),
                "baseline_match_count": baseline_match,
                "current_match_count": current_match,
                "match_delta": current_match - baseline_match,
                "baseline_inlier_count": parse_float(tie_baseline.get("tiepoint_inlier_count")) or 0.0,
                "current_inlier_count": parse_float(tie_current.get("tiepoint_inlier_count")) or 0.0,
                "baseline_rmse": parse_float(tie_baseline.get("tiepoint_xy_error_rmse_m")),
                "current_rmse": parse_float(tie_current.get("tiepoint_xy_error_rmse_m")),
                "baseline_p90": parse_float(tie_baseline.get("tiepoint_xy_error_p90_m")),
                "current_p90": parse_float(tie_current.get("tiepoint_xy_error_p90_m")),
                "baseline_horizontal": parse_float(pose_baseline.get("horizontal_error_m")),
                "current_horizontal": parse_float(pose_current.get("horizontal_error_m")),
                "baseline_view_dir": parse_float(pose_baseline.get("view_dir_angle_error_deg")),
                "current_view_dir": parse_float(pose_current.get("view_dir_angle_error_deg")),
                "truth_crop_path": tie_current.get("truth_crop_path", ""),
                "pred_crop_path": tie_current.get("pred_crop_path", ""),
            }
        )
    return rows


def choose_success_queries(comparison_rows: list[dict[str, object]], count: int = 2) -> list[str]:
    ordered = sorted(
        comparison_rows,
        key=lambda row: (
            row["current_rmse"] if row["current_rmse"] is not None else float("inf"),
            row["current_horizontal"] if row["current_horizontal"] is not None else float("inf"),
        ),
    )
    return [row["query_id"] for row in ordered[:count]]


def choose_anomaly_queries(comparison_rows: list[dict[str, object]], count: int = 2) -> list[str]:
    baseline_matches = [row["baseline_match_count"] for row in comparison_rows]
    threshold = statistics.quantiles(baseline_matches, n=4, method="inclusive")[0]
    low_match = [row for row in comparison_rows if row["baseline_match_count"] <= threshold]
    if len(low_match) < count:
        low_match = comparison_rows
    ordered = sorted(
        low_match,
        key=lambda row: (
            -(row["current_rmse"] if row["current_rmse"] is not None else -1.0),
            -(row["current_horizontal"] if row["current_horizontal"] is not None else -1.0),
        ),
    )
    picked: list[str] = []
    for row in ordered:
        if row["query_id"] not in picked:
            picked.append(row["query_id"])
        if len(picked) >= count:
            break
    return picked


def build_case_assets(
    comparison_rows: list[dict[str, object]],
    assets_dir: Path,
    *,
    success_queries: list[str],
    anomaly_queries: list[str],
    suite_root: Path,
) -> list[dict[str, object]]:
    case_dir = assets_dir / "sample_cases"
    ensure_dir(case_dir)
    row_map = {row["query_id"]: row for row in comparison_rows}
    output: list[dict[str, object]] = []
    for label, query_ids in (("success", success_queries), ("anomaly", anomaly_queries)):
        for query_id in query_ids:
            row = row_map[query_id]
            truth_path = to_local_path(row["truth_crop_path"])
            pred_path = to_local_path(row["pred_crop_path"])
            truth_rgb, truth_alpha = read_rgba_or_rgb(truth_path)
            pred_rgb, pred_alpha = read_rgba_or_rgb(pred_path)
            truth_rgb = normalize_rgb(truth_rgb)
            pred_rgb = normalize_rgb(pred_rgb)
            alpha = pred_alpha if pred_alpha is not None else truth_alpha
            overlay_rgb = overlay_truth_and_pred(truth_rgb, pred_rgb, alpha)
            mask_rgb = mask_to_rgb(alpha, pred_rgb)

            truth_png = case_dir / f"{query_id}_truth.png"
            pred_png = case_dir / f"{query_id}_pred.png"
            overlay_png = case_dir / f"{query_id}_overlay.png"
            mask_png = case_dir / f"{query_id}_mask.png"
            tie_overlay_src = (
                suite_root
                / "tiepoint_ground_error"
                / "viz_tiepoints"
                / f"{query_id}_tiepoints_overlay.png"
            )
            tie_overlay_png = case_dir / f"{query_id}_tiepoints_overlay.png"
            panel_png = case_dir / f"{query_id}_case_panel.png"

            save_rgb_png(truth_rgb, truth_png)
            save_rgb_png(pred_rgb, pred_png)
            save_rgb_png(overlay_rgb, overlay_png)
            save_rgb_png(mask_rgb, mask_png)
            if tie_overlay_src.exists():
                shutil.copyfile(tie_overlay_src, tie_overlay_png)

            figure, axes = plt.subplots(2, 2, figsize=(10, 10))
            axes = axes.ravel()
            images = [
                (truth_png, "Satellite truth patch"),
                (pred_png, "Predicted ortho"),
                (overlay_png, "Truth/pred overlay"),
                (tie_overlay_png if tie_overlay_png.exists() else mask_png, "RoMa v2 tiepoints"),
            ]
            for axis, (image_path, title) in zip(axes, images):
                axis.imshow(plt.imread(image_path))
                axis.set_title(title, fontsize=10)
                axis.axis("off")
            figure.suptitle(
                f"{query_id} | {label} | RMSE={fmt(row['current_rmse'], 3)} m | "
                f"horizontal={fmt(row['current_horizontal'], 3)} m",
                fontsize=11,
            )
            figure.tight_layout()
            figure.savefig(panel_png, dpi=180, bbox_inches="tight")
            plt.close(figure)

            output.append(
                {
                    "query_id": query_id,
                    "label": label,
                    "truth_png": truth_png,
                    "pred_png": pred_png,
                    "overlay_png": overlay_png,
                    "mask_png": mask_png,
                    "tie_overlay_png": tie_overlay_png if tie_overlay_png.exists() else mask_png,
                    "panel_png": panel_png,
                    "current_rmse": row["current_rmse"],
                    "current_horizontal": row["current_horizontal"],
                    "current_match_count": row["current_match_count"],
                    "baseline_match_count": row["baseline_match_count"],
                }
            )
    return output


def build_overall_metrics_chart(data: dict[str, object], assets_dir: Path) -> Path:
    path = assets_dir / "overall_metrics_comparison.png"
    labels = [
        "L1 phase_corr",
        "L1 ortho_iou",
        "L1 ssim",
        "L2 horiz err",
        "L2 view_dir",
        "L3 RMSE",
        "L3 p90",
    ]
    baseline = [
        mean_metric(data["baseline_suite_summary"], "phase_corr_error_m"),
        mean_metric(data["baseline_suite_summary"], "ortho_iou"),
        mean_metric(data["baseline_suite_summary"], "ssim"),
        mean_metric(data["baseline_suite_summary"], "horizontal_error_m"),
        mean_metric(data["baseline_suite_summary"], "view_dir_angle_error_deg"),
        float(data["baseline_tie_overall"]["tiepoint_xy_error_rmse_m"]),
        float(data["baseline_tie_overall"]["tiepoint_xy_error_p90_m"]),
    ]
    current = [
        mean_metric(data["current_suite_summary"], "phase_corr_error_m"),
        mean_metric(data["current_suite_summary"], "ortho_iou"),
        mean_metric(data["current_suite_summary"], "ssim"),
        mean_metric(data["current_suite_summary"], "horizontal_error_m"),
        mean_metric(data["current_suite_summary"], "view_dir_angle_error_deg"),
        float(data["current_tie_overall"]["tiepoint_xy_error_rmse_m"]),
        float(data["current_tie_overall"]["tiepoint_xy_error_p90_m"]),
    ]
    x = np.arange(len(labels))
    width = 0.38
    fig, ax = plt.subplots(figsize=(12, 5.5))
    ax.bar(x - width / 2, baseline, width, label="Baseline", color="#9DB4C0")
    ax.bar(x + width / 2, current, width, label="Current", color="#D36B4A")
    ax.set_xticks(x)
    ax.set_xticklabels(labels, rotation=20, ha="right")
    ax.set_title("Baseline vs current overall metrics")
    ax.legend()
    ax.grid(axis="y", alpha=0.25)
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return path


def build_layer_metric_chart(path: Path, labels: list[str], baseline: list[float], current: list[float], title: str) -> Path:
    x = np.arange(len(labels))
    width = 0.38
    fig, ax = plt.subplots(figsize=(8.8, 4.8))
    ax.bar(x - width / 2, baseline, width, label="Baseline", color="#9DB4C0")
    ax.bar(x + width / 2, current, width, label="Current", color="#D36B4A")
    ax.set_xticks(x)
    ax.set_xticklabels(labels)
    ax.set_title(title)
    ax.legend()
    ax.grid(axis="y", alpha=0.25)
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return path


def build_low_match_chart(comparison_rows: list[dict[str, object]], assets_dir: Path) -> Path:
    path = assets_dir / "low_match_queries_improvement.png"
    threshold = statistics.quantiles(
        [row["baseline_match_count"] for row in comparison_rows],
        n=4,
        method="inclusive",
    )[0]
    low_rows = [row for row in comparison_rows if row["baseline_match_count"] <= threshold]
    low_rows = sorted(low_rows, key=lambda row: row["match_delta"], reverse=True)
    labels = [row["query_id"] for row in low_rows]
    baseline = [row["baseline_match_count"] for row in low_rows]
    current = [row["current_match_count"] for row in low_rows]
    x = np.arange(len(labels))
    width = 0.38
    fig, ax = plt.subplots(figsize=(10.8, 4.8))
    ax.bar(x - width / 2, baseline, width, label="Baseline", color="#9DB4C0")
    ax.bar(x + width / 2, current, width, label="Current", color="#D36B4A")
    ax.set_xticks(x)
    ax.set_xticklabels(labels)
    ax.set_title("Low-match queries improvement")
    ax.legend()
    ax.grid(axis="y", alpha=0.25)
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return path


def build_runtime_status_chart(data: dict[str, object], assets_dir: Path) -> Path:
    path = assets_dir / "runtime_status_comparison.png"
    baseline_counter = data["baseline_pose_summary"]["score_status_counts"]
    current_counter = data["current_pose_summary"]["score_status_counts"]
    labels = sorted(set(baseline_counter) | set(current_counter))
    baseline = [baseline_counter.get(label, 0) for label in labels]
    current = [current_counter.get(label, 0) for label in labels]
    x = np.arange(len(labels))
    width = 0.38
    fig, ax = plt.subplots(figsize=(7.8, 4.6))
    ax.bar(x - width / 2, baseline, width, label="Baseline", color="#9DB4C0")
    ax.bar(x + width / 2, current, width, label="Current", color="#D36B4A")
    ax.set_xticks(x)
    ax.set_xticklabels(labels)
    ax.set_title("Runtime score status counts")
    ax.legend()
    ax.grid(axis="y", alpha=0.25)
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return path


def build_pipeline_overview_chart(assets_dir: Path) -> Path:
    path = assets_dir / "pipeline_overview.png"
    fig, ax = plt.subplots(figsize=(11.5, 2.4))
    ax.axis("off")
    steps = [
        "UAV query image",
        "DINOv2 coarse retrieval",
        "RoMa v2 rerank",
        "SRTM-backed pose / PnP",
        "Layer-1 satellite truth",
        "Layer-2 pose_vs_at",
        "Layer-3 RoMa tiepoints",
    ]
    x_positions = np.linspace(0.08, 0.92, len(steps))
    for idx, (x_pos, text) in enumerate(zip(x_positions, steps)):
        rect = plt.Rectangle((x_pos - 0.065, 0.38), 0.13, 0.24, color="#F1E3D3", ec="#7D5A50", lw=1.5)
        ax.add_patch(rect)
        ax.text(x_pos, 0.5, text, ha="center", va="center", fontsize=10)
        if idx < len(steps) - 1:
            ax.annotate("", xy=(x_positions[idx + 1] - 0.07, 0.5), xytext=(x_pos + 0.07, 0.5), arrowprops=dict(arrowstyle="->", lw=1.4))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return path


def build_report_assets(data: dict[str, object], assets_dir: Path) -> dict[str, object]:
    ensure_dir(assets_dir)
    comparison_rows = build_comparison_rows(data)
    success_queries = choose_success_queries(comparison_rows)
    anomaly_queries = choose_anomaly_queries(comparison_rows)
    case_assets = build_case_assets(
        comparison_rows,
        assets_dir,
        success_queries=success_queries,
        anomaly_queries=anomaly_queries,
        suite_root=data["suite_root"],
    )
    charts = {
        "overall_metrics_comparison": build_overall_metrics_chart(data, assets_dir),
        "layer1_metrics_bar": build_layer_metric_chart(
            assets_dir / "layer1_metrics_bar.png",
            ["phase_corr_error_m", "ortho_iou", "ssim"],
            [
                mean_metric(data["baseline_suite_summary"], "phase_corr_error_m"),
                mean_metric(data["baseline_suite_summary"], "ortho_iou"),
                mean_metric(data["baseline_suite_summary"], "ssim"),
            ],
            [
                mean_metric(data["current_suite_summary"], "phase_corr_error_m"),
                mean_metric(data["current_suite_summary"], "ortho_iou"),
                mean_metric(data["current_suite_summary"], "ssim"),
            ],
            "Layer-1 metrics",
        ),
        "layer2_metrics_bar": build_layer_metric_chart(
            assets_dir / "layer2_metrics_bar.png",
            ["horizontal_error_m", "view_dir_angle_error_deg"],
            [
                mean_metric(data["baseline_suite_summary"], "horizontal_error_m"),
                mean_metric(data["baseline_suite_summary"], "view_dir_angle_error_deg"),
            ],
            [
                mean_metric(data["current_suite_summary"], "horizontal_error_m"),
                mean_metric(data["current_suite_summary"], "view_dir_angle_error_deg"),
            ],
            "Layer-2 metrics",
        ),
        "layer3_metrics_bar": build_layer_metric_chart(
            assets_dir / "layer3_metrics_bar.png",
            [
                "match_count_median",
                "inlier_count_median",
                "rmse_m",
                "p90_m",
            ],
            [
                median_from_rows(data["baseline_tie_rows"], "tiepoint_match_count"),
                median_from_rows(data["baseline_tie_rows"], "tiepoint_inlier_count"),
                float(data["baseline_tie_overall"]["tiepoint_xy_error_rmse_m"]),
                float(data["baseline_tie_overall"]["tiepoint_xy_error_p90_m"]),
            ],
            [
                median_from_rows(data["current_tie_rows"], "tiepoint_match_count"),
                median_from_rows(data["current_tie_rows"], "tiepoint_inlier_count"),
                float(data["current_tie_overall"]["tiepoint_xy_error_rmse_m"]),
                float(data["current_tie_overall"]["tiepoint_xy_error_p90_m"]),
            ],
            "Layer-3 metrics",
        ),
        "low_match_queries_improvement": build_low_match_chart(comparison_rows, assets_dir),
        "runtime_status_comparison": build_runtime_status_chart(data, assets_dir),
        "pipeline_overview": build_pipeline_overview_chart(assets_dir),
    }
    return {
        "comparison_rows": comparison_rows,
        "success_queries": success_queries,
        "anomaly_queries": anomaly_queries,
        "case_assets": case_assets,
        "charts": charts,
    }


def build_markdown(
    data: dict[str, object],
    assets: dict[str, object],
    out_md: Path,
) -> str:
    query_summary = summarize_queries(data["current_selected_summary"], data["current_query_manifest"])
    current_pose_summary = data["current_pose_summary"]
    current_suite_summary = data["current_suite_summary"]
    current_tie_overall = data["current_tie_overall"]
    baseline_pose_summary = data["baseline_pose_summary"]
    baseline_suite_summary = data["baseline_suite_summary"]
    baseline_tie_overall = data["baseline_tie_overall"]
    comparison_rows = assets["comparison_rows"]
    case_assets = assets["case_assets"]
    charts = assets["charts"]

    low_match_rows = sorted(
        comparison_rows,
        key=lambda row: row["baseline_match_count"],
    )[:8]
    top_rmse_rows = sorted(
        comparison_rows,
        key=lambda row: row["current_rmse"] if row["current_rmse"] is not None else float("inf"),
        reverse=True,
    )[:5]

    def img_line(key: str, alt: str) -> str:
        path = charts[key]
        rel = relative_to(out_md.parent, path)
        return f"![{alt}]({rel})"

    lines: list[str] = []
    lines.append("# 最终实验报告：Satellite Truth + SRTM + RoMa v2 Tiepoints")
    lines.append("")
    lines.append("## 1. 实验目的")
    lines.append(
        "本实验的目标是验证：在运行时仅提供单张不带地理元数据的 UAV 图像、并仅依赖固定遥感正射卫星影像库作为候选源的条件下，"
        "是否能够完成 UAV 图像的初始地理定位。当前路线保持 runtime 检索与位姿估计主链不变，仅将 validation truth 替换为 satellite truth patch，"
        "将 DSM 固定回 SRTM，并将 layer-3 的 tie-point matcher 替换为 RoMa v2。"
    )
    lines.append("")
    lines.append("## 2. 评估方法与评估指标")
    lines.append(
        "本实验沿用三层验证体系。layer-1 用于回答“预测正射结果与 satellite truth patch 的图像级几何对齐程度如何”；"
        "layer-2 用于回答“估计 pose 与参考航空三角测量/ODM pose 的相对一致性如何”；"
        "layer-3 用于回答“在同一地面区域内，局部 tie-point 的地面 XY 误差是否稳定”。"
    )
    lines.append("- layer-1 指标。`phase_corr_error_m` 越小越好，表示全局平移残差；`ortho_iou` 越大越好，表示预测正射与真值有效覆盖的重叠程度；`ssim` 越大越好，但在跨源影像条件下同时受到纹理、光照与成像外观差异影响。")
    lines.append("- layer-2 指标。`horizontal_error_m` 越小越好，衡量估计相机中心与参考相机中心的平面误差；`view_dir_angle_error_deg` 越小越好，衡量相机视线方向差异。该层仍使用 `pose_vs_at` 语义，因此与 baseline 具有直接可比性。")
    lines.append("- layer-3 指标。`tiepoint_match_count` 与 `tiepoint_inlier_count` 越大越好，表示可用于地面误差估计的 RoMa v2 对应点支持更充分；`tiepoint_inlier_ratio` 越大越好；`tiepoint_xy_error_rmse_m` 与 `tiepoint_xy_error_p90_m` 越小越好，分别反映整体误差能量与高误差尾部。")
    lines.append("")
    lines.append(img_line("pipeline_overview", "pipeline overview"))
    lines.append("")
    lines.append("## 3. 实验流程与数据准备")
    lines.append(
        f"本次 full-run 固定在 009/010 两条 nadir 航线的 40 张 query 上，每条航线 20 张。"
        f"query 经过 metadata 去除后进入 runtime；当前 `query_count={query_summary['query_count']}`，"
        f"`metadata_removed_count={query_summary['metadata_removed_count']}`，"
        f"`gimbal_pitch_degree` 范围为 {fmt(query_summary['gimbal_pitch_min'], 1)} 到 {fmt(query_summary['gimbal_pitch_max'], 1)}。"
    )
    lines.append("- 离线数据准备。复用 baseline 的 selected_queries、query_inputs、retrieval 与 RoMa v2 rerank 资产；validation truth 改为从 source satellite GeoTIFF 裁切出的 truth patch；PnP 使用 SRTM DSM；query reference pose 仍来自 `odm_report/shots.geojson`，缺失时才退回 seed。")
    lines.append("- Runtime 主链。单张 UAV query 先经过 DINOv2 coarse retrieval，再由 RoMa v2 rerank 生成候选排序，随后在 SRTM 支持下完成对应关系、2D-3D 构建、PnP 与 best-pose 选优。")
    lines.append("- Validation 主链。layer-1 对比 predicted ortho 与 satellite truth patch；layer-2 运行 `pose_vs_at`；layer-3 在 truth/pred common valid mask 内用 RoMa v2 重新做 dense/semi-dense matching，并计算地面 XY 误差。")
    lines.append("")
    lines.append("## 4. 实验结果")
    lines.append("### 4.1 Runtime 结果")
    lines.append(
        f"当前 full-run `best_status_counts` 为 `{current_pose_summary['best_status_counts']}`，"
        f"`score_status_counts` 为 `{current_pose_summary['score_status_counts']}`；"
        f"baseline 分别为 `{baseline_pose_summary['best_status_counts']}` 与 `{baseline_pose_summary['score_status_counts']}`。"
        "这表明当前路线在保持 40/40 query 全部产出可用 best pose 的同时，runtime 主链定义并未变化。"
    )
    lines.append(img_line("runtime_status_comparison", "runtime status comparison"))
    lines.append("")
    lines.append("### 4.2 Layer-1 结果")
    lines.append(
        f"当前 layer-1 平均 `phase_corr_error_m={fmt(mean_metric(current_suite_summary, 'phase_corr_error_m'))}`，"
        f"`ortho_iou={fmt(mean_metric(current_suite_summary, 'ortho_iou'))}`，"
        f"`ssim={fmt(mean_metric(current_suite_summary, 'ssim'))}`。"
        f"对比 baseline，对应值分别为 `{fmt(mean_metric(baseline_suite_summary, 'phase_corr_error_m'))}`、"
        f"`{fmt(mean_metric(baseline_suite_summary, 'ortho_iou'))}`、`{fmt(mean_metric(baseline_suite_summary, 'ssim'))}`。"
    )
    lines.append(
        "从几何一致性角度看，`phase_corr_error_m` 明显下降、`ortho_iou` 上升，说明以 satellite truth patch 为真值时，预测正射结果与最终 truth 的平移一致性和有效覆盖重叠更强；"
        "`ssim` 下降则主要反映跨源影像在纹理、色调与成像条件上的外观差异，不应直接解释为几何退化。"
    )
    lines.append(img_line("layer1_metrics_bar", "layer1 metrics"))
    lines.append("")
    lines.append("### 4.3 Layer-2 结果")
    lines.append(
        f"当前 layer-2 平均 `horizontal_error_m={fmt(mean_metric(current_suite_summary, 'horizontal_error_m'))}`，"
        f"`view_dir_angle_error_deg={fmt(mean_metric(current_suite_summary, 'view_dir_angle_error_deg'))}`；"
        f"baseline 分别为 `{fmt(mean_metric(baseline_suite_summary, 'horizontal_error_m'))}` 与 "
        f"`{fmt(mean_metric(baseline_suite_summary, 'view_dir_angle_error_deg'))}`。"
        "由于该层仍是 `pose_vs_at`，因此可以直接理解为估计 pose 与参考 pose 的一致性对比。"
    )
    lines.append(
        "结果显示 layer-2 没有出现明显改善，整体上略有回退但幅度有限。这说明将 validation truth 切换为 satellite truth patch，并不会自动带来 camera pose 与参考 pose 的同步提升。"
    )
    lines.append(img_line("layer2_metrics_bar", "layer2 metrics"))
    lines.append("")
    lines.append("### 4.4 Layer-3 结果")
    lines.append(
        f"当前 layer-3 `status_counts={tie_status_counts(current_tie_overall)}`，"
        f"`tiepoint_match_count_mean={fmt(current_tie_overall['tiepoint_match_count_mean'], 3)}`，"
        f"`tiepoint_inlier_ratio_mean={fmt(current_tie_overall['tiepoint_inlier_ratio_mean'], 4)}`，"
        f"`tiepoint_xy_error_rmse_m={fmt(current_tie_overall['tiepoint_xy_error_rmse_m'])}`，"
        f"`tiepoint_xy_error_p90_m={fmt(current_tie_overall['tiepoint_xy_error_p90_m'])}`。"
        f"baseline 对应 `status_counts={tie_status_counts(baseline_tie_overall)}`，"
        f"`tiepoint_xy_error_rmse_m={fmt(baseline_tie_overall['tiepoint_xy_error_rmse_m'])}`，"
        f"`tiepoint_xy_error_p90_m={fmt(baseline_tie_overall['tiepoint_xy_error_p90_m'])}`。"
    )
    lines.append(
        f"若按中位数比较，`tiepoint_match_count` 从 baseline 的 `{fmt(median_from_rows(data['baseline_tie_rows'], 'tiepoint_match_count'))}` "
        f"提升到 `{fmt(median_from_rows(data['current_tie_rows'], 'tiepoint_match_count'))}`，"
        f"`tiepoint_inlier_count` 从 `{fmt(median_from_rows(data['baseline_tie_rows'], 'tiepoint_inlier_count'))}` "
        f"提升到 `{fmt(median_from_rows(data['current_tie_rows'], 'tiepoint_inlier_count'))}`。"
        "这说明 RoMa v2 作为 layer-3 matcher 后，局部几何评估获得了显著更强的 tiepoint support。"
    )
    lines.append(img_line("layer3_metrics_bar", "layer3 metrics"))
    lines.append("")
    lines.append("### 4.5 baseline 低匹配 query 改善情况")
    lines.append(
        "以 baseline 的 layer-3 `tiepoint_match_count` 下四分位作为 low-match 集合，可以看到多数 query 的匹配支持数量显著提升，但这并不等价于所有 query 的几何误差都会同步下降。"
    )
    lines.append(img_line("low_match_queries_improvement", "low match improvement"))
    lines.append("")
    lines.append("| Query | Baseline Matches | Current Matches | Delta | Baseline RMSE | Current RMSE |")
    lines.append("| --- | ---: | ---: | ---: | ---: | ---: |")
    for row in low_match_rows:
        lines.append(
            f"| {row['query_id']} | {int(row['baseline_match_count'])} | {int(row['current_match_count'])} | "
            f"{int(row['match_delta'])} | {fmt(row['baseline_rmse'])} | {fmt(row['current_rmse'])} |"
        )
    lines.append("")
    lines.append("## 5. 结论与结果分析")
    lines.append(
        "综合 full-run 结果，可以得出以下结论。第一，runtime 任务定义保持不变的前提下，当前路线仍然实现了 `best_status_counts={ok: 40}`，因此 satellite truth + SRTM + RoMa layer-3 方案不会破坏原有定位主链。"
    )
    lines.append(
        "第二，主要收益集中在 layer-1 与 layer-3。layer-1 的 `phase_corr_error_m` 和 `ortho_iou` 均优于 baseline，说明以 satellite truth patch 作为 formal truth 时，预测正射与最终真值之间的图像级几何对齐更稳定；layer-3 的 tiepoint 数量和 inlier 数量显著增加，且 RMSE 从 baseline 的 "
        f"`{fmt(baseline_tie_overall['tiepoint_xy_error_rmse_m'])}` 下降到 `{fmt(current_tie_overall['tiepoint_xy_error_rmse_m'])}`。"
    )
    lines.append(
        "第三，layer-2 没有明显改善，`horizontal_error_m` 与 `view_dir_angle_error_deg` 基本持平或略有回退。这表明 validation truth 的更换主要改善的是 orthophoto-level alignment 与局部 tiepoint support，而不是直接改善相机 pose 与参考 pose 的差距。"
    )
    lines.append(
        "第四，`ssim` 低于 baseline 不能直接解释为几何退化。由于当前 layer-1 属于跨源对比，外观差异会显著拉低 `ssim`，因此应将其与 `phase_corr_error_m`、`ortho_iou` 联合解释。"
    )
    lines.append(
        "第五，tiepoint 数量大增并不意味着所有 query 的地面几何误差都同步改善。部分 query 虽然已经获得高密度匹配，但 `RMSE` 仍然偏高，说明后续仍需关注局部几何失真、遮挡、跨源外观差异以及 truth/pred 有效覆盖不完全一致等问题。"
    )
    lines.append("")
    lines.append(img_line("overall_metrics_comparison", "overall metrics comparison"))
    lines.append("")
    lines.append("## 6. 后续想法")
    lines.append("- 扩展到更一般的非 nadir / arbitrary UAV query，验证当前路线是否能从受控 nadir 集迁移到更真实的任意视角输入。")
    lines.append("- 继续研究 satellite truth 作为正式主验证口径的稳定性，尤其是其与 UAV DOM/ODM truth 在空间分辨率、外观域差异上的系统偏差。")
    lines.append("- 分析“高匹配但 RMSE 仍偏高”的 query，区分是 RoMa 匹配分布问题、局部区域形变问题，还是 truth/pred 几何边界不一致。")
    lines.append("- 在 layer-3 中加入更强几何约束与分区域评估，例如分块 RMSE、边界区域剔除、显式不确定性建模，以减少仅靠 match count 解读结果的偏差。")
    lines.append("- 进一步比较 runtime 与 validation 的差异，明确哪些改动应当只停留在验证链，哪些改动有可能反向迁移到运行时主链。")
    lines.append("")
    lines.append("## 7. 关键实验设置与变量说明")
    lines.append(f"- 实验对象：`{relative_to_cwd(data['experiment_root'])}`。")
    lines.append(f"- baseline 对比对象：`{relative_to_cwd(data['baseline_root'])}`。")
    lines.append("- runtime 保持不变：DINOv2 coarse retrieval + RoMa v2 rerank + fixed satellite library + SRTM-backed pose/PnP。")
    lines.append("- validation 变更点：truth 改为 satellite truth patch；layer-2 继续使用 `pose_vs_at`；layer-3 matcher 改为 RoMa v2。")
    lines.append(f"- 当前 full-run runtime 状态：`best_status_counts={current_pose_summary['best_status_counts']}`，`score_status_counts={current_pose_summary['score_status_counts']}`。")
    lines.append(f"- 当前 layer-3 状态：`status_counts={tie_status_counts(current_tie_overall)}`。")
    lines.append(
        f"- query 基本信息：航线分布 `{counter_to_text(query_summary['flight_counts'])}`，"
        f"平均 absolute altitude `{fmt(query_summary['absolute_altitude_mean'], 3)}`，"
        f"平均 relative altitude `{fmt(query_summary['relative_altitude_mean'], 3)}`。"
    )
    lines.append("")
    lines.append("## 8. 典型样例与异常样例分析")
    lines.append("成功样例优先选取 layer-3 RMSE 低且 layer-2 平面误差较低的 query；异常样例优先选取 baseline low-match 集合中当前 RMSE 仍偏高的 query。")
    for case in case_assets:
        rel = relative_to(out_md.parent, case["panel_png"])
        lines.append("")
        lines.append(
            f"### {case['query_id']} ({'代表性成功样例' if case['label'] == 'success' else '异常样例'})"
        )
        lines.append(
            f"- `current_match_count={int(case['current_match_count'])}`，"
            f"`baseline_match_count={int(case['baseline_match_count'])}`，"
            f"`current_rmse_m={fmt(case['current_rmse'])}`，"
            f"`current_horizontal_error_m={fmt(case['current_horizontal'])}`。"
        )
        if case["label"] == "success":
            lines.append("- 该样例说明在 satellite truth patch 与 RoMa v2 tiepoints 支持下，truth/pred 的局部几何关系能够保持较高稳定性。")
        else:
            lines.append("- 该样例说明即使匹配点数量已经显著增加，局部几何误差仍可能偏高，后续需要重点排查区域形变、边界覆盖与跨源差异。")
        lines.append(f"![{case['query_id']} case]({rel})")
    lines.append("")
    lines.append("## 附加说明")
    lines.append("- 本报告只基于已存在的 full-run 结果重组与可视化，不引入新的实验口径。")
    lines.append("- `validation truth` 仅用于离线评估，不参与 runtime 候选选择或定位推理。")
    return "\n".join(lines) + "\n"


def build_docx(data: dict[str, object], assets: dict[str, object], out_docx: Path) -> None:
    doc = Document()
    add_paragraph(doc, "最终实验报告：Satellite Truth + SRTM + RoMa v2 Tiepoints", size=16, center=True, bold=True)
    add_paragraph(doc, "实验对象：009/010 nadir 40-query full run", size=11, center=True)

    markdown_text = build_markdown(data, assets, out_docx.with_suffix(".md"))
    sections = [part.strip() for part in markdown_text.split("\n## ") if part.strip()]
    for index, section in enumerate(sections):
        lines = section.splitlines()
        title = lines[0] if index == 0 else "## " + lines[0]
        heading_text = title.lstrip("# ").strip()
        add_heading(doc, heading_text, level=1)
        body_lines = lines[1:] if index == 0 else lines[1:]
        for line in body_lines:
            if not line:
                continue
            if line.startswith("![") and "](" in line:
                image_rel = line.split("](", 1)[1].rstrip(")")
                image_path = (out_docx.parent / image_rel).resolve()
                add_picture_with_caption(doc, image_path, image_path.stem, width_inch=6.0)
            elif line.startswith("| "):
                continue
            elif line.startswith("- "):
                add_bullets(doc, [line[2:]])
            elif line.startswith("### "):
                add_heading(doc, line[4:], level=2)
            else:
                add_paragraph(doc, line)

    comparison_rows = assets["comparison_rows"]
    low_match_rows = sorted(comparison_rows, key=lambda row: row["baseline_match_count"])[:8]
    add_heading(doc, "附表：Low-match Query 对比", level=1)
    add_table(
        doc,
        ["Query", "Baseline Matches", "Current Matches", "Delta", "Baseline RMSE", "Current RMSE"],
        [
            [
                row["query_id"],
                str(int(row["baseline_match_count"])),
                str(int(row["current_match_count"])),
                str(int(row["match_delta"])),
                fmt(row["baseline_rmse"]),
                fmt(row["current_rmse"]),
            ]
            for row in low_match_rows
        ],
    )
    ensure_dir(out_docx.parent)
    doc.save(str(out_docx))


def main() -> None:
    args = parse_args()
    experiment_root = Path(args.experiment_root).resolve()
    baseline_root = Path(args.baseline_root).resolve()
    suite_root = Path(args.suite_root).resolve()
    out_md = Path(args.out_md).resolve() if args.out_md else experiment_root / "reports" / "final_experiment_report_sattruth_srtm_romatie.md"
    out_docx = Path(args.out_docx).resolve() if args.out_docx else experiment_root / "reports" / "final_experiment_report_sattruth_srtm_romatie.docx"
    assets_dir = Path(args.assets_dir).resolve() if args.assets_dir else experiment_root / "reports" / "final_experiment_report_assets"

    data = load_required_inputs(experiment_root, baseline_root, suite_root)
    assets = build_report_assets(data, assets_dir)
    markdown = build_markdown(data, assets, out_md)
    write_text(out_md, markdown)
    build_docx(data, assets, out_docx)


if __name__ == "__main__":
    main()
