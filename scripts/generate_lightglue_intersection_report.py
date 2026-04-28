#!/usr/bin/env python3
"""Generate a formal Chinese report for LightGlue reranking under intersection-truth evaluation."""

from __future__ import annotations

import argparse
import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--result-dir", required=True)
    parser.add_argument("--out-docx", required=True)
    parser.add_argument("--out-md", required=True)
    return parser.parse_args()


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def load_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", newline="", encoding="utf-8-sig") as f:
        return list(csv.DictReader(f))


def set_cn_font(run, size: int = 11, bold: bool = False) -> None:
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(size)
    run.bold = bold


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 10, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    set_cn_font(r, size=size, bold=bold)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.style = "Heading 1" if level == 1 else "Heading 2"
    r = p.add_run(text)
    set_cn_font(r, size=14 if level == 1 else 12, bold=True)


def add_paragraph(doc: Document, text: str, size: int = 11, center: bool = False) -> None:
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn_font(r, size=size)


def add_bullets(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(line)
        set_cn_font(r, size=11)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn_font(r, size=10)


def add_picture(doc: Document, path: Path, width_inch: float) -> None:
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(width_inch))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def short_flight_name(flight_id: str) -> str:
    parts = flight_id.split("_")
    return parts[2] if len(parts) >= 3 else flight_id


def format_seconds(seconds: float) -> str:
    return f"{seconds:.2f}s ({seconds / 60.0:.2f} min)"


def stage_seconds(timing: dict) -> dict[str, float]:
    return {item["stage"]: float(item["elapsed_seconds"]) for item in timing["stages"]}


def build_metric_definition_table(doc: Document) -> None:
    rows = [
        ("Intersection Recall@1", "首位候选命中 intersection truth 的比例，是本轮重排最直接的首位粗定位指标。"),
        ("Intersection Recall@5", "前 5 名候选中命中 intersection truth 的比例，用于衡量前排覆盖能力。"),
        ("Intersection Recall@10", "前 10 名候选中命中 intersection truth 的比例，用于衡量区域级初步定位能力。"),
        ("Intersection Recall@20/50", "用于观察粗检索窗口扩展上限，以及 LightGlue 在 Top-50 内的可挽救空间。"),
        ("Intersection MRR", "首个 intersection truth 排名倒数平均值，越高表示正确区域越靠前。"),
        ("Top-1 error mean (m)", "首位候选中心与 query 参考位置之间的平均距离，用于衡量首位候选空间偏差。"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, title in enumerate(["指标", "含义"]):
        set_cell_text(table.cell(0, i), title, bold=True)
        shade_cell(table.cell(0, i), "D9EAF7")
    for name, desc in rows:
        row = table.add_row().cells
        set_cell_text(row[0], name)
        row[1].text = ""
        p = row[1].paragraphs[0]
        r = p.add_run(desc)
        set_cn_font(r, size=10)


def build_overall_table(doc: Document, overall: dict) -> None:
    table = doc.add_table(rows=1, cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["方法", "R@1", "R@5", "R@10", "R@20", "R@50", "MRR", "Top-1误差均值(m)"]
    for i, title in enumerate(headers):
        set_cell_text(table.cell(0, i), title, bold=True)
        shade_cell(table.cell(0, i), "D9EAF7")
    rows = [
        (
            "DINOv2 Baseline",
            overall["baseline_intersection_recall@1"],
            overall["baseline_intersection_recall@5"],
            overall["baseline_intersection_recall@10"],
            overall["baseline_intersection_recall@20"],
            "-",
            overall["baseline_intersection_mrr"],
            overall["baseline_top1_error_m_mean"],
        ),
        (
            "Coarse Top50",
            "-",
            "-",
            "-",
            "-",
            overall["coarse_intersection_recall@50"],
            "-",
            "-",
        ),
        (
            "SuperPoint + LightGlue",
            overall["lightglue_intersection_recall@1"],
            overall["lightglue_intersection_recall@5"],
            overall["lightglue_intersection_recall@10"],
            overall["lightglue_intersection_recall@20"],
            overall["lightglue_intersection_recall@50"],
            overall["lightglue_intersection_mrr"],
            overall["lightglue_top1_error_m_mean"],
        ),
    ]
    for vals in rows:
        row = table.add_row().cells
        for i, val in enumerate(vals):
            if isinstance(val, float):
                set_cell_text(row[i], f"{val:.3f}")
            else:
                set_cell_text(row[i], str(val))


def build_per_flight_table(doc: Document, rows: list[dict[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["航线", "Baseline R@10", "LightGlue R@10", "Delta", "Baseline MRR", "LightGlue MRR", "LightGlue Top1Err(m)"]
    for i, title in enumerate(headers):
        set_cell_text(table.cell(0, i), title, bold=True)
        shade_cell(table.cell(0, i), "D9EAF7")
    for item in rows:
        baseline_r10 = float(item["baseline_intersection_recall@10"])
        lightglue_r10 = float(item["lightglue_intersection_recall@10"])
        row = table.add_row().cells
        set_cell_text(row[0], item["flight_tag"])
        set_cell_text(row[1], f"{baseline_r10:.3f}")
        set_cell_text(row[2], f"{lightglue_r10:.3f}")
        set_cell_text(row[3], f"{lightglue_r10 - baseline_r10:+.3f}")
        set_cell_text(row[4], f"{float(item['baseline_intersection_mrr']):.3f}")
        set_cell_text(row[5], f"{float(item['lightglue_intersection_mrr']):.3f}")
        set_cell_text(row[6], f"{float(item['lightglue_top1_error_m_mean']):.3f}")


def build_timing_table(doc: Document, timing_map: dict[str, float]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, title in enumerate(["阶段", "耗时"]):
        set_cell_text(table.cell(0, i), title, bold=True)
        shade_cell(table.cell(0, i), "D9EAF7")
    rows = [
        ("coarse Top-50 导出", format_seconds(timing_map.get("coarse_topk_export", 0.0))),
        ("输入准备", format_seconds(timing_map.get("input_preparation", 0.0))),
        ("LightGlue 重排", format_seconds(timing_map.get("lightglue_rerank", 0.0))),
        ("结果汇总", format_seconds(timing_map.get("summary_aggregation", 0.0))),
        ("可视化", format_seconds(timing_map.get("visualization", 0.0))),
    ]
    for key, value in rows:
        row = table.add_row().cells
        set_cell_text(row[0], key)
        set_cell_text(row[1], value)


def choose_cases(comp_rows: list[dict[str, str]]) -> list[dict[str, str]]:
    improved = [r for r in comp_rows if r["baseline_first_truth_rank"] and r["lightglue_first_truth_rank"] and int(r["lightglue_first_truth_rank"]) < int(r["baseline_first_truth_rank"])]
    improved.sort(key=lambda r: int(r["baseline_first_truth_rank"]) - int(r["lightglue_first_truth_rank"]), reverse=True)

    promoted = [r for r in comp_rows if int(r["promoted_11_50_to_top10"]) == 1]
    promoted.sort(key=lambda r: r["query_id"])

    still_miss10 = [r for r in comp_rows if r["coarse_intersection_hit@50"] == "1" and r["lightglue_intersection_hit@10"] == "0"]
    still_miss10.sort(key=lambda r: ((int(r["coarse_first_truth_rank"]) if r["coarse_first_truth_rank"] else 999), r["query_id"]))

    selected: list[dict[str, str]] = []
    if promoted:
        case = dict(promoted[0])
        case["title"] = "从 11..50 拉回 Top-10 的改进样例"
        selected.append(case)
    if improved:
        cand = next((dict(r) for r in improved if r["query_id"] not in {x["query_id"] for x in selected}), dict(improved[0]))
        cand["title"] = "前排排名进一步提升样例"
        selected.append(cand)
    if still_miss10:
        cand = next((dict(r) for r in still_miss10 if r["query_id"] not in {x["query_id"] for x in selected}), dict(still_miss10[0]))
        cand["title"] = "Top-50 内存在真值但仍未进入 Top-10 的样例"
        selected.append(cand)
    return selected[:3]


def build_md(result_dir: Path, overall: dict, per_flight_rows: list[dict[str, str]], timing_map: dict[str, float], k_full: dict, k_unique: dict, cases: list[dict[str, str]]) -> str:
    lines: list[str] = []
    lines.append("# SuperPoint + LightGlue 在 Intersection Truth 口径下的重排结果说明")
    lines.append("")
    lines.append("## 1. 任务定义与实验设置")
    lines.append("本组实验用于回答：在 `intersection truth` 新真值口径下，若先用 `DINOv2 + FAISS` 对全库粗检索，再对前 `Top-50` 候选使用 `SuperPoint + LightGlue` 做局部重排，是否能够进一步改善正式指标。")
    lines.append("")
    lines.extend(
        [
            "- 数据范围：4 条航线，共 `40` 个 query。",
            "- 固定候选库：`1029` 张卫片，来源于四条航线总体范围外扩 `250m` 后构建的固定库。",
            "- 粗检索：DINOv2 pooler 特征 + FAISS `IndexFlatIP`。",
            "- 重排方法：SuperPoint 局部特征 + LightGlue 匹配 + RANSAC 几何一致性 + 融合分数排序。",
            "- 重排窗口：每个 query 的 coarse `Top-50`。",
            "- Top-K 曲线口径：`1..50` 使用 LightGlue 结果，`51..1029` 保持 baseline 原顺序。",
        ]
    )
    lines.append("")
    lines.append("## 2. Intersection Truth 定义")
    lines.append("本轮正式真值定义为：只要 query 覆盖范围与卫片存在非零面积相交，该卫片就记为 `intersection truth`。")
    lines.append("")
    lines.append("## 3. 指标定义")
    lines.extend(
        [
            "- `Intersection Recall@1/5/10/20/50`：前 K 名中是否命中 intersection truth。",
            "- `Intersection MRR`：首个 intersection truth 排名倒数的平均值。",
            "- `Top-1 error mean (m)`：首位候选中心与 query 参考位置之间的平均距离。",
        ]
    )
    lines.append("")
    lines.append("## 4. 总体定量结果")
    lines.append("")
    lines.append(f"- Baseline：`R@1={overall['baseline_intersection_recall@1']:.3f}`，`R@5={overall['baseline_intersection_recall@5']:.3f}`，`R@10={overall['baseline_intersection_recall@10']:.3f}`，`R@20={overall['baseline_intersection_recall@20']:.3f}`，`MRR={overall['baseline_intersection_mrr']:.3f}`")
    lines.append(f"- Coarse Top50 上限：`R@50={overall['coarse_intersection_recall@50']:.3f}`")
    lines.append(f"- LightGlue：`R@1={overall['lightglue_intersection_recall@1']:.3f}`，`R@5={overall['lightglue_intersection_recall@5']:.3f}`，`R@10={overall['lightglue_intersection_recall@10']:.3f}`，`R@20={overall['lightglue_intersection_recall@20']:.3f}`，`R@50={overall['lightglue_intersection_recall@50']:.3f}`，`MRR={overall['lightglue_intersection_mrr']:.3f}`")
    lines.append(f"- 指标变化：`ΔR@1={overall['delta_intersection_recall@1']:+.3f}`，`ΔR@5={overall['delta_intersection_recall@5']:+.3f}`，`ΔR@10={overall['delta_intersection_recall@10']:+.3f}`，`ΔMRR={overall['delta_intersection_mrr']:+.3f}`")
    lines.append(f"- Top-1 误差：`{overall['baseline_top1_error_m_mean']:.3f}m -> {overall['lightglue_top1_error_m_mean']:.3f}m`")
    lines.append("")
    lines.append("## 5. 分航线结果")
    lines.append("")
    for row in per_flight_rows:
        delta = float(row["lightglue_intersection_recall@10"]) - float(row["baseline_intersection_recall@10"])
        lines.append(
            f"- `{row['flight_tag']}`：Baseline `R@10={float(row['baseline_intersection_recall@10']):.3f}`，LightGlue `R@10={float(row['lightglue_intersection_recall@10']):.3f}`，`Δ={delta:+.3f}`"
        )
    lines.append("")
    lines.append("## 6. 时间开销统计")
    lines.append("")
    for key, label in [
        ("coarse_topk_export", "coarse Top-50 导出"),
        ("input_preparation", "输入准备"),
        ("lightglue_rerank", "LightGlue 重排"),
        ("summary_aggregation", "结果汇总"),
        ("visualization", "可视化"),
    ]:
        lines.append(f"- {label}：`{format_seconds(timing_map.get(key, 0.0))}`")
    lines.append("")
    lines.append("## 7. Top-K 曲线结果")
    lines.append("")
    lines.append(
        f"- full-truth：`40/40` 个 query 都能达到真值饱和，`mean={k_full['overall']['mean']:.3f}`，`median={k_full['overall']['median']}`，`p95={k_full['overall']['p95']}`。"
    )
    lines.append(
        f"- unique-tile：唯一真值 tile 数为 `{k_unique['overall']['total_truth_unique_tiles']}`，`k_full_truth={k_unique['overall']['k_full_truth']}`，候选唯一 tile 总数为 `{k_unique['overall']['candidate_unique_tiles']}`。"
    )
    lines.append("- 这说明 LightGlue 虽然改变了前 50 名内部顺序，但如果目标是把全部真值 tile 找全，仍然需要接近全库深度。")
    lines.append("")
    lines.append("## 8. 代表性样例")
    lines.append("")
    for case in cases:
        lines.append(
            f"- `{case['query_id']}` / `{short_flight_name(case['flight_id'])}`：{case['title']}；Baseline 首个真值 rank=`{case['baseline_first_truth_rank'] or 'miss'}`，Coarse Top50 rank=`{case['coarse_first_truth_rank'] or 'miss'}`，LightGlue rank=`{case['lightglue_first_truth_rank'] or 'miss'}`。"
        )
    lines.append("")
    lines.append("## 9. 结论")
    lines.append("")
    lines.extend(
        [
            f"- 本轮 `LightGlue` 没有提升 `Recall@1`，但把 `Recall@10` 从 `{overall['baseline_intersection_recall@10']:.3f}` 提升到 `{overall['lightglue_intersection_recall@10']:.3f}`。",
            f"- `Recall@5` 与 `MRR` 略有下降，说明当前局部重排并未稳定改善最前排排序质量。",
            f"- `Top-1 error mean` 从 `{overall['baseline_top1_error_m_mean']:.3f}m` 降到 `{overall['lightglue_top1_error_m_mean']:.3f}m`，说明首位候选的空间误差有一定改善。",
            "- 更准确的结论是：当前 LightGlue 更像是在局部改善前十名覆盖，而不是已经稳定提升首位候选最优性。",
        ]
    )
    lines.append("")
    return "\n".join(lines)


def build_docx(result_dir: Path, out_docx: Path, overall: dict, per_flight_rows: list[dict[str, str]], timing_map: dict[str, float], k_full: dict, k_unique: dict, cases: list[dict[str, str]]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("SuperPoint + LightGlue 在 Intersection Truth 口径下的重排结果说明")
    set_cn_font(r, size=16, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("面向无人机单张图像初步地理定位的正式技术报告")
    set_cn_font(r, size=11)

    add_heading(doc, "1. 任务定义与实验设置", 1)
    add_paragraph(
        doc,
        "本组实验用于回答：在 `intersection truth` 新真值定义下，若先用 `DINOv2 + FAISS` 对全库执行粗检索，再对前 `Top-50` 候选使用 `SuperPoint + LightGlue` 做局部匹配与融合重排，是否能够进一步提升正式主指标。"
    )
    add_bullets(
        doc,
        [
            "数据范围：4 条航线，共 40 个 query。",
            "固定候选库：1029 张卫片，来源于四条航线总体范围外扩 250m 后构建的固定库。",
            "粗检索方法：DINOv2 pooler 特征 + FAISS IndexFlatIP。",
            "重排方法：SuperPoint 局部特征 + LightGlue 匹配 + RANSAC 几何一致性 + 融合分数排序。",
            "执行窗口：coarse Top-50。",
            "Top-K 曲线口径：1..50 使用 LightGlue 结果，51..1029 保持 baseline 原顺序。",
        ],
    )

    add_heading(doc, "2. Intersection Truth 定义", 1)
    add_paragraph(
        doc,
        "本轮正式真值定义为：只要 query 覆盖范围与卫片存在非零面积相交，该卫片就记为 intersection truth。该定义仍以真实地理交集为准，而不是只依赖 query 中心点。",
    )

    add_heading(doc, "3. 指标定义", 1)
    build_metric_definition_table(doc)

    add_heading(doc, "4. 方法说明", 1)
    add_paragraph(
        doc,
        "本轮不改变 DINOv2 粗检索本身，只在其 Top-50 候选集合上增加局部几何验证与融合排序。因此它更准确地说是一轮“扩窗后的局部重排”实验，而不是重新训练新的全局检索器。"
    )

    add_heading(doc, "5. 总体定量结果", 1)
    add_paragraph(doc, "表 1 给出 baseline、coarse Top-50 上限和 LightGlue 重排后的 overall 指标对照。")
    build_overall_table(doc, overall)
    add_caption(doc, "表 1  Baseline、Coarse Top-50 与 LightGlue 的 overall 指标对照")
    add_paragraph(
        doc,
        f"从总体结果看，LightGlue 后 `R@1` 保持 `{overall['lightglue_intersection_recall@1']:.3f}` 不变，`R@10` 从 `{overall['baseline_intersection_recall@10']:.3f}` 提升到 `{overall['lightglue_intersection_recall@10']:.3f}`，但 `R@5` 和 `MRR` 略有下降。与此同时，`Top-1 error mean` 从 `{overall['baseline_top1_error_m_mean']:.3f}m` 下降到 `{overall['lightglue_top1_error_m_mean']:.3f}m`。"
    )

    add_heading(doc, "6. 分航线结果", 1)
    add_paragraph(doc, "表 2 给出四条航线上的 `R@10`、`MRR` 和 Top-1 误差对照。")
    build_per_flight_table(doc, per_flight_rows)
    add_caption(doc, "表 2  四条航线上的 LightGlue 重排结果对照")

    add_heading(doc, "7. 时间开销统计", 1)
    add_paragraph(doc, "表 3 总结本次 run 中 coarse 导出、输入准备、LightGlue 重排、结果汇总与可视化的时间开销。")
    build_timing_table(doc, timing_map)
    add_caption(doc, "表 3  本次 LightGlue run 的时间开销统计")

    add_heading(doc, "8. 汇总图解读", 1)
    add_paragraph(
        doc,
        "图 1 到图 6 分别展示 overall 指标、分航线 Recall、baseline 与 LightGlue 的 overall 对照、以及本轮新增的两张 Top-K 对照图，用于从主指标和真值覆盖深度两个角度理解当前重排效果。"
    )
    figures = [
        (result_dir / "figures" / "_aggregate" / "overall_metrics_bar.png", "图 1  Intersection Truth 口径下的 LightGlue overall 指标"),
        (result_dir / "figures" / "_aggregate" / "multi_flight_recall.png", "图 2  Intersection Truth 口径下的分航线 Recall"),
        (result_dir / "figures" / "_compare" / "baseline_vs_lightglue_compare.png", "图 3  Baseline vs LightGlue 的 overall 指标对照"),
        (result_dir / "figures_topk_fulltruth" / "_aggregate" / "overall_topk_truth_count_curve_all.png", "图 4  LightGlue 的 full-truth Top-K 曲线"),
        (result_dir / "figures_topk_unique_tile" / "_aggregate" / "overall_topk_unique_truth_count_curve.png", "图 5  LightGlue 的 unique-tile Top-K 曲线"),
        (result_dir / "figures_topk_compare" / "baseline_vs_lightglue_full_truth_curve.png", "图 6  Baseline vs LightGlue 的 full-truth Top-K 对照"),
        (result_dir / "figures_topk_compare" / "baseline_vs_lightglue_unique_truth_curve.png", "图 7  Baseline vs LightGlue 的 unique-tile Top-K 对照"),
    ]
    for path, caption in figures:
        add_picture(doc, path, 6.2)
        add_caption(doc, caption)

    add_paragraph(
        doc,
        f"从 Top-K 曲线看，LightGlue 的 full-truth `k_full` 统计为 `mean={k_full['overall']['mean']:.3f}`、`median={k_full['overall']['median']}`、`p95={k_full['overall']['p95']}`；"
        f"而 unique-tile 口径下，唯一真值 tile 数为 `{k_unique['overall']['total_truth_unique_tiles']}`，`k_full_truth={k_unique['overall']['k_full_truth']}`。"
        "这说明当前 LightGlue 虽然改变了前 50 名内部顺序，但若目标是把所有真值 tile 全部找全，仍然需要接近全库深度。"
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "9. 代表性样例", 1)
    for idx, case in enumerate(cases, start=1):
        qid = case["query_id"]
        flight_id = case["flight_id"]
        add_heading(doc, f"9.{idx} {case['title']}：{short_flight_name(flight_id)} / {qid}", 2)
        add_paragraph(
            doc,
            f"Baseline 首个 intersection truth 排名：{case['baseline_first_truth_rank'] or '未命中'}；"
            f"Coarse Top50 首个 intersection truth 排名：{case['coarse_first_truth_rank'] or '未命中'}；"
            f"LightGlue 首个 intersection truth 排名：{case['lightglue_first_truth_rank'] or '未命中'}。"
        )
        add_picture(doc, result_dir / "figures" / flight_id / f"{qid}_top10.png", 5.8)
        add_caption(doc, f"{qid} 的 LightGlue Top-10 联系图")

    add_heading(doc, "10. 结论", 1)
    add_bullets(
        doc,
        [
            f"在 intersection truth 口径下，LightGlue 没有提升 `Recall@1`，但把 `Recall@10` 从 `{overall['baseline_intersection_recall@10']:.3f}` 提升到了 `{overall['lightglue_intersection_recall@10']:.3f}`。",
            f"`Recall@5` 与 `MRR` 略有下降，说明当前局部重排尚未稳定改善最前排排序质量。",
            f"`Top-1 error mean` 从 `{overall['baseline_top1_error_m_mean']:.3f}m` 下降到 `{overall['lightglue_top1_error_m_mean']:.3f}m`，说明首位候选的空间偏差有所缓解。",
            "因此，更准确的结论是：当前 LightGlue 更像是在局部改善前十名覆盖，而不是已经稳定提升首位候选最优性。",
        ],
    )

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))


def main() -> None:
    args = parse_args()
    result_dir = Path(args.result_dir)
    out_docx = Path(args.out_docx)
    out_md = Path(args.out_md)

    overall = load_json(result_dir / "overall_summary.json")
    per_flight_rows = load_csv(result_dir / "per_flight_comparison.csv")
    per_query_rows = load_csv(result_dir / "per_query_comparison.csv")
    timing = load_json(result_dir / "timing" / "timing_summary.json")
    k_full = load_json(result_dir / "figures_topk_fulltruth" / "k_full_truth_summary.json")
    k_unique = load_json(result_dir / "figures_topk_unique_tile" / "k_full_truth_unique_tile_summary.json")

    timing_map = stage_seconds(timing)
    cases = choose_cases(per_query_rows)

    out_md.parent.mkdir(parents=True, exist_ok=True)
    out_md.write_text(build_md(result_dir, overall, per_flight_rows, timing_map, k_full, k_unique, cases), encoding="utf-8")
    build_docx(result_dir, out_docx, overall, per_flight_rows, timing_map, k_full, k_unique, cases)
    print(out_md)
    print(out_docx)


if __name__ == "__main__":
    main()
