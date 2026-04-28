#!/usr/bin/env python3
"""Generate a Chinese comparison report for DINOv2 vs DINOv3 intersection-truth results."""

from __future__ import annotations

import argparse
import json
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--dinov2-dir", required=True)
    parser.add_argument("--dinov3-dir", required=True)
    parser.add_argument("--out-md", required=True)
    parser.add_argument("--out-docx", required=True)
    return parser.parse_args()


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def set_cn_font(run, size: int = 11, bold: bool = False) -> None:
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(size)
    run.bold = bold


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 10, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    set_cn_font(r, size=size, bold=bold)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.style = "Heading 1" if level == 1 else "Heading 2"
    r = p.add_run(text)
    set_cn_font(r, size=14 if level == 1 else 12, bold=True)


def add_paragraph(doc: Document, text: str, size: int = 11, center: bool = False) -> None:
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn_font(r, size=size)


def add_bullets(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(line)
        set_cn_font(r, size=11)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn_font(r, size=10)


def short_flight_name(flight_id: str) -> str:
    parts = flight_id.split("_")
    return parts[2] if len(parts) >= 3 else flight_id


def metric_bool(value) -> int:
    if isinstance(value, bool):
        return int(value)
    if isinstance(value, (int, float)):
        return int(value)
    if isinstance(value, str):
        return 1 if value.lower() in {"1", "true", "yes"} else 0
    return 0


def per_flight_metrics(summary: dict) -> list[dict[str, float | str | int]]:
    grouped: dict[str, list[dict]] = defaultdict(list)
    for row in summary["per_query"]:
        grouped[row["flight_id"]].append(row)

    out: list[dict[str, float | str | int]] = []
    for flight_id in sorted(grouped):
        rows = grouped[flight_id]
        total = len(rows)
        errors = [float(row["top1_error_m"]) for row in rows]
        out.append(
            {
                "flight_id": flight_id,
                "query_count": total,
                "intersection_recall@1": sum(metric_bool(row["intersection_hit@1"]) for row in rows) / total,
                "intersection_recall@5": sum(metric_bool(row["intersection_hit@5"]) for row in rows) / total,
                "intersection_recall@10": sum(metric_bool(row["intersection_hit@10"]) for row in rows) / total,
                "intersection_recall@20": sum(metric_bool(row["intersection_hit@20"]) for row in rows) / total,
                "intersection_mrr": sum(float(row["intersection_reciprocal_rank"]) for row in rows) / total,
                "top1_error_m_mean": sum(errors) / len(errors),
            }
        )
    return out


def format_seconds(seconds: float) -> str:
    return f"{seconds:.2f}s ({seconds / 60.0:.2f} min)"


def format_delta(value: float, invert_good: bool = False) -> str:
    sign = "+" if value > 0 else ""
    rendered = f"{sign}{value:.3f}"
    if invert_good:
        if value < 0:
            return f"{rendered} (DINOv3更低)"
        if value > 0:
            return f"{rendered} (DINOv3更高)"
        return f"{rendered} (持平)"
    if value > 0:
        return f"{rendered} (DINOv3更高)"
    if value < 0:
        return f"{rendered} (DINOv3更低)"
    return f"{rendered} (持平)"


def read_dinov3_timing(dinov3_dir: Path, dinov2_dir: Path) -> dict[str, float]:
    timing_map: dict[str, float] = {}

    primary_files = {
        "satellite_feature_extraction": dinov2_dir / "timing" / "satellite_feature_extraction_dinov3.json",
        "query_feature_extraction": dinov2_dir / "timing" / "query_feature_extraction_dinov3.json",
        "faiss_index_build": dinov2_dir / "timing" / "faiss_index_build_dinov3_satellite.json",
    }
    for key, path in primary_files.items():
        if path.exists():
            timing_map[key] = float(load_json(path)["elapsed_seconds"])

    retrieval_files = {
        "query_retrieval_top20": dinov3_dir / "timing" / "query_retrieval_dinov3_top20.json",
        "query_retrieval_all": dinov3_dir / "timing" / "query_retrieval_dinov3_all.json",
        "plot_topk_intersection": dinov3_dir / "timing" / "plot_topk_intersection_dinov3.json",
        "plot_topk_fulltruth": dinov3_dir / "timing" / "plot_topk_fulltruth_dinov3.json",
        "visualization": dinov3_dir / "timing" / "visualize_intersection_retrieval_dinov3.json",
        "report_generation": dinov3_dir / "timing" / "report_generation_dinov3.json",
    }
    for key, path in retrieval_files.items():
        if path.exists():
            timing_map[key] = float(load_json(path)["elapsed_seconds"])
    return timing_map


def metric_rows(dinov2_top20: dict, dinov3_top20: dict, dinov2_all: dict, dinov3_all: dict) -> list[tuple[str, float, float, float, bool]]:
    return [
        ("Top-20 R@1", float(dinov2_top20["intersection_recall@1"]), float(dinov3_top20["intersection_recall@1"]), float(dinov3_top20["intersection_recall@1"]) - float(dinov2_top20["intersection_recall@1"]), False),
        ("Top-20 R@5", float(dinov2_top20["intersection_recall@5"]), float(dinov3_top20["intersection_recall@5"]), float(dinov3_top20["intersection_recall@5"]) - float(dinov2_top20["intersection_recall@5"]), False),
        ("Top-20 R@10", float(dinov2_top20["intersection_recall@10"]), float(dinov3_top20["intersection_recall@10"]), float(dinov3_top20["intersection_recall@10"]) - float(dinov2_top20["intersection_recall@10"]), False),
        ("Top-20 R@20", float(dinov2_top20["intersection_recall@20"]), float(dinov3_top20["intersection_recall@20"]), float(dinov3_top20["intersection_recall@20"]) - float(dinov2_top20["intersection_recall@20"]), False),
        ("Top-20 MRR", float(dinov2_top20["intersection_mrr"]), float(dinov3_top20["intersection_mrr"]), float(dinov3_top20["intersection_mrr"]) - float(dinov2_top20["intersection_mrr"]), False),
        ("Top-20 Top-1误差均值(m)", float(dinov2_top20["top1_error_m_mean"]), float(dinov3_top20["top1_error_m_mean"]), float(dinov3_top20["top1_error_m_mean"]) - float(dinov2_top20["top1_error_m_mean"]), True),
        ("全库 MRR", float(dinov2_all["intersection_mrr"]), float(dinov3_all["intersection_mrr"]), float(dinov3_all["intersection_mrr"]) - float(dinov2_all["intersection_mrr"]), False),
    ]


def build_md(
    dinov2_dir: Path,
    dinov3_dir: Path,
    dinov2_top20: dict,
    dinov3_top20: dict,
    dinov2_all: dict,
    dinov3_all: dict,
    dinov2_flights: list[dict],
    dinov3_flights: list[dict],
    dinov3_timing: dict[str, float],
) -> str:
    d2_by_flight = {str(item["flight_id"]): item for item in dinov2_flights}
    d3_by_flight = {str(item["flight_id"]): item for item in dinov3_flights}
    lines: list[str] = []
    lines.append("# DINOv3 vs DINOv2 在 Intersection Truth v2 口径下的对比说明")
    lines.append("")
    lines.append("## 1. 对比目标")
    lines.append("本报告用于对比当前最新 `query v2` 数据口径下，`DINOv2 + FAISS` 与 `DINOv3 + FAISS` 两套全局特征基线在 `intersection truth` 口径下的区域级检索表现。")
    lines.append("")
    lines.append("## 2. 对比前提与数据来源")
    lines.append("- 对比口径统一为 `intersection truth`。")
    lines.append("- 对比 query 统一为 `query_reselect_2026-03-26_v2` 这一批 4 航线共 `40` 张 query。")
    lines.append(f"- DINOv2 来源目录：`{dinov2_dir}`")
    lines.append(f"- DINOv3 来源目录：`{dinov3_dir}`")
    lines.append("- 当前结论只比较正式输出结果，不重新跑实验。")
    lines.append("")
    lines.append("## 3. 总体指标对比")
    lines.append("| 指标 | DINOv2 | DINOv3 | Delta (v3-v2) |")
    lines.append("| --- | ---: | ---: | ---: |")
    for name, d2, d3, delta, invert_good in metric_rows(dinov2_top20, dinov3_top20, dinov2_all, dinov3_all):
        lines.append(f"| {name} | {d2:.3f} | {d3:.3f} | {format_delta(delta, invert_good)} |")
    lines.append("")
    lines.append(
        f"总体上，DINOv2 在当前 `query v2` 口径下优于 DINOv3：DINOv2 的 `R@1={float(dinov2_top20['intersection_recall@1']):.3f}`，DINOv3 的 `R@1={float(dinov3_top20['intersection_recall@1']):.3f}`，两者相差 `{float(dinov3_top20['intersection_recall@1']) - float(dinov2_top20['intersection_recall@1']):+.3f}`；`MRR` 也从 DINOv2 的 `{float(dinov2_top20['intersection_mrr']):.3f}` 降到 DINOv3 的 `{float(dinov3_top20['intersection_mrr']):.3f}`。"
    )
    lines.append(
        f"误差方面，DINOv3 的 `Top-1 error mean` 为 `{float(dinov3_top20['top1_error_m_mean']):.3f}m`，高于 DINOv2 的 `{float(dinov2_top20['top1_error_m_mean']):.3f}m`，说明这次升级没有把首位候选的空间偏差压低。"
    )
    lines.append("")
    lines.append("## 4. 分航线总览对比（Top-20）")
    lines.append("| 航线 | DINOv2 R@1 | DINOv3 R@1 | Delta | DINOv2 MRR | DINOv3 MRR | DINOv2误差(m) | DINOv3误差(m) |")
    lines.append("| --- | ---: | ---: | ---: | ---: | ---: | ---: | ---: |")
    for flight_id in sorted(d2_by_flight):
        d2 = d2_by_flight[flight_id]
        d3 = d3_by_flight[flight_id]
        lines.append(
            f"| {short_flight_name(flight_id)} | {float(d2['intersection_recall@1']):.3f} | {float(d3['intersection_recall@1']):.3f} | {float(d3['intersection_recall@1']) - float(d2['intersection_recall@1']):+.3f} | {float(d2['intersection_mrr']):.3f} | {float(d3['intersection_mrr']):.3f} | {float(d2['top1_error_m_mean']):.3f} | {float(d3['top1_error_m_mean']):.3f} |"
        )
    lines.append("")
    lines.append("- `009` 航线：DINOv2 `R@1=0.800`，DINOv3 `R@1=0.700`，DINOv3 更弱。")
    lines.append("- `010` 航线：两者都较强，但 DINOv2 `R@1=0.900` 仍高于 DINOv3。")
    lines.append("- `011` 航线：DINOv2 `R@1=0.900`，DINOv3 `R@1=0.700`，差距明显。")
    lines.append("- `012` 航线：DINOv2 `R@1=0.700`，DINOv3 `R@1=0.800`，这是 DINOv3 唯一更强的一条航线。")
    lines.append("")
    lines.append("## 5. 时间说明")
    lines.append("当前 `query v2` 目录下未发现 DINOv2 同批次同口径的 timing 文件，因此本报告不做严格公平的 DINOv2 vs DINOv3 耗时横向结论。")
    lines.append("目前能确认的只有 DINOv3 本轮实测时间：")
    lines.append("")
    lines.append("| 阶段 | DINOv3耗时 |")
    lines.append("| --- | ---: |")
    for label, key in [
        ("卫片特征提取", "satellite_feature_extraction"),
        ("FAISS 建库", "faiss_index_build"),
        ("Query 特征提取", "query_feature_extraction"),
        ("检索评估（Top-20）", "query_retrieval_top20"),
    ]:
        if key in dinov3_timing:
            lines.append(f"| {label} | {format_seconds(dinov3_timing[key])} |")
    lines.append("")
    lines.append("补充说明：`new1output/query_reselect_2026-03-26_v2/timing/` 下的 3 个 json 文件名都显式标记为 `dinov3`，应视为 DINOv3 的前处理与建库时间，不属于 DINOv2。")
    lines.append("")
    lines.append("## 6. 结论")
    lines.append("- 在当前最新 `intersection truth v2` 口径下，DINOv2 是更强的全局特征基线。")
    lines.append("- DINOv3 没有带来召回率、MRR 或 Top-1 平均误差的改善。")
    lines.append("- DINOv3 的主要问题不是“Top-20 覆盖不够”，因为两者都已在 `R@10/R@20` 上接近饱和；问题在于前排排序判别力没有优于 DINOv2。")
    lines.append("- 若后续还要比较速度，需要为 DINOv2 `query v2` 这套结果补齐同批次 timing，再做严格耗时对比。")
    lines.append("")
    return "\n".join(lines)


def build_overall_table(doc: Document, dinov2_top20: dict, dinov3_top20: dict, dinov2_all: dict, dinov3_all: dict) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["指标", "DINOv2", "DINOv3", "Delta(v3-v2)"]
    for i, title in enumerate(headers):
        set_cell_text(table.cell(0, i), title, bold=True)
        shade_cell(table.cell(0, i), "D9EAF7")
    for name, d2, d3, delta, invert_good in metric_rows(dinov2_top20, dinov3_top20, dinov2_all, dinov3_all):
        row = table.add_row().cells
        set_cell_text(row[0], name)
        set_cell_text(row[1], f"{d2:.3f}")
        set_cell_text(row[2], f"{d3:.3f}")
        set_cell_text(row[3], format_delta(delta, invert_good))


def build_per_flight_table(doc: Document, dinov2_flights: list[dict], dinov3_flights: list[dict]) -> None:
    d2_by_flight = {str(item["flight_id"]): item for item in dinov2_flights}
    d3_by_flight = {str(item["flight_id"]): item for item in dinov3_flights}
    table = doc.add_table(rows=1, cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["航线", "DINOv2 R@1", "DINOv3 R@1", "Delta", "DINOv2 MRR", "DINOv3 MRR", "DINOv2误差", "DINOv3误差"]
    for i, title in enumerate(headers):
        set_cell_text(table.cell(0, i), title, bold=True)
        shade_cell(table.cell(0, i), "D9EAF7")
    for flight_id in sorted(d2_by_flight):
        d2 = d2_by_flight[flight_id]
        d3 = d3_by_flight[flight_id]
        row = table.add_row().cells
        set_cell_text(row[0], short_flight_name(flight_id))
        set_cell_text(row[1], f"{float(d2['intersection_recall@1']):.3f}")
        set_cell_text(row[2], f"{float(d3['intersection_recall@1']):.3f}")
        set_cell_text(row[3], f"{float(d3['intersection_recall@1']) - float(d2['intersection_recall@1']):+.3f}")
        set_cell_text(row[4], f"{float(d2['intersection_mrr']):.3f}")
        set_cell_text(row[5], f"{float(d3['intersection_mrr']):.3f}")
        set_cell_text(row[6], f"{float(d2['top1_error_m_mean']):.3f}")
        set_cell_text(row[7], f"{float(d3['top1_error_m_mean']):.3f}")


def build_timing_table(doc: Document, dinov3_timing: dict[str, float]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, title in enumerate(["阶段", "DINOv3耗时"]):
        set_cell_text(table.cell(0, i), title, bold=True)
        shade_cell(table.cell(0, i), "D9EAF7")
    for label, key in [
        ("卫片特征提取", "satellite_feature_extraction"),
        ("FAISS 建库", "faiss_index_build"),
        ("Query 特征提取", "query_feature_extraction"),
        ("检索评估（Top-20）", "query_retrieval_top20"),
    ]:
        if key not in dinov3_timing:
            continue
        row = table.add_row().cells
        set_cell_text(row[0], label)
        set_cell_text(row[1], format_seconds(dinov3_timing[key]))


def build_docx(
    out_docx: Path,
    dinov2_dir: Path,
    dinov3_dir: Path,
    dinov2_top20: dict,
    dinov3_top20: dict,
    dinov2_all: dict,
    dinov3_all: dict,
    dinov2_flights: list[dict],
    dinov3_flights: list[dict],
    dinov3_timing: dict[str, float],
) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("DINOv3 vs DINOv2 在 Intersection Truth v2 口径下的对比说明")
    set_cn_font(r, size=16, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("面向 query v2 数据集的正式结果对比报告")
    set_cn_font(r, size=11)

    add_heading(doc, "1. 对比目标", 1)
    add_paragraph(doc, "本报告用于对比当前最新 query v2 数据口径下，DINOv2 + FAISS 与 DINOv3 + FAISS 两套全局特征基线在 intersection truth 口径下的区域级检索表现。")

    add_heading(doc, "2. 对比前提与数据来源", 1)
    add_bullets(
        doc,
        [
            "对比口径统一为 intersection truth。",
            "对比 query 统一为 query_reselect_2026-03-26_v2 这一批 4 航线共 40 张 query。",
            f"DINOv2 来源目录：{dinov2_dir}",
            f"DINOv3 来源目录：{dinov3_dir}",
            "本报告只消费正式输出结果，不重新跑实验。",
        ],
    )

    add_heading(doc, "3. 总体指标对比", 1)
    add_paragraph(doc, "表 1 汇总两套基线在当前 v2 口径下的总体指标差异，Delta 按 DINOv3 减去 DINOv2 计算。")
    build_overall_table(doc, dinov2_top20, dinov3_top20, dinov2_all, dinov3_all)
    add_caption(doc, "表 1  DINOv3 vs DINOv2 总体指标对比")
    add_paragraph(
        doc,
        f"总体上，DINOv2 在当前 v2 query 集上优于 DINOv3。DINOv2 的 Top-20 主口径结果为 `R@1={float(dinov2_top20['intersection_recall@1']):.3f}`、`MRR={float(dinov2_top20['intersection_mrr']):.3f}`、`Top-1误差均值={float(dinov2_top20['top1_error_m_mean']):.3f}m`；"
        f"DINOv3 对应结果为 `R@1={float(dinov3_top20['intersection_recall@1']):.3f}`、`MRR={float(dinov3_top20['intersection_mrr']):.3f}`、`Top-1误差均值={float(dinov3_top20['top1_error_m_mean']):.3f}m`。"
    )
    add_paragraph(doc, "这说明在当前最新正式口径下，DINOv3 没有带来召回率、MRR 或首位候选空间偏差的改善。")

    add_heading(doc, "4. 分航线总览对比", 1)
    add_paragraph(doc, "表 2 用于观察不同航线上的差异是否一致。")
    build_per_flight_table(doc, dinov2_flights, dinov3_flights)
    add_caption(doc, "表 2  分航线 Top-20 结果对比")
    add_bullets(
        doc,
        [
            "009 航线：DINOv2 R@1=0.800，DINOv3 R@1=0.700，DINOv3 更弱。",
            "010 航线：两者都较强，但 DINOv2 R@1=0.900 仍高于 DINOv3。",
            "011 航线：DINOv2 R@1=0.900，DINOv3 R@1=0.700，差距明显。",
            "012 航线：DINOv2 R@1=0.700，DINOv3 R@1=0.800，这是 DINOv3 唯一更强的一条航线。",
        ],
    )

    add_heading(doc, "5. 时间说明", 1)
    add_paragraph(doc, "当前 query v2 目录下未发现 DINOv2 同批次同口径的 timing 文件，因此本报告不对两模型做严格公平的耗时横向结论。")
    add_paragraph(doc, "表 3 仅列出当前已确认的 DINOv3 本轮实测时间。")
    build_timing_table(doc, dinov3_timing)
    add_caption(doc, "表 3  当前可确认的 DINOv3 实测时间")
    add_paragraph(doc, "补充说明：new1output/query_reselect_2026-03-26_v2/timing/ 下的 3 个 json 文件名都显式标记为 dinov3，应视为 DINOv3 的前处理与建库时间，不属于 DINOv2。")

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "6. 结论", 1)
    add_bullets(
        doc,
        [
            "在当前最新 intersection truth v2 口径下，DINOv2 是更强的全局特征基线。",
            "DINOv3 没有带来召回率、MRR 或 Top-1 平均误差的改善。",
            "DINOv3 的主要问题不是 Top-20 覆盖不够，因为两者都在 R@10 和 R@20 上接近饱和；问题在于前排排序判别力没有优于 DINOv2。",
            "若后续还要比较速度，需要为 DINOv2 query v2 这套结果补齐同批次 timing，再做严格耗时对比。",
        ],
    )

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))


def main() -> None:
    args = parse_args()
    dinov2_dir = Path(args.dinov2_dir)
    dinov3_dir = Path(args.dinov3_dir)
    out_md = Path(args.out_md)
    out_docx = Path(args.out_docx)

    dinov2_top20 = load_json(dinov2_dir / "retrieval" / "summary_top20.json")
    dinov2_all = load_json(dinov2_dir / "retrieval" / "summary_all.json")
    dinov3_top20 = load_json(dinov3_dir / "retrieval" / "summary_top20.json")
    dinov3_all = load_json(dinov3_dir / "retrieval" / "summary_all.json")

    dinov2_flights = per_flight_metrics(dinov2_top20)
    dinov3_flights = per_flight_metrics(dinov3_top20)
    dinov3_timing = read_dinov3_timing(dinov3_dir, dinov2_dir)

    out_md.parent.mkdir(parents=True, exist_ok=True)
    out_md.write_text(
        build_md(
            dinov2_dir,
            dinov3_dir,
            dinov2_top20,
            dinov3_top20,
            dinov2_all,
            dinov3_all,
            dinov2_flights,
            dinov3_flights,
            dinov3_timing,
        ),
        encoding="utf-8",
    )
    build_docx(
        out_docx,
        dinov2_dir,
        dinov3_dir,
        dinov2_top20,
        dinov3_top20,
        dinov2_all,
        dinov3_all,
        dinov2_flights,
        dinov3_flights,
        dinov3_timing,
    )
    print(out_md)
    print(out_docx)


if __name__ == "__main__":
    main()
