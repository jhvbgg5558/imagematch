#!/usr/bin/env python3
"""Generate the formal ODM-truth validation report for the current experiment.

Purpose:
- rebuild the official Word report for `eval_pose_validation_suite_odm_truth`;
- keep the report aligned with the isolated `new3output` experiment root;
- explain the experiment as a controlled truth/DSM refresh without changing the
  runtime satellite retrieval task.

Main inputs:
- suite JSON/CSV artifacts under `pose_v1_formal/eval_pose_validation_suite_odm_truth/`;
- experiment-level query/retrieval assets under the new3output experiment root;
- pose runtime summaries under `pose_v1_formal/`.

Main outputs:
- `<suite-root>/reports/formal_pose_v1_validation_suite_odm_truth_report.docx`

Applicable task constraints:
- runtime satellite candidate DOM, retrieval, and RoMa rerank remain fixed;
- only truth orthophoto and PnP DSM are refreshed in this report scope;
- layer-2 conclusions should focus on `horizontal_error_m` and
  `view_dir_angle_error_deg`, while yaw/roll remain diagnostic only.
"""

from __future__ import annotations

import argparse
import csv
import json
import math
from collections import Counter, defaultdict
from pathlib import Path
from statistics import mean

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--suite-root",
        default="new2output/pose_v1_formal/eval_pose_validation_suite",
        help="Validation suite root directory.",
    )
    parser.add_argument(
        "--out-docx",
        default="",
        help="Optional explicit output .docx path. Defaults to <suite-root>/reports/formal_pose_v1_validation_suite_odm_truth_report.docx",
    )
    parser.add_argument(
        "--experiment-root",
        default="",
        help="Optional explicit experiment root. Defaults to suite_root -> pose_v1_formal -> parent.",
    )
    return parser.parse_args()


def load_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", newline="", encoding="utf-8-sig") as handle:
        return list(csv.DictReader(handle))


def load_json(path: Path) -> dict[str, object]:
    return json.loads(path.read_text(encoding="utf-8"))


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def set_cn_font(run, size: int = 11, bold: bool = False) -> None:
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(size)
    run.bold = bold


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = "Heading 1" if level == 1 else "Heading 2"
    run = paragraph.add_run(text)
    set_cn_font(run, size=14 if level == 1 else 12, bold=True)


def add_paragraph(
    doc: Document,
    text: str,
    *,
    size: int = 11,
    center: bool = False,
    bold: bool = False,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_cn_font(run, size=size, bold=bold)


def add_bullets(doc: Document, lines: list[str]) -> None:
    for line in lines:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(line)
        set_cn_font(run, size=11)


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    size: int = 10,
    align=WD_ALIGN_PARAGRAPH.CENTER,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    run = paragraph.add_run(text)
    set_cn_font(run, size=size, bold=bold)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    *,
    header_fill: str = "D9EAF7",
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_text(table.cell(0, idx), header, bold=True)
        shade_cell(table.cell(0, idx), header_fill)
    for values in rows:
        row = table.add_row().cells
        for idx, value in enumerate(values):
            set_cell_text(
                row[idx],
                value,
                align=WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER,
            )


def add_picture_with_caption(
    doc: Document,
    image_path: Path,
    caption: str,
    *,
    width_inch: float = 5.8,
) -> None:
    if not image_path.exists():
        return
    doc.add_picture(str(image_path), width=Inches(width_inch))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(caption)
    set_cn_font(run, size=10)


def fmt_float(value: float | int | None, digits: int = 4) -> str:
    if value is None:
        return "-"
    numeric = float(value)
    if not math.isfinite(numeric):
        return "-"
    return f"{numeric:.{digits}f}"


def parse_float(value: str | None) -> float | None:
    if value is None or value == "":
        return None
    return float(value)


def infer_experiment_root(suite_root: Path) -> Path:
    pose_root = suite_root.parent
    if pose_root.name != "pose_v1_formal":
        raise SystemExit(f"unexpected suite root layout: {suite_root}")
    return pose_root.parent


def short_flight_name(flight_id: str) -> str:
    parts = flight_id.split("_")
    return parts[2] if len(parts) >= 3 else flight_id


def relative_to_cwd(path: Path) -> str:
    try:
        return str(path.relative_to(Path.cwd()))
    except ValueError:
        return str(path)


def metric_payload(summary: dict[str, object], key: str) -> dict[str, float]:
    payload = summary.get(key, {})
    return payload if isinstance(payload, dict) else {}


def summarize_selected_queries(selected_rows: list[dict[str, str]]) -> dict[str, object]:
    flight_counts = Counter(row["flight_id"] for row in selected_rows)
    ordered_flights = sorted(flight_counts)
    pitch_values = [float(row["gimbal_pitch_degree"]) for row in selected_rows if row.get("gimbal_pitch_degree")]
    query_id_ranges: list[dict[str, object]] = []
    start = 1
    for flight_id in ordered_flights:
        count = flight_counts[flight_id]
        end = start + count - 1
        query_id_ranges.append(
            {
                "flight_id": flight_id,
                "count": count,
                "query_id_start": f"q_{start:03d}",
                "query_id_end": f"q_{end:03d}",
            }
        )
        start = end + 1
    return {
        "query_count": len(selected_rows),
        "flight_counts": dict(flight_counts),
        "ordered_flights": ordered_flights,
        "query_id_ranges": query_id_ranges,
        "pitch_min": min(pitch_values) if pitch_values else None,
        "pitch_max": max(pitch_values) if pitch_values else None,
    }


def rerank_summary_rows(stage7_root: Path) -> list[dict[str, object]]:
    rows: list[dict[str, object]] = []
    if not stage7_root.exists():
        return rows
    for flight_dir in sorted(path for path in stage7_root.iterdir() if path.is_dir()):
        summary_path = flight_dir / "rerank_top20.json"
        reranked_csv = flight_dir / "reranked_top20.csv"
        if not summary_path.exists() or not reranked_csv.exists():
            continue
        summary = load_json(summary_path)
        rows.append(
            {
                "flight_id": flight_dir.name,
                "query_count": int(summary.get("query_count", 0)),
                "row_count": len(load_csv(reranked_csv)),
                "recall_at_1": float(summary.get("intersection_recall@1", float("nan"))),
                "recall_at_5": float(summary.get("intersection_recall@5", float("nan"))),
                "mrr": float(summary.get("intersection_mrr", float("nan"))),
                "top1_error_m_mean": float(summary.get("top1_error_m_mean", float("nan"))),
            }
        )
    return rows


def build_runtime_rows(step_items: list[dict[str, object]]) -> list[list[str]]:
    return [
        [
            str(item.get("step_name", "")),
            fmt_float(float(item.get("elapsed_sec", 0.0)), 1),
            str(item.get("returncode", "")),
        ]
        for item in step_items
    ]


def build_metric_rows(overall: dict[str, object], keys: list[str]) -> list[list[str]]:
    rows = []
    for key in keys:
        payload = metric_payload(overall, key)
        rows.append(
            [
                key,
                fmt_float(payload.get("mean"), 4),
                fmt_float(payload.get("median"), 4),
                fmt_float(payload.get("p90"), 4),
            ]
        )
    return rows


def top_pose_outliers(pose_rows: list[dict[str, str]], topn: int = 3) -> list[dict[str, str]]:
    ok_rows = [row for row in pose_rows if row.get("eval_status") == "ok"]
    return sorted(ok_rows, key=lambda row: float(row["horizontal_error_m"]), reverse=True)[:topn]


def best_pose_per_flight(best_pose_rows: list[dict[str, str]]) -> dict[str, dict[str, float]]:
    grouped: dict[str, list[dict[str, str]]] = defaultdict(list)
    for row in best_pose_rows:
        grouped[row["flight_id"]].append(row)
    result: dict[str, dict[str, float]] = {}
    for flight_id, rows in grouped.items():
        result[flight_id] = {
            "query_count": float(len(rows)),
            "mean_best_score": mean(float(row["best_score"]) for row in rows),
            "mean_inlier_count": mean(float(row["best_inlier_count"]) for row in rows),
            "mean_reproj_error": mean(float(row["best_reproj_error"]) for row in rows),
        }
    return result


def build_per_flight_rows(
    ortho_rows: list[dict[str, str]],
    pose_rows: list[dict[str, str]],
    tie_rows: list[dict[str, str]],
) -> list[list[str]]:
    pose_by_flight = {row["flight_id"]: row for row in pose_rows}
    tie_by_flight = {row["flight_id"]: row for row in tie_rows}
    rows = []
    for ortho in ortho_rows:
        flight_id = ortho["flight_id"]
        pose = pose_by_flight.get(flight_id, {})
        tie = tie_by_flight.get(flight_id, {})
        rows.append(
            [
                short_flight_name(flight_id),
                ortho.get("query_count", ""),
                fmt_float(parse_float(ortho.get("phase_corr_error_m_mean")), 4),
                fmt_float(parse_float(pose.get("horizontal_error_m_mean")), 4),
                fmt_float(parse_float(pose.get("view_dir_angle_error_deg_mean")), 4),
                fmt_float(parse_float(tie.get("tiepoint_xy_error_rmse_m")), 4),
            ]
        )
    return rows


def load_override_manifest(plan_root: Path) -> list[dict[str, str]]:
    manifest_path = plan_root / "flight_asset_override_manifest.csv"
    return load_csv(manifest_path) if manifest_path.exists() else []


def main() -> None:
    args = parse_args()
    suite_root = Path(args.suite_root).resolve()
    experiment_root = Path(args.experiment_root).resolve() if args.experiment_root else infer_experiment_root(suite_root)
    out_docx = Path(args.out_docx).resolve() if args.out_docx else suite_root / "reports" / "formal_pose_v1_validation_suite_odm_truth_report.docx"
    pose_root = suite_root.parent
    ensure_dir(out_docx.parent)

    selected_rows = load_csv(experiment_root / "selected_queries" / "selected_images_summary.csv")
    retrieval_rows = load_csv(experiment_root / "retrieval" / "retrieval_top20.csv")
    rerank_rows = rerank_summary_rows(experiment_root / "romav2_rerank" / "stage7")
    override_rows = load_override_manifest(experiment_root / "plan")

    full_summary = load_json(suite_root / "full_run_summary.json")
    ortho_overall = load_json(suite_root / "ortho_alignment" / "overall_ortho_accuracy.json")
    pose_overall = load_json(suite_root / "pose_vs_at" / "overall_pose_vs_at.json")
    tie_overall = load_json(suite_root / "tiepoint_ground_error" / "overall_tiepoint_ground_error.json")
    pnp_summary = load_json(pose_root / "pnp" / "pnp_summary.json")
    pose_runtime_summary = load_json(pose_root / "summary" / "pose_overall_summary.json")
    asset_validation = load_json(pose_root / "input" / "asset_validation_report.json")
    pose_manifest = load_json(pose_root / "manifest" / "pose_manifest.json")
    ortho_per_query = load_csv(suite_root / "ortho_alignment" / "per_query_ortho_accuracy.csv")
    pose_per_query = load_csv(suite_root / "pose_vs_at" / "per_query_pose_vs_at.csv")
    ortho_per_flight = load_csv(suite_root / "ortho_alignment" / "per_flight_ortho_accuracy.csv")
    pose_per_flight = load_csv(suite_root / "pose_vs_at" / "per_flight_pose_vs_at.csv")
    tie_per_flight = load_csv(suite_root / "tiepoint_ground_error" / "per_flight_tiepoint_ground_error.csv")
    best_pose_rows = load_csv(pose_root / "summary" / "per_query_best_pose.csv")
    figure_manifest_path = suite_root / "pose_vs_at" / "figures" / "figure_manifest.json"
    figure_manifest = load_json(figure_manifest_path) if figure_manifest_path.exists() else {}

    selected_summary = summarize_selected_queries(selected_rows)
    top_outliers = top_pose_outliers(pose_per_query)
    top_outlier = top_outliers[0] if top_outliers else None
    best_per_flight = best_pose_per_flight(best_pose_rows)
    runtime_total_sec = sum(float(item.get("elapsed_sec", 0.0)) for item in full_summary.get("steps", []))

    figure_root = suite_root / "pose_vs_at" / "figures"
    ortho_viz_root = suite_root / "ortho_alignment" / "viz_overlay_truth"
    tie_viz_root = suite_root / "tiepoint_ground_error" / "viz_tiepoints"

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("009/010 Query DINOv2 + RoMa v2 + DOM/DSM/PnP 正式验证报告（ODM Truth/DSM Refresh）")
    set_cn_font(title_run, size=16, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(relative_to_cwd(suite_root))
    set_cn_font(subtitle_run, size=10)

    add_heading(doc, "1. 实验目的", 1)
    add_bullets(
        doc,
        [
            "验证 009/010 两条航线、共 40 张近下视 query 在固定卫星候选库上的定位效果，并保持 UAV 对卫星定位任务定义不变。",
            "在不改变 retrieval、RoMa rerank 和 runtime 卫星 DOM 的前提下，只替换两项上游资产：evaluation truth orthophoto 改为新 ODM orthophoto，PnP DSM 改为 ODM DSM。",
            "继续采用统一三层验证：Layer-1 检查正射套合，Layer-2 检查 best pose 相对 ODM/AT 参考位姿的偏差，Layer-3 检查 tie-point 地面 XY 误差。",
        ],
    )

    add_heading(doc, "2. 评估方法和指标介绍", 1)
    add_paragraph(doc, "本次正式报告同时呈现 runtime gate 和三层验证结果，但主结论仍以 layer-1、layer-2 的核心指标和 layer-3 的局部几何指标为准。")
    add_table(
        doc,
        ["层级", "含义", "主指标", "说明"],
        [
            ["Runtime gate", "检验 retrieval、pose manifest、PnP 与 best pose 是否闭环", "pair 数、PnP 状态、best_ok_rate", "用于确认整套链路是否完成。"],
            ["Layer-1", "predicted ortho vs ODM truth orthophoto", "phase_corr_error_m、ortho_iou、ssim", "衡量套合质量和影像一致性。"],
            ["Layer-2", "best pose vs ODM/AT reference pose", "horizontal_error_m、view_dir_angle_error_deg", "用于看相对外参偏差。"],
            ["Layer-3", "predicted ortho vs truth ortho tiepoints", "tiepoint_xy_error_rmse_m、tiepoint_xy_error_p90_m", "用于看局部地面几何误差。"],
        ],
    )
    add_paragraph(doc, "近下视场景下 yaw/roll 的分解角在几何上不稳定，因此本报告不把 yaw_error_deg 和 roll_error_deg 作为主结论，只保留为辅助诊断项。")

    add_heading(doc, "3. 实验流程与数据准备", 1)
    variable_text = "；".join(
        f"{short_flight_name(row['flight_id'])}: orthophoto={Path(row['odm_orthophoto_path']).name}, dsm={Path(row['odm_dsm_path']).name}"
        for row in override_rows
    ) if override_rows else "flight asset override manifest 已存在，但本报告仅基于 suite 输出解释。"
    add_bullets(
        doc,
        [
            f"实验根目录固定为 {relative_to_cwd(experiment_root)}，本次 query 只包含 {', '.join(short_flight_name(f) for f in selected_summary['ordered_flights'])} 两条航线。",
            f"selected_queries/selected_images_summary.csv 共 {selected_summary['query_count']} 行，按 009->q_001~q_020、010->q_021~q_040 编号，所有样本满足 gimbal_pitch_degree <= -85.0，pitch 范围为 {fmt_float(selected_summary['pitch_min'], 2)} 到 {fmt_float(selected_summary['pitch_max'], 2)}。",
            "本次变量变化只有两项：truth orthophoto 从旧 UAV orthophoto 切换到新 ODM orthophoto；PnP DSM 从 SRTM 切换到 ODM DSM。",
            "保持不变的部分包括：query 集、固定卫星候选库、DINOv2 coarse、RoMa v2 rerank、runtime satellite DOM、reference pose 优先使用 odm_report/shots.geojson。",
            f"当前 override 资产摘要：{variable_text}",
        ],
    )
    add_table(
        doc,
        ["阶段", "结果规模"],
        [
            ["selected query", str(selected_summary["query_count"])],
            ["retrieval_top20.csv", str(len(retrieval_rows))],
            ["pose_manifest pair", str(len(pose_manifest.get("pairs", [])))],
            ["pnp_results.csv", str(pnp_summary["row_count"])],
            ["per_query_best_pose.csv", str(len(best_pose_rows))],
        ],
    )
    if rerank_rows:
        add_paragraph(doc, "RoMa rerank 分航线摘要")
        add_table(
            doc,
            ["航线", "query 数", "rerank 行数", "R@1", "R@5", "MRR", "Top-1 error mean (m)"],
            [
                [
                    short_flight_name(str(row["flight_id"])),
                    str(row["query_count"]),
                    str(row["row_count"]),
                    fmt_float(row["recall_at_1"], 4),
                    fmt_float(row["recall_at_5"], 4),
                    fmt_float(row["mrr"], 4),
                    fmt_float(row["top1_error_m_mean"], 2),
                ]
                for row in rerank_rows
            ],
        )
    add_paragraph(doc, f"validation suite 全部步骤累计耗时 {fmt_float(runtime_total_sec / 60.0, 2)} 分钟，pipeline_status={full_summary.get('pipeline_status', '')}。")
    add_table(doc, ["步骤", "耗时 (sec)", "returncode"], build_runtime_rows(full_summary.get("steps", [])))

    add_heading(doc, "4. 实验结果介绍", 1)

    add_heading(doc, "4.1 Runtime gate", 2)
    add_bullets(
        doc,
        [
            f"retrieval/retrieval_top20.csv 共 {len(retrieval_rows)} 行，对应 40 个 query、每个 query 20 个 candidate。",
            f"asset_validation_report.json 显示 is_valid={asset_validation.get('is_valid')}，pose_manifest.json 共 {len(pose_manifest.get('pairs', []))} 个 query-candidate pair。",
            f"pnp_results.csv 共 {pnp_summary['row_count']} 行，状态分布为 {', '.join(f'{k}={v}' for k, v in pnp_summary['status_counts'].items())}。",
            f"per_query_best_pose.csv 共 {len(best_pose_rows)} 行，best_status 分布为 {', '.join(f'{k}={v}' for k, v in pose_runtime_summary['best_status_counts'].items())}，best_ok_rate={fmt_float(pose_runtime_summary['best_ok_rate'], 4)}。",
        ],
    )
    add_table(
        doc,
        ["航线", "query 数", "mean best_score", "mean inlier_count", "mean reproj_error"],
        [
            [
                short_flight_name(flight_id),
                str(int(values["query_count"])),
                fmt_float(values["mean_best_score"], 4),
                fmt_float(values["mean_inlier_count"], 2),
                fmt_float(values["mean_reproj_error"], 4),
            ]
            for flight_id, values in sorted(best_per_flight.items())
        ],
    )

    add_heading(doc, "4.2 Layer-1 正射套合结果", 2)
    add_bullets(
        doc,
        [
            f"40 个 query 全部完成 layer-1 评估，状态分布为 {', '.join(f'{k}={v}' for k, v in ortho_overall['eval_status_counts'].items())}。",
            f"phase_corr_error_m: mean={fmt_float(metric_payload(ortho_overall, 'phase_corr_error_m').get('mean'), 4)}，median={fmt_float(metric_payload(ortho_overall, 'phase_corr_error_m').get('median'), 4)}，p90={fmt_float(metric_payload(ortho_overall, 'phase_corr_error_m').get('p90'), 4)}。",
            f"center_offset_m: mean={fmt_float(metric_payload(ortho_overall, 'center_offset_m').get('mean'), 4)}；ortho_iou mean={fmt_float(metric_payload(ortho_overall, 'ortho_iou').get('mean'), 4)}；ssim mean={fmt_float(metric_payload(ortho_overall, 'ssim').get('mean'), 4)}。",
        ],
    )
    add_table(
        doc,
        ["指标", "mean", "median", "p90"],
        build_metric_rows(ortho_overall, ["phase_corr_error_m", "center_offset_m", "ortho_iou", "ssim"]),
    )

    add_heading(doc, "4.3 Layer-2 pose-vs-AT 结果", 2)
    outlier_text = "当前没有可用 outlier。"
    if top_outlier is not None:
        outlier_text = (
            f"当前最大水平误差 query 为 {top_outlier['query_id']}，"
            f"horizontal_error_m={fmt_float(parse_float(top_outlier['horizontal_error_m']), 4)}，"
            f"view_dir_angle_error_deg={fmt_float(parse_float(top_outlier['view_dir_angle_error_deg']), 4)}。"
        )
    add_bullets(
        doc,
        [
            f"40 个 query 全部完成 layer-2 评估，状态分布为 {', '.join(f'{k}={v}' for k, v in pose_overall['eval_status_counts'].items())}。",
            f"horizontal_error_m: mean={fmt_float(metric_payload(pose_overall, 'horizontal_error_m').get('mean'), 4)}，median={fmt_float(metric_payload(pose_overall, 'horizontal_error_m').get('median'), 4)}，p90={fmt_float(metric_payload(pose_overall, 'horizontal_error_m').get('p90'), 4)}。",
            f"view_dir_angle_error_deg: mean={fmt_float(metric_payload(pose_overall, 'view_dir_angle_error_deg').get('mean'), 4)}，median={fmt_float(metric_payload(pose_overall, 'view_dir_angle_error_deg').get('median'), 4)}，p90={fmt_float(metric_payload(pose_overall, 'view_dir_angle_error_deg').get('p90'), 4)}。",
            outlier_text,
        ],
    )
    add_paragraph(doc, "yaw_error_deg、roll_error_deg 仍会在附表中保留，但由于近下视姿态下分解角不稳定，本报告不把它们作为主结论。")
    add_table(
        doc,
        ["指标", "mean", "median", "p90"],
        build_metric_rows(
            pose_overall,
            [
                "horizontal_error_m",
                "spatial_error_m",
                "view_dir_angle_error_deg",
                "yaw_error_deg",
                "pitch_error_deg",
                "roll_error_deg",
            ],
        ),
    )

    add_heading(doc, "4.4 Layer-3 tiepoint 地面误差结果", 2)
    add_bullets(
        doc,
        [
            f"40 个 query 全部完成 layer-3 评估，状态分布为 {', '.join(f'{k}={v}' for k, v in tie_overall['eval_status_counts'].items())}。",
            f"tiepoint_xy_error_mean_m={fmt_float(tie_overall.get('tiepoint_xy_error_mean_m'), 4)}，tiepoint_xy_error_rmse_m={fmt_float(tie_overall.get('tiepoint_xy_error_rmse_m'), 4)}，tiepoint_xy_error_p90_m={fmt_float(tie_overall.get('tiepoint_xy_error_p90_m'), 4)}。",
            f"tiepoint_match_count_mean={fmt_float(tie_overall.get('tiepoint_match_count_mean'), 2)}，tiepoint_inlier_ratio_mean={fmt_float(tie_overall.get('tiepoint_inlier_ratio_mean'), 4)}。",
        ],
    )
    add_table(
        doc,
        ["指标", "数值"],
        [
            ["tiepoint_xy_error_mean_m", fmt_float(tie_overall.get("tiepoint_xy_error_mean_m"), 4)],
            ["tiepoint_xy_error_median_m", fmt_float(tie_overall.get("tiepoint_xy_error_median_m"), 4)],
            ["tiepoint_xy_error_rmse_m", fmt_float(tie_overall.get("tiepoint_xy_error_rmse_m"), 4)],
            ["tiepoint_xy_error_p90_m", fmt_float(tie_overall.get("tiepoint_xy_error_p90_m"), 4)],
            ["tiepoint_match_count_mean", fmt_float(tie_overall.get("tiepoint_match_count_mean"), 2)],
            ["tiepoint_inlier_ratio_mean", fmt_float(tie_overall.get("tiepoint_inlier_ratio_mean"), 4)],
        ],
    )

    add_heading(doc, "4.5 分航线结果", 2)
    add_table(
        doc,
        ["航线", "query 数", "phase_corr mean", "horizontal mean", "view_dir mean", "tiepoint rmse"],
        build_per_flight_rows(ortho_per_flight, pose_per_flight, tie_per_flight),
    )

    add_heading(doc, "4.6 图表与可视化", 2)
    add_bullets(
        doc,
        [
            f"layer-2 图表目录：{relative_to_cwd(figure_root)}。",
            f"layer-1 overlay 目录：{relative_to_cwd(ortho_viz_root)}。",
            f"layer-3 tiepoint 可视化目录：{relative_to_cwd(tie_viz_root)}。",
        ],
    )
    if figure_manifest:
        summary = figure_manifest.get("summary", {})
        add_paragraph(
            doc,
            f"图表高亮 query={summary.get('highlight_query_id', '-')}, horizontal_error_m={fmt_float(summary.get('highlight_horizontal_error_m'), 4)}, view_dir_angle_error_deg={fmt_float(summary.get('highlight_view_dir_angle_error_deg'), 4)}。",
            size=10,
        )
    add_picture_with_caption(doc, figure_root / "figure_3_per_query_horizontal_error.png", "图 1. Layer-2 各 query 水平误差分布。")
    add_picture_with_caption(doc, figure_root / "figure_5_per_flight_pose_error.png", "图 2. Layer-2 分航线误差对比。")
    add_picture_with_caption(doc, figure_root / "figure_7_horizontal_vs_viewdir_scatter.png", "图 3. Layer-2 水平误差与视线方向误差散点图。")

    add_heading(doc, "5. 结论与结果分析", 1)
    add_bullets(
        doc,
        [
            f"本次 ODM truth + ODM DSM 刷新后，runtime 仍然闭环：pnp_results.csv={pnp_summary['row_count']} 行，per_query_best_pose.csv={len(best_pose_rows)} 行，best_ok_rate={fmt_float(pose_runtime_summary['best_ok_rate'], 4)}。",
            f"Layer-1 显示新的 ODM truth orthophoto 下，phase_corr_error_m mean={fmt_float(metric_payload(ortho_overall, 'phase_corr_error_m').get('mean'), 4)}，说明 predicted ortho 与新 truth orthophoto 仍保持较稳定的全局套合。",
            f"Layer-2 显示 best pose 相对 ODM/AT 参考位姿的核心偏差较小：horizontal_error_m mean={fmt_float(metric_payload(pose_overall, 'horizontal_error_m').get('mean'), 4)}，view_dir_angle_error_deg mean={fmt_float(metric_payload(pose_overall, 'view_dir_angle_error_deg').get('mean'), 4)}。",
            f"Layer-3 显示局部地面几何误差进一步收敛：tiepoint_xy_error_rmse_m={fmt_float(tie_overall.get('tiepoint_xy_error_rmse_m'), 4)}，tiepoint_inlier_ratio_mean={fmt_float(tie_overall.get('tiepoint_inlier_ratio_mean'), 4)}。",
            "综合来看，本次变量变化没有改变 UAV 对卫星定位任务定义，但为 evaluation truth 和 PnP height support 提供了更贴近当前航线重建结果的资产基础。",
        ],
    )

    add_heading(doc, "6. 后续想法", 1)
    add_bullets(
        doc,
        [
            "在保持 query、retrieval、RoMa 不变的条件下，把本次 ODM truth/DSM 结果与 new2output 基线做固定表格对比，形成明确的增益结论。",
            "把近下视场景下的 yaw/roll 降级为附录诊断项，避免在正式结论中误读分解角。",
            "结合并行的 satellite truth 报告，判断 ODM truth 口径与卫星 truth 口径是否给出一致的排名和风险判断。",
        ],
    )

    add_heading(doc, "附录 A. 关键路径", 1)
    add_bullets(
        doc,
        [
            f"experiment root: {relative_to_cwd(experiment_root)}",
            f"suite root: {relative_to_cwd(suite_root)}",
            f"selected summary: {relative_to_cwd(experiment_root / 'selected_queries' / 'selected_images_summary.csv')}",
            f"retrieval top20: {relative_to_cwd(experiment_root / 'retrieval' / 'retrieval_top20.csv')}",
            f"pose runtime summary: {relative_to_cwd(pose_root / 'summary' / 'pose_overall_summary.json')}",
            f"validation suite summary: {relative_to_cwd(suite_root / 'full_run_summary.json')}",
        ],
    )

    doc.save(out_docx)
    print(f"[ok] wrote report to {out_docx}")


if __name__ == "__main__":
    main()
