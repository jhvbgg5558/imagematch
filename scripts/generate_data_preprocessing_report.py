#!/usr/bin/env python3
"""Generate a formal Chinese report about current data inputs and preprocessing."""

from __future__ import annotations

import argparse
import csv
import json
import shutil
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--dataset-summary-json", required=True)
    parser.add_argument("--selected-summary-csv", required=True)
    parser.add_argument("--query-manifest-csv", required=True)
    parser.add_argument("--query-seed-csv", required=True)
    parser.add_argument("--query-truth-tiles-csv", required=True)
    parser.add_argument("--tile-metadata-csv", required=True)
    parser.add_argument("--range-fig-dir", required=True)
    parser.add_argument("--scale-note-dir", required=True)
    parser.add_argument("--out-docx", required=True)
    parser.add_argument("--out-md", required=True)
    return parser.parse_args()


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def load_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", newline="", encoding="utf-8-sig") as f:
        return list(csv.DictReader(f))


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def set_cn_font(run, size: int = 11, bold: bool = False) -> None:
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(size)
    run.bold = bold


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 10, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    set_cn_font(r, size=size, bold=bold)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.style = "Heading 1" if level == 1 else "Heading 2"
    r = p.add_run(text)
    set_cn_font(r, size=14 if level == 1 else 12, bold=True)


def add_paragraph(doc: Document, text: str, size: int = 11, center: bool = False) -> None:
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn_font(r, size=size)


def add_bullets(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(line)
        set_cn_font(r, size=11)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn_font(r, size=10)


def add_picture(doc: Document, path: Path, width_inch: float) -> None:
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(width_inch))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def copy_figures(paths: list[Path], out_dir: Path) -> dict[str, Path]:
    ensure_dir(out_dir)
    copied: dict[str, Path] = {}
    for path in paths:
        dst = out_dir / path.name
        shutil.copy2(path, dst)
        copied[path.name] = dst
    return copied


def km2(value_m2: float) -> float:
    return value_m2 / 1_000_000.0


def build_stats(
    dataset_summary: dict,
    selected_rows: list[dict[str, str]],
    query_manifest_rows: list[dict[str, str]],
    query_seed_rows: list[dict[str, str]],
    truth_tile_rows: list[dict[str, str]],
    tile_rows: list[dict[str, str]],
) -> dict:
    selected_by_flight = Counter(row["flight_id"] for row in selected_rows)
    selected_tag_counts = Counter()
    for row in selected_rows:
        tags = row["tags"]
        if tags == "downview":
            selected_tag_counts["downview"] += 1
        elif "tilted" in tags:
            selected_tag_counts["tilted"] += 1
        else:
            selected_tag_counts["diverse_extra"] += 1

    truth_by_query: dict[str, list[dict[str, str]]] = defaultdict(list)
    truth_scale_counts = Counter()
    truth_tile_unique = set()
    for row in truth_tile_rows:
        truth_by_query[row["query_id"]].append(row)
        truth_scale_counts[f"{int(float(row['tile_size_m']))}m"] += 1
        truth_tile_unique.add(row["tile_id"])

    truth_counts = [len(truth_by_query[row["query_id"]]) for row in query_seed_rows]
    contain_center_count = sum(int(row["contains_query_center"]) for row in truth_tile_rows)

    gsd_values = [float(row["gsd_x_m_per_px"]) for row in tile_rows]
    native_widths = [int(row["native_width"]) for row in tile_rows]
    native_heights = [int(row["native_height"]) for row in tile_rows]

    return {
        "query_count": len(query_seed_rows),
        "flight_count": len(selected_by_flight),
        "selected_by_flight": dict(sorted(selected_by_flight.items())),
        "selected_near": selected_tag_counts["downview"],
        "selected_downview": selected_tag_counts["downview"],
        "selected_tilt": selected_tag_counts["tilted"],
        "selected_extra": selected_tag_counts["diverse_extra"],
        "sanitized_query_count": sum(int(row["has_metadata_removed"]) for row in query_manifest_rows),
        "roi_area_km2": km2(float(dataset_summary["roi_area_m2"])),
        "raw_area_km2": km2(float(dataset_summary["raw_bbox_area_m2"])),
        "buffer_gain_km2": km2(float(dataset_summary["buffer_area_gain_m2"])),
        "tile_count_total": len(tile_rows),
        "truth_pair_total": len(truth_tile_rows),
        "truth_unique_tile_count": len(truth_tile_unique),
        "truth_count_mean": sum(truth_counts) / len(truth_counts) if truth_counts else 0.0,
        "truth_count_min": min(truth_counts) if truth_counts else 0,
        "truth_count_max": max(truth_counts) if truth_counts else 0,
        "truth_contain_center_ratio": contain_center_count / len(truth_tile_rows) if truth_tile_rows else 0.0,
        "truth_scale_counts": {
            key: truth_scale_counts.get(key, 0) for key in ["200m", "300m", "500m", "700m"]
        },
        "footprint_area_mean_ha": float(dataset_summary["footprint_area_m2_mean"]) / 1e4,
        "footprint_area_min_ha": float(dataset_summary["footprint_area_m2_min"]) / 1e4,
        "footprint_area_max_ha": float(dataset_summary["footprint_area_m2_max"]) / 1e4,
        "satellite_gsd_mean_m": sum(gsd_values) / len(gsd_values) if gsd_values else 0.0,
        "tile_native_width_min": min(native_widths) if native_widths else 0,
        "tile_native_width_max": max(native_widths) if native_widths else 0,
        "tile_native_height_min": min(native_heights) if native_heights else 0,
        "tile_native_height_max": max(native_heights) if native_heights else 0,
    }


def build_input_table(doc: Document, dataset_summary: dict, stats: dict) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["数据类型", "来源/位置", "空间属性", "当前用途"]
    for i, title in enumerate(headers):
        set_cell_text(table.cell(0, i), title, bold=True)
        shade_cell(table.cell(0, i), "D9EAF7")

    rows = [
        (
            "无人机原始影像",
            r"D:\数据\武汉影像\无人机0.1m",
            f"4条航线，原始经纬度可从DJI元数据读取；当前统一转换到 {dataset_summary['query_crs']}",
            "作为 query 候选来源，并用于真值几何构建",
        ),
        (
            "query 选中集",
            r"D:\数据\武汉影像\挑选无人机0.1m",
            f"共 {stats['query_count']} 张，按 4 条航线组织；目录名表征原始分辨率约为 0.1m",
            "形成正式 query 清单与人工可复核样本集",
        ),
        (
            "去元数据 query",
            r"D:\aiproject\imagematch\output\query_sanitized_40_v2",
            "保留像素内容，移除 EXIF/XMP/GPS/DJI 元数据",
            "作为检索模型的正式 query 输入",
        ),
        (
            "卫星大图与切片库",
            r"D:\数据\武汉影像\卫星0.5m_32650\武汉_大图\L19",
            f"卫片坐标系 {dataset_summary['satellite_crs']}；当前切成 {stats['tile_count_total']} 张多尺度 tile",
            "构建固定候选库与后续特征索引",
        ),
    ]
    for values in rows:
        row = table.add_row().cells
        for i, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.LEFT if i in {1, 2, 3} else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(row[i], value, align=align)


def build_asset_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["资产", "关键文件", "说明"]
    for i, title in enumerate(headers):
        set_cell_text(table.cell(0, i), title, bold=True)
        shade_cell(table.cell(0, i), "D9EAF7")
    rows = [
        ("query 选中清单", "selected_images_summary.csv", "记录 40 张 query 的原始来源、元数据标签与筛选原因。"),
        ("query 实验清单", "query_manifest.csv", "记录原图与去元数据副本的对应关系。"),
        ("query 几何主表", "queries_truth_seed.csv", "记录 query 中心点、footprint、多边形范围与坐标系。"),
        ("卫片候选库索引", "tiles.csv / roi_summary.json", "记录 tile 的尺度、范围、中心点、原始 GeoTIFF 来源。"),
        ("交集真值表", "query_truth_tiles.csv", "记录 query 与卫片 tile 的交集关系，是当前正式真值来源。"),
    ]
    for values in rows:
        row = table.add_row().cells
        set_cell_text(row[0], values[0])
        set_cell_text(row[1], values[1], align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(row[2], values[2], align=WD_ALIGN_PARAGRAPH.LEFT)


def build_markdown(dataset_summary: dict, stats: dict, fig_names: dict[str, str]) -> str:
    lines: list[str] = []
    lines.append("# 当前任务数据输入与预处理说明")
    lines.append("")
    lines.append("## 1. 文档目的")
    lines.append("本报告用于系统说明当前项目在正式检索实验之前完成的数据整理与预处理工作，重点回答以下问题：")
    lines.append("")
    lines.extend(
        [
            "- 当前实验到底使用了哪些航片与卫片数据；",
            "- 40 张 query 是如何从原始无人机影像中得到的；",
            "- 卫片候选库是如何围绕 4 条航线构建的；",
            "- 为什么 query 在检索时需要去除坐标元数据；",
            "- 为什么卫片按 `200 / 300 / 500 / 700m` 多尺度划分，并设置 `overlap = 0.25`；",
            "- 当前正式真值为什么定义为 `intersection truth`。",
        ]
    )
    lines.append("")
    lines.append("## 2. 输入数据概况")
    lines.append("")
    lines.append(f"- 航线数量：`{stats['flight_count']}`")
    lines.append(f"- Query 数量：`{stats['query_count']}`")
    lines.append(f"- Query/卫片统一坐标系：`{dataset_summary['query_crs']}`")
    lines.append(f"- 原始航线总体 bbox 面积：`{stats['raw_area_km2']:.3f} km^2`")
    lines.append(f"- 外扩 `250m` 后 ROI 面积：`{stats['roi_area_km2']:.3f} km^2`")
    lines.append(f"- 外扩新增面积：`{stats['buffer_gain_km2']:.3f} km^2`")
    lines.append(f"- 固定候选卫片总数：`{stats['tile_count_total']}`")
    lines.append(f"- 候选尺度分布：`200m={dataset_summary['tile_count_by_scale']['200m']}`、`300m={dataset_summary['tile_count_by_scale']['300m']}`、`500m={dataset_summary['tile_count_by_scale']['500m']}`、`700m={dataset_summary['tile_count_by_scale']['700m']}`")
    lines.append(f"- 卫片切片像素分辨率近似保持原始 `0.5m` 数据源尺度，当前 tile 的平均 GSD 约为 `{stats['satellite_gsd_mean_m']:.3f} m/px`")
    lines.append(f"- query footprint 面积统计：均值 `{stats['footprint_area_mean_ha']:.2f} ha`，最小 `{stats['footprint_area_min_ha']:.2f} ha`，最大 `{stats['footprint_area_max_ha']:.2f} ha`")
    lines.append("")
    lines.append(f"![航片整体范围]({fig_names['航片整体范围.png']})")
    lines.append("")
    lines.append(f"![卫片整体范围]({fig_names['卫片整体范围.png']})")
    lines.append("")
    lines.append(f"![航片与外扩卫片叠加显示]({fig_names['航片与外扩卫片叠加显示.png']})")
    lines.append("")
    lines.append("## 3. 40 个 query 的生成过程")
    lines.append("")
    lines.append("当前 query 不是从正射影像中固定裁切得到，而是直接从原始无人机单图中筛选。筛选逻辑来自 `select_raw_uav_images.py`，每条航线目标选取 `10` 张图，共 `40` 张：")
    lines.append("")
    lines.extend(
        [
            "- 候选倾斜角窗口固定为 `-85 <= GimbalPitchDegree <= -40`；",
            "- 以 `-62.5` 度左右作为分割，尽量让下视样本与倾斜样本各占约一半；",
            "- 选取时优先覆盖不同地物类型，并保留时间间隔与空间间隔约束，避免样本过于密集；",
            "- 如果不足，再使用多样性补充策略，从剩余帧里补足。"
        ]
    )
    lines.append("")
    lines.append(f"- 当前统计结果：`downview={stats['selected_downview']}`，`tilted={stats['selected_tilt']}`。")
    lines.append("- 各航线最终都是 `10` 张，因此 4 条航线总计 `40` 张。")
    lines.append("- 选中样本先复制到 `D:\\数据\\武汉影像\\挑选无人机0.1m`，形成可复核的 query 候选集。")
    lines.append("")
    lines.append("## 4. Query 检索前的预处理")
    lines.append("")
    lines.append("正式检索时并不直接使用保留 DJI 元数据的原图，而是对选中 query 做去元数据处理。对应脚本为 `sanitize_query_images.py`。")
    lines.append("")
    lines.extend(
        [
            "- 处理方法是重新编码图像像素，移除 EXIF/XMP/GPS/DJI 元数据；",
            f"- 当前 `query_manifest.csv` 中 `{stats['sanitized_query_count']}/{stats['query_count']}` 张 query 已标记为 `has_metadata_removed=1`；",
            "- 这样做的原因不是为了改变图像内容，而是为了保证检索只能依赖视觉内容，不能从文件中的地理坐标直接恢复位置；",
            "- query 的真实空间位置与 footprint 仍然保存在 `queries_truth_seed.csv` 中，但这些信息只用于真值生成与评估，不进入检索模型输入。",
        ]
    )
    lines.append("")
    lines.append("## 5. 卫片候选库构建")
    lines.append("")
    lines.append("卫片候选库不是针对单张 query 临时裁出来的，而是一个固定离线库。其构建顺序如下：")
    lines.append("")
    lines.extend(
        [
            "1. 读取 4 条航线原始影像中的 GPS/MRK 点；",
            "2. 在统一坐标系 `EPSG:32650` 下求 4 条航线总体活动区域的包围框；",
            "3. 在包围框外固定外扩 `250m`，得到正式 ROI；",
            "4. 在 ROI 内对卫星大图按多尺度切片，形成固定候选库；",
            "5. 后续所有方法都在这同一候选库上做检索或重排。"
        ]
    )
    lines.append("")
    lines.append("这里需要强调两点：")
    lines.append("")
    lines.extend(
        [
            "- 外扩 `250m` 的作用是给航线边缘区域留出缓冲，避免 query footprint 接近边界时没有足够的候选 tile；",
            "- 先切 tile，再做 DINOv2 特征提取和 FAISS 建库，因此当前特征库是 tile 级特征库，而不是整幅卫片特征库。"
        ]
    )
    lines.append("")
    lines.append(f"![航片外扩250m的卫片候选库]({fig_names['航片外扩250m的卫片候选库.png']})")
    lines.append("")
    lines.append(f"![卫片整体与卫片候选库叠加显示]({fig_names['卫片整体与卫片候选库叠加显示.png']})")
    lines.append("")
    lines.append("## 6. 多尺度与 overlap 设计")
    lines.append("")
    lines.append("当前正式候选库采用 `200 / 300 / 500 / 700m` 四个尺度，并设置 `overlap = 0.25`。")
    lines.append("")
    lines.extend(
        [
            "- `200m` 保留较细粒度局部区域，适合覆盖更紧凑的真实范围；",
            "- `300m` 是从小尺度向中尺度过渡的补充档位，历史对比中它曾带来新增 Top-1 命中；",
            "- `500m` 与 `700m` 是因为当前 query 已经不是同尺度正射小裁块，而是覆盖更大、视角更复杂的无人机单图，需要更大的区域窗口把正确地理区域送入 coarse Top-K；",
            "- `overlap = 0.25` 表示相邻 tile 在地面范围上保留 25% 重叠，因此步长为 `0.75 * tile_size`；",
            "- overlap 的主要作用是降低边界截断风险，提高候选覆盖连续性，但代价是候选库规模与真值数量都会增加。"
        ]
    )
    lines.append("")
    lines.append(f"![query_satellite_overview]({fig_names['query_satellite_overview.png']})")
    lines.append("")
    lines.append(f"![scale_tile_count_bar]({fig_names['scale_tile_count_bar.png']})")
    lines.append("")
    lines.append("## 7. 真值生成与当前正式定义")
    lines.append("")
    lines.append("当前正式真值已经从更严格的旧口径切换为 `intersection truth`。对应脚本为 `generate_query_truth_by_intersection.py`。")
    lines.append("")
    lines.extend(
        [
            "- 对每张 query，先利用无人机影像的位姿元数据近似恢复其地面 footprint 多边形；",
            "- 再与固定候选库中的每张卫片 tile 做几何相交判断；",
            "- 只要 query footprint 与 tile 存在非零面积交集，该 tile 就记为当前 query 的真值；",
            "- 这一定义不再要求 query 中心点必须落入某张 tile，也不要求只保留极少数“唯一正确块”，而是把真实空间交集作为正样本依据。"
        ]
    )
    lines.append("")
    lines.append(f"- 当前交集真值总记录数：`{stats['truth_pair_total']}`")
    lines.append(f"- 去重后的唯一真值 tile 数：`{stats['truth_unique_tile_count']}`")
    lines.append(f"- 每个 query 的真值数量：均值 `{stats['truth_count_mean']:.2f}`，最小 `{stats['truth_count_min']}`，最大 `{stats['truth_count_max']}`")
    lines.append(f"- 真值中 `contains_query_center=1` 的比例约为 `{stats['truth_contain_center_ratio']:.3f}`，说明当前真值并不等价于“只看中心点所在 tile”。")
    lines.append(f"- 按尺度统计的真值记录数：`200m={stats['truth_scale_counts']['200m']}`、`300m={stats['truth_scale_counts']['300m']}`、`500m={stats['truth_scale_counts']['500m']}`、`700m={stats['truth_scale_counts']['700m']}`")
    lines.append("")
    lines.append("## 8. 当前预处理产物")
    lines.append("")
    lines.extend(
        [
            "- `selected_images_summary.csv`：40 张 query 的选取清单；",
            "- `query_manifest.csv`：原图与去元数据副本的映射；",
            "- `queries_truth_seed.csv`：query 的中心点、坐标系和 footprint；",
            "- `tiles.csv` / `roi_summary.json`：卫片候选库空间索引；",
            "- `query_truth_tiles.csv`：当前 intersection truth 主表。"
        ]
    )
    lines.append("")
    lines.append("## 9. 结论")
    lines.append("")
    lines.append("当前数据预处理链路已经把原始无人机单图和原始卫星大图，整理成了一套可以稳定支撑后续检索实验的统一输入资产。它的核心特点是：query 侧去除了地理元数据，卫片侧采用固定离线多尺度候选库，评估侧则以真实空间交集定义正式真值。这一设计更符合当前“跨视角区域级检索定位”的任务目标。")
    lines.append("")
    return "\n".join(lines)


def build_docx(dataset_summary: dict, stats: dict, fig_paths: dict[str, Path], out_docx: Path) -> None:
    doc = Document()
    add_paragraph(doc, "当前任务数据输入与预处理说明", size=16, center=True)
    add_paragraph(doc, "日期：2026-03-24", size=10, center=True)

    add_heading(doc, "1. 文档目的", 1)
    add_paragraph(doc, "本报告用于系统梳理当前项目在正式检索实验之前完成的数据来源整理与预处理工作，给后续 DINOv2、FAISS、LightGlue 等方法实验提供统一的数据口径说明。")
    add_bullets(
        doc,
        [
            "明确当前使用的航片与卫片数据来源、范围、坐标系和分辨率背景。",
            "说明 40 张 query 的形成方式，以及为什么检索时必须使用去元数据副本。",
            "说明固定卫片候选库的构建方式、多尺度设计与当前正式真值定义。",
        ],
    )

    add_heading(doc, "2. 输入数据概况", 1)
    add_paragraph(doc, "表 1 汇总当前实验使用的主要输入数据资产及其作用。")
    build_input_table(doc, dataset_summary, stats)
    add_paragraph(
        doc,
        f"当前 4 条航线统一到 {dataset_summary['query_crs']} 下进行空间分析。原始总体 bbox 面积为 {stats['raw_area_km2']:.3f} km^2，外扩 250m 后 ROI 面积为 {stats['roi_area_km2']:.3f} km^2，新增缓冲面积为 {stats['buffer_gain_km2']:.3f} km^2。",
    )
    add_paragraph(
        doc,
        f"40 张 query 的地面 footprint 面积均值约为 {stats['footprint_area_mean_ha']:.2f} ha，最小 {stats['footprint_area_min_ha']:.2f} ha，最大 {stats['footprint_area_max_ha']:.2f} ha；当前卫片候选库总规模为 {stats['tile_count_total']} 张。",
    )
    add_picture(doc, fig_paths["航片整体范围.png"], 5.8)
    add_caption(doc, "图 1  四条无人机航片的总体空间范围")
    add_picture(doc, fig_paths["卫片整体范围.png"], 5.8)
    add_caption(doc, "图 2  当前参与构库的卫片整体空间范围")
    add_picture(doc, fig_paths["航片与外扩卫片叠加显示.png"], 6.1)
    add_caption(doc, "图 3  航片总体范围与外扩后卫片 ROI 的叠加关系")

    add_heading(doc, "3. 40 个 Query 的生成过程", 1)
    add_paragraph(doc, "当前 query 直接来自原始无人机单图，不再沿用旧任务中的同尺度正射裁块。对应筛选脚本为 `select_raw_uav_images.py`。")
    add_bullets(
        doc,
        [
            "每条航线目标选取 10 张图，共 4 条航线，总计 40 张。",
            "候选倾斜角窗口固定为 `-85 <= GimbalPitchDegree <= -40`。",
            "以 `-62.5` 度左右作为分割，尽量让下视与倾斜样本各占约 50%。",
            "选取时优先覆盖不同地物类型，并同时施加时间间隔和空间间隔约束，避免连续帧过密。",
            "若样本不足，再使用多样性补充策略从剩余帧中补齐。",
        ],
    )
    add_paragraph(
        doc,
        f"当前统计结果为：downview 样本 {stats['selected_downview']} 张，tilted 样本 {stats['selected_tilt']} 张。各航线最终都稳定得到 10 张，因此形成当前正式 query 清单。",
    )

    add_heading(doc, "4. Query 检索前的预处理", 1)
    add_paragraph(doc, "为了避免模型直接利用图像文件中的 GPS/EXIF/XMP 元数据完成“作弊式定位”，正式检索时不使用原始 query，而使用去元数据副本。对应脚本为 `sanitize_query_images.py`。")
    add_bullets(
        doc,
        [
            "处理方式是重新编码像素内容，移除 EXIF/XMP/GPS/DJI 元数据，但不改变图像语义内容。",
            f"当前 query_manifest 中 {stats['sanitized_query_count']}/{stats['query_count']} 张均已标记为去元数据完成。",
            "检索模型只读取 sanitized query；query 的真实空间位置仍保留在 truth seed 文件中，仅用于真值构造与评估。",
        ],
    )

    add_heading(doc, "5. 卫片候选库构建", 1)
    add_paragraph(doc, "当前卫片候选库是一个固定离线库，不会针对单个 query 在推理时临时裁库。对应构库脚本为 `build_fixed_satellite_library.py`。")
    add_bullets(
        doc,
        [
            "先读取 4 条航线原始影像中的 GPS 或 MRK 点，得到整体活动区域。",
            "在总体包围框外固定外扩 250m，形成正式 ROI。",
            "在 ROI 内对卫星大图按多尺度切片，形成固定候选库。",
            "候选 tile 保持原始裁块分辨率，不预先统一为单一固定尺寸。",
            "只有在 tile 切完之后，才进行 DINOv2 特征提取和 FAISS 建库，因此当前是 tile 级特征库。",
        ],
    )
    add_paragraph(
        doc,
        f"当前固定候选库总数为 {stats['tile_count_total']}，其中 200m={dataset_summary['tile_count_by_scale']['200m']}、300m={dataset_summary['tile_count_by_scale']['300m']}、500m={dataset_summary['tile_count_by_scale']['500m']}、700m={dataset_summary['tile_count_by_scale']['700m']}。tile 像素尺寸随尺度变化，当前宽度范围约为 {stats['tile_native_width_min']} 到 {stats['tile_native_width_max']} px。",
    )
    add_picture(doc, fig_paths["航片外扩250m的卫片候选库.png"], 6.0)
    add_caption(doc, "图 4  四条航线外扩 250m 后形成的固定卫片候选库")
    add_picture(doc, fig_paths["卫片整体与卫片候选库叠加显示.png"], 6.0)
    add_caption(doc, "图 5  卫片整体范围与当前候选库范围对比")

    add_heading(doc, "6. 多尺度与 Overlap 设计", 1)
    add_paragraph(doc, "当前正式配置为 `200 / 300 / 500 / 700m` 四尺度，且相邻 tile 采用 `overlap = 0.25`。这部分设计已经在独立说明中验证过，这里给出与当前任务直接相关的解释。")
    add_bullets(
        doc,
        [
            "200m 负责较细粒度局部覆盖，适合更紧凑的真实范围。",
            "300m 是从小尺度到中尺度的过渡档位，历史对比中它带来过新增 Top-1 命中。",
            "500m 和 700m 是针对当前 query 覆盖范围更大、视角更复杂的特点而加入的大窗口候选。",
            "overlap = 0.25 表示相邻 tile 在地面范围上保留 25% 重叠，对应步长为 0.75 * tile_size。",
            "这一设置的核心作用是减轻边界截断、提升候选覆盖连续性，但也会提高库规模和真值数量。",
        ],
    )
    add_picture(doc, fig_paths["query_satellite_overview.png"], 6.0)
    add_caption(doc, "图 6  Query、ROI 与多尺度卫片中心点的空间关系")
    add_picture(doc, fig_paths["scale_tile_count_bar.png"], 5.6)
    add_caption(doc, "图 7  四个尺度对应的卫片候选数量")

    add_heading(doc, "7. 真值生成与当前正式定义", 1)
    add_paragraph(doc, "当前正式真值已经切换为 `intersection truth`。对应脚本为 `generate_query_truth_by_intersection.py`。")
    add_bullets(
        doc,
        [
            "首先利用 query 的位置、姿态和相机参数近似恢复其地面 footprint。",
            "随后将 footprint 与固定候选库中每张卫片 tile 的矩形范围做几何相交。",
            "只要存在非零面积交集，该 tile 就记为当前 query 的真值。",
            "因此当前真值不再依赖 query 中心点是否落入 tile，也不再把正样本压缩成过窄的单点式定义。",
        ],
    )
    add_paragraph(
        doc,
        f"当前交集真值总记录数为 {stats['truth_pair_total']}，去重后的唯一真值 tile 数为 {stats['truth_unique_tile_count']}。平均每个 query 对应 {stats['truth_count_mean']:.2f} 张真值 tile，最少 {stats['truth_count_min']} 张，最多 {stats['truth_count_max']} 张。真值中仍有约 {stats['truth_contain_center_ratio']:.3f} 的样本同时包含 query 中心点，但它已不再是必要条件。",
    )
    add_paragraph(
        doc,
        f"按尺度统计，当前真值记录数为 200m={stats['truth_scale_counts']['200m']}、300m={stats['truth_scale_counts']['300m']}、500m={stats['truth_scale_counts']['500m']}、700m={stats['truth_scale_counts']['700m']}。这也说明尺度越大、overlap 越高时，单个 query 的可接受正样本通常越多。",
    )

    add_heading(doc, "8. 当前预处理产物汇总", 1)
    add_paragraph(doc, "表 2 列出当前后续检索实验会直接依赖的关键数据资产。")
    build_asset_table(doc)

    add_heading(doc, "9. 结论", 1)
    add_paragraph(
        doc,
        "当前数据预处理链路已经把原始无人机单图、卫星大图和空间元数据整理成了一套统一、可复用的实验输入资产。它的关键原则是：query 侧屏蔽地理元数据，候选侧采用固定离线多尺度卫片库，评估侧以真实空间交集作为正式真值依据。这一设计更符合当前跨视角、区域级检索定位任务的目标。",
    )

    ensure_dir(out_docx.parent)
    doc.save(out_docx)


def main() -> None:
    args = parse_args()
    dataset_summary = load_json(Path(args.dataset_summary_json))
    selected_rows = load_csv(Path(args.selected_summary_csv))
    query_manifest_rows = load_csv(Path(args.query_manifest_csv))
    query_seed_rows = load_csv(Path(args.query_seed_csv))
    truth_tile_rows = load_csv(Path(args.query_truth_tiles_csv))
    tile_rows = load_csv(Path(args.tile_metadata_csv))

    out_docx = Path(args.out_docx)
    out_md = Path(args.out_md)
    out_dir = out_md.parent
    ensure_dir(out_dir)

    range_fig_dir = Path(args.range_fig_dir)
    scale_note_dir = Path(args.scale_note_dir)
    figure_paths = [
        range_fig_dir / "航片整体范围.png",
        range_fig_dir / "卫片整体范围.png",
        range_fig_dir / "航片与外扩卫片叠加显示.png",
        range_fig_dir / "航片外扩250m的卫片候选库.png",
        range_fig_dir / "卫片整体与卫片候选库叠加显示.png",
        scale_note_dir / "query_satellite_overview.png",
        scale_note_dir / "scale_tile_count_bar.png",
    ]
    copied = copy_figures(figure_paths, out_dir / "figures")
    stats = build_stats(dataset_summary, selected_rows, query_manifest_rows, query_seed_rows, truth_tile_rows, tile_rows)

    md_text = build_markdown(dataset_summary, stats, {name: f"figures/{name}" for name in copied})
    out_md.write_text(md_text, encoding="utf-8")
    build_docx(dataset_summary, stats, copied, out_docx)


if __name__ == "__main__":
    main()
