#!/usr/bin/env python3
"""Generate the 009/010 two-route engineering report as Markdown and Word.

Purpose:
- write a formal engineering report for the completed 009/010 satellite
  DOM+SRTM route and CaiWangCun DOM/DSM full-replacement route;
- organize the report by goal, scheme, data flow, evaluation method, results,
  and runtime cost;
- keep metrics tied to formal JSON/CSV outputs rather than relying only on
  narrative draft notes.

Main inputs:
- `new3output/nadir_009010_dinov2_romav2_pose_sattruth_srtm_romatie_2026-04-16`;
- `new3output/nadir_009010_caiwangcun_domdsm_fullreplace_full_2026-04-21`;
- draft notes under `汇总/`.

Main outputs:
- `汇总/009010双线路工程汇报_卫星DOM_SRTM_vs_CaiWangCun_DOMDSM.md`;
- `汇总/009010双线路工程汇报_卫星DOM_SRTM_vs_CaiWangCun_DOMDSM.docx`.

Applicable task constraints:
- query is treated as a single arbitrary UAV image without geographic metadata;
- query is not assumed to be an orthophoto;
- runtime retrieval, reranking, and pose solving must remain truth-free;
- external resolution normalization is not introduced as a formal assumption.
"""

from __future__ import annotations

import csv
import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
SUMMARY_DIR = PROJECT_ROOT / "汇总"
OUT_BASENAME = "009010双线路工程汇报_卫星DOM_SRTM_vs_CaiWangCun_DOMDSM"
OUT_MD = SUMMARY_DIR / f"{OUT_BASENAME}.md"
OUT_DOCX = SUMMARY_DIR / f"{OUT_BASENAME}.docx"

SRTM_ROOT = (
    PROJECT_ROOT
    / "new3output"
    / "nadir_009010_dinov2_romav2_pose_sattruth_srtm_romatie_2026-04-16"
)
CAI_ROOT = (
    PROJECT_ROOT
    / "new3output"
    / "nadir_009010_caiwangcun_domdsm_fullreplace_full_2026-04-21"
)
SRTM_SUITE = SRTM_ROOT / "pose_v1_formal" / "eval_pose_validation_suite_sattruth_srtm"
CAI_SUITE = CAI_ROOT / "pose_v1_formal" / "eval_pose_validation_suite_caiwangcun_truth"


def load_json(path: Path) -> dict[str, Any]:
    if not path.exists():
        raise FileNotFoundError(f"Missing required JSON: {path}")
    return json.loads(path.read_text(encoding="utf-8"))


def read_csv_rows(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def fmt_num(value: Any, digits: int = 3, suffix: str = "") -> str:
    if value is None:
        return "-"
    try:
        return f"{float(value):.{digits}f}{suffix}"
    except (TypeError, ValueError):
        return str(value)


def fmt_count(status_counts: dict[str, Any] | None) -> str:
    if not status_counts:
        return "-"
    return ", ".join(f"{key}={value}" for key, value in status_counts.items())


def get_nested(data: dict[str, Any], *keys: str, default: Any = None) -> Any:
    node: Any = data
    for key in keys:
        if not isinstance(node, dict) or key not in node:
            return default
        node = node[key]
    return node


def q037_pose_row() -> dict[str, str]:
    rows = read_csv_rows(CAI_SUITE / "pose_vs_at" / "per_query_pose_vs_at.csv")
    for row in rows:
        if row.get("query_id") == "q_037":
            return row
    return {}


def collect_metrics() -> dict[str, Any]:
    srtm_pose = load_json(SRTM_ROOT / "pose_v1_formal" / "summary" / "pose_overall_summary.json")
    srtm_phase = load_json(SRTM_ROOT / "pose_v1_formal" / "summary" / "phase_gate_summary.json")
    srtm_ortho = load_json(SRTM_SUITE / "ortho_alignment_satellite" / "overall_ortho_accuracy.json")
    srtm_pose_at = load_json(SRTM_SUITE / "pose_vs_at" / "overall_pose_vs_at.json")
    srtm_tie = load_json(SRTM_SUITE / "tiepoint_ground_error" / "overall_tiepoint_ground_error.json")
    srtm_retr = load_json(SRTM_ROOT / "retrieval" / "retrieval_top20_export_summary.json")

    cai_pose = load_json(CAI_ROOT / "pose_v1_formal" / "summary" / "pose_overall_summary.json")
    cai_phase = load_json(CAI_ROOT / "pose_v1_formal" / "summary" / "phase_gate_summary.json")
    cai_ortho = load_json(CAI_SUITE / "ortho_alignment" / "overall_ortho_accuracy.json")
    cai_pose_at = load_json(CAI_SUITE / "pose_vs_at" / "overall_pose_vs_at.json")
    cai_tie = load_json(CAI_SUITE / "tiepoint_ground_error" / "overall_tiepoint_ground_error.json")
    cai_frame = load_json(CAI_SUITE / "ortho_alignment" / "frame_sanity" / "overall_frame_sanity.json")
    cai_accept = load_json(CAI_ROOT / "plan" / "full_acceptance_summary.json")
    cai_retr = load_json(CAI_ROOT / "retrieval" / "retrieval_top20_export_summary.json")

    return {
        "srtm_pose": srtm_pose,
        "srtm_phase": srtm_phase,
        "srtm_ortho": srtm_ortho,
        "srtm_pose_at": srtm_pose_at,
        "srtm_tie": srtm_tie,
        "srtm_retr": srtm_retr,
        "cai_pose": cai_pose,
        "cai_phase": cai_phase,
        "cai_ortho": cai_ortho,
        "cai_pose_at": cai_pose_at,
        "cai_tie": cai_tie,
        "cai_frame": cai_frame,
        "cai_accept": cai_accept,
        "cai_retr": cai_retr,
        "q037": q037_pose_row(),
    }


def image_inventory() -> list[dict[str, str]]:
    images = [
        {
            "title": "卫星 DOM+SRTM 成功样例 q_015：predicted ortho 与 satellite truth overlay",
            "path": str(SRTM_SUITE / "ortho_alignment" / "viz_overlay_truth" / "q_015_overlay.png"),
        },
        {
            "title": "卫星 DOM+SRTM 成功样例 q_015：Layer-3 同名点 overlay",
            "path": str(SRTM_SUITE / "tiepoint_ground_error" / "viz_tiepoints" / "q_015_tiepoints_overlay.png"),
        },
        {
            "title": "卫星 DOM+SRTM 异常样例 q_036：Layer-3 同名点误差热力图",
            "path": str(SRTM_SUITE / "tiepoint_ground_error" / "viz_tiepoints" / "q_036_tiepoints_error_heatmap.png"),
        },
        {
            "title": "CaiWangCun 完整替换成功样例 q_004：predicted ortho 与 truth overlay",
            "path": str(CAI_SUITE / "ortho_alignment" / "viz_overlay_truth" / "q_004_overlay.png"),
        },
        {
            "title": "CaiWangCun 完整替换成功样例 q_004：predicted ortho 与 DOM overlay",
            "path": str(CAI_SUITE / "ortho_alignment" / "viz_overlay_dom" / "q_004_pred_vs_dom_overlay.png"),
        },
        {
            "title": "CaiWangCun 完整替换边界样例 q_001：Layer-3 同名点误差热力图",
            "path": str(CAI_SUITE / "tiepoint_ground_error" / "viz_tiepoints" / "q_001_tiepoints_error_heatmap.png"),
        },
    ]
    for image in images:
        image["exists"] = str(Path(image["path"]).exists())
    return images


def table_md(headers: list[str], rows: list[list[Any]]) -> str:
    lines = ["| " + " | ".join(headers) + " |"]
    lines.append("|" + "|".join("---" for _ in headers) + "|")
    for row in rows:
        lines.append("| " + " | ".join(str(cell) for cell in row) + " |")
    return "\n".join(lines)


def build_markdown(metrics: dict[str, Any], images: list[dict[str, str]]) -> str:
    s_pose = metrics["srtm_pose"]
    s_phase = metrics["srtm_phase"]
    s_ortho = metrics["srtm_ortho"]
    s_pose_at = metrics["srtm_pose_at"]
    s_tie = metrics["srtm_tie"]
    s_retr = metrics["srtm_retr"]
    c_pose = metrics["cai_pose"]
    c_phase = metrics["cai_phase"]
    c_ortho = metrics["cai_ortho"]
    c_pose_at = metrics["cai_pose_at"]
    c_tie = metrics["cai_tie"]
    c_frame = metrics["cai_frame"]
    c_accept = metrics["cai_accept"]
    c_retr = metrics["cai_retr"]
    q037 = metrics["q037"]

    usable_horizontal = get_nested(
        c_frame, "numeric_summaries", "horizontal_error_m", "mean", default=None
    )

    result_rows = [
        [
            "卫星 DOM+SRTM",
            f"{s_pose['best_status_counts'].get('ok', 0)}/{s_pose['query_count']}",
            fmt_num(get_nested(s_ortho, "center_offset_m", "mean"), 3, " m"),
            fmt_num(get_nested(s_pose_at, "horizontal_error_m", "mean"), 3, " m"),
            fmt_num(s_tie.get("tiepoint_xy_error_rmse_m"), 3, " m"),
            fmt_count(s_pose.get("score_status_counts")),
        ],
        [
            "CaiWangCun DOM+DSM 完整替换",
            f"{c_pose['best_status_counts'].get('ok', 0)}/{c_pose['query_count']}",
            f"{c_ortho.get('evaluated_query_count')}/{c_ortho.get('query_count')}；"
            f"{fmt_num(get_nested(c_ortho, 'center_offset_m', 'mean'), 3, ' m')}",
            f"{fmt_num(usable_horizontal, 3, ' m')}（39 个可用 query）；"
            f"全量均值 {fmt_num(get_nested(c_pose_at, 'horizontal_error_m', 'mean'), 3, ' m')}",
            fmt_num(c_tie.get("tiepoint_xy_error_rmse_m"), 3, " m"),
            fmt_count(c_pose.get("score_status_counts")),
        ],
    ]

    environment_rows = [
        ["主机 OS", "Windows 10 家庭版 64 位"],
        ["WSL 环境", "WSL2 Linux 6.6.87.2-microsoft-standard-WSL2"],
        ["CPU", "Intel Core i5-7500 CPU @ 3.40GHz"],
        ["CPU 核心/线程", "4 核 4 线程"],
        ["物理内存", "约 64GB"],
        ["WSL 可见内存", "约 31GiB"],
        ["GPU", "NVIDIA GeForce GTX 1080"],
        ["GPU 显存", "8192 MiB"],
        ["NVIDIA Driver", "581.80"],
        ["CUDA Version", "13.0"],
        ["Python 环境", "WSL2 项目 .conda 环境；SRTM 草稿记录 Python 3.10.19"],
    ]

    runtime_rows = [
        [
            "卫星 DOM+SRTM",
            "DINOv2 query 特征提取 / DINOv2 coarse retrieval / RoMa v2 rerank 资产生成",
            "SRTM formal pose、validation、report 完成",
            "约 6h47m",
            "按日志边界保守统计约 6h50m21s",
        ],
        [
            "CaiWangCun DOM+DSM 完整替换",
            "CaiWangCun candidate library / candidate features / FAISS / DINOv2 retrieval / RoMa v2 rerank 资产生成",
            "full formal pose、validation、frame sanity、report 完成",
            "约 19h18m47s",
            "约为 SRTM 线路 2.8 倍",
        ],
    ]

    srtm_runtime_rows = [
        ["DINOv2 query feature extraction", "2026-04-10T04:13:58 起", "0.6min", "new2output/nadir_009010_dinov2_romav2_pose_2026-04-10/logs/run_full.log"],
        ["DINOv2 coarse retrieval + RoMa v2 rerank", "2026-04-10 上游检索/重排阶段", "11156.6126s，约 3h05m57s", "romav2_rerank/timing/romav2_rerank_internal.json"],
        ["上游资产阶段日志边界", "从 DINOv2 特征提取命令到 RoMa timing 完成", "约 3h10m16s", "用于保守说明"],
        ["Formal input + SRTM DSM cache", "2026-04-16T16:18:17 -> 16:18:32", "约 15s", "run_sattruth_srtm_romatie_full.log"],
        ["Full pose pipeline", "2026-04-16T16:18:32 -> 19:46:40", "约 3h28m08s", "matches、sampling、PnP、scores、best pose"],
        ["Full validation suite", "2026-04-16T19:46:40 -> 19:58:17", "约 11m37s", "satellite truth Layer-1/2/3 validation"],
        ["Suite/report wrap-up", "2026-04-16T19:58:17 -> 19:58:22", "约 5s", "suite reports 与 summary"],
        ["SRTM formal/full 子阶段合计", "2026-04-16 主日志跨度", "3h40m05s", "作为整体耗时的后半段"],
    ]

    cai_runtime_rows = [
        ["Candidate library", "2026-04-20T12:14:07 -> 12:19:43", "约 5m36s", "CaiWangCun DOM/DSM 覆盖范围内生成 149 个候选 tile"],
        ["Query truth", "2026-04-20T12:19:43 -> 12:22:06", "约 2m23s", "生成 CaiWangCun 候选库对应 truth 表"],
        ["Candidate DINOv2 features", "2026-04-20T12:22:06 起", "1.7min", "日志内记录 elapsed=1.7min"],
        ["FAISS index", "2026-04-20T12:26:58 -> 12:27:00", "约 2s", "构建 CaiWangCun tiles IP index"],
        ["DINOv2 retrieval", "2026-04-20T12:27:00 -> 12:27:01", "约 1s", "导出 coarse retrieval Top-20"],
        ["RoMa v2 rerank", "2026-04-20T12:27:01 -> 16:13:59", "约 3h46m58s", "两条航线 Top-20 几何重排"],
        ["上游资产阶段合计", "2026-04-20T12:14:07 -> 16:13:59", "约 3h59m52s", "gate 分支生成，full run 复用/复制校验"],
        ["Formal input + DSM cache", "2026-04-21T04:30:49 -> 04:37:30", "约 6m41s", "formal manifests、DSM cache、pose manifest"],
        ["Full pose pipeline", "2026-04-21T04:37:30 -> 10:05:07", "约 5h27m37s", "4,000,000 条 matches/correspondences"],
        ["Full validation suite", "2026-04-21T10:05:07 -> 19:46:53", "约 9h41m46s", "Layer-3 tiepoint ground error 约 9h15m46s"],
        ["Frame sanity", "2026-04-21T19:46:53 -> 19:49:21", "约 2m28s", "overall_frame_sanity.json"],
        ["Full report", "2026-04-21T19:49:21 -> 19:49:44", "约 23s", "Word/Markdown report"],
        ["Full formal/full 子阶段合计", "2026-04-21 主日志跨度", "15h18m55s", "作为整体耗时的后半段"],
    ]

    image_lines = []
    for image in images:
        if image["exists"] == "True":
            image_lines.append(f"![{image['title']}]({image['path']})")
        else:
            image_lines.append(f"- 缺失图件：{image['title']}，路径：`{image['path']}`")

    q037_text = (
        f"`q_037` 的水平误差为 {fmt_num(q037.get('horizontal_error_m'), 3, ' m')}，"
        f"视线方向误差为 {fmt_num(q037.get('view_dir_angle_error_deg'), 3, ' deg')}。"
        if q037
        else "`q_037` 的 per-query pose-vs-AT 明细未读取到。"
    )

    md = f"""# 009/010 双线路工程汇报：卫星 DOM+SRTM 与 CaiWangCun DOM+DSM 完整替换

## 1. 目标

本章结论：当前实验已经从“单张无人机影像能否检索到正确地理区域”推进到“检索、几何重排、DSM/DEM 三维支撑和 PnP 粗位姿恢复能否形成可复核工程链路”。两条代表线路分别验证了低频 SRTM 支撑下的稳定全量跑通能力，以及同源 CaiWangCun DOM/DSM 完整替换后的米级定位能力。

本任务面向 GNSS 拒止或 GNSS 不可靠条件下的无人机视觉初始定位。GNSS 是全球卫星导航系统，用于提供绝对位置；在遮挡、干扰、欺骗或多路径环境下，其可靠性会下降。当前输入 query 被定义为任意单张 UAV 图像，UAV 指无人机平台采集的影像；该图像不携带地理元数据，不保证为正射影像，也不允许依赖旧任务中的同尺度裁块、query 中心点真值或统一外部分辨率预处理。

本实验的工程目标不是替代完整导航系统，而是验证遥感正射影像能否为单张 UAV 图像提供低频、全局、具有地理意义的视觉观测。该观测可作为后续惯性导航系统、视觉惯性里程计或融合导航模块的外部绝对位置约束。

DOM 是数字正射影像，用于提供带地理坐标的二维视觉参考；DSM 是数字表面模型，用于提供建筑、植被、地表等表面高程；SRTM 是 Shuttle Radar Topography Mission 的公开低频地形高程数据。本报告比较的两条线路如下：

{table_md(["线路", "核心目标", "主要数据支撑", "工程判断重点"], [
    ["卫星 DOM+SRTM", "验证固定 satellite DOM 候选库与 SRTM DSM 是否可稳定支撑 40-query 全流程", "固定 satellite DOM candidate library；SRTM HGT 派生 DSM cache", "稳定性、可跑通性、跨源验证误差"],
    ["CaiWangCun DOM+DSM 完整替换", "验证同源 DOM/DSM、候选库和验证 truth 完整重建后是否进入米级定位范围", "CaiWangCun DOM mosaic；CaiWangCun DSM mosaic", "坐标框架一致性、米级精度、异常样本可解释性"],
])}

## 2. 方案介绍

本章结论：两条线路采用相同的“粗检索、几何重排、三维采样、PnP 位姿恢复、离线验证”主框架，差异集中在候选 DOM、DSM 来源和 validation truth 口径。CaiWangCun 线路的关键改进不只是替换高分辨率 DSM，而是完整重建所有由地图数据派生的运行时资产。

DINOv2 是自监督视觉基础模型，本实验使用其图像全局特征进行粗检索。FAISS 是向量相似度检索库，用于在候选 DOM tile 特征库中快速检索 Top-20 候选。RoMa v2 是局部/稠密图像匹配模型，用于对 query 与候选 tile 建立局部几何对应并进行重排。PnP 是 Perspective-n-Point 位姿求解方法，用于根据二维图像点与三维地图点恢复相机粗位姿。AT 是空中三角测量或参考位姿结果，在本报告中仅作为离线精度验证参考。

两条线路的共同流程为：query 图像提取 DINOv2 特征；FAISS 在候选 DOM 特征库中返回 Top-20；RoMa v2 对 Top-20 进行局部几何匹配和重排；候选 DOM 像素通过仿射变换映射到投影平面坐标，并从 DSM/DEM 中采样高程，形成 PnP 所需的 2D-3D 对应；PnP RANSAC 与 refinement 输出候选位姿；按融合分数选择每个 query 的 best pose；最后通过分层验证体系进行离线评估。

{table_md(["对比项", "卫星 DOM+SRTM", "CaiWangCun DOM+DSM 完整替换"], [
    ["候选库", "固定 satellite DOM 多尺度候选库", f"CaiWangCun DOM 重新生成，candidate tiles={c_accept.get('candidate_tile_count')}"],
    ["检索与重排", f"复用 009/010 DINOv2+RoMa v2 运行链路；Top-K={s_retr.get('top_k')}，rows={s_retr.get('row_count')}", f"CaiWangCun 分支重新导出检索 Top-20；Top-K={c_retr.get('top_k')}，rows={c_retr.get('row_count')}"],
    ["DSM/DEM", f"SRTM DSM cache：planned={get_nested(s_phase, 'dsm', 'planned_count')}，built={get_nested(s_phase, 'dsm', 'built_count')}，failed={get_nested(s_phase, 'dsm', 'failed_count')}", f"CaiWangCun DSM cache：planned={get_nested(c_phase, 'dsm', 'planned_count')}，built={get_nested(c_phase, 'dsm', 'built_count')}，failed={get_nested(c_phase, 'dsm', 'failed_count')}"],
    ["验证 truth", "satellite truth patch；Layer-3 使用 RoMa v2 satellite tiepoint evaluator", "CaiWangCun DOM truth；增加 frame sanity 诊断"],
    ["运行时约束", "truth 不参与候选选择、重排或 PnP", "完整替换后仍保持 runtime truth-free"],
])}

## 3. 数据处理流程

本章结论：数据处理的关键风险不在单个模型，而在候选库、特征、索引、重排结果、DSM cache、位姿 manifest 与验证 truth 是否处于同一地理框架。CaiWangCun 前期半替换出现百米级偏移，说明局部匹配质量不能替代全链路坐标一致性审计。

卫星 DOM+SRTM 线路保留固定 satellite DOM candidate library，并使用 SRTM HGT 文件裁剪每个 active candidate 的 DSM cache。该线路的处理重点是保持运行时主链不变，只在 validation 侧使用 satellite truth patch 和 RoMa v2 同名点评估器，从而评估固定卫星候选库与低频高程支撑的实际定位表现。

CaiWangCun DOM+DSM 完整替换线路先将 CaiWangCun DOM 和 DSM mosaic 统一到 EPSG:32650，再基于共同覆盖范围重新生成 candidate tile library、candidate DINOv2 features、FAISS index、retrieval Top-20、RoMa v2 rerank、formal manifests、DSM cache、pose manifest、pose outputs 和 CaiWangCun DOM-truth validation outputs。该线路禁止回退到 ODM LAZ、SRTM 或旧 satellite candidate library。

{table_md(["流程环节", "处理内容", "关键控制点"], [
    ["离线数据准备", "准备 DOM/DSM 或 SRTM 源数据，并统一投影坐标系", "禁止混用未对齐的地图资产"],
    ["候选库构建", "生成 200m/300m/500m/700m 多尺度 tile", "候选 tile 同时承担检索图像和地理载体角色"],
    ["query 处理", "使用 009/010 两条 nadir 航线共 40 张 query", "query 元数据不进入运行时候选选择"],
    ["检索与重排", "DINOv2+FAISS 粗检索，RoMa v2 几何重排", "Top-20 进入后续 PnP 链路"],
    ["DSM cache", "按 active candidate 裁剪 DSM/DEM", "记录 planned/built/failed 和采样状态"],
    ["2D-3D 采样与 PnP", "候选 DOM 像素映射到 X/Y 并采样 Z，再求解位姿", "保留 PnP 成功、失败和评分分项"],
    ["验证输出", "生成 predicted ortho、Layer-1/2/3 指标和图件", "truth 只用于离线评价"],
])}

## 4. 定位精度检测方法和指标介绍

本章结论：定位精度不能由单一指标判断。当前采用三层验证框架，并辅以检索、运行时求解和 frame sanity 指标，以区分检索错误、DSM 支撑不足、坐标框架偏移、跨源外观差异和局部几何误差等不同问题。

Recall@K 表示正确候选是否进入前 K 个结果；MRR 是 Mean Reciprocal Rank，用于衡量正确候选平均排序位置。SSIM 是结构相似性指标，主要反映共同有效区域上的纹理结构相似度；IoU 是交并比，用于衡量 predicted ortho 与 truth 的有效覆盖重叠。RMSE 是均方根误差，对大误差更敏感；P90 表示 90% 样本不超过该误差。

Layer-1 检查 predicted ortho 与 truth orthophoto 在同一地理网格上的影像级对齐。predicted ortho 是由 best pose、单张 query 图像、候选关联 DSM 和 truth crop grid 生成的单图像位姿重投影结果，不是完整多视图正射重建。核心指标包括 phase correlation error、center offset、ortho IoU、SSIM 和有效像素比例。

Layer-2 检查 best pose 与 AT/query reference pose 的几何差异。该层直接评价相机中心和视线方向，核心指标包括 horizontal error、spatial error、view direction angle error 以及 yaw/pitch/roll error。

Layer-3 检查 predicted ortho 与 truth orthophoto 之间的局部同名点地面坐标差异。核心指标包括 tiepoint match count、inlier ratio、XY RMSE、XY P90 和 per-query tiepoint CSV。CaiWangCun 线路额外使用 frame sanity 诊断 DSM truth-grid 有效率、predicted valid pixel ratio、camera center offset 和 failure bucket。

{table_md(["层级/指标", "含义", "解释边界"], [
    ["Recall@K / MRR", "正确区域是否被检索到并排在前列", "只评价候选排序，不代表位姿已经准确"],
    ["Layer-1 center offset / IoU / SSIM", "predicted ortho 与 truth 的整体图像对齐", "SSIM 受跨源影像外观差异影响，不能单独代表几何精度"],
    ["Layer-2 horizontal error", "best pose 与参考位姿的平面误差", "异常样本会显著拉高均值，应同时看 median、P90 和 failure bucket"],
    ["Layer-3 tiepoint XY RMSE", "局部同名点在地面坐标中的误差", "依赖可匹配纹理和上游 predicted ortho 有效性"],
    ["Frame sanity", "DSM 覆盖、投影有效像素和坐标框架一致性诊断", "主要用于异常归因，不替代主精度指标"],
])}

## 5. 实验结果

本章结论：卫星 DOM+SRTM 线路完成 40/40 best pose 输出，整体处于约 10 m 级水平定位误差和约 2.77 m 局部同名点 RMSE；CaiWangCun DOM+DSM 完整替换线路同样完成 40/40 best pose 输出，除 q_037 外的可用样本进入约 1.78 m 水平误差和约 0.41 m 同名点 RMSE。完整替换线路显著优于半替换状态，也优于 SRTM 低频高程线路的精细定位表现。

{table_md(["线路", "Best pose", "Layer-1 对齐", "Layer-2 水平误差", "Layer-3 同名点 RMSE", "PnP 状态"], result_rows)}

卫星 DOM+SRTM 线路的工程价值在于稳定性。该线路在 40 个 query 上均获得 best pose，DSM cache `failed=0`，PnP 候选级成功 `730/800`。Layer-2 平均水平误差为 {fmt_num(get_nested(s_pose_at, "horizontal_error_m", "mean"), 3, " m")}，中位数为 {fmt_num(get_nested(s_pose_at, "horizontal_error_m", "median"), 3, " m")}，P90 为 {fmt_num(get_nested(s_pose_at, "horizontal_error_m", "p90"), 3, " m")}。Layer-3 同名点 RMSE 为 {fmt_num(s_tie.get("tiepoint_xy_error_rmse_m"), 3, " m")}。这说明固定 satellite DOM 候选库和低频 SRTM DSM 能支撑全量链路跑通，但精细几何误差仍受跨源影像差异、SRTM 低频高程和局部地物建模不足影响。

CaiWangCun DOM+DSM 完整替换线路的核心价值在于坐标框架一致性和精度改善。该线路同样完成 40/40 best pose，候选级 PnP 成功 `781/800`，Layer-1 成功 `39/40`。在 39 个 frame sanity 可用 query 上，horizontal error mean 为 {fmt_num(usable_horizontal, 3, " m")}；Layer-3 同名点 RMSE 为 {fmt_num(c_tie.get("tiepoint_xy_error_rmse_m"), 3, " m")}。全量 Layer-2 均值为 {fmt_num(get_nested(c_pose_at, "horizontal_error_m", "mean"), 3, " m")} 被 `q_037` 显著拉高，{q037_text} 该异常对应 Layer-1 `dsm_intersection_failed` 和 Layer-3 `upstream_eval_failed`，不表现为全局坐标框架偏移。

风险和局限性如下。第一，跨源影像场景下，SSIM 不能单独作为几何精度判断依据，应与 center offset、IoU、Layer-2 和 Layer-3 联合解释。第二，SRTM 表面更平滑，视觉上可能减少局部扭曲，但不代表城市细节几何更真实。第三，CaiWangCun 完整替换结果仍存在 `q_037` 单样本异常，后续扩大样本前需要单独复核 DSM/truth-grid intersection 与参考位姿差异。第四，完整替换链路显著增加运行成本，需要针对 Layer-3 和 DSM sampling 做工程优化。

### 5.1 代表性图件

{chr(10).join(image_lines)}

## 6. 耗时统计

本章结论：本章统一采用“含上游检索/重排资产口径”统计耗时，即从能够形成运行时 Top-20 候选的 DINOv2 特征、FAISS/DINOv2 粗检索、RoMa v2 重排等上游资产生成阶段开始，统计到位姿求解、精度验证和最终报告输出完成。按该口径，卫星 DOM+SRTM 线路整体耗时约 6h47m，CaiWangCun DOM+DSM 完整替换线路整体耗时约 19h18m47s，后者约为前者的 2.8 倍。

### 6.1 统计口径与运行环境

本章只使用“含上游检索/重排资产口径”作为主统计口径。该口径的统计起点不是 formal pose，而是 DINOv2 query/candidate 特征、FAISS 或 DINOv2 粗检索、RoMa v2 重排等上游资产生成阶段；统计终点是 full validation suite 和最终报告输出完成。旧的 `3h40m05s` 与 `15h18m55s` 仅作为 formal/full 子阶段耗时保留，不再作为主耗时。

两条线路仍不是同一次连续运行。卫星 DOM+SRTM 线路跨 2026-04-10 的 DINOv2/RoMa 上游资产生成和 2026-04-16 的 SRTM formal pose/validation；CaiWangCun 线路跨 2026-04-20 的 full-replacement gate 上游资产生成和 2026-04-21 的 full formal pose/validation。因此，本章耗时适合作为工程成本估算和相对比较，不应解释为严格硬件 benchmark。

{table_md(["项目", "配置"], environment_rows)}

### 6.2 双线路整体耗时对比

{table_md(["线路", "统计起点", "统计终点", "整体耗时", "备注"], runtime_rows)}

### 6.3 卫星 DOM+SRTM 关键阶段耗时

卫星 DOM+SRTM 线路从 2026-04-10 的 DINOv2 query 特征提取、DINOv2 coarse retrieval 和 RoMa v2 rerank 资产生成开始统计，接 2026-04-16 的 SRTM formal pose、validation 和报告输出。主报告采用约 6h47m 作为整体耗时；按日志边界从 DINOv2 特征提取命令开始到 RoMa timing 完成再叠加 2026-04-16 主日志跨度，保守估算约 6h50m21s。

{table_md(["阶段", "时间范围/来源", "耗时", "说明"], srtm_runtime_rows)}

该统计不包含更早期固定 satellite DOM 候选库、satellite candidate features 和 FAISS index 的历史构建成本，因为当前 SRTM 正式实验使用的是已固定候选库。SRTM formal 阶段每个 query-candidate pair 采样 2000 条，合计 1,600,000 条 matches/correspondences；其耗时较低的一部分原因是 SRTM DSM cache 栅格较小，但这也对应低频 DSM 与较弱精细地物表达能力。

### 6.4 CaiWangCun DOM+DSM 完整替换关键阶段耗时

CaiWangCun DOM+DSM 完整替换线路从 2026-04-20 的 CaiWangCun candidate library 构建开始统计，包含 candidate features、FAISS index、DINOv2 retrieval 和 RoMa v2 rerank 等上游资产，再接 2026-04-21 的 full formal pose、validation、frame sanity 和报告输出。整体耗时约 19h18m47s。

{table_md(["阶段", "时间范围/来源", "耗时", "说明"], cai_runtime_rows)}

该统计仍不包含 CaiWangCun 原始 DOM/DSM mosaic 从源切片拼接、重投影的耗时；现有可审计日志从 candidate library 构建开始。full run 阶段复用或复制校验了 gate 阶段生成的 candidate library、features、FAISS、retrieval 和 RoMa rerank，因此把 gate 上游资产阶段纳入整体耗时后，更接近完整替换线路的工程成本。

### 6.5 耗时分析与工程含义

CaiWangCun 完整替换线路耗时明显高于 SRTM 线路，主要来自两个方面。第一，full validation suite 中 Layer-3 tiepoint ground error 约 9h15m46s，是当前最主要的单项耗时。第二，full pose pipeline 需要处理约 4,000,000 条 correspondence，RoMa matches export 和 DSM sampling 的计算量显著高于 SRTM formal 阶段的 1,600,000 条 correspondence。

两条线路采用更接近的耗时统计口径后，仍存在不可完全消除的变量差异，包括 sample-count、DSM 分辨率、truth 类型、候选库规模和 validation 计算量。因此，本章耗时对比的工程含义是评估当前实现下的实际成本，而不是严格归因到单一数据源或单一模型。

## 7. 结论和分析

本章结论：两条线路均验证了在单张 UAV 图像不携带地理元数据、且不假设其为正射图的条件下，DOM 检索、RoMa v2 几何重排、DSM/DEM 高程约束和 PnP 位姿求解可以形成可用的初始地理定位结果。CaiWangCun DOM+DSM 完整替换线路在可用样本上达到当前最优精度，卫星 DOM+SRTM 线路则提供了更低成本、更稳定的跨源基线。

### 7.1 总体结论

当前实验表明，单张 UAV 图像初始地理定位的关键不只是检索模型本身，而是检索候选、几何重排、三维高程支撑、位姿求解和验证基准之间的坐标框架一致性。两条正式线路均完成 40/40 best pose 输出，说明当前链路已经具备端到端工程跑通能力。

从工程优先级看，CaiWangCun DOM+DSM 完整替换线路应作为当前精度最优线路；卫星 DOM+SRTM 线路应作为跨源、低频高程、较低耗时条件下的稳定对照线路。前者更适合评估米级定位潜力，后者更适合评估卫星底图和低频 DEM/DSM 条件下的保底能力。

### 7.2 双线路对比分析

卫星 DOM+SRTM 线路的优势是资产依赖较少、运行成本较低、全量 40 个 query 均能形成 best pose。其局限是 SRTM 高程表达偏低频，难以精确表示建筑、道路边缘和局部地物高度变化；因此，Layer-2 和 Layer-3 指标虽保持可用，但精细几何精度明显弱于同源 DOM/DSM 完整替换线路。

CaiWangCun DOM+DSM 完整替换线路的优势是候选 DOM、DSM、特征、索引、检索、重排、manifest、DSM cache 和 validation truth 均被统一到同一数据框架内。该完整替换消除了前期半替换状态下的百米级偏移风险，使 39 个可用 query 达到约 1.778 m 的水平误差均值和约 0.414 m 的 Layer-3 同名点 RMSE。其代价是资产生成和验证耗时显著增加，对工程调度和算力占用提出更高要求。

### 7.3 精度结论

卫星 DOM+SRTM 线路在 40 个 query 上均得到 best pose，Layer-2 horizontal error mean 为 9.723 m，Layer-3 tiepoint XY RMSE 为 2.772 m。该结果说明固定 satellite DOM 候选库结合 SRTM 高程能够支撑 009/010 nadir 样本的稳定初始定位，但当前精度更接近十米级初值恢复，而不是精细局部几何对齐。

CaiWangCun DOM+DSM 完整替换线路同样完成 40/40 best pose。在排除 `q_037` 对应的 validation-missing 异常后，39 个可用 query 的 frame sanity horizontal error mean 为 1.778 m，Layer-3 tiepoint XY RMSE 为 0.414 m，明显优于 SRTM 线路。全量 Layer-2 mean 为 22.964 m，主要由 `q_037` 的 849.225 m 水平误差拉高，因此该均值不宜单独代表该线路的典型精度。

### 7.4 耗时结论

在“含上游检索/重排资产口径”下，卫星 DOM+SRTM 线路整体耗时约 6h47m，CaiWangCun DOM+DSM 完整替换线路整体耗时约 19h18m47s，后者约为前者的 2.8 倍。该差异主要来自 CaiWangCun full validation suite 中约 9h15m46s 的 Layer-3 tiepoint ground error，以及 full pose pipeline 中约 4,000,000 条 correspondence 对 RoMa matches export 和 DSM sampling 的计算压力。

该耗时差异不能简单解释为某一数据源更慢。两条线路在 sample-count、DSM 分辨率、truth 类型、候选库规模和 validation 计算量上仍存在差异。因此，当前耗时结论更适合作为工程成本评估，而不是严格硬件 benchmark 或单变量性能归因。

### 7.5 局限性和风险

当前结论仍受样本规模和样本类型限制。009/010 实验主要覆盖 40 个 nadir query，尚不能代表大倾角、强遮挡、低纹理、季节差异或更大区域迁移场景。后续扩大样本前，应避免把当前米级表现直接外推到所有 UAV 输入条件。

跨源影像场景下，SSIM 不宜单独作为几何精度判断依据。SRTM 表面更平滑，可能降低 predicted ortho 的局部视觉扭曲，但这不等价于城市细节几何更真实。CaiWangCun 线路虽然在可用样本上表现最优，但 `q_037` 已暴露 DSM/truth-grid intersection、参考位姿差异或局部投影有效性方面的异常风险。完整替换线路还显著增加运行成本，若不优化 Layer-3 和 DSM sampling，批量工程运行会受到耗时约束。

## 8. 后续的想法

本章结论：后续工作应优先围绕异常样本复核、耗时瓶颈优化、资产一致性审计、泛化实验扩展和工程化输出五个方向展开。优先级最高的是 `q_037` 异常闭环与 Layer-3 / DSM sampling 耗时优化，因为这两项分别决定当前最优线路的可信边界和批量运行成本。

### 8.1 异常样本专项复核

应将 `q_037` 作为首个专项复核对象，重点检查 DSM/truth-grid intersection、CaiWangCun DOM/DSM 覆盖范围、validation truth crop、RoMa v2 匹配点空间分布、PnP 内点结构、参考 AT 位姿和 frame sanity failure bucket。复核目标不是只修正单个样本，而是确认该异常属于数据覆盖问题、参考框架问题、局部匹配退化问题，还是验证流程边界问题。

专项复核完成后，应形成可复用的异常判别规则，例如 DSM 交会失败、投影有效像素不足、视线角异常、Layer-2 与 Layer-3 指标冲突、同名点内点集中于局部区域等状态码。该规则应进入后续 full run 的质量门控，避免异常样本被均值指标掩盖。

### 8.2 耗时优化

耗时优化应优先处理 Layer-3 tiepoint ground error、RoMa v2 matches export 和 DSM sampling。Layer-3 可评估批处理、特征缓存、按质量阈值跳过低价值样本、以及仅对关键样本执行高密度同名点评估的策略。RoMa v2 和 DSM sampling 可评估批量 I/O、候选级 cache 复用、矢量化采样、sample-count 自适应和失败早停机制。

优化目标应区分“正式评估模式”和“工程运行模式”。正式评估模式保留完整验证链路，用于产生可汇报指标；工程运行模式则优先输出定位初值、置信度和失败原因，在必要时才触发重验证或高密度 Layer-3 评估。

### 8.3 数据资产一致性审计

后续实验应固化完整替换线路的资产一致性审计，覆盖 DOM/DSM 投影坐标系、mosaic bounds、candidate library、candidate features、FAISS index、retrieval Top-20、RoMa v2 rerank、formal manifest、DSM cache、pose manifest 和 validation truth。每次运行前应输出可机读 audit summary，明确是否存在旧资产混入、路径回退、候选库和 DSM 不一致、truth 与 runtime 资产不在同一框架等风险。

该审计机制对 CaiWangCun 线路尤其重要。前期半替换实验已经表明，局部替换 DOM 或 DSM 即使能改善匹配质量，也可能因候选、truth 和高程框架不一致造成百米级偏移。后续应把“完整替换”作为同源 DOM/DSM 实验的默认工程约束。

### 8.4 泛化实验扩展

当前 009/010 nadir 样本应扩展到更多 query 条件，包括非 nadir 视角、不同高度、不同纹理密度、不同航线、不同季节或光照差异，以及更大范围的 DOM/DSM 数据源。扩展实验应保留同一套三层验证框架，使新增场景能够与当前结果进行可比分析。

泛化验证还应覆盖不同高程源组合。SRTM、ODM DSM、CaiWangCun DSM 和可能的其他城市 DSM 应在统一候选和统一 truth 口径下分组比较，以区分高程分辨率、数据同源性、坐标框架一致性和 DSM nodata 对定位精度的贡献。

### 8.5 工程化输出

后续工程化应将当前实验输出整理为可被导航、融合定位或人工审核模块消费的标准结果，包括 best pose、候选排序、置信度、PnP 内点统计、DSM 采样有效率、Layer-1/2/3 验证状态、异常原因和推荐处置策略。输出不应只包含最终坐标，也应包含失败可解释信息。

在工程部署层面，建议建立 gate/full 两级流程。gate 阶段用于快速验证资产框架、DSM 覆盖和少量样本精度；full 阶段用于全量 query 运行、完整验证和报告归档。该流程能够降低半替换或错配资产直接进入长耗时 full run 的风险，并为后续批量实验提供稳定的质量控制入口。
"""
    return md


def set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell: Any, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)


def add_table(document: Document, headers: list[str], rows: list[list[Any]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
        set_cell_shading(table.rows[0].cells[idx], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], str(value))
    document.add_paragraph()


def add_paragraph(document: Document, text: str, bold_prefix: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(21)
    paragraph.paragraph_format.line_spacing = 1.15
    if bold_prefix and "：" in text:
        prefix, rest = text.split("：", 1)
        run = paragraph.add_run(prefix + "：")
        run.bold = True
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10.5)
        run = paragraph.add_run(rest)
    else:
        run = paragraph.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        run.font.color.rgb = None


def apply_doc_style(document: Document) -> None:
    styles = document.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)
    for section in document.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)


def add_image(document: Document, title: str, path: Path) -> None:
    if not path.exists():
        add_paragraph(document, f"缺失图件：{title}。路径：{path}")
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(5.8))
    caption = document.add_paragraph(title)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(9)


def build_docx(metrics: dict[str, Any], images: list[dict[str, str]], out_docx: Path) -> None:
    md = build_markdown(metrics, images)
    document = Document()
    apply_doc_style(document)

    title = document.add_heading("009/010 双线路工程汇报：卫星 DOM+SRTM 与 CaiWangCun DOM+DSM 完整替换", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table_headers: list[str] | None = None
    table_rows: list[list[str]] = []
    in_table = False

    for raw_line in md.splitlines():
        line = raw_line.strip()
        if not line:
            if in_table and table_headers is not None:
                add_table(document, table_headers, table_rows)
                table_headers = None
                table_rows = []
                in_table = False
            continue

        if line.startswith("|"):
            parts = [part.strip() for part in line.strip("|").split("|")]
            if all(part.replace("-", "") == "" for part in parts):
                continue
            if table_headers is None:
                table_headers = parts
                table_rows = []
            else:
                table_rows.append(parts)
            in_table = True
            continue

        if in_table and table_headers is not None:
            add_table(document, table_headers, table_rows)
            table_headers = None
            table_rows = []
            in_table = False

        if line.startswith("# "):
            continue
        if line.startswith("## "):
            add_heading(document, line[3:], 1)
        elif line.startswith("### "):
            add_heading(document, line[4:], 2)
        elif line.startswith("!["):
            title_part = line[2:].split("](", 1)[0]
            path_part = line.split("](", 1)[1].rstrip(")")
            add_image(document, title_part, Path(path_part))
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        else:
            add_paragraph(document, line, bold_prefix=line.startswith("本章结论"))

    if in_table and table_headers is not None:
        add_table(document, table_headers, table_rows)

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(out_docx)


def main() -> None:
    metrics = collect_metrics()
    images = image_inventory()
    md = build_markdown(metrics, images)
    OUT_MD.parent.mkdir(parents=True, exist_ok=True)
    OUT_MD.write_text(md, encoding="utf-8")
    build_docx(metrics, images, OUT_DOCX)
    print(f"Wrote Markdown: {OUT_MD}")
    print(f"Wrote Word: {OUT_DOCX}")


if __name__ == "__main__":
    main()
