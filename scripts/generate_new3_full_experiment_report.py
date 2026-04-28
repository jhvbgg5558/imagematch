#!/usr/bin/env python3
"""Generate an integrated Word report for the new3output ODM-refresh experiment.

Purpose:
- summarize the full `new3output` 009/010 nadir experiment in one standalone
  Word report under the experiment `reports/` directory;
- combine runtime closure, ODM-truth validation, satellite-truth validation,
  cross-suite comparison, and the predicted-ortho missing-coverage analysis;
- generate companion charts and sample image panels used directly by the Word
  report so the output is not limited to text and tables.

Main inputs:
- experiment-level selection and retrieval outputs under the experiment root;
- formal pose runtime summaries under `pose_v1_formal/summary/` and `pnp/`;
- ODM-truth suite outputs under
  `pose_v1_formal/eval_pose_validation_suite_odm_truth/`;
- satellite-truth suite outputs under
  `pose_v1_formal/eval_pose_validation_suite_satellite_truth/`;
- baseline comparison assets under the historical `new2output` root.

Main outputs:
- `<experiment-root>/reports/nadir_009010_odmrefresh_sattruth_full_experiment_report.docx`;
- `<experiment-root>/reports/full_experiment_report_assets/*.png`.

Applicable task constraints:
- keep the runtime task definition unchanged as UAV-to-satellite localization;
- describe ODM orthophoto truth and satellite-truth patches as evaluation-only
  assets;
- explicitly explain that predicted orthophotos are valid projected coverage on
  the truth grid rather than complete orthophoto reconstructions.
"""

from __future__ import annotations

import argparse
import csv
import json
import math
from collections import Counter
from pathlib import Path
from statistics import mean

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import rasterio
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_BASELINE_ROOT = PROJECT_ROOT / "new2output" / "nadir_009010_dinov2_romav2_pose_2026-04-10"
DEFAULT_EXPERIMENT_ROOT = (
    PROJECT_ROOT / "new3output" / "nadir_009010_dinov2_romav2_pose_odmrefresh_sattruth_2026-04-16"
)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--experiment-root", default=str(DEFAULT_EXPERIMENT_ROOT))
    parser.add_argument("--baseline-root", default=str(DEFAULT_BASELINE_ROOT))
    parser.add_argument(
        "--out-docx",
        default="",
        help="Optional explicit output path. Defaults to <experiment-root>/reports/nadir_009010_odmrefresh_sattruth_full_experiment_report.docx",
    )
    parser.add_argument(
        "--assets-dirname",
        default="full_experiment_report_assets",
        help="Directory name under reports/ used for generated figures.",
    )
    return parser.parse_args()


def load_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", newline="", encoding="utf-8-sig") as handle:
        return list(csv.DictReader(handle))


def load_json(path: Path) -> dict[str, object]:
    return json.loads(path.read_text(encoding="utf-8"))


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def set_cn_font(run, size: int = 11, bold: bool = False) -> None:
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(size)
    run.bold = bold


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = "Heading 1" if level == 1 else "Heading 2"
    run = paragraph.add_run(text)
    set_cn_font(run, size=14 if level == 1 else 12, bold=True)


def add_paragraph(
    doc: Document,
    text: str,
    *,
    size: int = 11,
    center: bool = False,
    bold: bool = False,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_cn_font(run, size=size, bold=bold)


def add_bullets(doc: Document, lines: list[str]) -> None:
    for line in lines:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(line)
        set_cn_font(run, size=11)


def set_cell_text(cell, text: str, *, bold: bool = False, size: int = 10) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_cn_font(run, size=size, bold=bold)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], *, header_fill: str = "D9EAF7") -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_text(table.cell(0, idx), header, bold=True)
        shade_cell(table.cell(0, idx), header_fill)
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            set_cell_text(cells[idx], value)


def add_picture_with_caption(doc: Document, image_path: Path, caption: str, *, width_inch: float = 6.2) -> None:
    if not image_path.exists():
        add_paragraph(doc, f"图像缺失：{image_path}", size=10)
        return
    doc.add_picture(str(image_path), width=Inches(width_inch))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(caption)
    set_cn_font(run, size=10)


def fmt(value: object, digits: int = 4) -> str:
    if value in ("", None):
        return "-"
    try:
        numeric = float(value)
    except (TypeError, ValueError):
        return str(value)
    if not math.isfinite(numeric):
        return "-"
    return f"{numeric:.{digits}f}"


def parse_float(value: str | None) -> float | None:
    if value in ("", None):
        return None
    try:
        return float(value)
    except ValueError:
        return None


def relative_to_cwd(path: Path) -> str:
    try:
        return str(path.relative_to(Path.cwd()))
    except ValueError:
        return str(path)


def short_flight_name(flight_id: str) -> str:
    parts = flight_id.split("_")
    return parts[2] if len(parts) >= 3 else flight_id


def read_rgba_or_rgb(path: Path) -> tuple[np.ndarray, np.ndarray | None]:
    with rasterio.open(path) as ds:
        data = ds.read()
    if data.shape[0] >= 4:
        rgb = np.moveaxis(data[:3], 0, 2)
        alpha = data[3]
    else:
        rgb = np.moveaxis(data[:3], 0, 2)
        alpha = None
    return rgb.astype(np.uint8), alpha


def normalize_rgb(rgb: np.ndarray) -> np.ndarray:
    if rgb.dtype != np.uint8:
        rgb = np.clip(rgb, 0, 255).astype(np.uint8)
    return rgb


def overlay_truth_and_pred(truth_rgb: np.ndarray, pred_rgb: np.ndarray, alpha: np.ndarray | None) -> np.ndarray:
    truth = truth_rgb.astype(np.float32) / 255.0
    pred = pred_rgb.astype(np.float32) / 255.0
    if alpha is None:
        mask = np.any(pred_rgb > 0, axis=2).astype(np.float32)
    else:
        mask = (alpha > 0).astype(np.float32)
    mask3 = np.repeat(mask[:, :, None], 3, axis=2)
    overlay = truth * (1.0 - 0.55 * mask3) + pred * (0.55 * mask3)
    return np.clip(overlay, 0.0, 1.0)


def mask_to_rgb(alpha: np.ndarray | None, fallback_rgb: np.ndarray) -> np.ndarray:
    if alpha is None:
        mask = np.any(fallback_rgb > 0, axis=2).astype(np.uint8) * 255
    else:
        mask = np.where(alpha > 0, 255, 0).astype(np.uint8)
    return np.repeat(mask[:, :, None], 3, axis=2)


def summarize_selected_queries(rows: list[dict[str, str]]) -> dict[str, object]:
    counts = Counter(row["flight_id"] for row in rows)
    pitch_values = [float(row["gimbal_pitch_degree"]) for row in rows if row.get("gimbal_pitch_degree")]
    ordered = sorted(counts)
    query_ranges = []
    start = 1
    for flight_id in ordered:
        count = counts[flight_id]
        query_ranges.append(
            {
                "flight_id": flight_id,
                "count": count,
                "query_id_start": f"q_{start:03d}",
                "query_id_end": f"q_{start + count - 1:03d}",
            }
        )
        start += count
    return {
        "query_count": len(rows),
        "flight_counts": dict(counts),
        "ordered_flights": ordered,
        "query_ranges": query_ranges,
        "pitch_min": min(pitch_values) if pitch_values else None,
        "pitch_max": max(pitch_values) if pitch_values else None,
    }


def choose_odm_sample_queries(odm_rows: list[dict[str, str]]) -> tuple[str, str, str]:
    rows = [row for row in odm_rows if row.get("eval_status") == "ok"]
    rows_sorted = sorted(rows, key=lambda row: float(row["common_valid_ratio"]))
    low = rows_sorted[0]["query_id"]
    high = rows_sorted[-1]["query_id"]
    ratios = [float(row["common_valid_ratio"]) for row in rows_sorted]
    median_ratio = np.median(ratios)
    mid = min(rows_sorted, key=lambda row: abs(float(row["common_valid_ratio"]) - median_ratio))["query_id"]
    return high, mid, low


def build_metric_bar_chart(
    csv_rows: list[dict[str, str]],
    columns: list[tuple[str, str]],
    title: str,
    out_path: Path,
    *,
    rotate_xticks: bool = True,
) -> None:
    query_ids = [row["query_id"] for row in csv_rows]
    fig, axes = plt.subplots(len(columns), 1, figsize=(14, 3.2 * len(columns)), sharex=True)
    if len(columns) == 1:
        axes = [axes]
    for axis, (column, label) in zip(axes, columns):
        values = [float(row[column]) if row.get(column) not in ("", None) else float("nan") for row in csv_rows]
        axis.bar(query_ids, values, color="#4C78A8")
        axis.set_ylabel(label)
        axis.grid(axis="y", alpha=0.25)
    axes[0].set_title(title)
    axes[-1].set_xlabel("query_id")
    if rotate_xticks:
        axes[-1].tick_params(axis="x", rotation=90)
    fig.tight_layout()
    fig.savefig(out_path, dpi=180)
    plt.close(fig)


def build_baseline_comparison_chart(
    baseline_overall: dict[str, object],
    odm_overall: dict[str, object],
    sat_ortho_overall: dict[str, object],
    sat_tie_overall: dict[str, object],
    sat_ortho_rows: list[dict[str, str]],
    sat_geom_rows: list[dict[str, str]],
    out_path: Path,
) -> None:
    def mean_payload(src: dict[str, object], key: str) -> float:
        payload = src.get(key, {})
        if isinstance(payload, dict):
            return float(payload.get("mean", float("nan")))
        return float("nan")

    geom_offsets = [
        float(row["camera_center_offset_m"])
        for row in sat_geom_rows
        if row.get("eval_status") == "ok" and row.get("camera_center_offset_m") not in ("", None)
    ]
    labels = ["baseline_uav_truth", "odm_truth_refresh", "satellite_truth"]
    metrics = {
        "layer1_ortho_iou_mean": [
            mean_payload(baseline_overall["ortho"], "ortho_iou"),
            mean_payload(odm_overall["ortho"], "ortho_iou"),
            mean_payload(sat_ortho_overall, "ortho_iou"),
        ],
        "layer1_common_valid_ratio_mean": [
            mean(float(row["common_valid_ratio"]) for row in baseline_overall["odm_rows"]),
            mean(float(row["common_valid_ratio"]) for row in odm_overall["odm_rows"]),
            mean(float(row["common_valid_ratio"]) for row in sat_ortho_rows),
        ],
        "layer3_tiepoint_rmse_m": [
            float(baseline_overall["tie"]["tiepoint_xy_error_rmse_m"]),
            float(odm_overall["tie"]["tiepoint_xy_error_rmse_m"]),
            float(sat_tie_overall["tiepoint_xy_error_rmse_m"]),
        ],
        "satellite_geometry_center_offset_m": [
            float("nan"),
            float("nan"),
            mean(geom_offsets) if geom_offsets else float("nan"),
        ],
    }

    fig, axes = plt.subplots(2, 2, figsize=(12, 8))
    axes = axes.ravel()
    for axis, (metric, values) in zip(axes, metrics.items()):
        axis.bar(labels, values, color=["#4C78A8", "#F58518", "#54A24B"])
        axis.set_title(metric)
        axis.grid(axis="y", alpha=0.25)
        axis.tick_params(axis="x", rotation=12)
    fig.suptitle("Baseline vs ODM-truth vs Satellite-truth key metrics", fontsize=13)
    fig.tight_layout()
    fig.savefig(out_path, dpi=180)
    plt.close(fig)


def build_sample_panel(
    truth_path: Path,
    pred_path: Path,
    out_path: Path,
    title: str,
    *,
    stats_lines: list[str],
) -> None:
    truth_rgb, _ = read_rgba_or_rgb(truth_path)
    pred_rgb, alpha = read_rgba_or_rgb(pred_path)
    truth_rgb = normalize_rgb(truth_rgb)
    pred_rgb = normalize_rgb(pred_rgb)
    overlay = overlay_truth_and_pred(truth_rgb, pred_rgb, alpha)
    mask_rgb = mask_to_rgb(alpha, pred_rgb)

    fig, axes = plt.subplots(1, 4, figsize=(16, 4.6))
    panels = [
        (truth_rgb, "truth"),
        (pred_rgb, "pred"),
        (mask_rgb, "alpha/mask"),
        (overlay, "overlay"),
    ]
    for axis, (image, label) in zip(axes, panels):
        if image.dtype == np.uint8:
            axis.imshow(image)
        else:
            axis.imshow(np.clip(image, 0.0, 1.0))
        axis.set_title(label)
        axis.axis("off")
    fig.suptitle(title, fontsize=13)
    footer = "\n".join(stats_lines)
    fig.text(0.01, 0.01, footer, fontsize=9, ha="left", va="bottom")
    fig.tight_layout(rect=(0, 0.08, 1, 0.94))
    fig.savefig(out_path, dpi=180)
    plt.close(fig)


def build_missing_issue_chart(rows: list[dict[str, str]], out_path: Path) -> None:
    sorted_rows = sorted(rows, key=lambda row: float(row["common_valid_ratio"]))
    query_ids = [row["query_id"] for row in sorted_rows]
    coverage = [float(row["common_valid_ratio"]) for row in sorted_rows]
    iou = [float(row["ortho_iou"]) for row in sorted_rows]
    center_offset = [float(row["center_offset_m"]) for row in sorted_rows]

    fig, axes = plt.subplots(3, 1, figsize=(14, 9), sharex=True)
    axes[0].bar(query_ids, coverage, color="#E45756")
    axes[0].set_ylabel("common_valid_ratio")
    axes[0].grid(axis="y", alpha=0.25)
    axes[1].bar(query_ids, iou, color="#4C78A8")
    axes[1].set_ylabel("ortho_iou")
    axes[1].grid(axis="y", alpha=0.25)
    axes[2].bar(query_ids, center_offset, color="#72B7B2")
    axes[2].set_ylabel("center_offset_m")
    axes[2].grid(axis="y", alpha=0.25)
    axes[2].set_xlabel("query_id")
    axes[2].tick_params(axis="x", rotation=90)
    fig.suptitle("ODM-truth predicted-ortho valid coverage and alignment", fontsize=13)
    fig.tight_layout()
    fig.savefig(out_path, dpi=180)
    plt.close(fig)


def build_runtime_rows(full_summary: dict[str, object]) -> list[list[str]]:
    rows = []
    for item in full_summary.get("steps", []):
        rows.append(
            [
                str(item.get("step_name", "")),
                fmt(item.get("elapsed_sec"), 1),
                str(item.get("returncode", "")),
            ]
        )
    return rows


def main() -> None:
    args = parse_args()
    experiment_root = Path(args.experiment_root).resolve()
    baseline_root = Path(args.baseline_root).resolve()
    reports_root = experiment_root / "reports"
    assets_root = reports_root / args.assets_dirname
    out_docx = (
        Path(args.out_docx).resolve()
        if args.out_docx
        else reports_root / "nadir_009010_odmrefresh_sattruth_full_experiment_report.docx"
    )
    ensure_dir(reports_root)
    ensure_dir(assets_root)

    pose_root = experiment_root / "pose_v1_formal"
    odm_suite_root = pose_root / "eval_pose_validation_suite_odm_truth"
    sat_suite_root = pose_root / "eval_pose_validation_suite_satellite_truth"
    baseline_pose_root = baseline_root / "pose_v1_formal"
    baseline_suite_root = baseline_pose_root / "eval_pose_validation_suite"

    selected_rows = load_csv(experiment_root / "selected_queries" / "selected_images_summary.csv")
    retrieval_rows = load_csv(experiment_root / "retrieval" / "retrieval_top20.csv")

    pose_summary = load_json(pose_root / "summary" / "pose_overall_summary.json")
    pnp_summary = load_json(pose_root / "pnp" / "pnp_summary.json")
    pose_runtime_steps = load_json(odm_suite_root / "full_run_summary.json")
    best_pose_rows = load_csv(pose_root / "summary" / "per_query_best_pose.csv")

    odm_ortho_overall = load_json(odm_suite_root / "ortho_alignment" / "overall_ortho_accuracy.json")
    odm_pose_overall = load_json(odm_suite_root / "pose_vs_at" / "overall_pose_vs_at.json")
    odm_tie_overall = load_json(odm_suite_root / "tiepoint_ground_error" / "overall_tiepoint_ground_error.json")
    odm_ortho_rows = load_csv(odm_suite_root / "ortho_alignment" / "per_query_ortho_accuracy.csv")
    odm_pose_rows = load_csv(odm_suite_root / "pose_vs_at" / "per_query_pose_vs_at.csv")
    odm_tie_rows = load_csv(odm_suite_root / "tiepoint_ground_error" / "per_query_tiepoint_ground_error.csv")

    sat_full_summary = load_json(sat_suite_root / "full_run_summary.json")
    sat_ortho_overall = load_json(sat_suite_root / "ortho_alignment_satellite" / "overall_ortho_accuracy.json")
    sat_geom_overall = load_json(sat_suite_root / "pose_vs_satellite_truth_geometry" / "overall_satellite_truth_geometry.json")
    sat_tie_overall = load_json(sat_suite_root / "tiepoint_ground_error_satellite" / "overall_tiepoint_ground_error.json")
    sat_ortho_rows = load_csv(sat_suite_root / "ortho_alignment_satellite" / "per_query_ortho_accuracy.csv")
    sat_geom_rows = load_csv(sat_suite_root / "pose_vs_satellite_truth_geometry" / "per_query_pose_vs_satellite_truth_geometry.csv")
    sat_tie_rows = load_csv(sat_suite_root / "tiepoint_ground_error_satellite" / "per_query_tiepoint_ground_error.csv")

    baseline_ortho_overall = load_json(baseline_suite_root / "ortho_alignment" / "overall_ortho_accuracy.json")
    baseline_pose_overall = load_json(baseline_suite_root / "pose_vs_at" / "overall_pose_vs_at.json")
    baseline_tie_overall = load_json(baseline_suite_root / "tiepoint_ground_error" / "overall_tiepoint_ground_error.json")
    baseline_ortho_rows = load_csv(baseline_suite_root / "ortho_alignment" / "per_query_ortho_accuracy.csv")

    query_summary = summarize_selected_queries(selected_rows)

    high_q, mid_q, low_q = choose_odm_sample_queries(odm_ortho_rows)
    sample_query_ids = []
    for query_id in [high_q, mid_q, low_q, "q_001"]:
        if query_id not in sample_query_ids:
            sample_query_ids.append(query_id)

    odm_chart_path = assets_root / "odm_truth_metrics.png"
    sat_chart_path = assets_root / "satellite_truth_metrics.png"
    comparison_chart_path = assets_root / "baseline_vs_new3_key_metrics.png"
    missing_chart_path = assets_root / "predicted_ortho_missing_coverage_analysis.png"

    build_metric_bar_chart(
        odm_ortho_rows,
        [
            ("ortho_iou", "ortho_iou"),
            ("common_valid_ratio", "common_valid_ratio"),
            ("center_offset_m", "center_offset_m"),
            ("ssim", "ssim"),
        ],
        "ODM-truth per-query metrics",
        odm_chart_path,
    )
    build_metric_bar_chart(
        sat_ortho_rows,
        [
            ("phase_corr_error_m", "phase_corr_error_m"),
            ("center_offset_m", "center_offset_m"),
        ],
        "Satellite-truth layer-1 per-query metrics",
        sat_chart_path,
    )
    build_missing_issue_chart(odm_ortho_rows, missing_chart_path)
    build_baseline_comparison_chart(
        {
            "ortho": baseline_ortho_overall,
            "tie": baseline_tie_overall,
            "odm_rows": baseline_ortho_rows,
        },
        {
            "ortho": odm_ortho_overall,
            "tie": odm_tie_overall,
            "odm_rows": odm_ortho_rows,
        },
        sat_ortho_overall,
        sat_tie_overall,
        sat_ortho_rows,
        sat_geom_rows,
        comparison_chart_path,
    )

    sample_panels: list[tuple[Path, str]] = []
    odm_by_query = {row["query_id"]: row for row in odm_ortho_rows}
    for query_id in sample_query_ids:
        row = odm_by_query[query_id]
        panel_path = assets_root / f"{query_id}_odm_truth_sample_panel.png"
        build_sample_panel(
            Path(row["truth_crop_path"]),
            Path(row["pred_crop_path"]),
            panel_path,
            f"{query_id} ODM-truth sample",
            stats_lines=[
                f"common_valid_ratio={fmt(row['common_valid_ratio'], 4)}",
                f"ortho_iou={fmt(row['ortho_iou'], 4)}",
                f"center_offset_m={fmt(row['center_offset_m'], 4)}",
                f"best_inlier_count={row['best_inlier_count']}",
            ],
        )
        sample_panels.append((panel_path, query_id))

    q001_row = odm_by_query["q_001"]
    low_row = odm_by_query[low_q]

    missing_summary_lines = [
        (
            f"q_001: common_valid_ratio={fmt(q001_row['common_valid_ratio'], 4)}, "
            f"ortho_iou={fmt(q001_row['ortho_iou'], 4)}, center_offset_m={fmt(q001_row['center_offset_m'], 4)}"
        ),
        (
            f"{low_q}: common_valid_ratio={fmt(low_row['common_valid_ratio'], 4)}, "
            f"ortho_iou={fmt(low_row['ortho_iou'], 4)}, center_offset_m={fmt(low_row['center_offset_m'], 4)}"
        ),
    ]

    runtime_elapsed_sec = sum(float(item.get("elapsed_sec", 0.0)) for item in pose_runtime_steps.get("steps", []))

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("009/010 ODM Refresh + Satellite Truth 综合实验报告")
    set_cn_font(title_run, size=16, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(relative_to_cwd(experiment_root))
    set_cn_font(subtitle_run, size=10)

    add_heading(doc, "1. 实验背景与目标", 1)
    add_bullets(
        doc,
        [
            "实验根固定在 new3output 下，保持任务定义不变：仍然是 UAV 查询图像对固定卫星候选库进行初始定位。",
            "本次不重跑 DINOv2 coarse 和 RoMa v2 rerank，直接复用 new2output 已完成的 retrieval / rerank 结果。",
            "主变量只有两项：评估 truth orthophoto 从旧 UAV orthophoto 切换为新的 ODM orthophoto；PnP DSM 从 SRTM 切换为 ODM DSM / LAZ 栅格化结果。",
            "在同一套 best pose 结果上，额外增加一套 satellite-truth 并行验证，用于提供独立 truth 口径交叉检查。",
        ],
    )

    add_heading(doc, "2. 数据范围与 query 选择规则", 1)
    add_bullets(
        doc,
        [
            f"query 总数为 {query_summary['query_count']}，仅来自航线 {', '.join(short_flight_name(fid) for fid in query_summary['ordered_flights'])}。",
            "每条航线固定选择 20 张 query，总计 40 张，约束为 gimbal_pitch_degree <= -85.0。",
            f"本次 query 的 pitch 范围为 {fmt(query_summary['pitch_min'], 2)} 到 {fmt(query_summary['pitch_max'], 2)}。",
            "009 航线对应 q_001-q_020，010 航线对应 q_021-q_040。",
        ],
    )
    add_table(
        doc,
        ["flight", "query_count", "query_id_range"],
        [
            [short_flight_name(item["flight_id"]), str(item["count"]), f"{item['query_id_start']} - {item['query_id_end']}"]
            for item in query_summary["query_ranges"]
        ],
    )

    add_heading(doc, "3. 本次实验链路与和 new2output 的差异", 1)
    add_bullets(
        doc,
        [
            "runtime candidate DOM 仍然来自固定卫星库，不替换为 UAV DOM。",
            "retrieval / RoMa rerank 结果复用 new2output，不改变 runtime 候选集合。",
            "query intrinsics 沿用现有 per-flight cameras.json 解析结果，本次不作为变量。",
            "ODM truth 仅用于 eval_pose_validation_suite_odm_truth；satellite truth 仅用于 eval_pose_validation_suite_satellite_truth；两者都不进入 runtime candidate 选择。",
            "PnP DSM 已从 SRTM 路线切换为 ODM DSM override；当缺少 raster DSM 时，使用 odm_georeferenced_model.laz 栅格化构建 candidate DSM。",
        ],
    )

    add_heading(doc, "4. Pose 主链运行结果", 1)
    add_table(
        doc,
        ["item", "value"],
        [
            ["retrieval_top20 rows", str(len(retrieval_rows))],
            ["pose score rows", str(pose_summary["score_row_count"])],
            ["best pose query count", str(pose_summary["scored_query_count"])],
            ["PnP rows", str(pnp_summary["row_count"])],
            ["PnP status counts", json.dumps(pnp_summary["status_counts"], ensure_ascii=False)],
            ["best status counts", json.dumps(pose_summary["best_status_counts"], ensure_ascii=False)],
            ["best score mean", fmt(pose_summary["best_score_mean"], 4)],
            ["best reproj error mean", fmt(pose_summary["best_success_reproj_error_mean"], 4)],
            ["runtime elapsed sec", fmt(runtime_elapsed_sec, 1)],
        ],
    )
    add_table(doc, ["step_name", "elapsed_sec", "returncode"], build_runtime_rows(pose_runtime_steps))

    add_heading(doc, "5. ODM-truth 三层验证结果", 1)
    add_table(
        doc,
        ["metric", "value"],
        [
            ["layer1 phase_corr_error_m mean", fmt(odm_ortho_overall["phase_corr_error_m"]["mean"], 4)],
            ["layer1 center_offset_m mean", fmt(odm_ortho_overall["center_offset_m"]["mean"], 4)],
            ["layer1 ortho_iou mean", fmt(odm_ortho_overall["ortho_iou"]["mean"], 4)],
            ["layer1 ssim mean", fmt(odm_ortho_overall["ssim"]["mean"], 4)],
            ["layer2 horizontal_error_m mean", fmt(odm_pose_overall["horizontal_error_m"]["mean"], 4)],
            ["layer2 view_dir_angle_error_deg mean", fmt(odm_pose_overall["view_dir_angle_error_deg"]["mean"], 4)],
            ["layer3 tiepoint_xy_error_rmse_m", fmt(odm_tie_overall["tiepoint_xy_error_rmse_m"], 4)],
            ["layer3 tiepoint_xy_error_p90_m", fmt(odm_tie_overall["tiepoint_xy_error_p90_m"], 4)],
        ],
    )
    add_picture_with_caption(doc, odm_chart_path, "图 1. ODM-truth per-query 指标图：ortho_iou、common_valid_ratio、center_offset_m、ssim。")

    add_heading(doc, "6. Satellite-truth 三层验证结果", 1)
    geom_offset_values = [
        float(row["camera_center_offset_m"])
        for row in sat_geom_rows
        if row.get("camera_center_offset_m") not in ("", None) and row.get("eval_status") == "ok"
    ]
    add_table(
        doc,
        ["metric", "value"],
        [
            ["layer1 phase_corr_error_m mean", fmt(sat_ortho_overall["phase_corr_error_m"]["mean"], 4)],
            ["layer1 center_offset_m mean", fmt(sat_ortho_overall["center_offset_m"]["mean"], 4)],
            ["layer1 ortho_iou mean", fmt(sat_ortho_overall["ortho_iou"]["mean"], 4)],
            ["layer1 ssim mean", fmt(sat_ortho_overall["ssim"]["mean"], 4)],
            ["layer2 camera_center_offset_m mean", fmt(mean(geom_offset_values) if geom_offset_values else None, 4)],
            ["layer2 status counts", json.dumps(sat_geom_overall["status_counts"], ensure_ascii=False)],
            ["layer3 tiepoint_xy_error_rmse_m", fmt(sat_tie_overall["tiepoint_xy_error_rmse_m"], 4)],
            ["layer3 tiepoint_xy_error_p90_m", fmt(sat_tie_overall["tiepoint_xy_error_p90_m"], 4)],
        ],
    )
    add_picture_with_caption(doc, sat_chart_path, "图 2. Satellite-truth per-query 指标图：phase_corr_error_m、center_offset_m。")

    add_heading(doc, "7. 两套 truth 口径结果对比", 1)
    add_bullets(
        doc,
        [
            f"baseline(new2output) layer1 ortho_iou mean = {fmt(baseline_ortho_overall['ortho_iou']['mean'], 4)}，new3 ODM-truth = {fmt(odm_ortho_overall['ortho_iou']['mean'], 4)}，satellite-truth = {fmt(sat_ortho_overall['ortho_iou']['mean'], 4)}。",
            f"baseline(new2output) layer3 tiepoint RMSE = {fmt(baseline_tie_overall['tiepoint_xy_error_rmse_m'], 4)}，new3 ODM-truth = {fmt(odm_tie_overall['tiepoint_xy_error_rmse_m'], 4)}，satellite-truth = {fmt(sat_tie_overall['tiepoint_xy_error_rmse_m'], 4)}。",
            "ODM-truth 反映的是替换了新的 ODM orthophoto truth 和 ODM DSM 后，在 UAV truth 口径下的表现；satellite-truth 则是独立的卫星影像 truth 验证视角。",
        ],
    )
    add_picture_with_caption(doc, comparison_chart_path, "图 3. baseline、ODM-truth、satellite-truth 关键指标对比。")

    add_heading(doc, "8. 预测图部分缺失问题分析", 1)
    add_bullets(
        doc,
        [
            "预测图不是完整正射重建图，而是把单张 query 按 best pose + candidate-linked DSM 投影到 truth grid 后得到的有效覆盖区域。",
            "当前渲染器不允许在 DSM 采样失败时回退到平面模型，因此 DSM 无效、投影出界、以及 query 视场外的区域都会直接留空。",
            "因此 pred_tile 中的空洞不等于 PnP 失败，而更接近“该 truth 网格位置在当前 pose + DSM 条件下没有有效可见投影”。",
            f"q_001 与低覆盖样例的定量证据：{missing_summary_lines[0]}；{missing_summary_lines[1]}。",
        ],
    )
    add_picture_with_caption(doc, missing_chart_path, "图 4. ODM-truth 下 predicted ortho 覆盖率、ortho_iou 与 center_offset_m 的逐 query 分布。")
    for panel_path, query_id in sample_panels:
        row = odm_by_query[query_id]
        add_picture_with_caption(
            doc,
            panel_path,
            (
                f"图 5-{query_id}. {query_id} 的 truth / pred / alpha-mask / overlay 样例；"
                f"common_valid_ratio={fmt(row['common_valid_ratio'], 4)}, "
                f"ortho_iou={fmt(row['ortho_iou'], 4)}, "
                f"center_offset_m={fmt(row['center_offset_m'], 4)}。"
            ),
        )

    add_heading(doc, "9. 结论与后续建议", 1)
    add_bullets(
        doc,
        [
            "本次 new3output 实验在不改变 runtime 卫星检索任务定义的前提下，完成了 ODM truth / ODM DSM 替换和 satellite-truth 并行验证。",
            "现有预测图空洞是系统性覆盖约束，而不是单张文件损坏；解释时必须结合 common_valid_ratio、ortho_iou 和 alpha/mask 一起看。",
            "后续最应该补的是：把 valid coverage / alpha mask 纳入正式报告默认图表，并为 satellite-truth 补充更强的整体指标图。",
            "如果目标是进一步降低空洞视觉影响，需要单独讨论是否允许平面回退、扩大 truth 裁剪策略，或增加 coverage-aware 渲染说明，而不是把当前 pred_tile 当作完整正射图使用。",
        ],
    )

    ensure_dir(out_docx.parent)
    doc.save(out_docx)
    print(out_docx)


if __name__ == "__main__":
    main()
