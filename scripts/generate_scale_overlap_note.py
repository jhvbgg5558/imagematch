#!/usr/bin/env python3
"""Generate a brief Word/Markdown note about satellite tile scales and overlap."""

from __future__ import annotations

import argparse
import json
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--dataset-summary-json", required=True)
    parser.add_argument("--overview-figure", required=True)
    parser.add_argument("--scale-figure", required=True)
    parser.add_argument("--out-docx", required=True)
    parser.add_argument("--out-md", required=True)
    return parser.parse_args()


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def set_cn_font(run, size: int = 11, bold: bool = False) -> None:
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(size)
    run.bold = bold


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.style = "Heading 1" if level == 1 else "Heading 2"
    r = p.add_run(text)
    set_cn_font(r, size=14 if level == 1 else 12, bold=True)


def add_paragraph(doc: Document, text: str, size: int = 11, center: bool = False) -> None:
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn_font(r, size=size)


def add_bullets(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(line)
        set_cn_font(r, size=11)


def add_picture(doc: Document, path: Path, width_inch: float) -> None:
    doc.add_picture(str(path), width=Inches(width_inch))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn_font(r, size=10)


def build_markdown(summary: dict, overview_figure: Path, scale_figure: Path) -> str:
    roi_area_km2 = summary["roi_area_m2"] / 1_000_000.0
    raw_area_km2 = summary["raw_bbox_area_m2"] / 1_000_000.0
    buffer_gain_km2 = summary["buffer_area_gain_m2"] / 1_000_000.0
    lines = [
        "# 卫片尺度与 Overlap 说明",
        "",
        "## 1. 当前处理顺序",
        "",
        "当前流程的先后顺序是：",
        "",
        "1. 用 4 条航线原始 GPS 点求整体活动区域的包围框。",
        f"2. 在包围框外固定外扩 `{summary['roi_buffer_meters']:.0f}m`，得到正式 ROI。",
        "3. 在 ROI 内按 `200 / 300 / 500 / 700m` 切出多尺度卫片 tile，且相邻 tile 采用 `overlap = 0.25`。",
        "4. 对切好的 tile 图像逐张提取 DINOv2 特征。",
        "5. 用这些 tile 特征建立 FAISS 索引，供 query 检索。",
        "",
        "结论是：尺度划分发生在 DINOv2 建库之前，当前特征库是 tile 级特征库，而不是整幅卫片特征库。",
        "",
        "## 2. 为什么选 200 / 300 / 500 / 700m",
        "",
        "- `200m` 保留较细粒度局部区域，适合覆盖更紧凑的真实范围。",
        "- `300m` 是从较小尺度过渡到中尺度的补充档位；历史结果里加入 `300m` 后曾带来新增 Top-1 命中，因此被保留。",
        "- `500m` 与 `700m` 的加入，是因为当前 query 已不再是同尺度正射小裁块，而是地面覆盖范围更大、视角更复杂的无人机单图；更大尺度有助于把正确区域送入 coarse Top-K。",
        "- 这组尺度不是理论最优解，而是从旧版 `80/120/200/300m` 演进到当前任务后的工程折中选择，目标是提升区域级召回和候选覆盖上限。",
        "",
        "## 3. 为什么 overlap = 0.25",
        "",
        "- overlap 的定义是相邻 tile 在地面范围上保留 25% 的重叠。",
        "- 对应步长 `stride = tile_size * (1 - overlap)`，因此当前步长是 `0.75 * tile_size`。",
        "- 这样做可以减少 query 刚好落在 tile 边界附近时被切断的情况，提高连续覆盖。",
        "- 同时，相邻候选之间会对同一真实区域保留一定冗余，有利于后续检索把真值送入 Top-K。",
        "- 但 overlap 不是越大越好；重叠增大虽然会提高覆盖连续性，也会让候选库规模和真值数量继续上升。",
        "",
        "## 4. 对召回和真值数量的影响",
        "",
        "- 更大尺度通常更有利于区域级召回，因为 query 覆盖范围较大时，较大的 tile 更容易与其发生有效交集。",
        "- overlap 提高后，边界附近的漏检风险会下降，因此 coarse Top-K 的覆盖能力通常会更稳定。",
        "- 代价是候选库规模变大，当前四尺度总 tile 数为 "
        f"`{summary['tile_count_total']}`，其中 `200m={summary['tile_count_by_scale']['200m']}`、"
        f"`300m={summary['tile_count_by_scale']['300m']}`、`500m={summary['tile_count_by_scale']['500m']}`、"
        f"`700m={summary['tile_count_by_scale']['700m']}`。",
        "- 在你当前改成“非零面积交集即为真值”之后，尺度越大、overlap 越高，单个 query 对应的真值 tile 数一般也会更多。",
        "- 因此，新口径下更容易出现“Top-K 里已经是 query 覆盖范围，但不是旧 strict truth”的情况，这也是你这次重定义真值的直接背景。",
        "",
        "## 5. 当前数据范围摘要",
        "",
        f"- Query 数：`{summary['query_count']}`",
        f"- ROI 面积：`{roi_area_km2:.3f} km^2`",
        f"- 原始 bbox 面积：`{raw_area_km2:.3f} km^2`",
        f"- 外扩增加面积：`{buffer_gain_km2:.3f} km^2`",
        f"- Query/卫片坐标系：`{summary['query_crs']}`",
        "",
        "## 图示",
        "",
        f"![query_satellite_overview]({overview_figure.name})",
        "",
        f"![scale_tile_count_bar]({scale_figure.name})",
        "",
        "## 结论",
        "",
        "当前 `200 / 300 / 500 / 700m + overlap=0.25` 的设置，本质上是在当前无人机单图 query 条件下，为区域级检索提供更稳的候选覆盖。它更偏向提升 coarse recall 和真值覆盖机会，而不是直接保证 Top-1 排序最优。",
        "",
    ]
    return "\n".join(lines)


def build_docx(
    summary: dict,
    overview_figure: Path,
    scale_figure: Path,
    out_docx: Path,
) -> None:
    roi_area_km2 = summary["roi_area_m2"] / 1_000_000.0
    raw_area_km2 = summary["raw_bbox_area_m2"] / 1_000_000.0
    buffer_gain_km2 = summary["buffer_area_gain_m2"] / 1_000_000.0

    doc = Document()
    add_paragraph(doc, "卫片尺度与 Overlap 说明", size=16, center=True)
    add_paragraph(doc, "日期：2026-03-23", size=10, center=True)

    add_heading(doc, "1. 当前处理顺序", 1)
    add_bullets(
        doc,
        [
            "先用 4 条航线原始 GPS 点求整体活动区域包围框，并固定外扩 250m。",
            "再在 ROI 内按 200 / 300 / 500 / 700m 切出多尺度卫片 tile，切片 overlap 为 0.25。",
            "随后对切好的 tile 图像逐张提取 DINOv2 特征。",
            "最后用这些 tile 特征建立 FAISS 索引，供 query 检索。",
        ],
    )
    add_paragraph(doc, "当前顺序的关键点是：尺度划分发生在 DINOv2 建库之前，当前建的是 tile 级特征库，而不是整幅卫片特征库。")

    add_heading(doc, "2. 为什么选 200 / 300 / 500 / 700m", 1)
    add_bullets(
        doc,
        [
            "200m 保留较细粒度局部区域，适合更紧凑的候选窗口。",
            "300m 是从小尺度向中尺度过渡的补充档位，历史结果里加入 300m 后曾出现新增 Top-1 命中，因此继续保留。",
            "500m 与 700m 是针对当前 query 不再是标准同尺度正射小裁块，而是范围更大、视角更复杂的无人机单图而加入的更大地面覆盖窗口。",
            "这组尺度不是理论最优结论，而是从旧版 80/120/200/300m 演进到当前任务后的工程折中，目标是提升 coarse candidate coverage。",
        ],
    )

    add_heading(doc, "3. 为什么 overlap = 0.25", 1)
    add_bullets(
        doc,
        [
            "overlap = 0.25 表示相邻 tile 在地面范围上保留 25% 重叠。",
            "对应步长 stride = tile_size * (1 - overlap)，因此当前 stride = 0.75 * tile_size。",
            "这样做可以降低 query 刚好落在 tile 边界附近时被切断的风险，提高相邻候选对同一区域的连续覆盖。",
            "但 overlap 增大也会让候选库规模和真值数量继续上升，因此 0.25 应理解为覆盖连续性与库规模之间的保守折中。",
        ],
    )

    add_heading(doc, "4. 对召回和真值数量的影响", 1)
    add_bullets(
        doc,
        [
            "更大尺度通常更有利于区域级召回，因为 query 覆盖范围较大时，较大的 tile 更容易与其发生有效交集。",
            "更高 overlap 有助于减少边界漏检，因此 coarse Top-K 的覆盖能力通常会更稳定。",
            "当前四尺度总 tile 数为 "
            f"{summary['tile_count_total']}，其中 200m={summary['tile_count_by_scale']['200m']}、"
            f"300m={summary['tile_count_by_scale']['300m']}、500m={summary['tile_count_by_scale']['500m']}、"
            f"700m={summary['tile_count_by_scale']['700m']}。",
            "在新的“非零面积交集即真值”口径下，尺度越大、overlap 越高，单个 query 对应的真值 tile 数通常也会更多。",
            "因此它们更偏向提升候选覆盖和召回机会，而不是直接保证 Top-1 排序更准。",
        ],
    )

    add_heading(doc, "5. 当前数据范围摘要", 1)
    add_bullets(
        doc,
        [
            f"Query 数：{summary['query_count']}",
            f"Query/卫片坐标系：{summary['query_crs']}",
            f"原始 bbox 面积：{raw_area_km2:.3f} km^2",
            f"外扩后 ROI 面积：{roi_area_km2:.3f} km^2",
            f"外扩增加面积：{buffer_gain_km2:.3f} km^2",
        ],
    )

    add_heading(doc, "6. 图示", 1)
    add_picture(doc, overview_figure, 6.2)
    add_caption(doc, "图 1  Query 活动范围、ROI 与卫片候选库空间关系示意")
    add_picture(doc, scale_figure, 5.8)
    add_caption(doc, "图 2  四个尺度对应的 tile 数量统计")

    add_heading(doc, "7. 结论", 1)
    add_paragraph(
        doc,
        "当前 200 / 300 / 500 / 700m 与 overlap=0.25 的设置，本质上是在当前无人机单图 query 条件下，为区域级检索提供更稳的候选覆盖。它主要服务于 coarse recall 和真值覆盖机会的提升，而不是直接保证首位候选最优。",
    )

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_docx)


def main() -> None:
    args = parse_args()
    summary = load_json(Path(args.dataset_summary_json))
    overview_figure = Path(args.overview_figure)
    scale_figure = Path(args.scale_figure)
    out_docx = Path(args.out_docx)
    out_md = Path(args.out_md)
    out_dir = out_md.parent

    out_dir.mkdir(parents=True, exist_ok=True)
    local_overview = out_dir / overview_figure.name
    local_scale = out_dir / scale_figure.name
    shutil.copy2(overview_figure, local_overview)
    shutil.copy2(scale_figure, local_scale)

    markdown = build_markdown(summary, local_overview, local_scale)
    out_md.write_text(markdown, encoding="utf-8")
    build_docx(summary, local_overview, local_scale, out_docx)


if __name__ == "__main__":
    main()
