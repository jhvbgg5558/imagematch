#!/usr/bin/env python3
"""Generate a formal Word report for the coverage-truth baseline experiment."""

from __future__ import annotations

import csv
import json
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path("/mnt/d/aiproject/imagematch")
RESULT_DIR = ROOT / "output" / "coverage_truth_200_300_500_700_dinov2_baseline"
FIG_DIR = RESULT_DIR / "figures"
OUT_PATH = RESULT_DIR / "DINOv2_coverage_truth_200_300_500_700_实验结果说明_2026-03-19.docx"


def load_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", newline="", encoding="utf-8-sig") as f:
        return list(csv.DictReader(f))


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def set_cn_font(run, size: int = 11, bold: bool = False) -> None:
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(size)
    run.bold = bold


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 10, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    set_cn_font(r, size=size, bold=bold)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.style = "Heading 1" if level == 1 else "Heading 2"
    r = p.add_run(text)
    set_cn_font(r, size=14 if level == 1 else 12, bold=True)


def add_paragraph(doc: Document, text: str, size: int = 11, center: bool = False) -> None:
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn_font(r, size=size)


def add_bullets(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(line)
        set_cn_font(r, size=11)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn_font(r, size=10)


def add_picture(doc: Document, path: Path, width_inch: float) -> None:
    doc.add_picture(str(path), width=Inches(width_inch))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def short_flight_name(flight_id: str) -> str:
    parts = flight_id.split("_")
    return parts[2] if len(parts) >= 3 else flight_id


def per_flight_metrics(summary: dict, seed_rows: dict[str, dict[str, str]]) -> list[dict[str, float | str]]:
    grouped: dict[str, list[dict]] = defaultdict(list)
    for row in summary["per_query"]:
        grouped[seed_rows[row["query_id"]]["flight_id"]].append(row)

    out = []
    for flight_id in sorted(grouped):
        rows = grouped[flight_id]
        total = len(rows)
        errors = [float(row["top1_error_m"]) for row in rows]
        out.append(
            {
                "flight_id": flight_id,
                "query_count": total,
                "coverage_recall@1": sum(int(row["coverage_hit@1"]) for row in rows) / total,
                "coverage_recall@5": sum(int(row["coverage_hit@5"]) for row in rows) / total,
                "coverage_recall@10": sum(int(row["coverage_hit@10"]) for row in rows) / total,
                "coverage_mrr": sum(float(row["coverage_reciprocal_rank"]) for row in rows) / total,
                "center_recall@1": sum(int(row["center_hit@1"]) for row in rows) / total,
                "center_recall@5": sum(int(row["center_hit@5"]) for row in rows) / total,
                "center_recall@10": sum(int(row["center_hit@10"]) for row in rows) / total,
                "top1_error_m_mean": sum(errors) / len(errors),
            }
        )
    return out


def truth_scale_stats(rows: list[dict[str, str]]) -> dict[str, float]:
    total = len(rows)
    return {
        "query_count": total,
        "coverage_threshold": float(rows[0]["coverage_threshold"]),
        "footprint_core_ratio": float(rows[0]["footprint_core_ratio"]),
        "truth_count_total_mean": sum(float(r["truth_count_total"]) for r in rows) / total,
        "truth_count_200m_mean": sum(float(r["truth_count_200m"]) for r in rows) / total,
        "truth_count_300m_mean": sum(float(r["truth_count_300m"]) for r in rows) / total,
        "truth_count_500m_mean": sum(float(r["truth_count_500m"]) for r in rows) / total,
        "truth_count_700m_mean": sum(float(r["truth_count_700m"]) for r in rows) / total,
        "footprint_area_m2_mean": sum(float(r["footprint_area_m2"]) for r in rows) / total,
    }


def build_metric_definition_table(doc: Document) -> None:
    rows = [
        ("Coverage Recall@1", "首位候选命中 coverage 真值的比例，衡量是否能直接给出正确区域候选。"),
        ("Coverage Recall@5", "前 5 名中命中 coverage 真值的比例，衡量前排候选覆盖能力。"),
        ("Coverage Recall@10", "前 10 名中命中 coverage 真值的比例，衡量区域级初步定位能力。"),
        ("Coverage MRR", "coverage 真值排名倒数的平均值，越高说明真值整体越靠前。"),
        ("Top-1 error mean (m)", "首位候选中心与 query 参考位置之间的平均距离，越低越好。"),
        ("Center Recall 系列", "辅助中心点口径指标，仅用于补充观察，不作为本轮主结论口径。"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["指标", "含义"]
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, bold=True)
        shade_cell(table.cell(0, i), "D9EAF7")
    for name, desc in rows:
        row = table.add_row().cells
        set_cell_text(row[0], name)
        row[1].text = ""
        p = row[1].paragraphs[0]
        r = p.add_run(desc)
        set_cn_font(r, size=10)


def build_overall_table(doc: Document, summary: dict) -> None:
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["口径", "R@1", "R@5", "R@10", "MRR", "Top-1误差均值(m)"]
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, bold=True)
        shade_cell(table.cell(0, i), "D9EAF7")

    coverage = table.add_row().cells
    vals = [
        "Coverage",
        f"{float(summary['coverage_recall@1']):.3f}",
        f"{float(summary['coverage_recall@5']):.3f}",
        f"{float(summary['coverage_recall@10']):.3f}",
        f"{float(summary['coverage_mrr']):.3f}",
        f"{float(summary['top1_error_m_mean']):.3f}",
    ]
    for i, val in enumerate(vals):
        set_cell_text(coverage[i], val)

    center = table.add_row().cells
    vals = [
        "Center",
        f"{float(summary['center_recall@1']):.3f}",
        f"{float(summary['center_recall@5']):.3f}",
        f"{float(summary['center_recall@10']):.3f}",
        "-",
        "-",
    ]
    for i, val in enumerate(vals):
        set_cell_text(center[i], val)


def build_per_flight_table(doc: Document, rows: list[dict[str, float | str]]) -> None:
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["航线", "Coverage R@1", "Coverage R@5", "Coverage R@10", "Coverage MRR", "Top-1误差均值(m)"]
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, bold=True)
        shade_cell(table.cell(0, i), "D9EAF7")
    for item in rows:
        row = table.add_row().cells
        set_cell_text(row[0], short_flight_name(str(item["flight_id"])))
        set_cell_text(row[1], f"{float(item['coverage_recall@1']):.3f}")
        set_cell_text(row[2], f"{float(item['coverage_recall@5']):.3f}")
        set_cell_text(row[3], f"{float(item['coverage_recall@10']):.3f}")
        set_cell_text(row[4], f"{float(item['coverage_mrr']):.3f}")
        set_cell_text(row[5], f"{float(item['top1_error_m_mean']):.3f}")


def build_truth_summary_table(doc: Document, stats: dict[str, float], roi_summary: dict) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["字段", "值"]
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, bold=True)
        shade_cell(table.cell(0, i), "D9EAF7")

    bounds = roi_summary["roi_bounds"]
    rows = [
        ("尺度集合", "200m / 300m / 500m / 700m"),
        ("Coverage 阈值", f"{stats['coverage_threshold']:.2f}"),
        ("Footprint core ratio", f"{stats['footprint_core_ratio']:.2f}"),
        ("Query 数量", f"{int(stats['query_count'])}"),
        ("卫星 tiles 数量", f"{int(roi_summary['point_count'])}"),
        ("ROI 范围", f"[{bounds[0]:.3f}, {bounds[1]:.3f}] - [{bounds[2]:.3f}, {bounds[3]:.3f}]"),
        ("平均真值数", f"{stats['truth_count_total_mean']:.2f}"),
        ("平均 200m 真值数", f"{stats['truth_count_200m_mean']:.2f}"),
        ("平均 300m 真值数", f"{stats['truth_count_300m_mean']:.2f}"),
        ("平均 500m 真值数", f"{stats['truth_count_500m_mean']:.2f}"),
        ("平均 700m 真值数", f"{stats['truth_count_700m_mean']:.2f}"),
        ("平均 footprint 面积(m2)", f"{stats['footprint_area_m2_mean']:.3f}"),
    ]
    for key, value in rows:
        row = table.add_row().cells
        set_cell_text(row[0], key)
        set_cell_text(row[1], value)


def choose_cases(summary: dict, seed_rows: dict[str, dict[str, str]]) -> list[dict]:
    rows = summary["per_query"]
    success = [r for r in rows if r["first_coverage_truth_rank"] == 1]
    success.sort(key=lambda r: (float(r["top1_error_m"]), r["query_id"]))

    boundary = [
        r
        for r in rows
        if (r["first_coverage_truth_rank"] in {2, 3, 4, 5}) or (r["coverage_hit@10"] and not r["coverage_hit@5"])
    ]
    boundary.sort(
        key=lambda r: (
            r["first_coverage_truth_rank"] if r["first_coverage_truth_rank"] is not None else 99,
            float(r["top1_error_m"]),
            r["query_id"],
        )
    )

    failure = [r for r in rows if not r["coverage_hit@10"]]
    failure.sort(key=lambda r: (-float(r["top1_error_m"]), r["query_id"]))

    selected: list[dict] = []
    used: set[str] = set()
    pools = [
        ("显著成功样例", success, "该 query 的首位候选已经直接命中 coverage 真值，说明在当前口径下该样本具备较强的跨视角区域辨识度。"),
        ("边界成功样例", boundary, "该 query 虽未在首位命中，但真值已进入前列候选，说明当前基线能够支持区域级初步定位，但前排排序仍有继续优化空间。"),
        ("残余失败样例", failure, "该 query 在 Top-10 内仍未命中 coverage 真值，说明当前尺度与表征组合在该类场景下仍存在稳定失败模式。"),
    ]
    for title, pool, note in pools:
        for row in pool:
            if row["query_id"] in used:
                continue
            case = dict(row)
            case["title"] = title
            case["note"] = note
            case["flight_id"] = seed_rows[row["query_id"]]["flight_id"]
            selected.append(case)
            used.add(row["query_id"])
            break
    return selected


def add_case_section(doc: Document, case_idx: int, case: dict) -> None:
    query_id = case["query_id"]
    flight_id = case["flight_id"]
    rank = case["first_coverage_truth_rank"]
    rank_text = "未命中 Top-10" if rank is None else str(rank)
    add_heading(doc, f"7.{case_idx} {case['title']}：{short_flight_name(flight_id)} / {query_id}", 2)
    add_paragraph(
        doc,
        f"首个 coverage 真值排名：{rank_text}；Coverage Truth Count：{case['coverage_truth_count']}；"
        f"Center Truth Count：{case['center_truth_count']}；Top-1 error：{float(case['top1_error_m']):.3f}m。{case['note']}",
    )
    img = FIG_DIR / flight_id / f"{query_id}_top10.png"
    add_picture(doc, img, 6.0)
    add_caption(doc, f"{case['title']}：{short_flight_name(flight_id)} / {query_id} 的 Top-10 检索结果")


def main() -> None:
    summary = load_json(RESULT_DIR / "retrieval" / "summary.json")
    roi_summary = load_json(RESULT_DIR / "fixed_satellite_library" / "roi_summary.json")
    truth_rows = load_csv(RESULT_DIR / "query_truth" / "query_truth.csv")
    seed_rows = {row["query_id"]: row for row in load_csv(RESULT_DIR / "query_truth" / "queries_truth_seed.csv")}
    flight_rows = per_flight_metrics(summary, seed_rows)
    truth_stats = truth_scale_stats(truth_rows)
    cases = choose_cases(summary, seed_rows)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("DINOv2 + FAISS 在 Coverage Truth 口径下的多尺度跨视角粗定位实验结果说明")
    set_cn_font(r, size=16, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("面向无人机单张图像初步地理定位的正式技术报告")
    set_cn_font(r, size=11)

    add_heading(doc, "1. 任务定义与实验设置", 1)
    add_paragraph(
        doc,
        "本组实验用于论证：在更贴近工程实际的输入条件下，仅依赖遥感正射影像，是否能够对任意单张无人机图像实现区域级初步地理定位。"
        "与此前基于中心点式真值和 `80/120/200/300m` 候选尺度的首轮结果不同，本轮实验进一步把卫星候选尺度调整为 `200/300/500/700m`，"
        "并采用 coverage 真值口径，以减少旧口径对中心点的过度依赖。",
    )
    add_bullets(
        doc,
        [
            f"数据范围：4 条航线，共 {summary['query_count']} 个 query。",
            f"卫星候选库：{roi_summary['point_count']} 个 tiles，坐标系为 {roi_summary['satellite_crs']}。",
            "查询图像：去元数据的原始无人机单张图像，不保证为正射视角。",
            "基线方法：DINOv2 pooler 全局特征 + FAISS IndexFlatIP 粗检索。",
        ],
    )

    add_heading(doc, "2. 真值定义与选取范围", 1)
    add_paragraph(
        doc,
        "本轮实验的关键变化在于真值定义不再采用“query 中心点落入某块卫星瓦片”的单点口径，"
        "而是先根据无人机图像元数据近似估计其地面覆盖 footprint，再计算该 footprint 与固定卫星库中各候选瓦片的地面覆盖框相交比例。"
        "当相交比例大于 `0.4` 时，该瓦片被记为 coverage 真值。为避免把边缘尾部覆盖误判为真值，footprint 计算时同时保留 `0.60` 的 core ratio 设置。"
    )
    add_paragraph(
        doc,
        "在该口径下，中心点式真值仍被保留为辅助统计项，用于和旧结果建立有限对应，但本轮正式结论统一以 coverage 指标为主。"
        "表 1 汇总了本轮真值构造与候选范围的关键参数。",
    )
    build_truth_summary_table(doc, truth_stats, roi_summary)
    add_caption(doc, "表 1  本轮 coverage 真值构造与固定卫星库范围概览")

    add_heading(doc, "3. 指标定义", 1)
    build_metric_definition_table(doc)

    add_heading(doc, "4. 方法说明", 1)
    add_heading(doc, "4.1 DINOv2 + FAISS 粗检索基线", 2)
    add_paragraph(
        doc,
        "本轮方法仍采用 DINOv2 `pooler` 输出作为统一全局特征，对 query 图像和固定卫星库瓦片分别提取表示，"
        "再经 L2 归一化后使用 FAISS `IndexFlatIP` 执行 Top-10 相似度检索。该方法不包含局部几何验证和重排，"
        "因此本轮结果可以被视为“新尺度口径 + 新真值口径”下的最小可运行正式基线。",
    )
    add_heading(doc, "4.2 与上一轮实验相比的变化", 2)
    add_paragraph(
        doc,
        "相较上一轮 `80/120/200/300m` 的首轮正式结果，本轮并未更换骨干模型或检索器，"
        "真正改变的是候选尺度集合与真值定义方式。因此，本轮结果更适合用于回答：当尺度扩展到 `500/700m`，且真值由中心点判定切换为 coverage 判定后，"
        "DINOv2 粗检索在区域级定位任务上的表现会发生怎样的变化。",
    )

    add_heading(doc, "5. 定量结果", 1)
    add_paragraph(doc, "表 2 给出当前实验的 overall 结果，其中 coverage 为正式主口径，center 为辅助口径。")
    build_overall_table(doc, summary)
    add_caption(doc, "表 2  当前实验的 overall 指标结果")
    add_paragraph(doc, "表 3 给出按航线拆分后的 coverage 主口径结果，用于观察不同航线上的稳定性差异。")
    build_per_flight_table(doc, flight_rows)
    add_caption(doc, "表 3  当前实验的分航线 coverage 指标结果")

    add_heading(doc, "6. 汇总图解读", 1)
    add_paragraph(
        doc,
        "为避免仅依据表格数字做静态判断，本文进一步给出本轮实验的汇总图。图 1 展示 overall 主指标，"
        "图 2 展示分航线 Recall，图 3 展示辅助中心点口径指标，图 4 展示 40 个 query 的 Top-1 误差分布。"
    )
    figures = [
        ("overall_metrics_bar.png", "图 1  当前实验的 overall coverage 指标图"),
        ("multi_flight_recall.png", "图 2  当前实验的分航线 Recall 图"),
        ("center_metrics_bar.png", "图 3  当前实验的辅助 center 指标图"),
        ("top1_error_distribution.png", "图 4  当前实验的 Top-1 误差分布图"),
    ]
    for fname, caption in figures:
        add_picture(doc, FIG_DIR / "_aggregate" / fname, 6.4)
        add_caption(doc, caption)

    add_paragraph(
        doc,
        "从 overall 结果可以看到，在当前 coverage 真值口径下，DINOv2 + FAISS 已经能够把接近一半的 query 在 Top-10 内检索到正确覆盖区域，"
        "`coverage Recall@10 = 0.475` 说明区域级初步定位能力已经形成；但 `coverage Recall@1 = 0.200` 仍然偏低，"
        "表明粗检索首位候选质量尚不足以支持直接落地。分航线结果进一步说明不同航线之间仍存在明显差异，当前性能并不均衡。"
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "7. 代表性案例", 1)
    add_paragraph(
        doc,
        "为使结论不局限于整体均值，本文自动选取显著成功、边界成功和残余失败三类代表性案例，"
        "并展示对应 query 的 Top-10 检索结果，以观察当前基线在候选排序层面的具体行为。"
    )
    for idx, case in enumerate(cases, start=1):
        add_case_section(doc, idx, case)

    add_heading(doc, "8. 结论", 1)
    add_paragraph(
        doc,
        "图 5 和图 6 进一步给出本轮实验最关键的两张结论图，分别对应 overall 指标和分航线 Recall 表现。"
    )
    add_picture(doc, FIG_DIR / "_aggregate" / "overall_metrics_bar.png", 6.2)
    add_caption(doc, "图 5  结论图：overall coverage 指标")
    add_picture(doc, FIG_DIR / "_aggregate" / "multi_flight_recall.png", 6.2)
    add_caption(doc, "图 6  结论图：分航线 Recall 表现")
    add_bullets(
        doc,
        [
            "在 `200/300/500/700m` 候选尺度和 coverage 真值口径下，DINOv2 + FAISS 已经能够支撑区域级初步地理定位，主命题在“区域召回”层面成立。",
            "更准确的结论表述是：当前基线已经具备把 query 检索到正确地理区域附近的能力，但首位候选质量仍不足，尚不能把粗检索结果等同于稳定可用的最终定位结果。",
            "coverage 真值定义比旧的中心点真值更贴近实际地面覆盖关系，也更适合评估大尺度候选在区域级检索中的真实贡献。",
            "当前仍存在稳定残余失败样本，说明后续若要进一步提升前排排序质量，仍有必要继续分析 `500/700m` 的收益边界，并考虑加入局部几何验证或重排步骤。",
        ],
    )

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(OUT_PATH)


if __name__ == "__main__":
    main()
