#!/usr/bin/env python3
"""Generate a formal Chinese report for an intersection-truth baseline run."""

from __future__ import annotations

import argparse
import json
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--result-dir", required=True)
    parser.add_argument("--out-docx", required=True)
    parser.add_argument("--out-md", required=True)
    parser.add_argument("--simplified", action="store_true", help="Generate a simplified report using only currently available outputs.")
    parser.add_argument("--model-label", default="DINOv2", help="Model label used in titles and method descriptions.")
    return parser.parse_args()


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def set_cn_font(run, size: int = 11, bold: bool = False) -> None:
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(size)
    run.bold = bold


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 10, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    set_cn_font(r, size=size, bold=bold)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.style = "Heading 1" if level == 1 else "Heading 2"
    r = p.add_run(text)
    set_cn_font(r, size=14 if level == 1 else 12, bold=True)


def add_paragraph(doc: Document, text: str, size: int = 11, center: bool = False) -> None:
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn_font(r, size=size)


def add_bullets(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(line)
        set_cn_font(r, size=11)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn_font(r, size=10)


def add_picture(doc: Document, path: Path, width_inch: float) -> None:
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(width_inch))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def short_flight_name(flight_id: str) -> str:
    parts = flight_id.split("_")
    return parts[2] if len(parts) >= 3 else flight_id


def metric_bool(value) -> int:
    if isinstance(value, bool):
        return int(value)
    if isinstance(value, (int, float)):
        return int(value)
    if isinstance(value, str):
        return 1 if value.lower() in {"1", "true", "yes"} else 0
    return 0


def per_flight_metrics(summary: dict) -> list[dict[str, float | str | int]]:
    grouped: dict[str, list[dict]] = defaultdict(list)
    for row in summary["per_query"]:
        grouped[row["flight_id"]].append(row)

    out: list[dict[str, float | str | int]] = []
    for flight_id in sorted(grouped):
        rows = grouped[flight_id]
        total = len(rows)
        errors = [float(row["top1_error_m"]) for row in rows]
        out.append(
            {
                "flight_id": flight_id,
                "query_count": total,
                "intersection_recall@1": sum(metric_bool(row["intersection_hit@1"]) for row in rows) / total,
                "intersection_recall@5": sum(metric_bool(row["intersection_hit@5"]) for row in rows) / total,
                "intersection_recall@10": sum(metric_bool(row["intersection_hit@10"]) for row in rows) / total,
                "intersection_recall@20": sum(metric_bool(row["intersection_hit@20"]) for row in rows) / total,
                "intersection_mrr": sum(float(row["intersection_reciprocal_rank"]) for row in rows) / total,
                "top1_error_m_mean": sum(errors) / len(errors),
            }
        )
    return out


def stage_seconds(timing: dict) -> dict[str, float]:
    return {item["stage"]: float(item["elapsed_seconds"]) for item in timing["stages"]}


def format_seconds(seconds: float) -> str:
    return f"{seconds:.2f}s ({seconds / 60.0:.2f} min)"


def metric_or_zero(value: object) -> float:
    return float(value) if value is not None else 0.0


def build_metric_definition_table(doc: Document) -> None:
    rows = [
        ("Intersection Recall@1", "首位候选命中 intersection truth 的比例，是本次 run 的首位粗定位指标。"),
        ("Intersection Recall@5", "前 5 名候选中命中 intersection truth 的比例，衡量前排候选覆盖能力。"),
        ("Intersection Recall@10", "前 10 名候选中命中 intersection truth 的比例，衡量区域级初步地理定位能力。"),
        ("Intersection Recall@20", "前 20 名候选中命中 intersection truth 的比例，用于观察候选窗口扩展上限。"),
        ("Intersection MRR", "首个 intersection truth 的排名倒数平均值，越高表示正确区域越靠前。"),
        ("Top-1 error mean (m)", "首位候选中心与 query 参考位置之间的平均距离，用于衡量首位候选的空间偏差。"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, title in enumerate(["指标", "含义"]):
        set_cell_text(table.cell(0, i), title, bold=True)
        shade_cell(table.cell(0, i), "D9EAF7")
    for name, desc in rows:
        row = table.add_row().cells
        set_cell_text(row[0], name)
        row[1].text = ""
        p = row[1].paragraphs[0]
        r = p.add_run(desc)
        set_cn_font(r, size=10)


def build_overall_table(doc: Document, summary_top20: dict, summary_all: dict) -> None:
    table = doc.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["口径", "R@1", "R@5", "R@10", "R@20", "MRR", "Top-1误差均值(m)"]
    for i, title in enumerate(headers):
        set_cell_text(table.cell(0, i), title, bold=True)
        shade_cell(table.cell(0, i), "D9EAF7")

    rows = [
        (
            "Top-20",
            f"{float(summary_top20['intersection_recall@1']):.3f}",
            f"{float(summary_top20['intersection_recall@5']):.3f}",
            f"{float(summary_top20['intersection_recall@10']):.3f}",
            f"{float(summary_top20['intersection_recall@20']):.3f}",
            f"{float(summary_top20['intersection_mrr']):.3f}",
            f"{float(summary_top20['top1_error_m_mean']):.3f}",
        ),
        (
            "全库排序",
            f"{float(summary_all['intersection_recall@1']):.3f}",
            f"{float(summary_all['intersection_recall@5']):.3f}",
            f"{float(summary_all['intersection_recall@10']):.3f}",
            f"{float(summary_all['intersection_recall@20']):.3f}",
            f"{float(summary_all['intersection_mrr']):.3f}",
            f"{float(summary_all['top1_error_m_mean']):.3f}",
        ),
    ]
    for values in rows:
        row = table.add_row().cells
        for i, value in enumerate(values):
            set_cell_text(row[i], value)


def build_per_flight_table(doc: Document, rows: list[dict[str, float | str | int]]) -> None:
    table = doc.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["航线", "R@1", "R@5", "R@10", "R@20", "MRR", "Top-1误差均值(m)"]
    for i, title in enumerate(headers):
        set_cell_text(table.cell(0, i), title, bold=True)
        shade_cell(table.cell(0, i), "D9EAF7")
    for item in rows:
        row = table.add_row().cells
        set_cell_text(row[0], short_flight_name(str(item["flight_id"])))
        set_cell_text(row[1], f"{float(item['intersection_recall@1']):.3f}")
        set_cell_text(row[2], f"{float(item['intersection_recall@5']):.3f}")
        set_cell_text(row[3], f"{float(item['intersection_recall@10']):.3f}")
        set_cell_text(row[4], f"{float(item['intersection_recall@20']):.3f}")
        set_cell_text(row[5], f"{float(item['intersection_mrr']):.3f}")
        set_cell_text(row[6], f"{float(item['top1_error_m_mean']):.3f}")


def build_timing_table(doc: Document, timing_map: dict[str, float]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, title in enumerate(["阶段", "耗时"]):
        set_cell_text(table.cell(0, i), title, bold=True)
        shade_cell(table.cell(0, i), "D9EAF7")
    rows = [
        ("卫片特征提取", format_seconds(timing_map.get("satellite_feature_extraction", 0.0))),
        ("FAISS 建库", format_seconds(timing_map.get("faiss_index_build", 0.0))),
        ("Query 特征提取", format_seconds(timing_map.get("query_feature_extraction", 0.0))),
        ("检索评估（Top-20）", format_seconds(timing_map.get("query_retrieval", 0.0))),
    ]
    for key, value in rows:
        row = table.add_row().cells
        set_cell_text(row[0], key)
        set_cell_text(row[1], value)


def topk_summary_lines(k_full: dict, k_unique: dict | None = None) -> list[str]:
    lines = [
        f"full-truth 曲线结果：`40/40` 个 query 都能达到真值饱和，整体 `mean={k_full['overall']['mean']:.3f}`、`median={k_full['overall']['median']}`、`p95={k_full['overall']['p95']}`。",
    ]
    if k_unique is not None:
        lines.append(
            f"unique-tile 曲线结果：全部唯一真值 tile 数为 `{k_unique['overall']['total_truth_unique_tiles']}`，而 `k_full_truth={k_unique['overall']['k_full_truth']}`，等于候选总量 `{k_unique['overall']['candidate_unique_tiles']}`。"
        )
        lines.append("这说明当前基线在前排候选上的命中能力已经较强，但若目标是把所有真值 tile 全部覆盖到，则仍需要接近全库深度。")
    else:
        lines.append("这说明当前基线在前排候选上的命中能力已经较强，但若希望把全部 intersection truth 全量覆盖到，仍需要接近全库深度。")
    return lines


def best_worst_flights(per_flight: list[dict]) -> tuple[str, str]:
    best = max(per_flight, key=lambda item: float(item["intersection_recall@1"]))
    worst = min(per_flight, key=lambda item: float(item["intersection_recall@1"]))
    return short_flight_name(str(best["flight_id"])), short_flight_name(str(worst["flight_id"]))


def build_md(
    result_dir: Path,
    summary_top20: dict,
    summary_all: dict,
    per_flight: list[dict],
    timing_map: dict[str, float] | None,
    k_full: dict,
    k_unique: dict | None,
    simplified: bool,
    model_label: str,
) -> str:
    lines: list[str] = []
    lines.append(f"# {model_label} + FAISS 在 Intersection Truth 口径下的基线实验结果说明")
    lines.append("")
    lines.append("## 1. 任务定义与实验设置")
    lines.append(f"本组实验用于回答：在新的 `intersection truth` 真值定义下，`{model_label} + FAISS` 作为跨视角区域级粗定位基线，能够达到怎样的检索表现。")
    lines.append("")
    lines.append("- 数据范围：4 条航线，共 `40` 个 query。")
    lines.append(f"- 卫片候选库：`{int(summary_all['top_k'])}` 张卫片，复用固定候选库与既有 FAISS 索引。")
    lines.append(f"- 方法：`{model_label}` 全局特征 + FAISS `IndexFlatIP`。")
    lines.append(f"- 主展示口径：`top_k=20`；辅助分析口径：`top_k={int(summary_all['top_k'])}`（全库排序）。")
    lines.append("")
    lines.append("## 2. 新的 Query 定义")
    lines.append("本轮实验中的 query 不再沿用旧版 query 集，而是基于 4 条航线重新筛选得到的新一组 `40` 张无人机图像。")
    lines.append("")
    lines.append("- 选取范围：4 条航线，每条航线 `10` 张，共 `40` 张。")
    lines.append("- 倾角范围：`-85° ~ -40°`。")
    lines.append("- 选取目标：尽量覆盖不同类型地物，同时保持样本在航线内分散，不集中于连续帧。")
    lines.append("- 数据使用方式：原始带地理信息图像仅用于 query footprint 恢复和真值生成；实际检索输入使用去元数据后的 query 副本。")
    lines.append(f"- 检索约束：{model_label} + FAISS 检索阶段不读取 query 的地理坐标、姿态或 EXIF/XMP 信息，只基于图像视觉特征排序。")
    lines.append("")
    lines.append("这一设置更贴近当前任务目标，即验证：在不给检索模型提供 query 地理坐标的前提下，仅依赖单张无人机图像内容，能否把候选区域检索到正确地理范围附近。")
    lines.append("")
    lines.append("## 3. Intersection Truth 定义")
    lines.append("本轮正式真值定义为：只要 query 覆盖范围与卫片存在非零面积相交，该卫片就记为 `intersection truth`。")
    lines.append("这一定义比原先的单点式真值更宽，也更贴近“只要检索到与查询范围有真实地理交集的区域即可视为有效候选”的任务目标。")
    lines.append("")
    lines.append("## 4. 指标定义")
    lines.append("- `Intersection Recall@1/5/10/20`：前 K 名中是否命中 intersection truth。")
    lines.append("- `Intersection MRR`：首个 intersection truth 排名倒数的平均值。")
    lines.append("- `Top-1 error mean (m)`：首位候选中心与 query 参考位置之间的平均距离。")
    lines.append("")
    lines.append("## 5. 总体定量结果")
    lines.append("| 口径 | R@1 | R@5 | R@10 | R@20 | MRR | Top-1误差均值(m) |")
    lines.append("| --- | ---: | ---: | ---: | ---: | ---: | ---: |")
    lines.append(
        f"| Top-20 | {float(summary_top20['intersection_recall@1']):.3f} | {float(summary_top20['intersection_recall@5']):.3f} | {float(summary_top20['intersection_recall@10']):.3f} | {float(summary_top20['intersection_recall@20']):.3f} | {float(summary_top20['intersection_mrr']):.3f} | {float(summary_top20['top1_error_m_mean']):.3f} |"
    )
    lines.append(
        f"| 全库排序 | {float(summary_all['intersection_recall@1']):.3f} | {float(summary_all['intersection_recall@5']):.3f} | {float(summary_all['intersection_recall@10']):.3f} | {float(summary_all['intersection_recall@20']):.3f} | {float(summary_all['intersection_mrr']):.3f} | {float(summary_all['top1_error_m_mean']):.3f} |"
    )
    lines.append("")
    lines.append(
        f"主口径 `Top-20` 下，当前基线达到 `R@1={float(summary_top20['intersection_recall@1']):.3f}`、`R@10={float(summary_top20['intersection_recall@10']):.3f}`、`R@20={float(summary_top20['intersection_recall@20']):.3f}`、`MRR={float(summary_top20['intersection_mrr']):.3f}`。"
    )
    lines.append(
        "与全库排序相比，`Top-20` 结果几乎没有损失，说明大部分有效命中已经集中在前 20 名候选中。"
    )
    lines.append("")
    lines.append("## 6. 分航线结果（Top-20）")
    lines.append("| 航线 | Query数 | R@1 | R@5 | R@10 | R@20 | MRR | Top-1误差均值(m) |")
    lines.append("| --- | ---: | ---: | ---: | ---: | ---: | ---: | ---: |")
    for item in per_flight:
        lines.append(
            f"| {short_flight_name(str(item['flight_id']))} | {int(item['query_count'])} | {float(item['intersection_recall@1']):.3f} | {float(item['intersection_recall@5']):.3f} | {float(item['intersection_recall@10']):.3f} | {float(item['intersection_recall@20']):.3f} | {float(item['intersection_mrr']):.3f} | {float(item['top1_error_m_mean']):.3f} |"
        )
    lines.append("")
    best_flight, worst_flight = best_worst_flights(per_flight)
    lines.append(f"可以看到，四条航线之间仍存在差异，其中 `{best_flight}` 航线表现最好，`{worst_flight}` 航线相对最难。")
    lines.append("")
    next_section = 7
    if timing_map is not None and not simplified:
        lines.append("## 7. 时间开销统计")
        lines.append("| 阶段 | 耗时 |")
        lines.append("| --- | ---: |")
        lines.append(f"| 卫片特征提取 | {format_seconds(timing_map.get('satellite_feature_extraction', 0.0))} |")
        lines.append(f"| FAISS 建库 | {format_seconds(timing_map.get('faiss_index_build', 0.0))} |")
        lines.append(f"| Query 特征提取 | {format_seconds(timing_map.get('query_feature_extraction', 0.0))} |")
        lines.append(f"| 检索评估（Top-20） | {format_seconds(timing_map.get('query_retrieval', 0.0))} |")
        lines.append("")
        next_section = 8
    lines.append(f"## {next_section}. Top-K 曲线结果解读")
    for line in topk_summary_lines(k_full, k_unique):
        lines.append(f"- {line}")
    lines.append("")
    lines.append(f"## {next_section + 1}. 关键图表")
    lines.append("- `figures_intersection_retrieval/_aggregate/overall_metrics_bar.png`：overall 指标图。")
    lines.append("- `figures_intersection_retrieval/_aggregate/multi_flight_recall.png`：分航线 Recall 对比。")
    lines.append("- `figures_intersection_retrieval/_aggregate/top1_error_distribution.png`：Top-1 误差分布。")
    lines.append("- `figures_topk_fulltruth/_aggregate/overall_topk_truth_count_curve_all.png`：full-truth 曲线。")
    if k_unique is not None and not simplified:
        lines.append("- `figures_topk_unique_tile/_aggregate/overall_topk_unique_truth_count_curve.png`：unique-tile 曲线。")
    lines.append("")
    lines.append(f"## {next_section + 2}. 结论")
    lines.append(
        f"- 在 `intersection truth` 口径下，`{model_label} + FAISS` 已经具备较强的区域级初步地理定位能力，`Top-20` 下可达到 `R@20={float(summary_top20['intersection_recall@20']):.3f}`。"
    )
    lines.append("- 当前方法的主要瓶颈不在于前 20 名候选覆盖不足，而在于若希望把所有真值 tile 全部找全，仍需要很深的检索深度。")
    lines.append("- 后续若继续优化，应优先提升排序判别力，降低达到全真值覆盖所需的 `K`，而不是仅扩大候选窗口。")
    lines.append("")
    return "\n".join(lines)


def build_docx(
    result_dir: Path,
    out_docx: Path,
    summary_top20: dict,
    summary_all: dict,
    per_flight: list[dict],
    timing_map: dict[str, float] | None,
    k_full: dict,
    k_unique: dict | None,
    simplified: bool,
    model_label: str,
) -> None:
    fig_dir = result_dir / "figures_intersection_retrieval"
    fig_full_dir = result_dir / "figures_topk_fulltruth"
    fig_unique_dir = result_dir / "figures_topk_unique_tile"

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(f"{model_label} + FAISS 在 Intersection Truth 口径下的基线实验结果说明")
    set_cn_font(r, size=16, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("面向无人机单张图像初步地理定位的正式技术报告")
    set_cn_font(r, size=11)

    add_heading(doc, "1. 任务定义与实验设置", 1)
    add_paragraph(
        doc,
        f"本组实验用于回答：在新的 intersection truth 真值定义下，{model_label} + FAISS 作为跨视角区域级粗定位基线，"
        "能够在四条航线、40 个 query、1029 张固定候选卫片的条件下达到怎样的检索表现。"
    )
    add_bullets(
        doc,
        [
            "数据范围：4 条航线，共 40 个 query。",
            f"卫片候选库：复用固定候选库，共 {int(summary_all['top_k'])} 张卫片。",
            f"基线方法：{model_label} 全局特征 + FAISS IndexFlatIP。",
            f"本轮主展示口径：Top-20；辅助分析口径：全库排序（Top-K={int(summary_all['top_k'])}）。",
        ],
    )

    add_heading(doc, "2. 新的 Query 定义", 1)
    add_paragraph(
        doc,
        "本轮实验中的 query 不再沿用旧版 query 集，而是基于四条航线重新筛选得到的新一组 40 张无人机图像。"
        "每条航线固定选取 10 张，候选倾角范围限定在 -85° 到 -40°，并尽量覆盖不同类型地物，避免样本集中在连续帧或局部重复场景。"
    )
    add_bullets(
        doc,
        [
            "数据规模：4 条航线，每条航线 10 张，共 40 张 query。",
            "倾角范围：-85° 到 -40°。",
            "选样目标：覆盖不同地物类型，并在航线内部保持时间与空间分散。",
            "地理信息使用：原始图像中的坐标与姿态仅用于 query footprint 恢复和 intersection truth 真值生成。",
            f"检索输入：实际送入 {model_label} + FAISS 的是去元数据后的 query 图像副本，不包含地理坐标信息。",
        ],
    )
    add_paragraph(
        doc,
        "因此，本轮结果反映的是“只依赖单张无人机图像视觉内容进行跨视角区域检索”的能力，而不是利用 query 地理先验进行定位。"
    )

    add_heading(doc, "3. Intersection Truth 定义", 1)
    add_paragraph(
        doc,
        "本轮正式真值定义为：只要 query 覆盖范围与卫片存在非零面积相交，该卫片就记为 intersection truth。"
        "这一定义不再要求 query 中心点必须落入候选瓦片，而是把“存在真实地理交集”作为判断依据，因此更贴近当前任务对区域级定位的目标描述。"
    )

    add_heading(doc, "4. 指标定义", 1)
    build_metric_definition_table(doc)

    add_heading(doc, "5. 方法说明", 1)
    add_paragraph(
        doc,
        f"本轮方法仍采用 {model_label} 全局特征作为统一图像表示，对 query 图像和固定卫片库分别提取表示，"
        "再使用 FAISS `IndexFlatIP` 执行相似度检索。该方法不包含局部几何验证和重排，因此本次结果可以视为 intersection truth 新口径下的正式基线。"
    )

    add_heading(doc, "6. 总体定量结果", 1)
    add_paragraph(doc, "表 1 给出 Top-20 与全库排序两种口径下的总体指标结果。")
    build_overall_table(doc, summary_top20, summary_all)
    add_caption(doc, "表 1  Intersection Truth 口径下的总体指标结果")

    add_heading(doc, "7. 分航线结果", 1)
    add_paragraph(doc, "表 2 给出按航线拆分后的 Top-20 主口径结果，用于观察不同航线上的稳定性差异。")
    build_per_flight_table(doc, per_flight)
    add_caption(doc, "表 2  Intersection Truth 口径下的分航线指标结果")

    figure_section_num = 8
    if timing_map is not None and not simplified:
        add_heading(doc, "8. 时间开销统计", 1)
        add_paragraph(doc, "表 3 总结本次 run 中卫片特征提取、FAISS 建库、query 特征提取与检索评估的时间开销。")
        build_timing_table(doc, timing_map)
        add_caption(doc, "表 3  本次 run 的时间开销统计")
        figure_section_num = 9

    add_heading(doc, f"{figure_section_num}. 汇总图解读", 1)
    add_paragraph(
        doc,
        "图 1 到图 4 分别展示 overall 指标、分航线 Recall、Top-1 误差分布与 full-truth 曲线，"
        "用于从整体命中能力与真值覆盖深度两个角度理解当前基线。"
    )
    figures = [
        (fig_dir / "_aggregate" / "overall_metrics_bar.png", "图 1  Intersection Truth 口径下的 overall 指标图"),
        (fig_dir / "_aggregate" / "multi_flight_recall.png", "图 2  Intersection Truth 口径下的分航线 Recall 图"),
        (fig_dir / "_aggregate" / "top1_error_distribution.png", "图 3  Intersection Truth 口径下的 Top-1 误差分布图"),
        (fig_full_dir / "_aggregate" / "overall_topk_truth_count_curve_all.png", "图 4  Full-truth Top-K 曲线"),
    ]
    if k_unique is not None and not simplified:
        figures.append((fig_unique_dir / "_aggregate" / "overall_topk_unique_truth_count_curve.png", "图 5  Unique-tile Top-K 曲线"))
    for path, caption in figures:
        add_picture(doc, path, 6.2)
        add_caption(doc, caption)

    add_paragraph(
        doc,
        f"从总体结果看，Top-20 口径下当前基线达到 `R@1={float(summary_top20['intersection_recall@1']):.3f}`、"
        f"`R@10={float(summary_top20['intersection_recall@10']):.3f}`、`R@20={float(summary_top20['intersection_recall@20']):.3f}`、"
        f"`MRR={float(summary_top20['intersection_mrr']):.3f}`。与全库排序相比，这些核心指标几乎没有变化，"
        "说明前 20 名候选已经覆盖了主要有效命中空间。"
    )
    for line in topk_summary_lines(k_full, k_unique):
        add_paragraph(doc, line)

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, f"{figure_section_num + 1}. 结论", 1)
    add_bullets(
        doc,
        [
            f"在 intersection truth 口径下，{model_label} + FAISS 基线在 Top-20 内已经取得较强的区域级初步定位能力，`R@20={float(summary_top20['intersection_recall@20']):.3f}`。",
            "当前方法的主要瓶颈不在于候选窗口太浅，而在于若希望把所有真值 tile 全部找全，仍需要接近全库深度。",
            "后续若继续优化，应优先提升排序判别力，降低达到全真值覆盖所需的 K，而不是仅扩大候选深度。",
        ],
    )

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))


def main() -> None:
    args = parse_args()
    result_dir = Path(args.result_dir)
    out_docx = Path(args.out_docx)
    out_md = Path(args.out_md)

    summary_top20 = load_json(result_dir / "retrieval" / "summary_top20.json")
    summary_all = load_json(result_dir / "retrieval" / "summary_all.json")
    k_full = load_json(result_dir / "figures_topk_fulltruth" / "k_full_truth_summary.json")
    timing_path = result_dir / "timing" / "timing_summary.json"
    k_unique_path = result_dir / "figures_topk_unique_tile" / "k_full_truth_unique_tile_summary.json"
    timing = load_json(timing_path) if timing_path.exists() and not args.simplified else None
    k_unique = load_json(k_unique_path) if k_unique_path.exists() and not args.simplified else None

    per_flight = per_flight_metrics(summary_top20)
    timing_map = stage_seconds(timing) if timing is not None else None

    out_md.parent.mkdir(parents=True, exist_ok=True)
    out_md.write_text(
        build_md(result_dir, summary_top20, summary_all, per_flight, timing_map, k_full, k_unique, args.simplified, args.model_label),
        encoding="utf-8",
    )
    build_docx(result_dir, out_docx, summary_top20, summary_all, per_flight, timing_map, k_full, k_unique, args.simplified, args.model_label)
    print(out_md)
    print(out_docx)


if __name__ == "__main__":
    main()
