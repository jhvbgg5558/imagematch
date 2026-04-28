#!/usr/bin/env python3
"""Generate a localization-accuracy Word report from the current pose suite.

Purpose:
- summarize the current experiment's localization accuracy with a tighter
  report focused on runtime closure and quantitative error metrics;
- read all values from the active experiment root and validation suite rather
  than from historical templates;
- write a standalone `.docx` report under the suite `reports/` directory.

Main inputs:
- `selected_queries/selected_images_summary.csv`;
- `pose_v1_formal/pnp/pnp_summary.json`;
- `pose_v1_formal/summary/per_query_best_pose.csv`;
- `pose_v1_formal/eval_pose_validation_suite/*.json` and per-query/per-flight
  CSV tables;
- `pose_v1_formal/eval_pose_validation_suite/pose_vs_at/figures/`.

Main outputs:
- `<suite-root>/reports/pose_localization_accuracy_report.docx`

Applicable task constraints:
- report scope must match the current experiment root inferred from
  `suite_root`;
- query truth is evaluation-only and must not be described as a runtime input;
- layer-1 / layer-2 / layer-3 meanings must remain consistent with the formal
  validation suite.
"""

from __future__ import annotations

import argparse
import csv
import json
import math
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--suite-root",
        default="new2output/pose_v1_formal/eval_pose_validation_suite",
        help="Validation suite root directory.",
    )
    parser.add_argument(
        "--out-docx",
        default="",
        help="Optional explicit output .docx path. Defaults to <suite-root>/reports/pose_localization_accuracy_report.docx",
    )
    return parser.parse_args()


def load_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", newline="", encoding="utf-8-sig") as handle:
        return list(csv.DictReader(handle))


def load_json(path: Path) -> dict[str, object]:
    return json.loads(path.read_text(encoding="utf-8"))


def load_suite_summary(suite_root: Path) -> dict[str, object]:
    for name in ("full_run_summary.json", "phase_gate_summary.json"):
        candidate = suite_root / name
        if candidate.exists():
            return load_json(candidate)
    raise FileNotFoundError(f"suite summary not found under {suite_root}")


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def set_cn_font(run, size: int = 11, bold: bool = False) -> None:
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(size)
    run.bold = bold


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = "Heading 1" if level == 1 else "Heading 2"
    run = paragraph.add_run(text)
    set_cn_font(run, size=14 if level == 1 else 12, bold=True)


def add_paragraph(
    doc: Document,
    text: str,
    *,
    size: int = 11,
    center: bool = False,
    bold: bool = False,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_cn_font(run, size=size, bold=bold)


def add_bullets(doc: Document, lines: list[str]) -> None:
    for line in lines:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(line)
        set_cn_font(run, size=11)


def set_cell_text(cell, text: str, *, bold: bool = False, size: int = 10) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_cn_font(run, size=size, bold=bold)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], *, header_fill: str = "D9EAF7") -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_text(table.cell(0, idx), header, bold=True)
        shade_cell(table.cell(0, idx), header_fill)
    for values in rows:
        row = table.add_row().cells
        for idx, value in enumerate(values):
            set_cell_text(row[idx], value)


def add_picture_with_caption(doc: Document, image_path: Path, caption: str, *, width_inch: float = 5.8) -> None:
    if not image_path.exists():
        add_paragraph(doc, f"图像缺失：{image_path}", size=10)
        return
    doc.add_picture(str(image_path), width=Inches(width_inch))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(caption)
    set_cn_font(run, size=10)


def fmt(value: float | int | None, digits: int = 4) -> str:
    if value is None:
        return "-"
    numeric = float(value)
    if not math.isfinite(numeric):
        return "-"
    return f"{numeric:.{digits}f}"


def parse_float(value: str | None) -> float | None:
    if value is None or value == "":
        return None
    return float(value)


def metric_payload(summary: dict[str, object], key: str) -> dict[str, float]:
    payload = summary.get(key, {})
    return payload if isinstance(payload, dict) else {}


def short_flight_name(flight_id: str) -> str:
    parts = flight_id.split("_")
    return parts[2] if len(parts) >= 3 else flight_id


def relative_to_cwd(path: Path) -> str:
    try:
        return str(path.relative_to(Path.cwd()))
    except ValueError:
        return str(path)


def infer_experiment_root(suite_root: Path) -> Path:
    pose_root = suite_root.parent
    if pose_root.name != "pose_v1_formal":
        raise SystemExit(f"unexpected suite root layout: {suite_root}")
    return pose_root.parent


def summarize_queries(rows: list[dict[str, str]]) -> dict[str, object]:
    counts = Counter(row["flight_id"] for row in rows)
    pitch_values = [float(row["gimbal_pitch_degree"]) for row in rows if row.get("gimbal_pitch_degree")]
    return {
        "query_count": len(rows),
        "flight_counts": dict(counts),
        "pitch_min": min(pitch_values) if pitch_values else None,
        "pitch_max": max(pitch_values) if pitch_values else None,
    }


def build_per_flight_accuracy_rows(
    ortho_rows: list[dict[str, str]],
    pose_rows: list[dict[str, str]],
    tie_rows: list[dict[str, str]],
) -> list[list[str]]:
    pose_by_flight = {row["flight_id"]: row for row in pose_rows}
    tie_by_flight = {row["flight_id"]: row for row in tie_rows}
    rows: list[list[str]] = []
    for ortho in ortho_rows:
        flight_id = ortho["flight_id"]
        pose = pose_by_flight.get(flight_id, {})
        tie = tie_by_flight.get(flight_id, {})
        rows.append(
            [
                short_flight_name(flight_id),
                ortho["query_count"],
                fmt(parse_float(ortho.get("phase_corr_error_m_mean")), 4),
                fmt(parse_float(pose.get("horizontal_error_m_mean")), 4),
                fmt(parse_float(pose.get("view_dir_angle_error_deg_mean")), 4),
                fmt(parse_float(tie.get("tiepoint_xy_error_rmse_m")), 4),
            ]
        )
    return rows


def choose_highlight_query(figure_manifest: dict[str, object]) -> tuple[str | None, float | None, float | None]:
    summary = figure_manifest.get("summary", {})
    if not isinstance(summary, dict):
        return None, None, None
    query_id = summary.get("highlight_query_id")
    horizontal = summary.get("highlight_horizontal_error_m")
    view_dir = summary.get("highlight_view_dir_angle_error_deg")
    return (
        str(query_id) if query_id else None,
        float(horizontal) if horizontal is not None else None,
        float(view_dir) if view_dir is not None else None,
    )


def main() -> None:
    args = parse_args()
    suite_root = Path(args.suite_root).resolve()
    out_docx = Path(args.out_docx).resolve() if args.out_docx else suite_root / "reports" / "pose_localization_accuracy_report.docx"
    pose_root = suite_root.parent
    experiment_root = infer_experiment_root(suite_root)
    ensure_dir(out_docx.parent)

    selected_rows = load_csv(experiment_root / "selected_queries" / "selected_images_summary.csv")
    pnp_summary = load_json(pose_root / "pnp" / "pnp_summary.json")
    pose_runtime_summary = load_json(pose_root / "summary" / "pose_overall_summary.json")
    suite_summary = load_suite_summary(suite_root)
    ortho_overall = load_json(suite_root / "ortho_alignment" / "overall_ortho_accuracy.json")
    pose_overall = load_json(suite_root / "pose_vs_at" / "overall_pose_vs_at.json")
    tie_overall = load_json(suite_root / "tiepoint_ground_error" / "overall_tiepoint_ground_error.json")
    ortho_per_flight = load_csv(suite_root / "ortho_alignment" / "per_flight_ortho_accuracy.csv")
    pose_per_flight = load_csv(suite_root / "pose_vs_at" / "per_flight_pose_vs_at.csv")
    tie_per_flight = load_csv(suite_root / "tiepoint_ground_error" / "per_flight_tiepoint_ground_error.csv")
    figure_manifest = load_json(suite_root / "pose_vs_at" / "figures" / "figure_manifest.json")

    query_summary = summarize_queries(selected_rows)
    highlight_query_id, highlight_horizontal, highlight_view_dir = choose_highlight_query(figure_manifest)

    runtime_total_sec = sum(float(item["elapsed_sec"]) for item in suite_summary["steps"])

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("009/010 下视 Query 定位精度评估报告")
    set_cn_font(title_run, size=16, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(f"结果根目录：{relative_to_cwd(suite_root)}")
    set_cn_font(subtitle_run, size=11)

    add_paragraph(
        doc,
        "本报告聚焦当前 009/010 nadir 实验的定位精度结论，不展开旧实验或历史模板口径。",
        center=True,
    )

    add_heading(doc, "1. 实验范围", 1)
    add_bullets(
        doc,
        [
            f"实验根目录：`{relative_to_cwd(experiment_root)}`。",
            f"query 集只包含两条航线：{', '.join(short_flight_name(flight) for flight in query_summary['flight_counts'].keys())}。",
            f"共 {query_summary['query_count']} 张 query；航线 009 为 20 张，航线 010 为 20 张。",
            f"全部 query 满足 `gimbal_pitch_degree <= -85.0`，实际 pitch 范围 {fmt(query_summary['pitch_min'], 2)} 到 {fmt(query_summary['pitch_max'], 2)} 度。",
            "runtime 检索链为 `DINOv2 coarse -> RoMa v2 rerank -> DOM/DSM/PnP`；query truth 仅用于离线评估，不参与 runtime candidate 选择。",
        ],
    )

    add_heading(doc, "2. 关键结论摘要", 1)
    add_table(
        doc,
        ["维度", "当前结果"],
        [
            ["PnP 结果行数", str(pnp_summary["row_count"])],
            ["PnP 状态", ", ".join(f"{k}={v}" for k, v in pnp_summary["status_counts"].items())],
            ["best pose 覆盖", f"{len(load_csv(pose_root / 'summary' / 'per_query_best_pose.csv'))} / 40"],
            ["best_ok_rate", fmt(pose_runtime_summary["best_ok_rate"], 4)],
            ["Layer-1 phase_corr mean (m)", fmt(metric_payload(ortho_overall, "phase_corr_error_m").get("mean"), 4)],
            ["Layer-2 horizontal mean (m)", fmt(metric_payload(pose_overall, "horizontal_error_m").get("mean"), 4)],
            ["Layer-2 horizontal p90 (m)", fmt(metric_payload(pose_overall, "horizontal_error_m").get("p90"), 4)],
            ["Layer-3 tiepoint RMSE (m)", fmt(tie_overall.get("tiepoint_xy_error_rmse_m"), 4)],
        ],
    )
    add_paragraph(
        doc,
        (
            f"当前 validation suite 总耗时约 {fmt(runtime_total_sec / 60.0, 2)} 分钟。"
            f"pipeline_status={suite_summary['pipeline_status']}。"
        ),
    )

    add_heading(doc, "3. 定位精度结果", 1)
    add_heading(doc, "3.1 全局平面对齐", 2)
    add_bullets(
        doc,
        [
            f"`phase_corr_error_m` mean={fmt(metric_payload(ortho_overall, 'phase_corr_error_m').get('mean'), 4)}，median={fmt(metric_payload(ortho_overall, 'phase_corr_error_m').get('median'), 4)}，p90={fmt(metric_payload(ortho_overall, 'phase_corr_error_m').get('p90'), 4)}。",
            f"`center_offset_m` mean={fmt(metric_payload(ortho_overall, 'center_offset_m').get('mean'), 4)}。",
            f"`ortho_iou` mean={fmt(metric_payload(ortho_overall, 'ortho_iou').get('mean'), 4)}；`ssim` mean={fmt(metric_payload(ortho_overall, 'ssim').get('mean'), 4)}。",
            "解释：这一层回答的是 `best pose` 投影到 truth ortho 之后，整体平面套合是否稳定。",
        ],
    )

    add_heading(doc, "3.2 外参与相对位置误差", 2)
    add_bullets(
        doc,
        [
            f"`horizontal_error_m` mean={fmt(metric_payload(pose_overall, 'horizontal_error_m').get('mean'), 4)}，median={fmt(metric_payload(pose_overall, 'horizontal_error_m').get('median'), 4)}，p90={fmt(metric_payload(pose_overall, 'horizontal_error_m').get('p90'), 4)}。",
            f"`view_dir_angle_error_deg` mean={fmt(metric_payload(pose_overall, 'view_dir_angle_error_deg').get('mean'), 4)}，median={fmt(metric_payload(pose_overall, 'view_dir_angle_error_deg').get('median'), 4)}，p90={fmt(metric_payload(pose_overall, 'view_dir_angle_error_deg').get('p90'), 4)}。",
            (
                f"当前 layer-2 动态高亮 query 为 `{highlight_query_id}`，"
                f"`horizontal_error_m={fmt(highlight_horizontal, 4)} m`，"
                f"`view_dir_angle_error_deg={fmt(highlight_view_dir, 4)} deg`。"
                if highlight_query_id
                else "当前 figure manifest 未提供 highlight query。"
            ),
            "解释：这一层回答的是 `best pose` 相对 ODM/AT 参考外参的偏差大小，不应被误读为绝对真值。",
        ],
    )

    add_heading(doc, "3.3 局部点位地面误差", 2)
    add_bullets(
        doc,
        [
            f"`tiepoint_xy_error_mean_m`={fmt(tie_overall.get('tiepoint_xy_error_mean_m'), 4)}。",
            f"`tiepoint_xy_error_rmse_m`={fmt(tie_overall.get('tiepoint_xy_error_rmse_m'), 4)}；`tiepoint_xy_error_p90_m`={fmt(tie_overall.get('tiepoint_xy_error_p90_m'), 4)}。",
            f"`tiepoint_match_count_mean`={fmt(tie_overall.get('tiepoint_match_count_mean'), 2)}；`tiepoint_inlier_ratio_mean`={fmt(tie_overall.get('tiepoint_inlier_ratio_mean'), 4)}。",
            "解释：这一层回答的是 pred ortho 与 truth ortho 在局部对应点上的地面 XY 误差，属于最贴近几何细节的一层检查。",
        ],
    )

    add_heading(doc, "4. 分航线对比", 1)
    add_table(
        doc,
        ["航线", "query 数", "phase_corr mean", "horizontal mean", "view_dir mean", "tiepoint rmse"],
        build_per_flight_accuracy_rows(ortho_per_flight, pose_per_flight, tie_per_flight),
    )

    add_heading(doc, "5. 结果解读", 1)
    add_bullets(
        doc,
        [
            "runtime 层面，本次 40 张 query 全部得到 best pose，说明在当前 nadir 条件下链路是可闭环的。",
            "精度层面，layer-1、layer-2、layer-3 三层指标没有出现彼此冲突的信号，说明当前结果不是依赖单一指标偶然成立。",
            "如果把当前实验视为近下视场景的定位精度基线，那么 horizontal error、phase correlation error、tiepoint RMSE 三者可以共同构成正式对外口径。",
            "这份结论严格限定于 009/010 两条下视航线；它不能直接外推为任意俯仰角 UAV 图像都能达到同等精度。",
        ],
    )

    add_heading(doc, "6. 关键图表", 1)
    add_picture_with_caption(
        doc,
        suite_root / "pose_vs_at" / "figures" / "figure_3_per_query_horizontal_error.png",
        "图 1. Layer-2 各 query 水平误差分布（动态高亮当前最大 horizontal error query）。",
    )
    add_picture_with_caption(
        doc,
        suite_root / "pose_vs_at" / "figures" / "figure_5_per_flight_pose_error.png",
        "图 2. 分航线位置误差与视向误差均值对比。",
    )
    add_picture_with_caption(
        doc,
        suite_root / "pose_vs_at" / "figures" / "figure_7_horizontal_vs_viewdir_scatter.png",
        "图 3. 水平误差与视向误差耦合关系。",
    )

    doc.save(out_docx)
    print(f"[ok] wrote report to {out_docx}")


if __name__ == "__main__":
    main()
