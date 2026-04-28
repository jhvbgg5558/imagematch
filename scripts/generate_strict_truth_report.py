#!/usr/bin/env python3
"""Generate a formal Word report for the strict-truth DINOv2 baseline experiment."""

from __future__ import annotations

import argparse
import csv
import json
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate a strict-truth formal Word report.")
    parser.add_argument("--result-dir", required=True)
    parser.add_argument("--baseline-summary-json", required=True)
    parser.add_argument("--out-docx", required=True)
    return parser.parse_args()


def load_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", newline="", encoding="utf-8-sig") as f:
        return list(csv.DictReader(f))


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def set_cn_font(run, size: int = 11, bold: bool = False) -> None:
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(size)
    run.bold = bold


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 10, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    set_cn_font(r, size=size, bold=bold)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.style = "Heading 1" if level == 1 else "Heading 2"
    r = p.add_run(text)
    set_cn_font(r, size=14 if level == 1 else 12, bold=True)


def add_paragraph(doc: Document, text: str, size: int = 11, center: bool = False) -> None:
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn_font(r, size=size)


def add_bullets(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(line)
        set_cn_font(r, size=11)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn_font(r, size=10)


def add_picture(doc: Document, path: Path, width_inch: float) -> None:
    doc.add_picture(str(path), width=Inches(width_inch))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def short_flight_name(flight_id: str) -> str:
    parts = flight_id.split("_")
    return parts[2] if len(parts) >= 3 else flight_id


def per_flight_metrics(summary: dict, seed_rows: dict[str, dict[str, str]]) -> list[dict[str, float | str]]:
    grouped: dict[str, list[dict]] = defaultdict(list)
    for row in summary["per_query"]:
        grouped[seed_rows[row["query_id"]]["flight_id"]].append(row)

    out = []
    for flight_id in sorted(grouped):
        rows = grouped[flight_id]
        total = len(rows)
        errors = [float(row["top1_error_m"]) for row in rows]
        out.append(
            {
                "flight_id": flight_id,
                "query_count": total,
                "strict_recall@1": sum(int(row["strict_hit@1"]) for row in rows) / total,
                "strict_recall@5": sum(int(row["strict_hit@5"]) for row in rows) / total,
                "strict_recall@10": sum(int(row["strict_hit@10"]) for row in rows) / total,
                "strict_mrr": sum(float(row["strict_reciprocal_rank"]) for row in rows) / total,
                "top1_error_m_mean": sum(errors) / len(errors),
            }
        )
    return out


def truth_scale_stats(rows: list[dict[str, str]]) -> dict[str, float]:
    total = len(rows)
    return {
        "query_count": total,
        "coverage_threshold": float(rows[0]["coverage_threshold"]),
        "footprint_core_ratio": float(rows[0]["footprint_core_ratio"]),
        "min_valid_ratio": float(rows[0]["min_valid_ratio"]),
        "truth_count_total_mean": sum(float(r["truth_count_total"]) for r in rows) / total,
        "strict_truth_count_total_mean": sum(float(r["strict_truth_count_total"]) for r in rows) / total,
        "soft_truth_count_total_mean": sum(float(r["soft_truth_count_total"]) for r in rows) / total,
        "strict_truth_count_200m_mean": sum(float(r["strict_truth_count_200m"]) for r in rows) / total,
        "strict_truth_count_300m_mean": sum(float(r["strict_truth_count_300m"]) for r in rows) / total,
        "strict_truth_count_500m_mean": sum(float(r["strict_truth_count_500m"]) for r in rows) / total,
        "strict_truth_count_700m_mean": sum(float(r["strict_truth_count_700m"]) for r in rows) / total,
        "footprint_area_m2_mean": sum(float(r["footprint_area_m2"]) for r in rows) / total,
    }


def build_metric_definition_table(doc: Document) -> None:
    rows = [
        ("Strict Recall@1", "首位候选命中 strict truth 的比例，是本轮正式主指标。"),
        ("Strict Recall@5", "前 5 名中命中 strict truth 的比例，衡量前排候选覆盖能力。"),
        ("Strict Recall@10", "前 10 名中命中 strict truth 的比例，衡量区域级初步定位能力。"),
        ("Strict MRR", "strict truth 排名倒数的平均值，越高说明有效真值整体越靠前。"),
        ("Top-1 error mean (m)", "首位候选中心与 query 参考位置之间的平均距离，检索排序未变时通常保持不变。"),
        ("Center-Strict Recall 系列", "辅助观察口径，仅统计既是 strict truth 又包含 query center 的 tile，不作为主结论。"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(["指标", "含义"]):
        set_cell_text(table.cell(0, i), h, bold=True)
        shade_cell(table.cell(0, i), "D9EAF7")
    for name, desc in rows:
        row = table.add_row().cells
        set_cell_text(row[0], name)
        row[1].text = ""
        p = row[1].paragraphs[0]
        r = p.add_run(desc)
        set_cn_font(r, size=10)


def build_overall_table(doc: Document, strict_summary: dict, baseline_summary: dict) -> None:
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(["口径", "R@1", "R@5", "R@10", "MRR", "Top-1误差均值(m)"]):
        set_cell_text(table.cell(0, i), h, bold=True)
        shade_cell(table.cell(0, i), "D9EAF7")

    rows = [
        (
            "Strict",
            f"{float(strict_summary['strict_recall@1']):.3f}",
            f"{float(strict_summary['strict_recall@5']):.3f}",
            f"{float(strict_summary['strict_recall@10']):.3f}",
            f"{float(strict_summary['strict_mrr']):.3f}",
            f"{float(strict_summary['top1_error_m_mean']):.3f}",
        ),
        (
            "Coverage(旧)",
            f"{float(baseline_summary['coverage_recall@1']):.3f}",
            f"{float(baseline_summary['coverage_recall@5']):.3f}",
            f"{float(baseline_summary['coverage_recall@10']):.3f}",
            f"{float(baseline_summary['coverage_mrr']):.3f}",
            f"{float(baseline_summary['top1_error_m_mean']):.3f}",
        ),
        (
            "Center-Strict",
            f"{float(strict_summary['center_strict_recall@1']):.3f}",
            f"{float(strict_summary['center_strict_recall@5']):.3f}",
            f"{float(strict_summary['center_strict_recall@10']):.3f}",
            "-",
            "-",
        ),
    ]
    for vals in rows:
        row = table.add_row().cells
        for i, val in enumerate(vals):
            set_cell_text(row[i], val)


def build_per_flight_table(doc: Document, rows: list[dict[str, float | str]]) -> None:
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(["航线", "Strict R@1", "Strict R@5", "Strict R@10", "Strict MRR", "Top-1误差均值(m)"]):
        set_cell_text(table.cell(0, i), h, bold=True)
        shade_cell(table.cell(0, i), "D9EAF7")
    for item in rows:
        row = table.add_row().cells
        set_cell_text(row[0], short_flight_name(str(item["flight_id"])))
        set_cell_text(row[1], f"{float(item['strict_recall@1']):.3f}")
        set_cell_text(row[2], f"{float(item['strict_recall@5']):.3f}")
        set_cell_text(row[3], f"{float(item['strict_recall@10']):.3f}")
        set_cell_text(row[4], f"{float(item['strict_mrr']):.3f}")
        set_cell_text(row[5], f"{float(item['top1_error_m_mean']):.3f}")


def build_truth_summary_table(doc: Document, stats: dict[str, float], roi_summary: dict) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(["字段", "值"]):
        set_cell_text(table.cell(0, i), h, bold=True)
        shade_cell(table.cell(0, i), "D9EAF7")

    bounds = roi_summary["roi_bounds"]
    rows = [
        ("尺度集合", "200m / 300m / 500m / 700m"),
        ("Coverage 阈值", f"{stats['coverage_threshold']:.2f}"),
        ("Footprint core ratio", f"{stats['footprint_core_ratio']:.2f}"),
        ("Min valid ratio", f"{stats['min_valid_ratio']:.2f}"),
        ("Query 数量", f"{int(stats['query_count'])}"),
        ("卫星 tiles 数量", f"{int(roi_summary['point_count'])}"),
        ("ROI 范围", f"[{bounds[0]:.3f}, {bounds[1]:.3f}] - [{bounds[2]:.3f}, {bounds[3]:.3f}]"),
        ("平均 truth 总数", f"{stats['truth_count_total_mean']:.2f}"),
        ("平均 strict truth 数", f"{stats['strict_truth_count_total_mean']:.2f}"),
        ("平均 soft truth 数", f"{stats['soft_truth_count_total_mean']:.2f}"),
        ("平均 strict 200m 数", f"{stats['strict_truth_count_200m_mean']:.2f}"),
        ("平均 strict 300m 数", f"{stats['strict_truth_count_300m_mean']:.2f}"),
        ("平均 strict 500m 数", f"{stats['strict_truth_count_500m_mean']:.2f}"),
        ("平均 strict 700m 数", f"{stats['strict_truth_count_700m_mean']:.2f}"),
        ("平均 footprint 面积(m2)", f"{stats['footprint_area_m2_mean']:.3f}"),
    ]
    for key, value in rows:
        row = table.add_row().cells
        set_cell_text(row[0], key)
        set_cell_text(row[1], value)


def choose_cases(summary: dict, seed_rows: dict[str, dict[str, str]]) -> list[dict]:
    rows = summary["per_query"]
    success = [r for r in rows if r["first_strict_truth_rank"] == 1]
    success.sort(key=lambda r: (float(r["top1_error_m"]), r["query_id"]))

    boundary = [
        r
        for r in rows
        if (r["first_strict_truth_rank"] in {2, 3, 4, 5}) or (r["strict_hit@10"] and not r["strict_hit@5"])
    ]
    boundary.sort(
        key=lambda r: (
            r["first_strict_truth_rank"] if r["first_strict_truth_rank"] is not None else 99,
            float(r["top1_error_m"]),
            r["query_id"],
        )
    )

    failure = [r for r in rows if not r["strict_hit@10"]]
    failure.sort(key=lambda r: (-float(r["top1_error_m"]), r["query_id"]))

    selected: list[dict] = []
    used: set[str] = set()
    pools = [
        ("显著成功样例", success, "该 query 的首位候选已经直接命中 strict truth，说明该样本不仅地理覆盖正确，而且候选内容也满足有效性约束。"),
        ("边界成功样例", boundary, "该 query 虽未在首位命中 strict truth，但有效真值已进入前列候选，说明当前基线具备区域级定位能力，但前排排序仍有优化空间。"),
        ("残余失败样例", failure, "该 query 在 Top-10 内仍未命中 strict truth，说明在内容过滤后，该类场景仍存在稳定失败模式。"),
    ]
    for title, pool, note in pools:
        for row in pool:
            if row["query_id"] in used:
                continue
            case = dict(row)
            case["title"] = title
            case["note"] = note
            case["flight_id"] = seed_rows[row["query_id"]]["flight_id"]
            selected.append(case)
            used.add(row["query_id"])
            break
    return selected


def add_case_section(doc: Document, case_idx: int, case: dict, fig_dir: Path) -> None:
    query_id = case["query_id"]
    flight_id = case["flight_id"]
    rank = case["first_strict_truth_rank"]
    rank_text = "未命中 Top-10" if rank is None else str(rank)
    add_heading(doc, f"7.{case_idx} {case['title']}：{short_flight_name(flight_id)} / {query_id}", 2)
    add_paragraph(
        doc,
        f"首个 strict truth 排名：{rank_text}；Strict Truth Count：{case['strict_truth_count']}；"
        f"Center-Strict Count：{case['center_strict_truth_count']}；Top-1 error：{float(case['top1_error_m']):.3f}m。{case['note']}",
    )
    img = fig_dir / flight_id / f"{query_id}_top10.png"
    add_picture(doc, img, 6.0)
    add_caption(doc, f"{case['title']}：{short_flight_name(flight_id)} / {query_id} 的 Top-10 检索结果")


def main() -> None:
    args = parse_args()
    result_dir = Path(args.result_dir)
    fig_dir = result_dir / "figures"
    out_path = Path(args.out_docx)

    strict_summary = load_json(result_dir / "retrieval" / "summary.json")
    baseline_summary = load_json(Path(args.baseline_summary_json))
    roi_summary = load_json(result_dir / "fixed_satellite_library" / "roi_summary.json")
    truth_rows = load_csv(result_dir / "query_truth" / "query_truth.csv")
    seed_rows = {row["query_id"]: row for row in load_csv(result_dir / "query_truth" / "queries_truth_seed.csv")}
    flight_rows = per_flight_metrics(strict_summary, seed_rows)
    truth_stats = truth_scale_stats(truth_rows)
    cases = choose_cases(strict_summary, seed_rows)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("DINOv2 + FAISS 在 Strict Truth 口径下的多尺度跨视角粗定位实验结果说明")
    set_cn_font(r, size=16, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("面向无人机单张图像初步地理定位的正式技术报告")
    set_cn_font(r, size=11)

    add_heading(doc, "1. 任务定义与实验设置", 1)
    add_paragraph(
        doc,
        "本组实验用于论证：在更贴近工程实际的输入条件下，仅依赖遥感正射影像，是否能够对任意单张无人机图像实现区域级初步地理定位。"
        "本轮不更换模型与索引，只在既有 `200/300/500/700m` 候选尺度与检索排序不变的前提下，把正式评估口径从 coverage truth 收紧为 strict truth，"
        "用于观察“几何覆盖 + 有效内容过滤”后的基线真实性能。",
    )
    add_bullets(
        doc,
        [
            f"数据范围：4 条航线，共 {strict_summary['query_count']} 个 query。",
            f"卫星候选库：{roi_summary['point_count']} 个 tiles，坐标系为 {roi_summary['satellite_crs']}。",
            "查询图像：去元数据的原始无人机单张图像，不保证为正射视角。",
            "基线方法：DINOv2 pooler 全局特征 + FAISS IndexFlatIP 粗检索。",
            "本轮变化：仅切换评估真值口径，不重跑特征提取和建库。",
        ],
    )

    add_heading(doc, "2. Strict Truth 定义与选取范围", 1)
    add_paragraph(
        doc,
        "strict truth 先沿用 coverage truth 的几何约束：根据无人机图像元数据近似估计其地面覆盖 footprint，"
        "再计算 footprint 与固定卫星库中各候选瓦片的地面覆盖框相交比例。"
        "当 `coverage_ratio >= 0.4` 时，该瓦片先进入 coverage 候选集合。"
        "随后进一步计算该卫星瓦片的有效内容比例 `valid_pixel_ratio`，只有满足 `valid_pixel_ratio >= 0.6` 的候选才进入 strict truth。"
        "不满足有效内容阈值但满足 coverage 几何条件的瓦片被降级为 soft truth，仅保留作诊断资产。"
    )
    add_paragraph(
        doc,
        "这一定义的目的，是把“几何上覆盖正确但几乎全是黑边或无效像素”的瓦片排除出正式主真值集合，从而让评估更贴近真实视觉检索监督。"
    )
    build_truth_summary_table(doc, truth_stats, roi_summary)
    add_caption(doc, "表 1  本轮 strict truth 构造与固定卫星库范围概览")

    add_heading(doc, "3. 指标定义", 1)
    build_metric_definition_table(doc)

    add_heading(doc, "4. 方法说明", 1)
    add_heading(doc, "4.1 DINOv2 + FAISS 粗检索基线", 2)
    add_paragraph(
        doc,
        "本轮方法仍采用 DINOv2 `pooler` 输出作为统一全局特征，对 query 图像和固定卫星库瓦片分别提取表示，"
        "再经 L2 归一化后使用 FAISS `IndexFlatIP` 执行 Top-10 相似度检索。该方法不包含局部几何验证和重排，因此本轮结果可以被视为“真值口径净化”后的最小可运行正式基线。",
    )
    add_heading(doc, "4.2 与 coverage 结果相比的变化", 2)
    add_paragraph(
        doc,
        "相较上一版 coverage truth 结果，本轮没有改变模型、索引或排序结果。真正改变的是正式真值集合："
        "将几何覆盖成立但有效内容不足的 tile 从主真值中移出。"
        "因此，本轮结果用于回答的问题是：在相同检索排序下，若只统计可作为有效视觉正样本的真值，当前基线性能会下降多少，且这种下降是否更真实。"
    )

    add_heading(doc, "5. 定量结果", 1)
    add_paragraph(doc, "表 2 给出 strict truth、旧 coverage truth 和 center-strict 辅助口径的 overall 对照结果。")
    build_overall_table(doc, strict_summary, baseline_summary)
    add_caption(doc, "表 2  Strict Truth 与旧 Coverage Truth 的 overall 对照结果")
    add_paragraph(doc, "表 3 给出按航线拆分后的 strict truth 主口径结果，用于观察不同航线上的稳定性差异。")
    build_per_flight_table(doc, flight_rows)
    add_caption(doc, "表 3  当前实验的分航线 strict truth 指标结果")

    add_heading(doc, "6. 汇总图解读", 1)
    add_paragraph(
        doc,
        "为避免仅依据表格数字做静态判断，本文进一步给出 strict truth 口径下的汇总图。图 1 展示 overall 主指标，图 2 展示分航线 Recall，图 3 展示辅助 center-strict 指标，图 4 展示 40 个 query 的 Top-1 误差分布。"
    )
    figures = [
        ("overall_metrics_bar.png", "图 1  当前实验的 overall strict truth 指标图"),
        ("multi_flight_recall.png", "图 2  当前实验的分航线 strict truth Recall 图"),
        ("center_metrics_bar.png", "图 3  当前实验的辅助 center-strict 指标图"),
        ("top1_error_distribution.png", "图 4  当前实验的 Top-1 误差分布图"),
    ]
    for fname, caption in figures:
        add_picture(doc, fig_dir / "_aggregate" / fname, 6.4)
        add_caption(doc, caption)

    add_paragraph(
        doc,
        "从 overall 对照结果可以看到，在把正式真值收紧为 strict truth 后，Recall 与 MRR 均有小幅下降，"
        "这说明旧 coverage 口径中确实存在一部分“几何上成立但内容上并不适合作为正样本”的 tile。"
        "另一方面，下降幅度有限，说明当前 DINOv2 + FAISS 粗检索仍保留了稳定的区域级检索能力，且原始结果并非完全依赖噪声真值支撑。"
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "7. 代表性案例", 1)
    add_paragraph(
        doc,
        "为使结论不局限于整体均值，本文自动选取显著成功、边界成功和残余失败三类代表性案例，并展示对应 query 的 Top-10 检索结果，以观察当前基线在 strict truth 口径下的具体行为。"
    )
    for idx, case in enumerate(cases, start=1):
        add_case_section(doc, idx, case, fig_dir)

    add_heading(doc, "8. 结论", 1)
    add_paragraph(
        doc,
        "图 5 和图 6 给出本轮实验最关键的两张结论图，分别对应 overall strict 指标和分航线 Recall 表现。"
    )
    add_picture(doc, fig_dir / "_aggregate" / "overall_metrics_bar.png", 6.2)
    add_caption(doc, "图 5  结论图：overall strict truth 指标")
    add_picture(doc, fig_dir / "_aggregate" / "multi_flight_recall.png", 6.2)
    add_caption(doc, "图 6  结论图：分航线 strict truth Recall 表现")
    add_bullets(
        doc,
        [
            "在 `200/300/500/700m` 候选尺度下，DINOv2 + FAISS 在 strict truth 口径上仍具备区域级初步地理定位能力，但指标相较旧 coverage 口径更克制，也更接近真实可用水平。",
            "Strict Truth 的下降主要来自对黑边或低有效内容 tile 的过滤，而不是模型或排序的退化；因此它更适合作为后续正式评估主口径。",
            "Top-1 error mean 基本不变，说明这次变化来自真值净化，而不是检索候选排序变化。",
            "后续若要进一步提升结果，应在 strict truth 口径下继续分析尺度贡献和残余失败样本，而不是回退到更宽松的真值定义。",
        ],
    )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(out_path)


if __name__ == "__main__":
    main()
