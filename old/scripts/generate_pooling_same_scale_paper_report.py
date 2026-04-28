#!/usr/bin/env python3
"""Generate a formal Word technical report for same-scale pooling comparison."""

from __future__ import annotations

import csv
import json
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path("/mnt/d/aiproject/imagematch")
RESULT_ROOT = ROOT / "方案" / "CLS token vs mean pooling vs GeM pooling_200m同尺度"
FIG_ROOT = RESULT_ROOT / "figures"
OUT_PATH = RESULT_ROOT / "DINOv2不同Pooling策略_200m同尺度跨视角粗定位实验结果说明_2026-03-17.docx"


@dataclass
class MethodSpec:
    key: str
    title: str
    summary_title: str


METHODS = [
    MethodSpec("pooler", "DINOv2 + pooler_output", "POOLER"),
    MethodSpec("cls", "DINOv2 + CLS token", "CLS"),
    MethodSpec("mean", "DINOv2 + mean pooling", "MEAN"),
    MethodSpec("gem", "DINOv2 + GeM pooling", "GeM"),
]


def set_cn_font(run, size: int = 11, bold: bool = False) -> None:
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(size)
    run.bold = bold


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 10, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    set_cn_font(r, size=size, bold=bold)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.style = "Heading 1" if level == 1 else "Heading 2"
    r = p.add_run(text)
    set_cn_font(r, size=14 if level == 1 else 12, bold=True)


def add_paragraph(doc: Document, text: str, size: int = 11, center: bool = False) -> None:
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn_font(r, size=size)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn_font(r, size=10)


def add_picture(doc: Document, path: Path, width_inch: float) -> None:
    doc.add_picture(str(path), width=Inches(width_inch))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def load_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def load_aggregate(method_key: str) -> list[dict]:
    return load_json(RESULT_ROOT / method_key / "aggregate_summary.json")["flights"]


def short_flight_name(flight_id: str) -> str:
    parts = flight_id.split("_")
    return parts[2] if len(parts) >= 3 else flight_id


def build_overall_table(doc: Document, overall_rows: list[dict[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=9)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["方法", "R@1", "R@5", "R@10", "MRR", "Top-1误差均值(m)", "特征耗时(ms)", "检索耗时(ms)", "总耗时(ms)"]
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, bold=True)
        shade_cell(table.cell(0, i), "D9EAF7")
    for row in overall_rows:
        cells = table.add_row().cells
        vals = [
            next(m.summary_title for m in METHODS if m.key == row["method"]),
            f"{float(row['recall@1']):.3f}",
            f"{float(row['recall@5']):.3f}",
            f"{float(row['recall@10']):.3f}",
            f"{float(row['mrr']):.3f}",
            f"{float(row['top1_error_m_mean']):.3f}",
            f"{float(row['feature_ms_mean']):.2f}",
            f"{float(row['retrieval_ms_mean']):.3f}",
            f"{float(row['total_ms_mean']):.2f}",
        ]
        for i, val in enumerate(vals):
            set_cell_text(cells[i], val)


def build_per_flight_table(doc: Document, method_data: dict[str, list[dict]]) -> None:
    flights = [item["flight_id"] for item in method_data["pooler"]]
    table = doc.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["航线", "方法", "R@1", "R@5", "R@10", "MRR", "Top-1误差均值(m)"]
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, bold=True)
        shade_cell(table.cell(0, i), "D9EAF7")
    for flight in flights:
        for idx, method in enumerate(METHODS):
            item = next(x for x in method_data[method.key] if x["flight_id"] == flight)
            row = table.add_row().cells
            set_cell_text(row[0], short_flight_name(flight) if idx == 0 else "")
            set_cell_text(row[1], method.summary_title)
            set_cell_text(row[2], f"{item['recall@1']:.3f}")
            set_cell_text(row[3], f"{item['recall@5']:.3f}")
            set_cell_text(row[4], f"{item['recall@10']:.3f}")
            set_cell_text(row[5], f"{item['mrr']:.3f}")
            set_cell_text(row[6], f"{item['top1_error_m_mean']:.3f}")


def first_truth_rank(summary_json: Path, query_id: str) -> int | None:
    data = load_json(summary_json)
    for item in data["per_query"]:
        if item["query_id"] != query_id:
            continue
        rr = float(item.get("reciprocal_rank", 0.0))
        return None if rr == 0 else int(round(1.0 / rr))
    return None


def build_case_rows() -> list[dict[str, str | int | None]]:
    cases = [
        {
            "flight_id": "DJI_202510311347_009_新建面状航线1",
            "query_id": "q_200m_05",
            "note": "稳定成功样例：四种方法均能在严格同尺度条件下把真值推到第 1 名，说明该查询块具有较强的跨视角辨识度。",
        },
        {
            "flight_id": "DJI_202510311500_012_新建面状航线1",
            "query_id": "q_200m_04",
            "note": "方法差异样例：POOLER/CLS 仅能把真值召回到前列，MEAN 与 GeM 可把真值提升到第 1 名，说明 patch-token 聚合在部分困难样本上具备优势。",
        },
        {
            "flight_id": "DJI_202510311435_011_新建面状航线1",
            "query_id": "q_200m_03",
            "note": "困难失败样例：四种方法均未把真值推到第 1 名，其中部分方法甚至未能在 Top-10 内稳定命中，表明严格同尺度条件下仍存在显著跨视角歧义。",
        },
    ]
    rows = []
    for case in cases:
        row = dict(case)
        for method in METHODS:
            summary_path = RESULT_ROOT / method.key / "stage4" / f"{case['flight_id']}_retrieval_top10.json"
            row[f"{method.key}_rank"] = first_truth_rank(summary_path, case["query_id"])
        rows.append(row)
    return rows


def add_case_section(doc: Document, case_idx: int, case: dict[str, str | int | None]) -> None:
    flight_id = str(case["flight_id"])
    query_id = str(case["query_id"])
    add_heading(doc, f"7.{case_idx} {short_flight_name(flight_id)} / {query_id}", 2)
    rank_text = "；".join(
        f"{m.summary_title} 首个真值排名：{case[f'{m.key}_rank'] if case[f'{m.key}_rank'] is not None else '未命中'}"
        for m in METHODS
    )
    add_paragraph(doc, f"{rank_text}。{case['note']}")

    for method in METHODS:
        img = FIG_ROOT / method.key / flight_id / f"{query_id}_top10.png"
        add_picture(doc, img, 5.8)
        add_caption(doc, f"{method.summary_title}：{short_flight_name(flight_id)} / {query_id} 的 Top-10 检索结果")


def main() -> None:
    overall_rows = load_csv(RESULT_ROOT / "overall_metrics.csv")
    method_data = {m.key: load_aggregate(m.key) for m in METHODS}
    case_rows = build_case_rows()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DINOv2 不同 Pooling 策略在 200m 同尺度条件下的跨视角粗定位实验结果说明")
    set_cn_font(r, size=16, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("面向 GNSS 拒止环境下无人机图像初步地理定位的技术报告")
    set_cn_font(r, size=11)

    add_heading(doc, "1. 任务定义", 1)
    add_paragraph(
        doc,
        "本组实验面向无人机影像与遥感正射影像之间的跨视角粗定位任务，目标是在不依赖 GNSS 的条件下，"
        "利用无人机图像在卫星候选库中进行检索，判断真值区域能否稳定进入前列候选。与此前“查询端固定 200m、卫星端允许多尺度候选”的口径不同，"
        "本轮实验进一步采用严格同尺度设置：无人机查询块固定为 200m，卫星候选瓦片亦固定为 200m，以便更清晰地评估不同全局特征聚合策略本身的影响。"
    )

    add_heading(doc, "2. 实验设置", 1)
    add_paragraph(
        doc,
        "实验继续复用既有 validation_round3_200m_fair 查询集，共包含 4 条航线、20 个 200m 查询块。"
        "卫星侧从原始多尺度候选库中筛选出全部 200m 瓦片，共 400 张，用于建立严格同尺度候选库。"
        "所有方法统一采用 DINOv2-base 作为骨干网络，统一使用 L2 归一化和 FAISS IndexFlatIP 检索；"
        "差异仅体现在全局特征的聚合方式，即 pooler_output、CLS token、mean pooling 与 GeM pooling 四种方案。"
    )

    add_heading(doc, "3. 指标定义", 1)
    add_paragraph(
        doc,
        "为保持与既有 PoC 一致的评估口径，本文继续采用 Recall@1、Recall@5、Recall@10、MRR、Top-1 定位误差均值，以及单次特征提取/检索/总耗时。"
        "其中 Recall 系列刻画真值进入前列候选的能力，MRR 描述真值排名靠前的总体趋势，Top-1 定位误差反映当前方法给出的首位候选与真实位置之间的几何偏差。"
    )

    add_heading(doc, "4. 方法说明", 1)
    add_heading(doc, "4.1 DINOv2 + pooler_output", 2)
    add_paragraph(
        doc,
        "该方法直接使用 DINOv2 模型输出中的 pooler_output 作为全局图像描述子，是既有粗检索 PoC 的主基线实现。"
        "其目的在于保留模型原生的全局语义表示能力，并以最少的附加设计完成无人机-卫星跨视角检索。"
    )
    add_heading(doc, "4.2 DINOv2 + CLS token", 2)
    add_paragraph(
        doc,
        "该方法直接取最后一层 Transformer 输出中的 CLS token 作为全局特征，用以检验显式 CLS 表示是否能够在严格同尺度条件下维持与 pooler_output 等价的检索能力。"
    )
    add_heading(doc, "4.3 DINOv2 + mean pooling", 2)
    add_paragraph(
        doc,
        "该方法对最后一层 patch tokens 做均值聚合，以构造全局描述子。其动机在于充分整合局部 token 的整体统计信息，观察此类平均式聚合是否能够缓解尺度固定后产生的局部歧义。"
    )
    add_heading(doc, "4.4 DINOv2 + GeM pooling", 2)
    add_paragraph(
        doc,
        "该方法对 patch tokens 进行 GeM 聚合，指数参数固定为 3.0。GeM 可视为位于均值池化与最大池化之间的一类广义池化方式，适合检验强调高响应局部区域是否有利于严格同尺度条件下的跨视角检索。"
    )

    add_heading(doc, "5. 定量结果", 1)
    add_paragraph(doc, "表 1 给出四种方法在 20 个查询块上的 overall 结果。")
    build_overall_table(doc, overall_rows)
    add_caption(doc, "表 1  200m 同尺度条件下四种 pooling 方法的 overall 结果")

    add_paragraph(doc, "表 2 按航线汇总了四种方法的分组结果，用以观察不同方法在不同航线上的稳定性差异。")
    build_per_flight_table(doc, method_data)
    add_caption(doc, "表 2  200m 同尺度条件下四种 pooling 方法的分航线结果")

    add_heading(doc, "6. 汇总图解读", 1)
    add_paragraph(
        doc,
        "为便于从整体上比较不同方法在召回能力、排序质量、定位误差与时延之间的差异，本文在表格基础上进一步给出汇总图。"
        "图 1 至图 4 分别对应 Recall@1、Recall@5、Recall@10 与 MRR；图 5 对应 Top-1 定位误差；图 6 与图 7 分别对应特征提取耗时与总耗时；图 8 则给出分航线召回表现。"
    )
    figure_names = [
        ("pooling_same_scale_recall1.png", "图 1  200m 同尺度条件下 Recall@1 对比"),
        ("pooling_same_scale_recall5.png", "图 2  200m 同尺度条件下 Recall@5 对比"),
        ("pooling_same_scale_recall10.png", "图 3  200m 同尺度条件下 Recall@10 对比"),
        ("pooling_same_scale_mrr.png", "图 4  200m 同尺度条件下 MRR 对比"),
        ("pooling_same_scale_top1_error.png", "图 5  200m 同尺度条件下 Top-1 定位误差对比"),
        ("pooling_same_scale_feature_ms.png", "图 6  200m 同尺度条件下特征提取耗时对比"),
        ("pooling_same_scale_total_ms.png", "图 7  200m 同尺度条件下总耗时对比"),
        ("multi_flight_recall.png", "图 8  200m 同尺度条件下分航线 Recall 对比"),
    ]
    for fname, cap in figure_names:
        add_picture(doc, FIG_ROOT / "_aggregate" / fname, 6.4)
        add_caption(doc, cap)

    add_paragraph(
        doc,
        "从 overall 结果可以看到，pooler 与 CLS 在严格同尺度条件下仍保持完全一致，这说明在当前任务设置中，"
        "二者的全局表示能力没有表现出可区分差异。与此前多尺度卫星候选库实验不同，本轮同尺度条件下，pooler/CLS 的 Recall@1 明显下降，"
        "表明多尺度候选库在前期实验中实际上对首位命中率提供了显著缓冲。与此同时，mean 与 GeM 的 Recall@1 提升到 0.25，超过了 pooler/CLS 的 0.05；"
        "但其 Recall@10 与整体稳定性并未全面占优，说明 patch-token 聚合在严格同尺度条件下更可能带来“部分样本更优、整体稳定性仍受限”的结果。"
        "进一步看定位误差，mean 的 Top-1 误差均值最低，说明其在成功命中时给出的首位候选更接近真实位置；然而结合 Recall@10 与分航线表现可知，这种优势并未转化为全面更强的粗检索能力。"
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "7. 典型案例", 1)
    add_paragraph(
        doc,
        "为避免仅依据 overall 指标作出过度概括，本文进一步选取稳定成功、方法差异和困难失败三类典型案例进行定性说明。"
        "每个案例均展示四种方法在同一查询块上的 Top-10 检索结果，以便从候选排序层面观察它们的差异。"
    )
    for idx, case in enumerate(case_rows, start=1):
        add_case_section(doc, idx, case)

    add_heading(doc, "8. 结论", 1)
    add_paragraph(
        doc,
        "本轮 200m 同尺度实验表明：当卫星候选库也被严格限制为 200m 时，四种 pooling 方法的行为与先前多尺度候选库口径下存在明显差异。"
        "首先，pooler 与 CLS 依旧保持等价，这一事实在同尺度条件下再次得到验证；其次，mean 与 GeM 在 Recall@1 上优于 pooler/CLS，"
        "说明 patch-token 聚合在若干查询样本上更容易直接给出正确首位候选；但从 Recall@10、分航线结果和困难案例看，这种优势并未扩展为全面更稳的检索性能。"
        "因此，更谨慎的结论是：在严格同尺度 200m 条件下，聚合策略确实会改变跨视角粗检索的表现，但当前证据更适合支持“不同 pooling 对首位排序和误差分布有显著影响”，"
        "而不足以支持“某一种 pooling 在所有维度上全面优于其余方案”。如需进一步形成更稳健的工程结论，后续仍有必要在更大样本量、更多城市区域或叠加几何重排条件下继续验证。"
    )

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(OUT_PATH)


if __name__ == "__main__":
    main()
