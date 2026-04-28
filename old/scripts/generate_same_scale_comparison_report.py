#!/usr/bin/env python3
"""Generate a Word report for strict same-scale three-method comparison."""

from __future__ import annotations

import csv
import json
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path("/mnt/d/aiproject/imagematch")
OUTPUT_ROOT = ROOT / "output"
PLAN_ROOT = ROOT / "方案"


@dataclass
class MethodSpec:
    key: str
    title: str
    result_root: Path
    summary_title: str


METHODS = [
    MethodSpec(
        key="baseline",
        title="DINOv2 + FAISS 粗检索",
        result_root=OUTPUT_ROOT / "validation_200m_same_scale",
        summary_title="基线方法",
    ),
    MethodSpec(
        key="sift",
        title="SIFT + RANSAC 保守门控重排",
        result_root=OUTPUT_ROOT / "validation_200m_same_scale_sift_gate3",
        summary_title="传统局部几何重排",
    ),
    MethodSpec(
        key="lightglue",
        title="SuperPoint + LightGlue 融合重排",
        result_root=OUTPUT_ROOT / "validation_200m_same_scale_lightglue_superpoint_fused_top10_k256",
        summary_title="学习型局部匹配融合重排",
    ),
]

METRICS = [
    ("Recall@1", "真值进入第 1 名的比例，衡量是否能直接给出最优候选。"),
    ("Recall@5", "真值进入前 5 名的比例，衡量前排候选覆盖能力。"),
    ("Recall@10", "真值进入前 10 名的比例，衡量区域级初步定位能力。"),
    ("MRR", "真值排名倒数的平均值，越高说明真值越靠前。"),
    ("Top-1 error mean (m)", "第 1 名候选中心与查询中心的平均距离，越低越好。"),
]


def set_cn_font(run, size: int = 11, bold: bool = False) -> None:
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(size)
    run.bold = bold


def set_cell_text(cell, text: str, bold: bool = False, size: int = 10) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn_font(r, size=size, bold=bold)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn_font(r, size=10)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def load_aggregate(method: MethodSpec) -> list[dict]:
    data = load_json(method.result_root / "aggregate_summary.json")
    return data["flights"]


def compute_overall(flights: list[dict]) -> dict[str, float]:
    total_q = sum(int(x["query_count"]) for x in flights)
    out = {"query_count": total_q}
    for metric in ["recall@1", "recall@5", "recall@10", "mrr", "top1_error_m_mean"]:
        out[metric] = sum(float(x[metric]) * int(x["query_count"]) for x in flights) / total_q
    return out


def first_truth_rank(csv_path: Path, query_id: str) -> int | None:
    with csv_path.open("r", encoding="utf-8", newline="") as f:
        for row in csv.DictReader(f):
            if row["query_id"] != query_id:
                continue
            if row["is_truth_hit"] == "1":
                return int(row["rank"])
    return None


def build_case_rows() -> list[dict[str, str | int]]:
    flights = [
        ("DJI_202510311347_009_新建面状航线1", "q_200m_03", "显著成功样例：基线真值第 6 名，LightGlue 融合后提升到第 1 名。"),
        ("DJI_202510311500_012_新建面状航线1", "q_200m_04", "弱航线改进样例：基线真值第 3 名，LightGlue 融合后提升到第 1 名。"),
        ("DJI_202510311500_012_新建面状航线1", "q_200m_01", "残余失败样例：LightGlue 仍未将真值推到第 1 名，说明跨视角差异仍然存在。"),
    ]
    rows = []
    for flight_id, query_id, note in flights:
        rows.append(
            {
                "flight_id": flight_id,
                "query_id": query_id,
                "note": note,
                "baseline_rank": first_truth_rank(
                    OUTPUT_ROOT / "validation_200m_same_scale" / "stage4" / flight_id / "retrieval_top10.csv",
                    query_id,
                ),
                "sift_rank": first_truth_rank(
                    OUTPUT_ROOT / "validation_200m_same_scale_sift_gate3" / "stage7" / flight_id / "reranked_top10.csv",
                    query_id,
                ),
                "lightglue_rank": first_truth_rank(
                    OUTPUT_ROOT
                    / "validation_200m_same_scale_lightglue_superpoint_fused_top10_k256"
                    / "stage7"
                    / flight_id
                    / "reranked_top10.csv",
                    query_id,
                ),
            }
        )
    return rows


def add_heading_text(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    style = "Heading 1" if level == 1 else "Heading 2"
    p.style = style
    r = p.add_run(text)
    set_cn_font(r, size=14 if level == 1 else 12, bold=True)


def add_normal_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.style = "Normal"
    r = p.add_run(text)
    set_cn_font(r, size=11)


def add_bullets(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(line)
        set_cn_font(r, size=11)


def add_metric_definition_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_cell_text(table.cell(0, 0), "指标", bold=True)
    set_cell_text(table.cell(0, 1), "含义", bold=True)
    shade_cell(table.cell(0, 0), "D9EAF7")
    shade_cell(table.cell(0, 1), "D9EAF7")
    for name, desc in METRICS:
        row = table.add_row().cells
        set_cell_text(row[0], name)
        row[1].text = ""
        p = row[1].paragraphs[0]
        r = p.add_run(desc)
        set_cn_font(r, size=10)


def add_overall_table(doc: Document, method_data: dict[str, list[dict]]) -> None:
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["方法", "Recall@1", "Recall@5", "Recall@10", "MRR", "Top-1 error mean (m)"]
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, bold=True)
        shade_cell(table.cell(0, i), "D9EAF7")
    for method in METHODS:
        overall = compute_overall(method_data[method.key])
        row = table.add_row().cells
        set_cell_text(row[0], method.title)
        set_cell_text(row[1], f"{overall['recall@1']:.3f}")
        set_cell_text(row[2], f"{overall['recall@5']:.3f}")
        set_cell_text(row[3], f"{overall['recall@10']:.3f}")
        set_cell_text(row[4], f"{overall['mrr']:.3f}")
        set_cell_text(row[5], f"{overall['top1_error_m_mean']:.3f}")


def add_per_flight_table(doc: Document, method_data: dict[str, list[dict]]) -> None:
    flights = [x["flight_id"] for x in method_data["baseline"]]
    table = doc.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["航线", "方法", "Recall@1", "Recall@5", "Recall@10", "MRR", "Top-1 error mean (m)"]
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, bold=True)
        shade_cell(table.cell(0, i), "D9EAF7")
    for flight in flights:
        for idx, method in enumerate(METHODS):
            item = next(x for x in method_data[method.key] if x["flight_id"] == flight)
            row = table.add_row().cells
            set_cell_text(row[0], flight if idx == 0 else "")
            set_cell_text(row[1], method.title)
            set_cell_text(row[2], f"{item['recall@1']:.3f}")
            set_cell_text(row[3], f"{item['recall@5']:.3f}")
            set_cell_text(row[4], f"{item['recall@10']:.3f}")
            set_cell_text(row[5], f"{item['mrr']:.3f}")
            set_cell_text(row[6], f"{item['top1_error_m_mean']:.3f}")


def add_picture(doc: Document, path: Path, width_inch: float) -> None:
    doc.add_picture(str(path), width=Inches(width_inch))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def main() -> None:
    method_data = {m.key: load_aggregate(m) for m in METHODS}
    case_rows = build_case_rows()

    out_path = PLAN_ROOT / "严格同尺度三方法对比实验结果解读_2026-03-17.docx"
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("严格同尺度三方法对比实验结果解读")
    set_cn_font(r, size=16, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("面向无人机-卫星跨视角初步地理定位（检索）的正式结果说明")
    set_cn_font(r, size=11)

    add_heading_text(doc, "1. 任务定义与实验设置", 1)
    add_normal_paragraph(
        doc,
        "本组实验用于论证：在跨视角条件下，仅依赖遥感正射影像，是否能够把无人机图像检索到正确地理区域附近。"
        "为避免尺度混杂带来的解释歧义，实验统一采用严格同尺度口径：无人机查询块固定为 200m，卫星候选库固定为 200m，"
        "并统一 resize 到同一网络输入分辨率；真值定义为查询中心点落入的 200m 卫星瓦片。",
    )
    add_bullets(
        doc,
        [
            "数据范围：4 条航线，共 20 个 200m 查询块。",
            "基线方法：DINOv2 全局特征 + FAISS 粗检索。",
            "重排方法 1：SIFT + RANSAC 保守门控重排。",
            "重排方法 2：SuperPoint + LightGlue，并采用全局分数与几何分数融合排序。",
        ],
    )

    add_heading_text(doc, "2. 指标定义", 1)
    add_metric_definition_table(doc)

    add_heading_text(doc, "3. 方法说明", 1)
    add_heading_text(doc, "3.1 DINOv2 + FAISS 粗检索", 2)
    add_normal_paragraph(
        doc,
        "首先对无人机查询块与 200m 卫星瓦片提取 DINOv2 全局特征，再使用 FAISS 进行 Top-K 相似度检索。"
        "该方法只利用全局表征，不做局部匹配与几何验证，是后续重排方法的统一基线。",
    )
    add_heading_text(doc, "3.2 SIFT + RANSAC 保守门控重排", 2)
    add_normal_paragraph(
        doc,
        "在基线粗检索结果上，对候选进行 SIFT 局部特征匹配与 RANSAC 几何验证。"
        "只有在候选位于前列且几何证据满足阈值时才允许前移，因此属于保守门控重排。"
        "这一方案主要用于验证传统局部几何验证在跨视角场景下是否足以改进排序。",
    )
    add_heading_text(doc, "3.3 SuperPoint + LightGlue 融合重排", 2)
    add_normal_paragraph(
        doc,
        "在同样的粗检索候选集上，使用 SuperPoint 提取学习型局部特征，并用 LightGlue 完成匹配，"
        "随后通过 RANSAC 估计几何一致性。与前一方案不同，本轮不是采用简单硬门控，而是把全局相似度、"
        "内点数、内点比例与重投影误差共同映射为融合得分，从而在保留粗检索先验的同时，引入更可靠的局部几何约束。",
    )

    add_heading_text(doc, "4. 定量结果", 1)
    add_normal_paragraph(doc, "表 1 给出三种方法在 20 个查询块上的 overall 对比结果。")
    add_overall_table(doc, method_data)
    add_caption(doc, "表 1  严格同尺度口径下三种方法的 overall 指标对比")

    add_normal_paragraph(doc, "表 2 给出按航线拆分的详细结果。")
    add_per_flight_table(doc, method_data)
    add_caption(doc, "表 2  四条航线上的分方法指标对比")

    add_heading_text(doc, "5. 汇总图解读", 1)
    add_picture(
        doc,
        OUTPUT_ROOT
        / "validation_200m_same_scale_lightglue_superpoint_fused_top10_k256"
        / "figures"
        / "_aggregate"
        / "baseline_vs_sift_vs_lightglue_fused_recall1.png",
        6.4,
    )
    add_caption(doc, "图 1  三方法 Recall@1 对比")
    add_picture(
        doc,
        OUTPUT_ROOT
        / "validation_200m_same_scale_lightglue_superpoint_fused_top10_k256"
        / "figures"
        / "_aggregate"
        / "baseline_vs_sift_vs_lightglue_fused_recall5.png",
        6.4,
    )
    add_caption(doc, "图 2  三方法 Recall@5 对比")
    add_picture(
        doc,
        OUTPUT_ROOT
        / "validation_200m_same_scale_lightglue_superpoint_fused_top10_k256"
        / "figures"
        / "_aggregate"
        / "baseline_vs_sift_vs_lightglue_fused_mrr.png",
        6.4,
    )
    add_caption(doc, "图 3  三方法 MRR 对比")

    add_normal_paragraph(
        doc,
        "从 overall 结果看，基线方法已经能够在严格同尺度条件下把真值稳定召回到前 10 名，"
        "说明仅依赖遥感正射影像完成区域级初步地理定位是可行的；但其 Recall@1 较低，表明粗检索本身难以稳定给出最优候选。"
        "SIFT + RANSAC 只在个别样例上起作用，整体提升有限。相比之下，SuperPoint + LightGlue 融合重排在 Recall@1、"
        "Recall@5、MRR 上均表现出显著优势，说明学习型局部匹配与全局分数融合可以有效提升跨视角检索排序质量。",
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading_text(doc, "6. 代表性案例", 1)
    for idx, case in enumerate(case_rows, start=1):
        add_heading_text(doc, f"6.{idx} {case['flight_id']} / {case['query_id']}", 2)
        add_normal_paragraph(
            doc,
            f"基线首个真值排名：{case['baseline_rank']}；"
            f"SIFT 首个真值排名：{case['sift_rank']}；"
            f"LightGlue 融合首个真值排名：{case['lightglue_rank']}。{case['note']}",
        )
        flight = str(case["flight_id"])
        qid = str(case["query_id"])
        base_img = OUTPUT_ROOT / "validation_200m_same_scale" / "figures" / flight / f"{qid}_top5.png"
        sift_img = OUTPUT_ROOT / "validation_200m_same_scale_sift_gate3" / "figures" / flight / f"{qid}_top5.png"
        lg_img = (
            OUTPUT_ROOT
            / "validation_200m_same_scale_lightglue_superpoint_fused_top10_k256"
            / "figures"
            / flight
            / f"{qid}_top10.png"
        )
        add_picture(doc, base_img, 5.8)
        add_caption(doc, f"图 {idx * 3 - 2}  基线结果：{flight} / {qid}")
        add_picture(doc, sift_img, 5.8)
        add_caption(doc, f"图 {idx * 3 - 1}  SIFT 重排结果：{flight} / {qid}")
        add_picture(doc, lg_img, 5.8)
        add_caption(doc, f"图 {idx * 3}  LightGlue 融合重排结果：{flight} / {qid}")

    add_heading_text(doc, "7. 结论", 1)
    add_bullets(
        doc,
        [
            "在严格 200m 同尺度条件下，DINOv2 + FAISS 已经能够把无人机图像检索到正确地理区域附近，主命题成立。",
            "单纯依赖传统 SIFT + RANSAC 保守门控重排，无法稳定优于全局粗检索基线。",
            "采用 SuperPoint + LightGlue，并使用全局分数与几何分数融合排序后，Top-1 与前排排序质量显著提升，成为当前最优方案。",
            "因此，更准确的结论是：遥感正射影像能够支撑无人机影像的区域级初步定位，而学习型局部几何重排可进一步提升最终候选质量。",
        ],
    )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(out_path)


if __name__ == "__main__":
    main()
