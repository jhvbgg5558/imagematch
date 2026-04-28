#!/usr/bin/env python3
"""Generate a Word report for the satellite-truth validation suite.

Purpose:
- summarize the satellite-truth validation run from the latest suite outputs;
- keep the report aligned with the isolated new3output experiment root;
- describe the satellite-truth evaluation as offline validation only.

Main inputs:
- satellite-truth suite summaries and per-query/per-flight CSV/JSON files;
- experiment-level query selection and selected-truth manifests;
- representative metrics from the orthophoto-alignment, geometry, and
  tie-point layers.

Main outputs:
- `<suite-root>/reports/formal_pose_v1_validation_suite_satellite_truth_report.docx`

Applicable task constraints:
- the report must describe satellite truth as a source-crop validation layer;
- fixed tiles are selection anchors only and must not be described as the final
  truth asset;
- layer-1 / layer-2 / layer-3 meanings must remain consistent with the new
  satellite-truth suite.
"""

from __future__ import annotations

import argparse
import csv
import json
import math
from collections import Counter
from pathlib import Path
from statistics import mean

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from satellite_truth_utils import DEFAULT_BUNDLE_ROOT, DEFAULT_SUITE_DIRNAME, shorten_flight_id, resolve_satellite_suite_root


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--suite-root", default=str(DEFAULT_BUNDLE_ROOT / DEFAULT_SUITE_DIRNAME))
    parser.add_argument(
        "--out-docx",
        default="",
        help="Optional explicit output .docx path. Defaults to <suite-root>/reports/formal_pose_v1_validation_suite_satellite_truth_report.docx",
    )
    return parser.parse_args()


def load_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", newline="", encoding="utf-8-sig") as handle:
        return list(csv.DictReader(handle))


def load_json(path: Path) -> dict[str, object]:
    return json.loads(path.read_text(encoding="utf-8"))


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def set_cn_font(run, size: int = 11, bold: bool = False) -> None:
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(size)
    run.bold = bold


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = "Heading 1" if level == 1 else "Heading 2"
    run = paragraph.add_run(text)
    set_cn_font(run, size=14 if level == 1 else 12, bold=True)


def add_paragraph(doc: Document, text: str, *, size: int = 11, center: bool = False, bold: bool = False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_cn_font(run, size=size, bold=bold)


def add_bullets(doc: Document, lines: list[str]) -> None:
    for line in lines:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(line)
        set_cn_font(run, size=11)


def set_cell_text(cell, text: str, *, bold: bool = False, size: int = 10) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_cn_font(run, size=size, bold=bold)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], *, header_fill: str = "D9EAF7") -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_text(table.cell(0, idx), header, bold=True)
        shade_cell(table.cell(0, idx), header_fill)
    for values in rows:
        row = table.add_row().cells
        for idx, value in enumerate(values):
            set_cell_text(row[idx], value)


def fmt_float(value: float | int | None, digits: int = 3) -> str:
    if value is None:
        return "-"
    numeric = float(value)
    if not math.isfinite(numeric):
        return "-"
    return f"{numeric:.{digits}f}"


def parse_float(value: str | None) -> float | None:
    if value is None or value == "":
        return None
    return float(value)


def infer_experiment_root(suite_root: Path) -> Path:
    pose_root = suite_root.parent
    if pose_root.name != "pose_v1_formal":
        raise SystemExit(f"unexpected suite root layout: {suite_root}")
    return pose_root.parent


def summarize_selected_queries(selected_rows: list[dict[str, str]]) -> dict[str, object]:
    flight_counts = Counter(row["flight_id"] for row in selected_rows)
    ordered_flights = sorted(flight_counts)
    pitch_values = [float(row["gimbal_pitch_degree"]) for row in selected_rows if row.get("gimbal_pitch_degree")]
    return {
        "query_count": len(selected_rows),
        "flight_counts": dict(flight_counts),
        "ordered_flights": ordered_flights,
        "pitch_min": min(pitch_values) if pitch_values else None,
        "pitch_max": max(pitch_values) if pitch_values else None,
    }


def metric_payload(summary: dict[str, object], key: str) -> dict[str, float]:
    payload = summary.get(key, {})
    return payload if isinstance(payload, dict) else {}


def main() -> None:
    args = parse_args()
    suite_root = Path(args.suite_root).resolve()
    out_docx = Path(args.out_docx).resolve() if args.out_docx else suite_root / "reports" / "formal_pose_v1_validation_suite_satellite_truth_report.docx"
    pose_root = suite_root.parent
    experiment_root = infer_experiment_root(suite_root)
    ensure_dir(out_docx.parent)

    selected_rows = []
    selected_summary_csv = experiment_root / "selected_queries" / "selected_images_summary.csv"
    if selected_summary_csv.exists():
        selected_rows = load_csv(selected_summary_csv)
    truth_manifest_csv = suite_root / "satellite_truth" / "query_satellite_truth_manifest.csv"
    truth_rows = load_csv(truth_manifest_csv) if truth_manifest_csv.exists() else []
    query_summary = summarize_selected_queries(selected_rows or truth_rows)

    suite_summary = load_json(suite_root / "full_run_summary.json") if (suite_root / "full_run_summary.json").exists() else {}
    ortho_overall = load_json(suite_root / "ortho_alignment_satellite" / "overall_ortho_accuracy.json") if (suite_root / "ortho_alignment_satellite" / "overall_ortho_accuracy.json").exists() else {}
    geom_overall = load_json(suite_root / "pose_vs_satellite_truth_geometry" / "overall_satellite_truth_geometry.json") if (suite_root / "pose_vs_satellite_truth_geometry" / "overall_satellite_truth_geometry.json").exists() else {}
    tie_overall = load_json(suite_root / "tiepoint_ground_error_satellite" / "overall_tiepoint_ground_error.json") if (suite_root / "tiepoint_ground_error_satellite" / "overall_tiepoint_ground_error.json").exists() else {}

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("009/010 卫星 Truth 验证报告")
    set_cn_font(title_run, size=16, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(str(suite_root))
    set_cn_font(subtitle_run, size=10)

    add_heading(doc, "1. 实验目的", 1)
    add_paragraph(
        doc,
        "本报告用于验证当前 UAV 对卫星候选库定位流程在卫星影像 truth 口径下的稳定性。truth 采用按 query footprint / 扩展 bbox 从原始卫星 GeoTIFF 裁出的 source-crop，不把 fixed tile 直接当作最终 truth，也不使用 top-k 拼接结果冒充 truth。",
    )

    add_heading(doc, "2. 评估方法和评估指标介绍", 1)
    add_bullets(
        doc,
        [
            "Layer-1: predicted ortho vs satellite truth patch，关注 phase correlation、center offset、IoU、NCC、SSIM。",
            "Layer-2: pose vs satellite truth geometry，关注 best pose 相机中心是否落在 truth crop 内，以及相机中心到 truth crop 中心的偏移。",
            "Layer-3: satellite truth patch 上的 tie-point ground XY error，关注 tie-point 的平均误差、RMSE、P90 和 inlier ratio。",
            "runtime 仍保持不变：query 仅用于检索与 PnP，卫星候选库仍来自固定 satellite library，truth 只用于离线验证。",
        ],
    )
    add_table(
        doc,
        ["指标", "说明"],
        [
            ["phase_corr_error_m", "truth 与 predicted ortho 的相位相关平移误差"],
            ["center_offset_m", "truth 与 predicted ortho 的质心偏移"],
            ["camera_center_offset_m", "best pose 相机中心到 truth crop 中心的偏移"],
            ["tiepoint_xy_error_rmse_m", "truth vs pred tie-point 的平面 XY RMSE"],
        ],
    )

    add_heading(doc, "3. 实验流程与数据准备", 1)
    add_bullets(
        doc,
        [
            f"query 规模：{query_summary['query_count']} 张，航线为 {', '.join(shorten_flight_id(f) for f in query_summary['flight_counts'])}。",
            f"query 选择：每条航线 20 张，`gimbal_pitch_degree <= -85.0`。",
            "truth 生成：先从 `query_truth_tiles.csv` 为每个 query 选择一个 canonical satellite source row，再用该 row 的 `source_tif` 按 query footprint + margin 裁 truth patch。",
            "truth 选择规则：strict truth 优先，其次 coverage_ratio、valid_pixel_ratio，最后 black_pixel_ratio / tile_size_m / tile_id 作为稳定排序。",
            "runtime 复用：DINOv2 coarse、RoMa rerank、PnP best pose 均直接复用现有 formal 结果，不在 truth 子链里重跑。",
            "输出根：所有 satellite truth 输出落在 `pose_v1_formal/eval_pose_validation_suite_satellite_truth/`。",
        ],
    )

    add_heading(doc, "4. 实验结果", 1)
    add_table(
        doc,
        ["项目", "结果"],
        [
            ["truth patch 数", str(len(truth_rows))],
            ["Layer-1 状态", ", ".join(f"{k}={v}" for k, v in (ortho_overall.get("status_counts", {}) or {}).items())],
            ["Layer-2 状态", ", ".join(f"{k}={v}" for k, v in (geom_overall.get("status_counts", {}) or {}).items())],
            ["Layer-3 状态", ", ".join(f"{k}={v}" for k, v in (tie_overall.get("status_counts", {}) or {}).items())],
            ["Phase corr mean (m)", fmt_float(metric_payload(ortho_overall, "phase_corr_error_m").get("mean"), 4)],
            ["Phase corr p90 (m)", fmt_float(metric_payload(ortho_overall, "phase_corr_error_m").get("p90"), 4)],
            ["Camera center offset mean (m)", fmt_float(metric_payload(geom_overall, "camera_center_offset_m").get("mean"), 4)],
            ["Tiepoint XY RMSE (m)", fmt_float(tie_overall.get("tiepoint_xy_error_rmse_m"), 4)],
        ],
    )
    if suite_summary:
        steps = suite_summary.get("steps", [])
        runtime_total_sec = sum(float(item.get("elapsed_sec", 0.0)) for item in steps) if isinstance(steps, list) else 0.0
        add_paragraph(doc, f"总运行时间（按步骤累积）约为 {fmt_float(runtime_total_sec / 60.0, 2)} 分钟。", size=10)

    add_heading(doc, "5. 结论与结果分析", 1)
    add_bullets(
        doc,
        [
            "satellite truth 子链已把 truth 定义从 fixed tile 级别下沉到 source GeoTIFF crop 级别，更适合做独立验证。",
            "Layer-1 和 Layer-3 直接反映最终地面落点质量；Layer-2 只作为几何诊断，不应被解释为绝对 pose 真值。",
            "如果后续要比较不同 truth 定义，应保持 retrieval / RoMa / PnP runtime 不变，只切换 truth 口径。",
        ],
    )

    add_heading(doc, "6. 后续的想法", 1)
    add_bullets(
        doc,
        [
            "继续对比新旧 ODM truth / SRTM DSM / satellite truth 三套口径的差异。",
            "在 satellite truth 上补一个异常 query 清单，优先检查 alignment 失败和 tie-point 退化样本。",
            "如果需要更强的判别力，再把 geometry layer 细分为 crop coverage、center offset 和 candidate anchoring 三类诊断。",
        ],
    )

    doc.save(str(out_docx))
    print(out_docx)


if __name__ == "__main__":
    main()
