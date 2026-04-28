#!/usr/bin/env python3
"""Generate comparison reports for the new3output ODM-truth and satellite-truth suites.

Purpose:
- compare the new ODM-truth validation outputs against the historical new2output
  baseline and the parallel satellite-truth validation outputs;
- produce one Markdown report and one Word report under the new3output
  experiment `reports/` directory;
- keep the comparison focused on metric deltas without changing runtime assets.

Main inputs:
- baseline experiment outputs under `new2output/nadir_009010_dinov2_romav2_pose_2026-04-10`;
- current experiment outputs under
  `new3output/nadir_009010_dinov2_romav2_pose_odmrefresh_sattruth_2026-04-16`;
- suite-level JSON/CSV summaries from the ODM-truth and satellite-truth suites.

Main outputs:
- `reports/odm_truth_vs_satellite_truth_comparison.md`
- `reports/odm_truth_vs_satellite_truth_comparison.docx`

Applicable task constraints:
- runtime satellite retrieval assets remain unchanged and must be described as
  fixed across the comparison;
- ODM truth replaces only evaluation orthophoto truth, and ODM DSM replaces
  only PnP height support;
- satellite truth remains an independent validation view rather than a runtime
  input.
"""

from __future__ import annotations

import argparse
import csv
import json
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_BASELINE_ROOT = PROJECT_ROOT / "new2output" / "nadir_009010_dinov2_romav2_pose_2026-04-10"
DEFAULT_EXPERIMENT_ROOT = (
    PROJECT_ROOT / "new3output" / "nadir_009010_dinov2_romav2_pose_odmrefresh_sattruth_2026-04-16"
)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--baseline-root", default=str(DEFAULT_BASELINE_ROOT))
    parser.add_argument("--experiment-root", default=str(DEFAULT_EXPERIMENT_ROOT))
    parser.add_argument("--out-md", default="")
    parser.add_argument("--out-docx", default="")
    return parser.parse_args()


def load_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", newline="", encoding="utf-8-sig") as handle:
        return list(csv.DictReader(handle))


def load_json(path: Path) -> dict[str, object]:
    return json.loads(path.read_text(encoding="utf-8"))


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def fmt(value: object, digits: int = 4) -> str:
    if value in ("", None):
        return "-"
    try:
        return f"{float(value):.{digits}f}"
    except (TypeError, ValueError):
        return str(value)


def set_cn_font(run, size: int = 11, bold: bool = False) -> None:
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(size)
    run.bold = bold


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_paragraph(doc: Document, text: str, *, bold: bool = False, size: int = 11) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_cn_font(run, size=size, bold=bold)


def add_heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_cn_font(run, size=14, bold=True)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        cell.text = ""
        run = cell.paragraphs[0].add_run(header)
        set_cn_font(run, size=10, bold=True)
        shade_cell(cell, "D9EAF7")
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].text = ""
            run = cells[idx].paragraphs[0].add_run(value)
            set_cn_font(run, size=10)


def summarize_pose_root(pose_root: Path) -> dict[str, object]:
    pnp_rows = load_csv(pose_root / "pnp" / "pnp_results.csv")
    best_rows = load_csv(pose_root / "summary" / "per_query_best_pose.csv")
    pnp_counts = Counter(row.get("status", "") for row in pnp_rows)
    best_counts = Counter(row.get("best_status", "") for row in best_rows)
    return {
        "pnp_row_count": len(pnp_rows),
        "best_row_count": len(best_rows),
        "pnp_status_counts": dict(pnp_counts),
        "best_status_counts": dict(best_counts),
    }


def suite_metric_rows(title: str, suite: dict[str, object], metrics: list[tuple[str, str]]) -> list[list[str]]:
    rows: list[list[str]] = []
    for key, label in metrics:
        rows.append([title, label, fmt(suite.get(key))])
    return rows


def main() -> None:
    args = parse_args()
    baseline_root = Path(args.baseline_root)
    experiment_root = Path(args.experiment_root)
    reports_root = experiment_root / "reports"
    ensure_dir(reports_root)

    out_md = Path(args.out_md) if args.out_md else reports_root / "odm_truth_vs_satellite_truth_comparison.md"
    out_docx = Path(args.out_docx) if args.out_docx else reports_root / "odm_truth_vs_satellite_truth_comparison.docx"

    baseline_pose_root = baseline_root / "pose_v1_formal"
    current_pose_root = experiment_root / "pose_v1_formal"
    baseline_suite_root = baseline_pose_root / "eval_pose_validation_suite"
    odm_suite_root = current_pose_root / "eval_pose_validation_suite_odm_truth"
    sat_suite_root = current_pose_root / "eval_pose_validation_suite_satellite_truth"

    baseline_pose = summarize_pose_root(baseline_pose_root)
    current_pose = summarize_pose_root(current_pose_root)

    baseline_ortho = load_json(baseline_suite_root / "ortho_alignment" / "overall_ortho_accuracy.json")
    baseline_pose_vs_at = load_json(baseline_suite_root / "pose_vs_at" / "overall_pose_vs_at.json")
    baseline_tie = load_json(baseline_suite_root / "tiepoint_ground_error" / "overall_tiepoint_ground_error.json")

    odm_ortho = load_json(odm_suite_root / "ortho_alignment" / "overall_ortho_accuracy.json")
    odm_pose_vs_at = load_json(odm_suite_root / "pose_vs_at" / "overall_pose_vs_at.json")
    odm_tie = load_json(odm_suite_root / "tiepoint_ground_error" / "overall_tiepoint_ground_error.json")

    sat_ortho = load_json(sat_suite_root / "ortho_alignment_satellite" / "overall_ortho_accuracy.json")
    sat_geom = load_json(sat_suite_root / "pose_vs_satellite_truth_geometry" / "overall_satellite_truth_geometry.json")
    sat_tie = load_json(sat_suite_root / "tiepoint_ground_error_satellite" / "overall_tiepoint_ground_error.json")

    md_lines = [
        "# ODM Truth vs Satellite Truth Comparison",
        "",
        "## Scope",
        "",
        f"- Baseline root: `{baseline_root}`",
        f"- Current experiment root: `{experiment_root}`",
        "- Runtime retrieval and candidate DOM library remained fixed to the satellite library.",
        "- Main variable changes were: ODM orthophoto truth replacement, ODM-derived DSM replacement, and an additional satellite-truth validation view.",
        "",
        "## Pose Runtime",
        "",
        f"- Baseline PnP rows: `{baseline_pose['pnp_row_count']}`, status counts: `{baseline_pose['pnp_status_counts']}`",
        f"- Current PnP rows: `{current_pose['pnp_row_count']}`, status counts: `{current_pose['pnp_status_counts']}`",
        f"- Baseline best-pose rows: `{baseline_pose['best_row_count']}`, status counts: `{baseline_pose['best_status_counts']}`",
        f"- Current best-pose rows: `{current_pose['best_row_count']}`, status counts: `{current_pose['best_status_counts']}`",
        "",
        "## Validation Summary",
        "",
        "| Suite | Metric | Value |",
        "| --- | --- | --- |",
    ]

    summary_rows = []
    summary_rows.extend(suite_metric_rows("baseline_uav_truth", baseline_ortho, [("phase_corr_error_m_mean", "layer1 phase_corr_error_m mean"), ("ortho_iou_mean", "layer1 ortho_iou mean"), ("ssim_mean", "layer1 ssim mean")]))
    summary_rows.extend(suite_metric_rows("baseline_uav_truth", baseline_pose_vs_at, [("horizontal_error_m_mean", "layer2 horizontal_error_m mean"), ("view_dir_angle_error_deg_mean", "layer2 view_dir_angle_error_deg mean")]))
    summary_rows.extend(suite_metric_rows("baseline_uav_truth", baseline_tie, [("tiepoint_xy_error_rmse_m", "layer3 tiepoint_xy_error_rmse_m"), ("tiepoint_xy_error_p90_m", "layer3 tiepoint_xy_error_p90_m")]))
    summary_rows.extend(suite_metric_rows("odm_truth_refresh", odm_ortho, [("phase_corr_error_m_mean", "layer1 phase_corr_error_m mean"), ("ortho_iou_mean", "layer1 ortho_iou mean"), ("ssim_mean", "layer1 ssim mean")]))
    summary_rows.extend(suite_metric_rows("odm_truth_refresh", odm_pose_vs_at, [("horizontal_error_m_mean", "layer2 horizontal_error_m mean"), ("view_dir_angle_error_deg_mean", "layer2 view_dir_angle_error_deg mean")]))
    summary_rows.extend(suite_metric_rows("odm_truth_refresh", odm_tie, [("tiepoint_xy_error_rmse_m", "layer3 tiepoint_xy_error_rmse_m"), ("tiepoint_xy_error_p90_m", "layer3 tiepoint_xy_error_p90_m")]))
    summary_rows.extend(suite_metric_rows("satellite_truth", sat_ortho, [("phase_corr_error_m_mean", "layer1 phase_corr_error_m mean"), ("ssim_mean", "layer1 ssim mean")]))
    summary_rows.extend(suite_metric_rows("satellite_truth", sat_geom, [("truth_patch_center_offset_m_mean", "layer2 truth_patch_center_offset_m mean"), ("truth_covering_tile_rank_mean", "layer2 truth_covering_tile_rank mean")]))
    summary_rows.extend(suite_metric_rows("satellite_truth", sat_tie, [("tiepoint_xy_error_rmse_m", "layer3 tiepoint_xy_error_rmse_m"), ("tiepoint_xy_error_p90_m", "layer3 tiepoint_xy_error_p90_m")]))

    for row in summary_rows:
        md_lines.append(f"| {row[0]} | {row[1]} | {row[2]} |")

    md_lines.extend(
        [
            "",
            "## Interpretation",
            "",
            "- `odm_truth_refresh` keeps the original three-layer suite semantics but swaps the truth orthophoto source and the runtime DSM source.",
            "- `satellite_truth` is an independent cross-check. Its layer-2 result is a geometry diagnostic relative to the truth patch rather than a pose-vs-AT comparison.",
            "",
        ]
    )
    out_md.write_text("\n".join(md_lines) + "\n", encoding="utf-8")

    doc = Document()
    add_heading(doc, "ODM Truth vs Satellite Truth Comparison")
    add_paragraph(doc, f"Baseline root: {baseline_root}")
    add_paragraph(doc, f"Current experiment root: {experiment_root}")
    add_paragraph(doc, "Runtime retrieval and candidate DOM library remained fixed to the satellite library.")
    add_heading(doc, "Pose Runtime")
    add_table(
        doc,
        ["Run", "PnP rows", "PnP status counts", "Best rows", "Best status counts"],
        [
            [
                "baseline",
                str(baseline_pose["pnp_row_count"]),
                json.dumps(baseline_pose["pnp_status_counts"], ensure_ascii=False),
                str(baseline_pose["best_row_count"]),
                json.dumps(baseline_pose["best_status_counts"], ensure_ascii=False),
            ],
            [
                "current",
                str(current_pose["pnp_row_count"]),
                json.dumps(current_pose["pnp_status_counts"], ensure_ascii=False),
                str(current_pose["best_row_count"]),
                json.dumps(current_pose["best_status_counts"], ensure_ascii=False),
            ],
        ],
    )
    add_heading(doc, "Validation Summary")
    add_table(doc, ["Suite", "Metric", "Value"], summary_rows)
    add_heading(doc, "Interpretation")
    add_paragraph(doc, "ODM truth refresh keeps the UAV orthophoto-truth suite structure but replaces the truth orthophoto source and the PnP DSM source.")
    add_paragraph(doc, "Satellite truth is an independent cross-check; its layer-2 is a geometry diagnostic against the truth patch rather than a pose-vs-AT metric.")
    ensure_dir(out_docx.parent)
    doc.save(out_docx)
    print(out_md)
    print(out_docx)


if __name__ == "__main__":
    main()
