#!/usr/bin/env python3
"""Generate the CaiWangCun full-replacement full-run Markdown and Word report.

Purpose:
- summarize the 40-query CaiWangCun DOM/DSM full-replacement run;
- compare full-run stability against the accepted 5-query gate;
- include full validation metrics, abnormal-query buckets, and representative
  predicted-ortho / tiepoint / frame-sanity images.

Main inputs:
- `<full-root>/pose_v1_formal/summary/pose_overall_summary.json`;
- `<full-root>/pose_v1_formal/eval_pose_validation_suite_caiwangcun_truth/full_run_summary.json`;
- gate summaries from the accepted CaiWangCun full-replacement gate root.

Main outputs:
- `<full-root>/reports/caiwangcun_fullreplace_full_report.md`;
- `<full-root>/reports/caiwangcun_fullreplace_full_report.docx`;
- `<full-root>/reports/assets/`.

Applicable task constraints:
- the report documents offline validation and must not describe truth as a
  runtime input;
- the report must keep the full-replacement constraint explicit: no ODM LAZ,
  SRTM, or old satellite candidate fallback.
"""

from __future__ import annotations

import argparse
import csv
import json
import math
import shutil
from pathlib import Path
from typing import Any

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

try:
    from PIL import Image
except ImportError:  # pragma: no cover
    Image = None


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_FULL_ROOT = (
    PROJECT_ROOT / "new3output" / "nadir_009010_caiwangcun_domdsm_fullreplace_full_2026-04-21"
)
DEFAULT_GATE_ROOT = (
    PROJECT_ROOT / "new3output" / "nadir_009010_caiwangcun_domdsm_fullreplace_gate_2026-04-20"
)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--experiment-root", default=str(DEFAULT_FULL_ROOT))
    parser.add_argument("--gate-root", default=str(DEFAULT_GATE_ROOT))
    parser.add_argument("--out-md", default="")
    parser.add_argument("--out-docx", default="")
    parser.add_argument("--assets-dir", default="")
    return parser.parse_args()


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def load_json(path: Path) -> dict[str, Any]:
    if not path.exists():
        raise FileNotFoundError(f"Missing required JSON: {path}")
    return json.loads(path.read_text(encoding="utf-8"))


def load_csv(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        raise FileNotFoundError(f"Missing required CSV: {path}")
    with path.open("r", newline="", encoding="utf-8-sig") as handle:
        return list(csv.DictReader(handle))


def count_csv_rows(path: Path) -> int:
    if not path.exists():
        return 0
    with path.open("r", encoding="utf-8-sig", errors="ignore") as handle:
        return max(sum(1 for _ in handle) - 1, 0)


def fmt(value: Any, digits: int = 3) -> str:
    if value in (None, ""):
        return "-"
    try:
        number = float(value)
    except (TypeError, ValueError):
        return str(value)
    if not math.isfinite(number):
        return "-"
    return f"{number:.{digits}f}"


def pct(value: Any, digits: int = 1) -> str:
    try:
        return f"{float(value) * 100:.{digits}f}%"
    except (TypeError, ValueError):
        return "-"


def nested(payload: dict[str, Any], *keys: str) -> Any:
    value: Any = payload
    for key in keys:
        if not isinstance(value, dict):
            return None
        value = value.get(key)
    return value


def mean_metric(payload: dict[str, Any], key: str) -> Any:
    value = payload.get(key)
    if isinstance(value, dict):
        return value.get("mean")
    return None


def status_text(counts: dict[str, Any] | None) -> str:
    if not counts:
        return "-"
    return ", ".join(f"{k}={v}" for k, v in counts.items())


def set_cn_font(run, size: int = 11, bold: bool = False) -> None:
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(size)
    run.bold = bold


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, size: int = 9) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_cn_font(run, size=size, bold=bold)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = "Heading 1" if level == 1 else "Heading 2"
    run = paragraph.add_run(text)
    set_cn_font(run, size=14 if level == 1 else 12, bold=True)


def add_paragraph(doc: Document, text: str, *, center: bool = False, bold: bool = False, size: int = 11) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_cn_font(run, size=size, bold=bold)


def add_bullets(doc: Document, lines: list[str]) -> None:
    for line in lines:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(line)
        set_cn_font(run)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        set_cell_text(table.cell(0, idx), header, bold=True)
        shade_cell(table.cell(0, idx), "D9EAF7")
    for values in rows:
        row = table.add_row().cells
        for idx, value in enumerate(values):
            set_cell_text(row[idx], value)


def add_picture(doc: Document, path: Path, caption: str, *, width_inch: float = 5.9) -> None:
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(width_inch))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(caption)
    set_cn_font(run, size=9)


def save_bar(path: Path, title: str, labels: list[str], values: list[float], ylabel: str) -> Path:
    ensure_dir(path.parent)
    plt.figure(figsize=(7.4, 4.1))
    plt.bar(labels, values, color=["#2f6f9f", "#45a778", "#d99b36", "#b85c5c", "#6c6c96", "#7b9e87"][: len(values)])
    plt.title(title)
    plt.ylabel(ylabel)
    plt.xticks(rotation=18, ha="right")
    plt.grid(axis="y", alpha=0.25)
    plt.tight_layout()
    plt.savefig(path, dpi=180)
    plt.close()
    return path


def save_grouped(path: Path, title: str, groups: list[str], series: dict[str, list[float]], ylabel: str) -> Path:
    ensure_dir(path.parent)
    plt.figure(figsize=(7.5, 4.2))
    x_positions = list(range(len(groups)))
    width = 0.8 / max(len(series), 1)
    for idx, (label, values) in enumerate(series.items()):
        offsets = [x + (idx - (len(series) - 1) / 2) * width for x in x_positions]
        plt.bar(offsets, values, width=width, label=label)
    plt.title(title)
    plt.ylabel(ylabel)
    plt.xticks(x_positions, groups, rotation=12, ha="right")
    plt.grid(axis="y", alpha=0.25)
    plt.legend()
    plt.tight_layout()
    plt.savefig(path, dpi=180)
    plt.close()
    return path


def copy_report_image(src: Path, assets_dir: Path, name: str, missing: list[str]) -> Path | None:
    if not src.exists():
        missing.append(str(src))
        return None
    dst = assets_dir / name
    if Image is not None and src.stat().st_size > 2 * 1024 * 1024:
        dst = dst.with_suffix(".jpg")
        with Image.open(src) as image:
            image = image.convert("RGB")
            max_side = 1800
            if max(image.size) > max_side:
                image.thumbnail((max_side, max_side))
            ensure_dir(dst.parent)
            image.save(dst, quality=88, optimize=True)
        return dst
    ensure_dir(dst.parent)
    shutil.copy2(src, dst)
    return dst


def collect(root: Path, gate_root: Path) -> dict[str, Any]:
    pose_root = root / "pose_v1_formal"
    suite_root = pose_root / "eval_pose_validation_suite_caiwangcun_truth"
    gate_suite = gate_root / "pose_v1_formal" / "eval_pose_validation_suite_caiwangcun_truth"
    return {
        "root": root,
        "gate_root": gate_root,
        "pose_root": pose_root,
        "suite_root": suite_root,
        "acceptance": load_json(root / "plan" / "full_acceptance_summary.json"),
        "pose": load_json(pose_root / "summary" / "pose_overall_summary.json"),
        "validation": load_json(suite_root / "full_run_summary.json"),
        "ortho": load_json(suite_root / "ortho_alignment" / "overall_ortho_accuracy.json"),
        "pose_vs_at": load_json(suite_root / "pose_vs_at" / "overall_pose_vs_at.json"),
        "tiepoint": load_json(suite_root / "tiepoint_ground_error" / "overall_tiepoint_ground_error.json"),
        "frame": load_json(suite_root / "ortho_alignment" / "frame_sanity" / "overall_frame_sanity.json"),
        "failure_rows": load_csv(root / "plan" / "full_failure_buckets.csv"),
        "gate_ortho": load_json(gate_suite / "ortho_alignment" / "overall_ortho_accuracy.json"),
        "gate_pose_vs_at": load_json(gate_suite / "pose_vs_at" / "overall_pose_vs_at.json"),
        "gate_tiepoint": load_json(gate_suite / "tiepoint_ground_error" / "overall_tiepoint_ground_error.json"),
        "gate_frame": load_json(gate_suite / "ortho_alignment" / "frame_sanity" / "overall_frame_sanity.json"),
    }


def build_assets(data: dict[str, Any], assets_dir: Path) -> tuple[dict[str, Path], list[str]]:
    if assets_dir.exists():
        shutil.rmtree(assets_dir)
    ensure_dir(assets_dir)
    missing: list[str] = []
    assets: dict[str, Path] = {}
    pose = data["pose"]
    ortho = data["ortho"]
    pose_vs_at = data["pose_vs_at"]
    tiepoint = data["tiepoint"]
    frame = data["frame"]
    frame_numeric = frame.get("numeric_summaries", {})
    gate_frame_numeric = data["gate_frame"].get("numeric_summaries", {})

    assets["pose_status"] = save_bar(
        assets_dir / "pose_best_status.png",
        "Full Run Best Pose Status",
        list(pose.get("best_status_counts", {}).keys()),
        [float(v) for v in pose.get("best_status_counts", {}).values()],
        "query count",
    )
    assets["validation_metrics"] = save_bar(
        assets_dir / "full_validation_metrics.png",
        "Full Run Validation Key Metrics",
        ["center_offset_m", "horizontal_error_m", "tiepoint_rmse_m"],
        [
            float(mean_metric(ortho, "center_offset_m") or 0.0),
            float(mean_metric(pose_vs_at, "horizontal_error_m") or 0.0),
            float(tiepoint.get("tiepoint_xy_error_rmse_m") or 0.0),
        ],
        "meter",
    )
    assets["gate_full_compare"] = save_grouped(
        assets_dir / "gate_vs_full_metrics.png",
        "Gate vs Full Key Metrics",
        ["center_offset_m", "horizontal_error_m", "tiepoint_rmse_m", "pred_valid_ratio"],
        {
            "gate": [
                float(mean_metric(data["gate_ortho"], "center_offset_m") or 0.0),
                float(mean_metric(data["gate_pose_vs_at"], "horizontal_error_m") or 0.0),
                float(data["gate_tiepoint"].get("tiepoint_xy_error_rmse_m") or 0.0),
                float(nested(gate_frame_numeric, "pred_valid_pixel_ratio", "mean") or 0.0),
            ],
            "full": [
                float(mean_metric(ortho, "center_offset_m") or 0.0),
                float(mean_metric(pose_vs_at, "horizontal_error_m") or 0.0),
                float(tiepoint.get("tiepoint_xy_error_rmse_m") or 0.0),
                float(nested(frame_numeric, "pred_valid_pixel_ratio", "mean") or 0.0),
            ],
        },
        "meter / ratio",
    )
    assets["frame_sanity"] = save_bar(
        assets_dir / "full_frame_sanity.png",
        "Full Run Frame Sanity",
        ["DSM valid", "Pred valid", "Camera offset", "BBox offset"],
        [
            float(nested(frame_numeric, "dsm_sample_valid_ratio_on_truth_grid", "mean") or 0.0),
            float(nested(frame_numeric, "pred_valid_pixel_ratio", "mean") or 0.0),
            float(nested(frame_numeric, "camera_center_offset_m", "mean") or 0.0),
            float(nested(frame_numeric, "bbox_center_delta_m", "mean") or 0.0),
        ],
        "ratio / meter",
    )

    suite_root = data["suite_root"]
    ortho_root = suite_root / "ortho_alignment"
    sample_queries = pick_visual_queries(data)
    for query_id in sample_queries:
        for key_suffix, src_suffix in [
            ("truth_overlay", ("viz_overlay_truth", f"{query_id}_overlay.png")),
            ("dom_overlay", ("viz_overlay_dom", f"{query_id}_pred_vs_dom_overlay.png")),
            ("tiepoints_overlay", ("tiepoint_ground_error/viz_tiepoints", f"{query_id}_tiepoints_overlay.png")),
        ]:
            base = ortho_root if not src_suffix[0].startswith("tiepoint") else suite_root
            src = base / src_suffix[0] / src_suffix[1]
            copied = copy_report_image(src, assets_dir, f"{query_id}_{key_suffix}.png", missing)
            if copied:
                assets[f"{query_id}_{key_suffix}"] = copied
    if sample_queries:
        qid = sample_queries[0]
        for name in [f"{qid}_frame_overlay.png", f"{qid}_offset_vectors.png", f"{qid}_dsm_valid_mask_on_truth_grid.png"]:
            copied = copy_report_image(ortho_root / "frame_sanity" / "figures" / name, assets_dir, name, missing)
            if copied:
                assets[name] = copied
    (assets_dir / "missing_images.json").write_text(
        json.dumps({"missing_images": missing}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return assets, missing


def pick_visual_queries(data: dict[str, Any]) -> list[str]:
    rows = []
    path = data["suite_root"] / "ortho_alignment" / "per_query_ortho_accuracy.csv"
    if path.exists():
        rows = load_csv(path)
    ok_rows = [row for row in rows if row.get("eval_status") == "ok"]
    if not ok_rows:
        return []
    best = max(ok_rows, key=lambda row: float(row.get("ortho_iou", "0") or 0))
    worst = max(ok_rows, key=lambda row: float(row.get("center_offset_m", "0") or 0))
    middle = ok_rows[len(ok_rows) // 2]
    selected: list[str] = []
    for row in (best, middle, worst):
        qid = row["query_id"]
        if qid not in selected:
            selected.append(qid)
    return selected


def rel(path: Path, base: Path) -> str:
    try:
        return path.relative_to(base).as_posix()
    except ValueError:
        return path.as_posix()


def build_markdown(data: dict[str, Any], assets: dict[str, Path], out_md: Path, missing: list[str]) -> str:
    pose = data["pose"]
    validation = data["validation"]
    ortho = data["ortho"]
    pose_vs_at = data["pose_vs_at"]
    tiepoint = data["tiepoint"]
    frame = data["frame"]
    frame_numeric = frame.get("numeric_summaries", {})

    def img(key: str, caption: str) -> str:
        path = assets.get(key)
        return f"\n![{caption}]({rel(path, out_md.parent)})\n\n*{caption}*\n" if path else ""

    rows = [
        ["query_count", str(pose.get("query_count"))],
        ["scored_query_count", str(pose.get("scored_query_count"))],
        ["best_status_counts", status_text(pose.get("best_status_counts"))],
        ["score_status_counts", status_text(pose.get("score_status_counts"))],
        ["best_score_mean", fmt(pose.get("best_score_mean"), 4)],
        ["best_success_inlier_ratio_mean", fmt(pose.get("best_success_inlier_ratio_mean"), 4)],
        ["best_success_reproj_error_mean", fmt(pose.get("best_success_reproj_error_mean"), 4)],
    ]
    validation_rows = [
        ["pipeline_status", str(validation.get("pipeline_status"))],
        ["Layer-1 center_offset_m mean", f"{fmt(mean_metric(ortho, 'center_offset_m'), 3)} m"],
        ["Layer-1 ortho_iou mean", fmt(mean_metric(ortho, "ortho_iou"), 4)],
        ["Layer-2 horizontal_error_m mean", f"{fmt(mean_metric(pose_vs_at, 'horizontal_error_m'), 3)} m"],
        ["Layer-2 view_dir_angle_error_deg mean", f"{fmt(mean_metric(pose_vs_at, 'view_dir_angle_error_deg'), 3)} deg"],
        ["Layer-3 tiepoint RMSE", f"{fmt(tiepoint.get('tiepoint_xy_error_rmse_m'), 3)} m"],
        ["Frame DSM valid ratio mean", pct(nested(frame_numeric, "dsm_sample_valid_ratio_on_truth_grid", "mean"), 2)],
        ["Frame pred valid pixel ratio mean", pct(nested(frame_numeric, "pred_valid_pixel_ratio", "mean"), 2)],
    ]
    lines = [
        "# CaiWangCun DOM/DSM 完整替换 Full Run 实验报告",
        "",
        "## 1. 实验目的",
        "本次 full run 将 gate 已验证的 CaiWangCun DOM/DSM 完整替换口径扩展到全部 40 张 009/010 query，验证该路线在全量 query 上是否仍保持米级定位误差和正常 predicted ortho 覆盖。",
        "",
        "## 2. 实验流程",
        "Full run 使用独立目录，不覆盖 gate。CaiWangCun mosaic、candidate library、candidate features、FAISS、retrieval 和 RoMa rerank 从 gate 复制并做路径重写与审计；formal input、DSM cache、pose manifest、full pose、full validation 和 frame sanity 在 full root 下重新生成。",
        "",
        "## 3. Full Pose 结果",
        "|指标|数值|",
        "|---|---:|",
    ]
    lines += [f"|{k}|{v}|" for k, v in rows]
    lines += [img("pose_status", "图 1. Full run best pose status。")]
    lines += ["", "## 4. Full Validation 结果", "|指标|数值|", "|---|---:|"]
    lines += [f"|{k}|{v}|" for k, v in validation_rows]
    lines += [
        img("validation_metrics", "图 2. Full run Layer-1/2/3 几何指标。"),
        img("frame_sanity", "图 3. Full run frame sanity 指标。"),
        img("gate_full_compare", "图 4. Gate 与 Full 关键指标对比。"),
        "",
        "## 5. 代表性可视化",
    ]
    for key, path in assets.items():
        if key.endswith("_truth_overlay"):
            qid = key.replace("_truth_overlay", "")
            lines += [
                img(key, f"{qid} predicted ortho 与 truth ortho overlay。"),
                img(f"{qid}_dom_overlay", f"{qid} predicted ortho 与 CaiWangCun DOM overlay。"),
                img(f"{qid}_tiepoints_overlay", f"{qid} tiepoint overlay。"),
            ]
    lines += [
        "",
        "## 6. 结论",
        "- 若 full run 的 center_offset_m 和 horizontal_error_m 仍保持米级，则 CaiWangCun 完整替换路线可扩展到全量 query。",
        "- 异常 query 应按 retrieval rank、RoMa inlier ratio、candidate tile scale、DSM footprint 与 frame sanity 分桶定位原因。",
        "- 本轮继续保持 no ODM LAZ / no SRTM / no old satellite candidate fallback 的约束。",
        "",
        "## 附录",
        f"- Full root: `{data['root']}`",
        f"- Gate root: `{data['gate_root']}`",
    ]
    if missing:
        lines += ["", "### 缺失图片"] + [f"- `{item}`" for item in missing]
    return "\n".join(lines)


def build_docx(data: dict[str, Any], assets: dict[str, Path], out_docx: Path, missing: list[str]) -> None:
    pose = data["pose"]
    validation = data["validation"]
    ortho = data["ortho"]
    pose_vs_at = data["pose_vs_at"]
    tiepoint = data["tiepoint"]
    frame = data["frame"]
    frame_numeric = frame.get("numeric_summaries", {})
    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CaiWangCun DOM/DSM 完整替换 Full Run 实验报告")
    set_cn_font(run, size=18, bold=True)
    add_paragraph(doc, "40-query full run", center=True)
    add_heading(doc, "1. 实验目的")
    add_paragraph(doc, "将 gate 已验证的 CaiWangCun DOM/DSM 完整替换口径扩展到全部 40 张 009/010 query，验证 predicted ortho、pose 和 tiepoint 质量是否保持稳定。")
    add_heading(doc, "2. 实验流程")
    add_bullets(
        doc,
        [
            "Full root 独立，不覆盖 gate root。",
            "候选库相关资产从 gate 复制并路径重写，formal input、DSM cache、pose、validation 在 full root 下重新生成。",
            "不使用 ODM LAZ、SRTM 或旧 satellite candidate library fallback。",
        ],
    )
    add_heading(doc, "3. Full Pose 结果")
    add_table(
        doc,
        ["指标", "数值"],
        [
            ["query_count", str(pose.get("query_count"))],
            ["scored_query_count", str(pose.get("scored_query_count"))],
            ["best_status_counts", status_text(pose.get("best_status_counts"))],
            ["score_status_counts", status_text(pose.get("score_status_counts"))],
            ["best_score_mean", fmt(pose.get("best_score_mean"), 4)],
            ["best_success_inlier_ratio_mean", fmt(pose.get("best_success_inlier_ratio_mean"), 4)],
            ["best_success_reproj_error_mean", fmt(pose.get("best_success_reproj_error_mean"), 4)],
        ],
    )
    add_picture(doc, assets["pose_status"], "图 1. Full run best pose status。")
    add_heading(doc, "4. Full Validation 结果")
    add_table(
        doc,
        ["指标", "数值"],
        [
            ["pipeline_status", str(validation.get("pipeline_status"))],
            ["Layer-1 center_offset_m mean", f"{fmt(mean_metric(ortho, 'center_offset_m'), 3)} m"],
            ["Layer-1 ortho_iou mean", fmt(mean_metric(ortho, "ortho_iou"), 4)],
            ["Layer-2 horizontal_error_m mean", f"{fmt(mean_metric(pose_vs_at, 'horizontal_error_m'), 3)} m"],
            ["Layer-2 view_dir_angle_error_deg mean", f"{fmt(mean_metric(pose_vs_at, 'view_dir_angle_error_deg'), 3)} deg"],
            ["Layer-3 tiepoint RMSE", f"{fmt(tiepoint.get('tiepoint_xy_error_rmse_m'), 3)} m"],
            ["Frame DSM valid ratio mean", pct(nested(frame_numeric, "dsm_sample_valid_ratio_on_truth_grid", "mean"), 2)],
            ["Frame pred valid pixel ratio mean", pct(nested(frame_numeric, "pred_valid_pixel_ratio", "mean"), 2)],
        ],
    )
    add_picture(doc, assets["validation_metrics"], "图 2. Full run Layer-1/2/3 几何指标。")
    add_picture(doc, assets["frame_sanity"], "图 3. Full run frame sanity 指标。")
    add_picture(doc, assets["gate_full_compare"], "图 4. Gate 与 Full 关键指标对比。")
    add_heading(doc, "5. 代表性可视化")
    shown: set[str] = set()
    for key in list(assets):
        if key.endswith("_truth_overlay"):
            qid = key.replace("_truth_overlay", "")
            if qid in shown:
                continue
            shown.add(qid)
            add_picture(doc, assets[key], f"{qid} predicted ortho 与 truth ortho overlay。")
            if f"{qid}_dom_overlay" in assets:
                add_picture(doc, assets[f"{qid}_dom_overlay"], f"{qid} predicted ortho 与 CaiWangCun DOM overlay。")
            if f"{qid}_tiepoints_overlay" in assets:
                add_picture(doc, assets[f"{qid}_tiepoints_overlay"], f"{qid} tiepoint overlay。")
    add_heading(doc, "6. 结论")
    add_bullets(
        doc,
        [
            "Full run 用于确认 gate 的完整替换口径是否可扩展到 40/40 query。",
            "异常 query 后续按 retrieval、RoMa、DSM footprint 和 frame sanity 分桶诊断。",
            "本轮继续保持 no ODM LAZ / no SRTM / no old satellite candidate fallback。",
        ],
    )
    if missing:
        add_heading(doc, "缺失图片记录", level=2)
        add_bullets(doc, missing)
    ensure_dir(out_docx.parent)
    doc.save(str(out_docx))


def main() -> None:
    args = parse_args()
    root = Path(args.experiment_root).resolve()
    gate_root = Path(args.gate_root).resolve()
    reports_dir = root / "reports"
    out_md = Path(args.out_md).resolve() if args.out_md else reports_dir / "caiwangcun_fullreplace_full_report.md"
    out_docx = Path(args.out_docx).resolve() if args.out_docx else reports_dir / "caiwangcun_fullreplace_full_report.docx"
    assets_dir = Path(args.assets_dir).resolve() if args.assets_dir else reports_dir / "assets"
    data = collect(root, gate_root)
    assets, missing = build_assets(data, assets_dir)
    ensure_dir(out_md.parent)
    out_md.write_text(build_markdown(data, assets, out_md, missing), encoding="utf-8")
    build_docx(data, assets, out_docx, missing)
    print(out_md)
    print(out_docx)


if __name__ == "__main__":
    main()
